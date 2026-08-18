from io import BytesIO

from fastapi.testclient import TestClient
import pytest
from docx import Document
from pypdf import PdfWriter

from backend.app.main import app


@pytest.fixture
def client():
    with TestClient(app) as test_client:
        yield test_client


def test_teacher_knowledge_graph_crud_and_persistence(client):
    created = client.post(
        "/api/teacher/knowledge-graphs",
        json={"title": "算法课程图谱", "description": "课程关系草稿", "target_classes": ["软件 1 班"]},
    )
    assert created.status_code == 201
    graph = created.json()["data"]
    assert graph["status"] == "draft"
    assert graph["nodes"][0]["label"] == "核心知识点"
    graph_id = graph["id"]

    extra_node = {
        "id": "node-custom-test",
        "label": "排序算法",
        "type": "方法",
        "description": "比较常见排序策略",
        "difficulty": 3,
        "x": 650,
        "y": 270,
        "color": "#0f766e",
        "source": "custom",
    }
    nodes = graph["nodes"] + [extra_node]
    edge = {
        "id": "edge-custom-test",
        "source": graph["nodes"][0]["id"],
        "target": extra_node["id"],
        "type": "前驱",
        "label": "前驱",
    }
    saved = client.put(
        f"/api/teacher/knowledge-graphs/{graph_id}",
        json={
            "title": "算法课程图谱",
            "description": "已编辑",
            "target_classes": ["软件 1 班", "软件 2 班"],
            "status": "draft",
            "nodes": nodes,
            "edges": [edge],
        },
    )
    assert saved.status_code == 200
    assert saved.json()["data"]["edges"] == [edge]

    published = client.post(f"/api/teacher/knowledge-graphs/{graph_id}/publish")
    assert published.status_code == 200
    assert published.json()["data"]["status"] == "published"

    listing = client.get("/api/teacher/knowledge-graphs")
    assert listing.status_code == 200
    row = next(item for item in listing.json()["data"] if item["id"] == graph_id)
    assert row["node_count"] == 2
    assert row["edge_count"] == 1
    assert "nodes" not in row

    restored = client.get(f"/api/teacher/knowledge-graphs/{graph_id}").json()["data"]
    assert restored["nodes"][1]["x"] == 650
    assert restored["edges"][0]["id"] == "edge-custom-test"

    deleted = client.delete(f"/api/teacher/knowledge-graphs/{graph_id}")
    assert deleted.status_code == 200
    assert client.get(f"/api/teacher/knowledge-graphs/{graph_id}").status_code == 404


def test_teacher_knowledge_graph_from_text_files_uses_fallback(client):
    response = client.post(
        "/api/teacher/knowledge-graphs/from-files",
        data={
            "title": "数据结构资料图谱",
            "description": "从讲义生成",
            "target_classes": "Python A 班，Python B 班\n实验班",
        },
        files=[
            ("files", ("outline.md", BytesIO("# 线性表\n## 栈与队列\n排序算法\n图的遍历\n递归方法".encode()), "text/markdown")),
            ("files", ("notes.txt", BytesIO("查找方法\n算法复杂度\n综合项目实践".encode()), "text/plain")),
        ],
    )
    assert response.status_code == 201
    graph = response.json()["data"]
    assert graph["status"] == "draft"
    assert graph["target_classes"] == ["Python A 班", "Python B 班", "实验班"]
    assert len(graph["source_files"]) == 2
    assert 6 <= graph["node_count"] <= 14
    assert graph["edge_count"] == graph["node_count"] - 1
    assert graph["source_summary"]


def test_teacher_knowledge_graph_validation_and_permissions(client):
    unsupported = client.post(
        "/api/teacher/knowledge-graphs/from-files",
        data={"title": "错误格式"},
        files={"files": ("archive.zip", b"not-a-document", "application/zip")},
    )
    assert unsupported.status_code == 415

    denied = client.get("/api/teacher/knowledge-graphs", headers={"X-User-Id": "student-03"})
    assert denied.status_code == 403


def test_teacher_knowledge_graph_accepts_docx_and_pdf(client):
    document_buffer = BytesIO()
    document = Document()
    document.add_heading("动态规划方法", level=1)
    document.add_paragraph("状态转移公式、递推步骤、综合案例实践和算法能力目标。")
    document.save(document_buffer)

    pdf_buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    writer.write(pdf_buffer)

    response = client.post(
        "/api/teacher/knowledge-graphs/from-files",
        data={"title": "Word PDF 格式验收"},
        files=[
            ("files", ("lesson.docx", document_buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
            ("files", ("appendix.pdf", pdf_buffer.getvalue(), "application/pdf")),
        ],
    )
    assert response.status_code == 201, response.text
    graph = response.json()["data"]
    assert [item["filename"] for item in graph["source_files"]] == ["lesson.docx", "appendix.pdf"]
    assert graph["node_count"] >= 6
    assert client.delete(f"/api/teacher/knowledge-graphs/{graph['id']}").status_code == 200
