from __future__ import annotations

from collections import Counter
from datetime import datetime
import hashlib
import json
from pathlib import Path
import sqlite3

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\teacher")
DB_PATH = ROOT / "codetrack.db"
OUTPUT_DIR = ROOT / "artifacts" / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "CodeTrack教师端数据库内容与前端接入说明.docx"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Preset: compact_reference_guide. Named overrides:
# 1. East Asian fallback uses Microsoft YaHei; ASCII remains Calibri.
# 2. CodeTrack green replaces preset heading blue consistently.
# 3. Dense schema tables use 8.2 pt text while preserving preset geometry.
GREEN = "2AAD16"
DARK_GREEN = "176B12"
INK = "1F2A24"
MUTED = "5F6B63"
LIGHT_GREEN = "EDF8EA"
LIGHT_GRAY = "F4F6F5"
TABLE_HEADER = "E8EEF5"
BORDER = "D8E0DA"
CAUTION = "FFF5D9"
RISK = "FCE8E6"
WHITE = "FFFFFF"

TABLE_PURPOSES = {
    "users": "教师、学生等平台用户主档案",
    "teacher_credentials": "教师登录凭据（盐值与密码哈希）",
    "teacher_preferences": "教师通知、AI 助教和邮件摘要偏好",
    "courses": "课程主数据及教师归属",
    "course_drafts": "每位教师一份课程创建草稿",
    "course_announcements": "课程公告正文、范围和发布时间",
    "announcement_reads": "用户维度的课程公告已读状态",
    "class_groups": "授课班级、邀请码与课表信息",
    "enrollments": "学生与班级的加入关系",
    "chapters": "课程章节、顺序与发布状态",
    "knowledge_points": "课程知识点及掌握度元数据",
    "materials": "教学材料、可见性和回收站状态",
    "material_folders": "教学材料目录及回收站层级",
    "material_knowledge_links": "材料与知识点的多对多关联",
    "teacher_knowledge_graphs": "教师端知识图谱快照和发布状态",
    "tasks": "作业、练习和编程任务主数据",
    "test_cases": "任务公开/隐藏测试用例及权重",
    "submissions": "学生提交版本与源代码记录",
    "evaluation_results": "自动评测通过数、耗时和得分",
    "diagnosis_results": "AI/规则诊断结果及置信度",
    "diagnosis_reviews": "教师对诊断结果的审核记录",
    "grades": "教师评分、维度分和发布状态",
    "teacher_feedback": "教师反馈内容及学生可见状态",
    "course_discussions": "课堂讨论主题、状态和参与统计",
    "discussion_replies": "学生对课堂讨论的回复",
    "notifications": "教师站内通知与已读状态",
    "audit_logs": "登录和关键写操作审计日志",
}

FIELD_DESCRIPTIONS = {
    "id": "主键标识",
    "user_id": "用户标识",
    "teacher_id": "教师标识",
    "student_id": "学生标识",
    "course_id": "课程标识",
    "class_id": "班级标识",
    "task_id": "任务标识",
    "chapter_id": "章节标识",
    "material_id": "材料标识",
    "folder_id": "目录标识",
    "knowledge_point_id": "知识点标识",
    "announcement_id": "课程公告标识",
    "discussion_id": "讨论主题标识",
    "submission_id": "提交记录标识",
    "diagnosis_id": "诊断记录标识",
    "actor_id": "操作人标识",
    "author_id": "作者标识",
    "name": "名称",
    "title": "标题",
    "number": "教师编号或学号",
    "email": "邮箱地址",
    "department": "所属院系",
    "role": "用户角色",
    "code": "业务代码",
    "term": "开课学期",
    "description": "说明或简介",
    "status": "业务状态",
    "position": "排序位置",
    "progress": "进度百分比",
    "student_visible": "是否对学生可见",
    "visibility": "可见范围",
    "created_at": "创建时间",
    "updated_at": "更新时间",
    "saved_at": "草稿保存时间",
    "published_at": "发布时间",
    "read_at": "阅读时间",
    "due_at": "截止时间",
    "payload_json": "课程草稿 JSON 载荷",
    "content_json": "结构化正文 JSON",
    "password_salt": "随机密码盐（敏感，不在文档展示值）",
    "password_hash": "PBKDF2 密码哈希（敏感，不在文档展示值）",
    "notifications_enabled": "是否启用站内通知",
    "ai_assistant_enabled": "是否启用 AI 助教",
    "email_digest": "是否启用邮件摘要",
    "summary": "摘要",
    "content": "正文内容",
    "audience": "接收范围",
    "pinned": "是否置顶",
    "read": "是否已读",
    "join_code": "班级邀请码",
    "schedule": "上课安排",
    "mentor": "授课教师名称",
    "grade": "年级",
    "major": "专业",
    "joined_at": "加入时间",
    "join_status": "加入状态",
    "join_method": "加入方式",
    "teaching_mode": "教学模式",
    "difficulty": "难度",
    "mastery": "掌握度",
    "type": "业务类型",
    "size": "文件大小描述",
    "content_url": "材料访问地址",
    "deleted_at": "进入回收站时间",
    "parent_id": "上级目录标识",
    "graph_json": "知识图谱 JSON 快照",
    "target_classes": "图谱目标班级",
    "starter_code": "任务起始代码",
    "publish_at": "计划发布时间",
    "total_score": "任务总分",
    "hidden": "是否为隐藏测试用例",
    "weight": "测试用例权重",
    "source_code": "学生源代码（敏感正文不在文档展示）",
    "version": "提交版本号",
    "hint_level": "最高提示等级",
    "submitted_at": "提交时间",
    "passed_tests": "通过测试数",
    "total_tests": "测试总数",
    "runtime_ms": "运行耗时（毫秒）",
    "score": "得分",
    "details_json": "评测明细 JSON",
    "explanation": "诊断解释",
    "confidence": "置信度",
    "source": "诊断来源",
    "fallback": "是否为规则兜底",
    "needs_teacher_review": "是否需要教师审核",
    "review_status": "审核状态",
    "reviewed_explanation": "教师修订解释",
    "comment": "教师评语",
    "dimensions_json": "评分维度 JSON",
    "student_visible": "是否对学生可见",
    "participant_count": "参与人数",
    "reply_count": "回复数量",
    "action": "审计动作",
    "resource_type": "资源类型",
    "resource_id": "资源标识",
    "detail": "审计补充说明",
}

GROUPS = [
    ("账号与个性化", ["users", "teacher_credentials", "teacher_preferences", "notifications", "audit_logs"]),
    ("课程、班级与章节", ["courses", "course_drafts", "course_announcements", "announcement_reads", "class_groups", "enrollments", "chapters", "knowledge_points"]),
    ("材料与知识图谱", ["materials", "material_folders", "material_knowledge_links", "teacher_knowledge_graphs"]),
    ("任务、提交与评价", ["tasks", "test_cases", "submissions", "evaluation_results", "diagnosis_results", "diagnosis_reviews", "grades", "teacher_feedback"]),
    ("课堂讨论", ["course_discussions", "discussion_replies"]),
]


def set_run_font(run, size=None, bold=None, color=None, italic=None, ascii_font="Calibri"):
    run.font.name = ascii_font
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    tr_pr.append(node)


def set_row_cant_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell(cell, bold=False, color=INK, size=8.8, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths_dxa, font_size=8.8, header_fill=TABLE_HEADER, aligns=None, keep_together=False):
    table = doc.add_table(rows=1, cols=len(headers))
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(value)
        set_cell_shading(cell, header_fill)
        style_cell(cell, bold=True, size=font_size, align=aligns[index] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            style_cell(cells[index], size=font_size, align=aligns[index] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
    for row_index, row in enumerate(table.rows):
        set_row_cant_split(row)
        if keep_together and row_index < len(table.rows) - 1:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return table


def add_field(paragraph, instruction, placeholder="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, keep=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.keep_with_next = keep
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    for run in paragraph.runs:
        set_run_font(run, size=11, color=INK)
    if not paragraph.runs:
        set_run_font(paragraph.add_run(text), size=11, color=INK)
    return paragraph


def add_callout(doc, title, text, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    heading = cell.paragraphs[0]
    heading.paragraph_format.space_after = Pt(4)
    set_run_font(heading.add_run(title), size=10.5, bold=True, color=DARK_GREEN)
    body = cell.add_paragraph()
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.line_spacing = 1.2
    set_run_font(body.add_run(text), size=9.5, color=INK)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=GREEN, size="8")
    add_para(doc, "", after=0)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    settings = {
        "Heading 1": (16, 18, 10, GREEN),
        "Heading 2": (13, 14, 7, GREEN),
        "Heading 3": (12, 10, 5, DARK_GREEN),
    }
    for name, (size, before, after, color) in settings.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = "CodeTrack 教师端"
    table.cell(0, 1).text = "数据库内容与前端接入说明"
    set_table_geometry(table, [3000, 6360], indent_dxa=0)
    for index, cell in enumerate(table.rows[0].cells):
        style_cell(cell, bold=index == 0, color=GREEN if index == 0 else MUTED, size=8.5,
                   align=WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT)
    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = "数据库快照：2026-08-14"
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.add_run("第 "), size=8.5, color=MUTED)
    add_field(right, "PAGE")
    set_run_font(right.add_run(" 页 / 共 "), size=8.5, color=MUTED)
    add_field(right, "NUMPAGES")
    set_run_font(right.add_run(" 页"), size=8.5, color=MUTED)
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    style_cell(table.cell(0, 0), size=8.5, color=MUTED)


def inspect_database(connection):
    tables = [row[0] for row in connection.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name"
    )]
    schemas = {}
    counts = {}
    for table in tables:
        safe = table.replace('"', '""')
        counts[table] = connection.execute(f'SELECT COUNT(*) FROM "{safe}"').fetchone()[0]
        columns = connection.execute(f'PRAGMA table_info("{safe}")').fetchall()
        foreign_keys = connection.execute(f'PRAGMA foreign_key_list("{safe}")').fetchall()
        fk_by_column = {row[3]: f"FK→{row[2]}.{row[4]}" for row in foreign_keys}
        unique_columns = set()
        for index_row in connection.execute(f'PRAGMA index_list("{safe}")').fetchall():
            if index_row[2]:
                index_name = index_row[1].replace('"', '""')
                fields = connection.execute(f'PRAGMA index_info("{index_name}")').fetchall()
                if len(fields) == 1:
                    unique_columns.add(fields[0][2])
        schema_rows = []
        for column in columns:
            cid, name, data_type, not_null, default, primary_key = column
            constraints = []
            if primary_key:
                constraints.append("PK")
            if name in unique_columns:
                constraints.append("UNIQUE")
            if not_null:
                constraints.append("NOT NULL")
            if default is not None:
                constraints.append(f"DEFAULT {default}")
            if name in fk_by_column:
                constraints.append(fk_by_column[name])
            schema_rows.append((name, data_type or "未声明", "；".join(constraints) or "—", FIELD_DESCRIPTIONS.get(name, "业务字段")))
        schemas[table] = schema_rows
    return tables, counts, schemas


def query_rows(connection, sql, params=()):
    return [tuple("—" if value is None else value for value in row) for row in connection.execute(sql, params).fetchall()]


def add_cover(doc, meta):
    add_para(doc, "CODETRACK / DATABASE REFERENCE", size=10, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, before=74, after=18)
    add_para(doc, "CodeTrack 教师端", size=17, bold=True, color=DARK_GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "数据库内容与前端接入说明", size=29, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "表结构 · 数据快照 · 页面/API 映射 · 运维与验证", size=14, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=44)
    add_table(doc, ["文档属性", "内容"], [
        ("数据库", str(DB_PATH)),
        ("数据库引擎", f"SQLite {meta['sqlite_version']}"),
        ("快照时间", meta["snapshot_time"]),
        ("文件大小", f"{meta['file_size']:,} 字节（{meta['file_size'] / 1024:.1f} KiB）"),
        ("数据表", f"{meta['table_count']} 张"),
        ("用途", "开发、测试、交接、验收与数据库维护"),
    ], [1900, 7460], font_size=9.5, header_fill=LIGHT_GREEN)
    add_para(doc, "本说明根据当前数据库自动生成。凭据盐值、密码哈希、学生源代码和大段正文不会写入文档。", size=9.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=28, after=0)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_document(doc, connection, tables, counts, schemas, meta):
    add_cover(doc, meta)

    doc.add_heading("阅读导航", level=1)
    add_table(doc, ["章节", "解决的问题"], [
        ("1. 接入结论与边界", "哪些前端数据已经进入数据库，哪些仍属于临时 UI 状态"),
        ("2. 数据库总览", "当前有哪些表、各有多少行、负责什么业务"),
        ("3. 新增持久化设计", "登录、草稿、公告已读、偏好如何落库"),
        ("4. 页面/API/数据表映射", "前端按钮最终调用哪个接口、读写哪些表"),
        ("5. 全量字段字典", "27 张表的字段、类型、约束和含义"),
        ("6. 当前数据内容", "教师、课程、公告和关键状态的代表性快照"),
        ("7. 运维、安全与验证", "启动、迁移、备份、隐私和测试结果"),
    ], [2100, 7260], font_size=9.3)

    doc.add_heading("1. 接入结论与边界", level=1)
    add_callout(doc, "接入结果", "教师登录凭据、课程创建草稿、课程公告、公告已读状态、教师通知/AI/邮件偏好均已由 FastAPI + SQLAlchemy 写入 SQLite；教师姓名、编号、院系和邮箱由 users 表返回。")
    doc.add_heading("1.1 数据流", level=2)
    add_table(doc, ["层级", "组件", "职责"], [
        ("浏览器", "React + Ant Design", "收集输入、渲染数据库返回值；不保存业务主数据"),
        ("前端数据层", "src/api.ts", "统一附加 X-User-Id，请求 /api/v1 接口"),
        ("服务层", "FastAPI", "鉴权、课程归属校验、业务校验、审计"),
        ("持久化层", "SQLAlchemy", "ORM 模型、事务提交、外键关系"),
        ("数据库", "codetrack.db / SQLite", "保存课程、教学活动及教师前端状态"),
    ], [1300, 2400, 5660], font_size=9.2)
    doc.add_heading("1.2 持久化边界", level=2)
    add_table(doc, ["保存到数据库", "仍保留在浏览器"], [
        ("教师账号与加盐密码哈希", "登录会话标识（localStorage）"),
        ("课程创建草稿及保存时间", "筛选条件、分页、抽屉/弹窗开关"),
        ("课程公告与每位教师的已读状态", "当前选中行、当前查看页签"),
        ("站内通知、AI 助教、邮件摘要偏好", "跨页面临时焦点（sessionStorage）"),
        ("教师档案、课程、班级、任务、成绩等业务数据", "URL 中的 discussion=1 等导航状态"),
    ], [4680, 4680], font_size=9.2)

    doc.add_heading("2. 数据库总览", level=1)
    add_para(doc, f"当前文件包含 {len(tables)} 张业务表，共 {sum(counts.values()):,} 行记录。行数为文档生成时快照，后续使用系统后会继续变化。", keep=True)
    inventory_rows = [(table, counts[table], TABLE_PURPOSES.get(table, "项目业务表")) for table in tables]
    add_table(doc, ["表名", "行数", "用途"], inventory_rows, [2800, 1000, 5560], font_size=8.6,
              aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT])

    doc.add_heading("3. 新增持久化设计", level=1)
    add_table(doc, ["新增表", "核心内容", "关键规则"], [
        ("teacher_credentials", "user_id、salt、hash、更新时间", "PBKDF2-HMAC-SHA256；不存明文密码"),
        ("teacher_preferences", "通知、AI、邮件摘要开关", "teacher_id 主键；每位教师一行"),
        ("course_drafts", "课程创建表单 JSON、保存时间", "teacher_id 唯一；每位教师最多一份"),
        ("course_announcements", "标题、摘要、正文 JSON、范围、置顶", "按课程归属校验访问"),
        ("announcement_reads", "公告、用户、阅读时间", "公告 + 用户唯一，重复点击幂等"),
    ], [2450, 3600, 3310], font_size=9.0)
    doc.add_heading("3.1 登录与凭据", level=2)
    add_bullet(doc, "登录接口同时接受教师姓名或教师编号；验证成功后返回 users 表中的教师档案。")
    add_bullet(doc, "演示初始化密码为 123456；数据库只保存 16 字节随机盐和 120,000 次 PBKDF2-HMAC-SHA256 结果。")
    add_bullet(doc, "成功登录写入 audit_logs，失败登录返回 HTTP 401，不回显凭据细节。")
    doc.add_heading("3.2 草稿、公告和偏好", level=2)
    add_bullet(doc, "课程草稿使用 PUT 覆盖保存；创建课程成功后 DELETE 清理草稿。")
    add_bullet(doc, "公告正文从 course_announcements 读取；打开公告后 PATCH 写入 announcement_reads。")
    add_bullet(doc, "设置页先 GET 读取 teacher_preferences，再通过 PUT 原子保存三个布尔偏好。")

    doc.add_heading("4. 页面 / API / 数据表映射", level=1)
    mapping_rows = [
        ("登录页 / 快速账号卡", "GET /teacher/auth/accounts", "users", "读取"),
        ("登录按钮", "POST /teacher/auth/login", "users、teacher_credentials、audit_logs", "读取 + 新增审计"),
        ("工作台初始化", "GET /teacher/bootstrap", "users、courses、class_groups、notifications", "读取"),
        ("创建课程 / 保存草稿", "PUT /teacher/course-draft", "course_drafts、audit_logs", "新增或更新"),
        ("创建课程 / 恢复草稿", "GET /teacher/course-draft", "course_drafts", "读取"),
        ("课程列表 / 删除草稿", "DELETE /teacher/course-draft", "course_drafts、audit_logs", "删除 + 新增审计"),
        ("课程工作空间 / 公告", "GET /teacher/courses/{course_id}/announcements", "course_announcements、announcement_reads、users", "读取"),
        ("打开公告", "PATCH /teacher/announcements/{id}/read", "announcement_reads", "幂等新增"),
        ("个人设置", "GET /teacher/preferences", "teacher_preferences", "读取"),
        ("保存偏好", "PUT /teacher/preferences", "teacher_preferences、audit_logs", "更新 + 新增审计"),
        ("上传材料", "POST /teacher/materials/upload", "materials、audit_logs", "新增"),
        ("从文件生成图谱", "POST /teacher/knowledge-graphs/from-files", "teacher_knowledge_graphs", "新增"),
    ]
    add_table(doc, ["前端位置/动作", "API", "主要数据表", "操作"], mapping_rows, [2200, 3000, 3000, 1160], font_size=8.2)

    doc.add_heading("5. 全量字段字典", level=1)
    add_para(doc, "约束来自 SQLite PRAGMA 实时检查。PK=主键，FK=外键，UNIQUE=唯一约束；默认值按数据库原始表达式记录。", size=9.5, color=MUTED)
    for group_name, group_tables in GROUPS:
        doc.add_heading(group_name, level=2)
        for table in group_tables:
            if table not in schemas:
                continue
            heading = doc.add_heading(f"{table}（{counts[table]} 行）", level=3)
            heading.paragraph_format.keep_with_next = True
            add_para(doc, TABLE_PURPOSES.get(table, "项目业务表"), size=9.2, color=MUTED, after=4, keep=True)
            add_table(doc, ["字段", "类型", "约束", "说明"], schemas[table], [1900, 1500, 3300, 2660], font_size=8.0, keep_together=True)

    doc.add_heading("6. 当前数据库内容快照", level=1)
    doc.add_heading("6.1 教师账号（非敏感字段）", level=2)
    teacher_rows = query_rows(connection, "SELECT id, name, number, department, email FROM users WHERE role='teacher' ORDER BY id")
    add_table(doc, ["ID", "姓名", "编号", "院系", "邮箱"], teacher_rows, [1500, 1100, 1300, 2500, 2960], font_size=8.4)
    doc.add_heading("6.2 教师课程", level=2)
    course_rows = query_rows(connection, """
        SELECT c.id, u.name, c.name, c.code, c.term, c.status
        FROM courses c JOIN users u ON u.id = c.teacher_id
        WHERE c.teacher_id IN ('teacher-01', 'teacher-02')
        ORDER BY c.teacher_id, c.created_at, c.id LIMIT 12
    """)
    add_table(doc, ["课程ID", "教师", "课程名称", "课程代码", "学期", "状态"], course_rows, [1450, 900, 2300, 1200, 2310, 1200], font_size=8.0)
    add_para(doc, "注：课程表记录较多，此处只展示前 12 行；完整行数见第 2 章。", size=9, color=MUTED, italic=True)
    doc.add_heading("6.3 课程公告", level=2)
    announcement_rows = query_rows(connection, """
        SELECT id, course_id, title, audience, pinned, published_at
        FROM course_announcements ORDER BY pinned DESC, published_at DESC
    """)
    add_table(doc, ["公告ID", "课程", "标题", "接收范围", "置顶", "发布时间"], announcement_rows, [1700, 1000, 2700, 1800, 600, 1560], font_size=8.0)
    doc.add_heading("6.4 新增表状态", level=2)
    credential_count = counts.get("teacher_credentials", 0)
    preference_rows = query_rows(connection, """
        SELECT p.teacher_id, u.name, p.notifications_enabled, p.ai_assistant_enabled, p.email_digest, p.updated_at
        FROM teacher_preferences p JOIN users u ON u.id=p.teacher_id ORDER BY p.teacher_id
    """)
    add_callout(doc, "凭据脱敏", f"teacher_credentials 当前 {credential_count} 行。文档只确认每个教师均已配置随机盐和哈希，不展示 password_salt 或 password_hash 的值。", fill=CAUTION)
    add_table(doc, ["教师ID", "姓名", "站内通知", "AI 助教", "邮件摘要", "更新时间"], preference_rows, [1400, 1000, 1100, 1100, 1100, 3660], font_size=8.3)

    doc.add_heading("7. 数据生命周期", level=1)
    lifecycle_rows = [
        ("登录", "读取账号 → PBKDF2 校验 → 写登录审计 → 返回教师档案"),
        ("草稿", "进入创建页读取 → 点击保存执行 upsert → 刷新恢复 → 创建成功或删除草稿后清除"),
        ("公告", "按课程读取 → 合并用户已读集合 → 打开公告幂等写入 → 刷新仍为已读"),
        ("偏好", "设置页读取 → 修改三个开关 → PUT 事务保存 → 写审计 → 刷新恢复"),
    ]
    add_table(doc, ["对象", "完整流程"], lifecycle_rows, [1600, 7760], font_size=9.2)

    doc.add_heading("8. 启动、迁移与备份", level=1)
    doc.add_heading("8.1 启动", level=2)
    add_table(doc, ["服务", "命令", "地址"], [
        ("后端", "python -m uvicorn backend.app.main:app --host 127.0.0.1 --port 8001", "http://127.0.0.1:8001"),
        ("前端", "pnpm dev:frontend", "http://127.0.0.1:5173"),
    ], [1200, 5860, 2300], font_size=8.5)
    doc.add_heading("8.2 迁移与兼容升级", level=2)
    add_bullet(doc, "迁移文件：backend/alembic/versions/20260814_0006_frontend_persistence.py。")
    add_bullet(doc, "标准环境执行：alembic -c backend/alembic.ini upgrade head。")
    add_bullet(doc, "应用启动还会执行 Base.metadata.create_all 和兼容种子，旧数据库可增量补齐缺失表与演示教师数据。")
    doc.add_heading("8.3 SQLite 备份", level=2)
    add_bullet(doc, "备份前停止后端写入，或使用 SQLite 在线备份 API，避免只复制主文件而遗漏 WAL 内容。")
    add_bullet(doc, "最小备份对象：codetrack.db；如存在 codetrack.db-wal / codetrack.db-shm，应在一致性快照中一并处理。")
    add_bullet(doc, "恢复前先保留当前数据库副本，再校验表数量、外键和核心行数。")

    doc.add_heading("9. 安全与隐私", level=1)
    add_table(doc, ["控制点", "当前实现", "运维要求"], [
        ("密码", "随机盐 + PBKDF2-HMAC-SHA256，120,000 次", "生产环境强制修改演示密码并增加重置流程"),
        ("权限", "X-User-Id + 教师角色 + 课程归属校验", "生产环境应替换为可信会话/JWT，禁止客户端自报身份"),
        ("审计", "登录、草稿、偏好等关键写操作记录 audit_logs", "定期归档并限制审计表访问"),
        ("敏感正文", "源代码、哈希、盐值和大段正文不进入本文档", "备份、日志、导出均按敏感数据管理"),
        ("SQLite", "已启用 PRAGMA foreign_keys=ON", "并发或多实例部署时评估迁移 PostgreSQL"),
    ], [1500, 4200, 3660], font_size=8.7)
    add_callout(doc, "生产环境提醒", "当前 X-User-Id 适合本地演示和联调，不等同于生产认证。正式上线前应由服务端签发并验证会话令牌，登录接口应增加限流、锁定和密码重置。", fill=RISK)

    doc.add_heading("10. 验证记录", level=1)
    add_table(doc, ["验证层级", "用例", "结果"], [
        ("后端", "账号列表、正确/错误登录、草稿 PUT/GET/DELETE、公告已读、偏好 GET/PUT", "通过"),
        ("后端全量", "backend/tests 全部 21 个测试", "21 passed"),
        ("前端类型", "TypeScript project build", "通过"),
        ("前端生产构建", "Vite production build", "通过；仅有大包体积警告"),
        ("浏览器", "双账号、草稿刷新恢复、公告已读刷新保持、偏好刷新保持", "4/4 通过"),
        ("视觉", "1440 × 1000 设置页截图检查", "无重叠、裁切或空白"),
    ], [1700, 5860, 1800], font_size=8.8)

    doc.add_heading("附录 A. 本次新增接口", level=1)
    endpoint_rows = [
        ("GET", "/api/v1/teacher/auth/accounts", "读取数据库教师账号卡"),
        ("POST", "/api/v1/teacher/auth/login", "数据库验密并记录登录审计"),
        ("GET", "/api/v1/teacher/course-draft", "读取当前教师草稿"),
        ("PUT", "/api/v1/teacher/course-draft", "新增或覆盖当前教师草稿"),
        ("DELETE", "/api/v1/teacher/course-draft", "删除当前教师草稿"),
        ("GET", "/api/v1/teacher/courses/{course_id}/announcements", "读取课程公告及已读状态"),
        ("PATCH", "/api/v1/teacher/announcements/{id}/read", "标记公告已读"),
        ("GET", "/api/v1/teacher/preferences", "读取教师偏好"),
        ("PUT", "/api/v1/teacher/preferences", "保存教师偏好"),
    ]
    add_table(doc, ["方法", "路径", "说明"], endpoint_rows, [1000, 5200, 3160], font_size=8.6)

def audit_document(doc):
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert abs(section.header_distance - Inches(0.492)) < 10000
    assert abs(section.footer_distance - Inches(0.492)) < 10000
    assert doc.styles["Normal"].paragraph_format.space_after == Pt(6)
    assert doc.styles["Heading 1"].font.size == Pt(16)
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) in (9360,)
        grid_total = sum(int(item.get(qn("w:w"))) for item in table._tbl.tblGrid.findall(qn("w:gridCol")))
        assert grid_total == 9360


def main():
    if not DB_PATH.exists():
        raise FileNotFoundError(DB_PATH)
    connection = sqlite3.connect(DB_PATH)
    connection.execute("PRAGMA foreign_keys=ON")
    tables, counts, schemas = inspect_database(connection)
    meta = {
        "sqlite_version": sqlite3.sqlite_version,
        "snapshot_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "file_size": DB_PATH.stat().st_size,
        "table_count": len(tables),
        "sha256": hashlib.sha256(DB_PATH.read_bytes()).hexdigest(),
        "journal_mode": connection.execute("PRAGMA journal_mode").fetchone()[0],
        "foreign_keys": connection.execute("PRAGMA foreign_keys").fetchone()[0],
    }

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        set_header_footer(section)
    doc.core_properties.title = "CodeTrack 教师端数据库内容与前端接入说明"
    doc.core_properties.subject = "数据库结构、内容快照与前端接入映射"
    doc.core_properties.author = "CodeTrack 项目组"
    doc.core_properties.keywords = "CodeTrack, SQLite, FastAPI, React, 数据库"
    add_document(doc, connection, tables, counts, schemas, meta)
    audit_document(doc)
    doc.save(OUTPUT_PATH)
    connection.close()
    print(json.dumps({
        "output": str(OUTPUT_PATH),
        "tables": len(tables),
        "rows": sum(counts.values()),
        "size": OUTPUT_PATH.stat().st_size,
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
