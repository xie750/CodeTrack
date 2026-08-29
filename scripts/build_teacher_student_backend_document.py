from __future__ import annotations

from datetime import datetime
import json
from pathlib import Path
import sqlite3

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw

from build_database_integration_document import (
    BORDER,
    DARK_GREEN,
    GREEN,
    INK,
    LIGHT_GRAY,
    LIGHT_GREEN,
    MUTED,
    RISK,
    ROOT,
    TABLE_HEADER,
    add_callout,
    add_field,
    add_para,
    add_table,
    audit_document,
    configure_section,
    configure_styles,
    set_cell_margins,
    set_cell_shading,
    set_header_footer,
    set_run_font,
    set_table_borders,
    set_table_geometry,
    style_cell,
)
from build_database_enterprise_document import (
    add_code_block,
    add_list_item,
    load_font,
)


OUTPUT_DIR = ROOT / "artifacts" / "deliverables"
ASSET_DIR = ROOT / "artifacts" / "teacher-student-backend-assets"
OUTPUT_PATH = OUTPUT_DIR / "CodeTrack教师端与学生端协同后端接入设计说明_企业版.docx"
DB_PATH = ROOT / "teacher_backend" / "codetrack.db"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)

DOC_ID = "CT-ARCH-TS-001"
DOC_VERSION = "V1.0"
BASELINE_DATE = "2026-08-16"


def add_numbering_definition(doc, bullet=False):
    """Create schema-valid numbering and keep abstract definitions before num instances."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.extend([start, num_fmt, level_text, suffix, p_pr])
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        level.append(r_pr)
    abstract.append(level)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def set_document_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = "CodeTrack 教学协同架构"
    table.cell(0, 1).text = "教师端与学生端协同后端接入设计说明"
    set_table_geometry(table, [3200, 6160], indent_dxa=0)
    style_cell(table.cell(0, 0), bold=True, color=GREEN, size=8.5)
    style_cell(table.cell(0, 1), color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = f"{DOC_ID}  |  {DOC_VERSION}  |  内部使用"
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.add_run("第 "), size=8.5, color=MUTED)
    add_field(right, "PAGE")
    set_run_font(right.add_run(" 页 / 共 "), size=8.5, color=MUTED)
    add_field(right, "NUMPAGES")
    set_run_font(right.add_run(" 页"), size=8.5, color=MUTED)
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    style_cell(table.cell(0, 0), size=8.5, color=MUTED)


def centered_text(draw, box, text, font, fill="#1F2A24"):
    left, top, right, bottom = box
    bounds = draw.multiline_textbbox((0, 0), text, font=font, spacing=7, align="center")
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.multiline_text(
        ((left + right - width) / 2, (top + bottom - height) / 2),
        text,
        font=font,
        fill=fill,
        spacing=7,
        align="center",
    )


def rounded_box(draw, box, title, detail, fill, outline="#2AAD16"):
    draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=3)
    left, top, right, bottom = box
    centered_text(draw, (left + 8, top + 8, right - 8, top + 58), title, load_font(24, True), "#176B12")
    centered_text(draw, (left + 12, top + 58, right - 12, bottom - 8), detail, load_font(18), "#1F2A24")


def arrow(draw, start, end, label=""):
    draw.line([start, end], fill="#2AAD16", width=5)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 10), (x - 18, y + 10)], fill="#2AAD16")
    if label:
        draw.text(((start[0] + end[0]) / 2 - 35, start[1] - 32), label, font=load_font(16), fill="#5F6B63")


def build_architecture_diagram():
    path = ASSET_DIR / "teacher-student-shared-backend.png"
    image = Image.new("RGB", (1600, 790), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "教师端与学生端共用后端架构", font=load_font(34, True), fill="#1F2A24")

    rounded_box(draw, (70, 150, 390, 350), "教师端 SPA", "创建 / 配置 / 发布\n审核 / 评分 / 分析", "#EDF8EA")
    rounded_box(draw, (70, 470, 390, 670), "学生端 SPA", "查看 / 加入 / 提交\n互动 / 接收结果", "#F4F6F5")
    rounded_box(draw, (515, 260, 850, 560), "统一 API 与认证", "Session / JWT\nRBAC + 对象级权限\n幂等 / 错误码 / 审计", "#EDF8EA")
    rounded_box(draw, (965, 150, 1285, 350), "共享业务服务", "课程 / 班级 / 内容\n任务 / 评价 / 互动", "#F4F6F5")
    rounded_box(draw, (965, 470, 1285, 670), "事件与通知", "Outbox / 消息\n通知 / 统计 / 预警", "#EDF8EA")
    rounded_box(draw, (1370, 260, 1550, 560), "数据库", "统一主数据\n状态机\n审计日志", "#F4F6F5")

    arrow(draw, (390, 250), (505, 350), "HTTPS")
    arrow(draw, (390, 570), (505, 470), "HTTPS")
    arrow(draw, (850, 350), (955, 250), "命令")
    arrow(draw, (850, 470), (955, 570), "事件")
    arrow(draw, (1285, 250), (1360, 350), "事务")
    arrow(draw, (1285, 570), (1360, 470), "投递")
    draw.text((70, 725), "原则：同一业务对象只保留一份权威数据；双端通过权限、可见性与状态机获得不同视图。", font=load_font(22), fill="#5F6B63")
    image.save(path)
    return path


def build_closed_loop_diagram():
    path = ASSET_DIR / "task-closed-loop.png"
    image = Image.new("RGB", (1600, 620), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "任务、提交与成绩闭环", font=load_font(34, True), fill="#1F2A24")
    steps = [
        ("教师建任务", "草稿 + 测试用例"),
        ("教师发布", "班级 / 时间 / 截止"),
        ("学生提交", "版本化源代码"),
        ("后端评测", "结果 + 诊断"),
        ("教师审核", "成绩 + 反馈"),
        ("学生查看", "发布后可见"),
    ]
    x = 50
    boxes = []
    for index, (title, detail) in enumerate(steps):
        box = (x, 190, x + 220, 420)
        rounded_box(draw, box, title, detail, "#EDF8EA" if index % 2 == 0 else "#F4F6F5")
        boxes.append(box)
        x += 260
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2] + 6, 305), (right[0] - 6, 305))
    draw.text((70, 505), "关键门禁：学生必须属于目标班级；隐藏测试不得下发；成绩仅在 published 状态对学生可见。", font=load_font(22), fill="#5F6B63")
    image.save(path)
    return path


def add_picture(doc, path, caption, alt_text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    shape = paragraph.add_run().add_picture(str(path), width=Inches(6.35))
    shape._inline.docPr.set("descr", alt_text)
    add_para(doc, caption, size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, keep=True)


def add_cover(doc, meta):
    add_para(doc, "ENTERPRISE ARCHITECTURE / INTERNAL", size=10, bold=True, color=GREEN, before=18, after=12)
    add_para(doc, "CodeTrack 教学协同平台", size=14, bold=True, color=DARK_GREEN, after=4)
    add_para(doc, "教师端与学生端协同后端接入设计说明", size=24, bold=True, color=INK, after=7)
    add_para(doc, "业务闭环、接口契约、数据模型、状态机与验收基线", size=14, color=MUTED, after=22)

    rows = [
        ("文档编号", DOC_ID),
        ("版本 / 状态", f"{DOC_VERSION} / 技术评审版"),
        ("适用系统", "CodeTrack 教师端、学生端及共享后端"),
        ("技术基线", "React / TypeScript / FastAPI / SQLAlchemy / SQLite"),
        ("编制日期", BASELINE_DATE),
        ("编制角色", "产品、前端、后端、测试、架构与安全"),
        ("密级", "内部使用"),
    ]
    for label, value in rows:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        set_run_font(paragraph.add_run(f"{label}："), size=10.5, bold=True, color=INK)
        set_run_font(paragraph.add_run(value), size=10.5, color=INK)

    add_callout(
        doc,
        "设计结论",
        "教师端与学生端必须共享课程、班级、内容、任务、提交、成绩、讨论和图谱等业务主数据。教师端负责创建、配置、审核与发布；学生端负责查看、参与、提交与接收结果；后端统一负责身份、权限、状态流转、事务一致性与审计。",
        fill=LIGHT_GREEN,
    )
    add_para(doc, f"当前数据库快照：{meta['tables']} 张业务表，{meta['rows']:,} 行记录。文中“建议接口”不代表已经实现。", size=9.2, color=MUTED, italic=True, before=14, after=0)
    add_page_break(doc)


def add_front_matter(doc):
    doc.add_heading("文档控制", level=1)
    add_table(doc, ["属性", "内容"], [
        ("文档名称", "CodeTrack 教师端与学生端协同后端接入设计说明"),
        ("文档编号 / 版本", f"{DOC_ID} / {DOC_VERSION}"),
        ("文档目的", "形成双端共享业务的开发、联调、测试、发布和运维共同基线"),
        ("目标读者", "产品经理、前端开发、后端开发、测试、架构、安全、运维"),
        ("实现口径", "已实现=代码可调用；部分实现=存在主链路但缺少闭环；待建设=目标设计"),
        ("保密要求", "不得外发真实凭据、学生源代码、成绩明细和个人信息"),
    ], [2300, 7060], font_size=9.2, header_fill=LIGHT_GREEN)

    doc.add_heading("版本修订记录", level=2)
    add_table(doc, ["版本", "日期", "变更摘要", "状态"], [
        ("V1.0", BASELINE_DATE, "建立教师端、学生端共享后端业务闭环及企业交付基线", "待评审"),
    ], [1100, 1600, 5160, 1500], font_size=8.8)

    doc.add_heading("评审与签署", level=2)
    add_table(doc, ["评审环节", "责任角色", "评审重点", "状态"], [
        ("产品评审", "产品负责人", "双端范围、角色动作、状态定义、异常体验", "待签署"),
        ("技术评审", "架构/后端负责人", "接口、事务、权限、幂等、数据模型", "待签署"),
        ("前端评审", "双端前端负责人", "页面跳转、按钮动作、状态刷新、错误处理", "待签署"),
        ("测试评审", "测试负责人", "场景覆盖、数据校验、权限负向、回归范围", "待签署"),
        ("安全评审", "安全负责人", "认证、越权、敏感数据、审计、滥用防护", "待签署"),
    ], [1450, 1850, 4660, 1400], font_size=8.6)

    doc.add_heading("阅读路径", level=1)
    add_table(doc, ["章节", "解决的问题"], [
        ("1-3", "为什么要共用后端、当前代码做到了什么、目标架构是什么"),
        ("4", "哪些业务必须与学生端连接，以及当前完成程度"),
        ("5", "教师按钮会影响学生哪个页面，学生动作如何回流教师端"),
        ("6-8", "状态机、接口契约、数据库与事务如何设计"),
        ("9-11", "权限、安全、前端接入、消息与一致性如何实现"),
        ("12-14", "如何测试、验收、分阶段上线和处理风险"),
        ("附录", "接口清单、数据表映射、错误码与交付清单"),
    ], [1700, 7660], font_size=9.0)


def add_scope_and_principles(doc, bullet_id):
    doc.add_heading("1. 文档说明", level=1)
    doc.add_heading("1.1 编制目的", level=2)
    add_para(doc, "本文件定义教师端与学生端需要共用后端的业务范围、页面交互、API、数据库、状态机、权限和验收标准，使双端开发不再依赖各自硬编码数据或互不一致的状态副本。")
    doc.add_heading("1.2 范围", level=2)
    for text in [
        "覆盖账号与会话、课程与班级、章节与知识点、教学资料、任务与测试、学生提交与自动评测、评分反馈、公告、课堂讨论、知识图谱、通知与学情分析。",
        "覆盖教师端按钮到学生端页面的影响、学生动作回流教师端的路径，以及中间的后端权限、事务和审计职责。",
        "同时记录当前项目已有实现和企业生产化目标，不把建议能力表述为现有能力。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("1.3 不在范围", level=2)
    for text in [
        "学生端完整视觉设计稿和所有页面组件实现。",
        "统一身份平台、短信、邮件、对象存储、消息中间件的厂商选型。",
        "教学制度、成绩申诉制度和数据保留年限的最终业务审批。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("1.4 术语", level=2)
    add_table(doc, ["术语", "定义"], [
        ("共享主数据", "教师端与学生端围绕同一业务对象读取和修改的唯一权威数据"),
        ("命令", "会改变状态的请求，例如发布任务、提交代码、发布成绩"),
        ("查询", "不改变状态的请求，例如学生查看任务、教师查看提交"),
        ("可见性", "对象是否可被特定课程、班级、角色或用户读取"),
        ("状态机", "对象允许的状态及合法迁移规则"),
        ("对象级权限", "不仅检查角色，还校验用户是否拥有或属于具体课程、班级、任务"),
        ("Outbox", "业务事务内记录待投递事件，事务提交后可靠发送通知或更新统计"),
    ], [1800, 7560], font_size=9.0)

    doc.add_heading("2. 设计原则", level=1)
    principles = [
        ("单一事实来源", "课程、班级、任务、提交、成绩等对象只保留一份权威记录，双端不得各自保存业务副本。"),
        ("后端裁决", "角色、归属、可见性、截止时间、状态迁移和成绩发布均由后端验证，不能只依赖前端按钮禁用。"),
        ("发布即授权", "教师执行发布操作时，后端同时完成状态变化、目标班级范围、可见性和审计。"),
        ("最小暴露", "学生只获得自身和已授权班级的数据；隐藏测试、其他学生源代码、教师草稿不得下发。"),
        ("事务一致性", "主记录、关联记录、审计和 Outbox 事件应处于同一事务，失败时整体回滚。"),
        ("幂等与可重试", "加入班级、标记已读、发布、通知消费等操作必须支持安全重试。"),
        ("兼容演进", "接口使用稳定业务码和版本；数据库结构通过 Alembic 迁移，不直接修改生产表。"),
    ]
    add_table(doc, ["原则", "开发要求"], principles, [2100, 7260], font_size=9.0, header_fill=LIGHT_GREEN)


def add_baseline_and_architecture(doc, meta, bullet_id):
    doc.add_heading("3. 当前基线与目标架构", level=1)
    add_callout(doc, "当前基线", f"项目为 React/Vite 前端、FastAPI/SQLAlchemy 后端和 SQLite 数据库；数据库当前 {meta['tables']} 张业务表、{meta['rows']:,} 行。部分学生端调用已存在，但仍以固定 student-01 / student-03 和 X-User-Id 模拟身份。", fill=LIGHT_GREEN)
    doc.add_heading("3.1 当前已验证的共享链路", level=2)
    for text in [
        "邀请码入班：教师生成邀请码，学生加入，enrollments 记录关系，教师查看班级学生。",
        "课程内容：教师发布章节并切换资料可见性，学生只读取已发布章节和 students 可见资料。",
        "任务提交：教师创建并发布任务，学生查看并提交，后端生成评测和诊断，教师查看提交。",
        "课堂讨论：教师创建/发布/结束讨论，已入班学生查看并回复，教师收到通知。",
        "成绩反馈：教师保存并发布成绩、追加反馈；学生端读取发布结果的专用接口仍需补齐。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("3.2 当前关键缺口", level=2)
    add_table(doc, ["优先级", "缺口", "影响", "企业要求"], [
        ("P0", "X-User-Id 由前端自报，学生 ID 固定", "可伪造身份、无法多用户真实联调", "接入服务端 Session/JWT，从令牌解析 user_id 与 role"),
        ("P0", "对象级权限覆盖不完整", "可能跨课程、跨班级读取或修改", "所有资源按 owner/enrollment/target scope 校验"),
        ("P1", "公告、成绩、知识图谱缺学生端读取闭环", "教师发布后学生无标准入口", "补学生列表/详情/已读接口并定义页面刷新"),
        ("P1", "班级加入无 pending/rejected/removed 状态", "无法审核申请和保留历史", "扩展 enrollment 状态机和审核接口"),
        ("P1", "通知仅局部实现", "跨端状态变化不能稳定触达", "统一通知表、Outbox、未读数与批量已读"),
        ("P2", "统计依赖同步查询", "数据量增长后页面慢、口径不稳定", "事件驱动增量聚合并保留重算能力"),
    ], [800, 2500, 2600, 3460], font_size=8.2)

    doc.add_heading("3.3 目标架构", level=2)
    add_picture(doc, build_architecture_diagram(), "图 1  教师端与学生端共用后端架构", "教师端和学生端通过统一认证 API 接入共享业务服务、事件通知和数据库")
    doc.add_heading("3.4 组件职责", level=2)
    add_table(doc, ["组件", "职责", "禁止事项"], [
        ("教师端", "创建、配置、审核、发布、撤回、评分、分析", "不得直接决定数据归属或绕过后端状态机"),
        ("学生端", "加入、查看、参与、提交、接收成绩和反馈", "不得自行传入可冒充他人的 user_id"),
        ("统一认证", "登录、令牌签发/刷新/注销、角色与会话管理", "不得信任 X-User-Id 作为生产身份"),
        ("业务服务", "权限、状态机、事务、幂等、错误码、审计", "不得由前端拼接跨表业务规则"),
        ("数据库", "权威主数据、唯一约束、外键、版本和审计", "不得让双端各建一套同义表"),
        ("事件/通知", "跨端提醒、未读数、统计增量和失败重试", "不得在主事务成功前发送不可撤销通知"),
    ], [1600, 4300, 3460], font_size=8.6)


def add_capability_map(doc):
    doc.add_heading("4. 必须连接学生端的功能清单", level=1)
    add_para(doc, "下表是本项目跨端功能的权威范围。状态依据当前代码与数据库检查：已实现表示主链路存在，部分实现表示缺少学生读取、完整状态或生产认证，待建设表示需要新增。", color=MUTED)
    rows = [
        ("身份与会话", "教师登录、角色校验", "学生登录、会话续期", "users / credentials；需统一 Session/JWT", "部分实现", "P0"),
        ("课程与班级", "建课、建班、邀请码、查看学生", "加入课程、查看所属班级", "courses / class_groups / enrollments", "已实现主链路", "P0"),
        ("加入审核", "审核、拒绝、移除、重新邀请", "申请、撤回、查看状态", "enrollments 状态与审核记录", "待建设", "P1"),
        ("章节与知识点", "创建、编辑、发布/撤回", "查看已发布章节和知识点", "chapters / knowledge_points", "已实现主链路", "P0"),
        ("教学资料", "上传、分类、设置可见性", "查看/下载授权资料", "materials / folders / links", "部分实现", "P0"),
        ("任务与测试", "建任务、测试用例、发布/定时/关闭", "查看任务、截止时间", "tasks / test_cases", "部分实现", "P0"),
        ("提交与评测", "监控提交、查看详情、重评", "提交、重交、查看公开评测", "submissions / evaluation_results", "已实现主链路", "P0"),
        ("诊断与审核", "审核 AI 诊断、修订建议", "查看教师确认后的诊断", "diagnosis_results / reviews", "部分实现", "P1"),
        ("成绩与反馈", "保存、发布成绩、反馈", "查看已发布成绩和反馈", "grades / teacher_feedback", "部分实现", "P0"),
        ("课程公告", "创建、发布、置顶、撤回", "查看、标记已读", "announcements / reads", "部分实现", "P1"),
        ("课堂讨论", "创建、发布、结束、查看回复", "参与讨论、回复", "discussions / replies", "已实现", "P1"),
        ("知识图谱", "编辑、指定班级、发布", "查看已发布图谱、关联资料", "teacher_knowledge_graphs；目标班级需规范化", "部分实现", "P1"),
        ("通知中心", "接收提交/回复/预警，标记已读", "接收发布/成绩/反馈通知", "notifications / outbox", "部分实现", "P1"),
        ("学情与预警", "查看聚合、风险、导出", "查看个人进度与建议", "提交/成绩/活跃聚合", "部分实现", "P2"),
    ]
    add_table(doc, ["业务域", "教师端", "学生端", "共享后端/数据", "当前状态", "优先级"], rows, [1300, 1900, 1800, 2550, 1150, 660], font_size=7.7, header_fill=LIGHT_GREEN)

    doc.add_heading("4.1 不应连接学生端的教师私有数据", level=2)
    add_table(doc, ["数据", "处理要求"], [
        ("课程草稿、未发布章节、未发布任务", "仅课程教师和授权协作者可见"),
        ("隐藏测试输入输出", "学生端只返回测试名称的脱敏结果或汇总，不返回内容"),
        ("其他学生源代码、诊断、成绩", "禁止向非本人学生下发"),
        ("AI 审核内部备注、风险处置记录", "教师/管理员可见，学生只看已发布结论"),
        ("教师账号安全设置、审计详情", "不得进入学生接口响应"),
    ], [2700, 6660], font_size=8.9)


def add_interaction_mapping(doc):
    heading = doc.add_heading("5. 双端页面与按钮交互映射", level=1)
    heading.paragraph_format.page_break_before = True
    add_para(doc, "页面路由中的学生端路径为建议命名；实际学生端若已有路由，应保持业务语义并在联调清单中映射。", color=MUTED)

    doc.add_heading("5.1 教师端动作对学生端的影响", level=2)
    teacher_rows = [
        ("班级管理 / 重新生成邀请码", "POST /teacher/classes/{id}/join-code", "学生加入班级页 /student/join", "旧码失效，新码可加入", "邀请码轮换、审计"),
        ("章节管理 / 发布章节", "PATCH /teacher/chapters/{id}", "课程学习页 /student/courses/{id}", "出现章节、知识点、可见资料和任务", "chapter=published"),
        ("章节管理 / 撤回章节", "PATCH /teacher/chapters/{id}", "课程学习页", "章节从学生列表移除", "chapter=draft；资料回 teacher"),
        ("资料库 / 设为学生可见", "PATCH /teacher/materials/{id}", "章节资料页", "学生可查看或下载", "visibility=students"),
        ("任务中心 / 发布任务", "POST /teacher/tasks/{id}/publish", "任务中心 /student/tasks", "目标班级学生看到任务", "draft/scheduled→published"),
        ("评分页 / 发布成绩", "POST /teacher/submissions/{id}/grade/publish", "提交详情 /student/submissions/{id}", "显示成绩、维度和教师评语", "grade=published"),
        ("反馈 / 发送给学生", "POST /teacher/submissions/{id}/feedback", "提交详情、通知中心", "显示反馈并增加未读通知", "feedback.student_visible=true"),
        ("公告 / 发布", "建议 POST /teacher/announcements/{id}/publish", "公告中心 /student/announcements", "目标学生看到公告并有未读状态", "announcement=published"),
        ("讨论 / 发布讨论", "POST /teacher/discussions/{id}/publish", "课堂讨论 /student/discussions", "目标班级可查看和回复", "discussion=published"),
        ("讨论 / 结束", "POST /teacher/discussions/{id}/end", "讨论详情", "保留历史但禁用回复", "discussion=ended"),
        ("知识图谱 / 发布", "POST /teacher/knowledge-graphs/{id}/publish", "知识图谱 /student/graphs", "目标班级可查看发布版本", "graph=published"),
    ]
    add_table(doc, ["教师页面/按钮", "调用接口", "学生页面", "学生端结果", "后端关键动作"], teacher_rows, [2050, 2350, 1900, 1800, 1260], font_size=7.7)

    doc.add_heading("5.2 学生端动作回流教师端", level=2)
    student_rows = [
        ("加入班级 / 确认加入", "POST /classes/{join_code}/join", "班级学生列表、加入状态", "新增/恢复 enrollment，更新人数"),
        ("任务详情 / 提交代码", "POST /student/tasks/{id}/submissions", "提交监控、AI 审核、通知", "创建版本、评测、诊断和必要审核"),
        ("讨论详情 / 发布回复", "POST /student/discussions/{id}/replies", "课堂讨论详情、通知中心", "创建回复，通知教师"),
        ("公告详情 / 标记已读", "建议 PATCH /student/announcements/{id}/read", "公告已读统计", "幂等写入 announcement_reads"),
        ("资料详情 / 查看或下载", "建议 POST /student/materials/{id}/view", "资料引用/学习统计", "记录访问事件，不改变资料主记录"),
        ("成绩详情 / 发起申诉", "建议 POST /student/grades/{id}/appeals", "教师成绩处理队列", "创建申诉状态与审计"),
    ]
    add_table(doc, ["学生页面/按钮", "调用接口", "教师端回流页面", "后端关键动作"], student_rows, [2300, 2700, 2200, 2160], font_size=8.0, header_fill=LIGHT_GREEN)

    doc.add_heading("5.3 交互反馈标准", level=2)
    add_table(doc, ["阶段", "前端要求", "后端要求"], [
        ("提交前", "按钮防重复、表单校验、展示影响范围", "不能依赖前端校验，必须二次验证"),
        ("处理中", "显示加载状态；长任务展示进度或轮询", "返回 request_id；异步任务返回 job_id"),
        ("成功", "更新本地缓存并重新拉取权威数据", "返回最新对象、版本号和状态"),
        ("失败", "保留用户输入，按业务码给出可行动提示", "稳定 error.code、message、field_errors、trace_id"),
        ("并发冲突", "提示数据已更新并提供重新加载", "使用 version/ETag，冲突返回 409"),
    ], [1300, 4000, 4060], font_size=8.7)


def add_state_machines(doc):
    heading = doc.add_heading("6. 核心状态机", level=1)
    heading.paragraph_format.page_break_before = True
    add_callout(doc, "强制规则", "状态迁移由后端执行。前端隐藏按钮只是体验优化，不能替代后端对当前状态、操作者、目标范围和时间条件的校验。")
    state_rows = [
        ("加入关系", "invited → pending → joined → removed", "申请、审核、移除；重复加入幂等", "当前仅用 enrollment 是否存在表示 joined"),
        ("章节", "draft ↔ published → archived", "发布后学生可见；撤回后立即隐藏", "当前支持 draft/published"),
        ("资料", "uploading → parsing → ready → deleted", "ready 且 visibility=students 才可下发", "当前支持 parsing/ready/deleted"),
        ("任务", "draft → scheduled → published → closed", "至少一个测试用例和知识点；到时发布", "当前有 draft/scheduled/published/closed"),
        ("提交", "submitted → evaluating → evaluated / failed", "每次重交产生递增 version", "当前创建时同步生成结果"),
        ("诊断", "generated → pending_review → confirmed/rejected", "低置信度需教师审核后再对学生展示", "当前有审核记录，学生展示待补"),
        ("成绩", "draft → published → amended", "published 前仅教师可见；修改需保留历史", "当前 draft/published，历史待补"),
        ("公告", "draft → scheduled → published → withdrawn", "按课程/班级授权并维护已读", "当前学生发布闭环待补"),
        ("讨论", "draft → published → ended", "published 可回复，ended 只读", "当前已实现"),
        ("图谱", "draft → published → archived", "发布版本面向目标班级；草稿不下发", "当前教师发布有，学生接口待补"),
    ]
    add_table(doc, ["对象", "目标状态", "迁移规则", "当前实现"], state_rows, [1300, 2400, 3300, 2360], font_size=8.1, header_fill=LIGHT_GREEN)

    doc.add_heading("6.1 非法迁移处理", level=2)
    add_table(doc, ["场景", "HTTP / 业务码", "处理"], [
        ("发布已关闭任务", "409 TASK_STATE_CONFLICT", "拒绝；返回 current_status=closed"),
        ("学生提交未发布任务", "404 TASK_NOT_AVAILABLE", "不暴露任务是否存在"),
        ("结束草稿讨论", "409 DISCUSSION_STATE_CONFLICT", "提示需先发布"),
        ("发布空知识图谱", "409 GRAPH_EMPTY", "保留草稿并提示至少一个节点"),
        ("发布成绩后再次修改", "409 GRADE_ALREADY_PUBLISHED 或创建 amended", "按业务制度二选一并留审计"),
    ], [2600, 2800, 3960], font_size=8.7)


def add_api_design(doc, bullet_id):
    doc.add_heading("7. API 契约与接口设计", level=1)
    doc.add_heading("7.1 统一约定", level=2)
    for text in [
        "基础路径使用 /api/v1；生产请求通过 Authorization: Bearer <token> 或 HttpOnly Session Cookie 认证。",
        "响应统一为 data、meta、request_id；失败统一为 error.code、message、field_errors、request_id。",
        "列表接口统一支持 page、page_size、sort、filter；默认 page_size=20，最大 100。",
        "写接口使用 Idempotency-Key；可并发编辑对象使用 version 或 If-Match/ETag。",
        "时间字段使用 ISO 8601 且含时区；数据库统一 UTC，前端按用户时区展示。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_code_block(doc, '{\n  "data": {"id": "task-01", "status": "published", "version": 4},\n  "meta": {"timestamp": "2026-08-16T10:30:00+08:00"},\n  "request_id": "req_01HXYZ"\n}')
    add_code_block(doc, '{\n  "error": {\n    "code": "TASK_STATE_CONFLICT",\n    "message": "已关闭任务不能重新发布",\n    "field_errors": []\n  },\n  "request_id": "req_01HXYZ"\n}')

    heading = doc.add_heading("7.2 当前接口与目标补齐", level=2)
    heading.paragraph_format.page_break_before = True
    endpoints = [
        ("认证", "POST /teacher/auth/login", "新增 POST /auth/login、/auth/refresh、/auth/logout、GET /me", "P0"),
        ("班级", "POST /teacher/classes/{id}/join-code；POST /classes/{code}/join", "新增申请审核、拒绝、移除、退出", "P1"),
        ("课程内容", "PATCH /teacher/chapters/{id}；GET /student/courses/{id}/content", "增加 ETag、增量同步、章节排序", "P1"),
        ("资料", "教师 CRUD；学生通过课程 content 获取资料", "新增学生资料详情/下载鉴权/访问记录", "P1"),
        ("任务", "教师创建/发布；学生列表/提交", "新增编辑、删除、关闭、定时发布调度、测试用例 CRUD", "P0"),
        ("成绩反馈", "教师保存/发布成绩、创建反馈", "新增 GET /student/submissions/{id}/result", "P0"),
        ("公告", "教师公告列表、已读", "新增教师 CRUD/发布/撤回；学生列表/详情/已读", "P1"),
        ("讨论", "教师创建/发布/结束；学生列表/回复", "新增分页、内容审核、删除/隐藏回复", "P2"),
        ("图谱", "教师图谱 CRUD/发布", "新增 GET /student/knowledge-graphs 与详情", "P1"),
        ("通知", "教师单条标记已读", "新增双端分页、全部已读、删除、未读数", "P1"),
        ("分析", "教师 analytics overview", "新增个人学习进度、导出任务、口径版本", "P2"),
    ]
    add_table(doc, ["业务", "当前接口", "目标补齐", "优先级"], endpoints, [1200, 3300, 4000, 860], font_size=8.0)

    doc.add_heading("7.3 关键目标接口", level=2)
    add_table(doc, ["方法", "路径", "调用方", "用途", "权限"], [
        ("POST", "/api/v1/auth/login", "双端", "登录并建立服务端会话", "公开+限流"),
        ("GET", "/api/v1/me", "双端", "恢复当前用户、角色和权限", "已登录"),
        ("POST", "/api/v1/teacher/classes/{id}/enrollments/{eid}/approve", "教师", "批准入班", "班级所属教师"),
        ("GET", "/api/v1/student/courses/{id}/content", "学生", "读取已发布课程内容", "已入班学生"),
        ("GET", "/api/v1/student/tasks", "学生", "读取目标班级已发布任务", "已入班学生"),
        ("POST", "/api/v1/student/tasks/{id}/submissions", "学生", "创建提交版本", "任务目标班级成员"),
        ("GET", "/api/v1/student/submissions/{id}/result", "学生", "读取本人公开评测、成绩和反馈", "提交本人"),
        ("GET", "/api/v1/student/announcements", "学生", "读取目标课程/班级公告", "已入班学生"),
        ("PATCH", "/api/v1/student/announcements/{id}/read", "学生", "幂等标记已读", "可见公告"),
        ("GET", "/api/v1/student/knowledge-graphs", "学生", "读取目标班级已发布图谱", "已入班学生"),
        ("GET", "/api/v1/notifications", "双端", "分页查询通知和未读数", "通知所有者"),
    ], [850, 3450, 900, 2200, 1960], font_size=7.7, header_fill=LIGHT_GREEN)


def add_database_design(doc, meta):
    doc.add_heading("8. 数据库与事务设计", level=1)
    add_para(doc, f"当前 SQLite 基线包含 {meta['tables']} 张业务表。双端共享不需要复制表，而是围绕现有表补充状态、权限范围、版本和事件记录。")
    doc.add_heading("8.1 现有共享表映射", level=2)
    add_table(doc, ["业务域", "现有表", "双端用途"], [
        ("账号", "users / teacher_credentials / preferences", "身份主数据；凭据模型需扩展为角色无关认证"),
        ("课程班级", "courses / class_groups / enrollments", "教师拥有课程；学生通过 enrollment 获得访问权"),
        ("内容", "chapters / knowledge_points / materials / material_folders / material_knowledge_links", "教师编辑，学生读取已发布和授权内容"),
        ("任务评价", "tasks / test_cases / submissions / evaluation_results", "教师发布任务，学生提交，后端评测"),
        ("成绩诊断", "grades / diagnosis_results / diagnosis_reviews / teacher_feedback", "教师审核发布，学生读取本人结果"),
        ("互动", "course_announcements / announcement_reads / course_discussions / discussion_replies", "发布、已读、回复和回流"),
        ("图谱通知", "teacher_knowledge_graphs / notifications / audit_logs", "目标班级图谱、通知与证据链"),
    ], [1500, 3800, 4060], font_size=8.3)

    heading = doc.add_heading("8.2 建议新增或调整", level=2)
    heading.paragraph_format.page_break_before = True
    add_table(doc, ["对象", "建议", "目的", "迁移要点"], [
        ("auth_sessions", "新增 user_id、refresh_hash、expires_at、revoked_at、device", "真实双端会话与注销", "迁移后停用 X-User-Id"),
        ("enrollments", "增加 status、join_method、reviewed_by、reviewed_at、removed_at", "申请审核和历史保留", "现有记录回填 joined"),
        ("course_announcements", "增加 status、target_scope、publish_at、withdrawn_at", "完整发布生命周期", "现有公告回填 published"),
        ("graph_targets", "将 target_classes JSON 规范化为 graph_id/class_id", "对象级权限和索引", "迁移 JSON 后双写验证"),
        ("grade_revisions", "新增成绩变更版本、原因、操作者", "已发布成绩可审计修订", "不覆盖历史值"),
        ("outbox_events", "新增 aggregate、event_type、payload、status、attempts", "可靠通知和统计更新", "与业务写同事务"),
        ("idempotency_keys", "新增 user_id、key、request_hash、response", "写请求重试防重", "按用户+key 唯一"),
        ("resource_access_logs", "可选记录资料查看/下载", "学情统计与安全审计", "设置保留期和脱敏"),
    ], [1650, 3000, 2600, 2110], font_size=8.0, header_fill=LIGHT_GREEN)

    doc.add_heading("8.3 关键约束与索引", level=2)
    add_table(doc, ["约束/索引", "要求"], [
        ("enrollments(class_id, student_id)", "唯一；重复加入返回已有关系，不创建重复行"),
        ("submissions(task_id, student_id, version)", "唯一；版本在事务中递增并防并发冲突"),
        ("announcement_reads(announcement_id, user_id)", "唯一；已读接口使用 upsert"),
        ("discussion_replies(discussion_id, created_at)", "支持详情分页和时间排序"),
        ("notifications(user_id, read, created_at)", "支持未读数和时间倒序分页"),
        ("graph_targets(graph_id, class_id)", "唯一；学生图谱按 class_id 快速过滤"),
        ("outbox_events(status, next_retry_at)", "支持失败重试扫描"),
        ("audit_logs(actor_id, action, created_at)", "支持用户与操作追溯"),
    ], [3200, 6160], font_size=8.6)

    doc.add_heading("8.4 事务边界", level=2)
    add_table(doc, ["命令", "同一事务内必须完成", "事务后动作"], [
        ("学生加入班级", "写 enrollment、审计、outbox", "通知教师、刷新班级统计"),
        ("教师发布任务", "校验测试/知识点、更新 task、审计、outbox", "通知目标学生、更新任务列表缓存"),
        ("学生提交", "写 submission、版本、评测任务/结果、审计、outbox", "异步评测时由 worker 更新并通知"),
        ("发布成绩", "更新 grade、写 revision、审计、outbox", "通知学生、更新学情聚合"),
        ("讨论回复", "写 reply、审计、outbox", "通知教师、更新参与人数"),
    ], [1900, 4560, 2900], font_size=8.4)


def add_security_and_frontend(doc, bullet_id):
    doc.add_heading("9. 身份、权限与安全", level=1)
    add_callout(doc, "生产门禁", "当前 X-User-Id 和固定 student-01 / student-03 只能用于本地演示。只要学生端正式接入，该方案必须替换，否则任何用户都可能伪造其他教师或学生身份。", fill=RISK)
    doc.add_heading("9.1 权限模型", level=2)
    add_table(doc, ["资源", "教师权限条件", "学生权限条件"], [
        ("课程/章节/资料", "course.teacher_id == current_user.id", "存在 joined enrollment 且对象已发布/可见"),
        ("班级/成员", "class.course.teacher_id == current_user.id", "仅查看本人加入关系和允许公开的班级信息"),
        ("任务", "任务所属课程归当前教师", "任务班级包含本人且状态 published、时间有效"),
        ("提交/成绩/反馈", "提交任务所属课程归当前教师", "submission.student_id == current_user.id 且结果已公开"),
        ("讨论", "讨论教师为当前教师", "本人加入目标班级且讨论 published"),
        ("图谱", "graph.user_id == current_user.id", "graph=published 且 graph_targets 包含本人班级"),
        ("通知", "notification.user_id == current_user.id", "notification.user_id == current_user.id"),
    ], [1800, 3800, 3760], font_size=8.5, header_fill=LIGHT_GREEN)
    doc.add_heading("9.2 安全控制", level=2)
    for text in [
        "认证：短期访问令牌 + 可撤销刷新会话，密码强度、首次改密、失败限流、设备和登录审计。",
        "授权：RBAC 检查教师/学生角色，同时执行课程所有权、班级成员关系、资源目标范围检查。",
        "越权防护：对不可见资源优先返回 404，避免暴露对象存在；所有 IDOR 场景纳入自动化测试。",
        "敏感数据：隐藏测试、源代码、成绩、个人信息最小化返回；日志不记录令牌和完整正文。",
        "上传安全：文件类型、大小、内容扫描、随机存储名、下载鉴权和响应头。",
        "审计：发布、撤回、评分、移除学生、权限变更、导出等高风险操作不可缺失。",
    ]:
        add_list_item(doc, text, bullet_id)

    doc.add_heading("10. 前端接入规范", level=1)
    doc.add_heading("10.1 双端统一数据层", level=2)
    add_table(doc, ["能力", "教师端与学生端共同要求"], [
        ("认证上下文", "应用启动调用 /me；401 统一刷新一次，失败后清理会话并跳转登录"),
        ("请求封装", "统一 base URL、token/cookie、request_id、超时、错误码映射和取消请求"),
        ("缓存键", "包含 role、user_id、course_id、class_id，退出登录必须清空"),
        ("状态刷新", "命令成功后使用返回对象更新缓存，并失效相关列表和未读数"),
        ("表单", "保留失败输入；422 映射字段；409 提示重新加载；403 返回可访问页面"),
        ("实时性", "第一阶段命令成功后轮询/重新拉取；第二阶段对通知和提交状态接入 SSE/WebSocket"),
        ("可观测性", "前端错误上报包含 route、action、request_id、build_version，不上传敏感正文"),
    ], [2000, 7360], font_size=8.7)

    doc.add_heading("10.2 页面加载与跳转规则", level=2)
    add_table(doc, ["触发", "加载顺序", "失败回退"], [
        ("登录成功", "/me → 可访问课程 → 最近课程/默认课程", "无课程进入空状态，不写死 course-ds"),
        ("切换课程", "清空旧课程缓存 → 拉课程权限 → 章节/班级/通知", "403 返回课程列表并提示权限变化"),
        ("学生打开任务", "任务详情 → 本人最新提交 → 公开结果", "任务不可用返回任务中心"),
        ("教师打开提交", "提交详情 → 评测/诊断 → 成绩/反馈", "404 返回当前任务提交列表"),
        ("打开通知", "标记已读 → 按 resource_type/id 跳转", "资源不可见时保留通知并提示已失效"),
    ], [1800, 4800, 2760], font_size=8.5)

    doc.add_heading("10.3 现有 src/api.ts 整改", level=2)
    add_table(doc, ["当前问题", "整改"], [
        ("_currentUserId 默认 teacher-01", "删除默认业务身份；由 /me 返回的服务端会话决定"),
        ("学生接口硬编码 student-01/student-03", "学生端使用当前会话，不允许调用方传 userId"),
        ("每个请求发送 X-User-Id", "改为 HttpOnly Cookie 或 Bearer token；后端忽略客户端用户 ID"),
        ("错误只含 status 和 message", "增加 code、fieldErrors、requestId、retryable"),
        ("缺少超时/取消/401 刷新", "统一 AbortController、刷新锁和一次重放"),
        ("大量 any", "按 OpenAPI 生成类型或维护稳定 DTO，禁止跨端复用数据库模型"),
    ], [3300, 6060], font_size=8.6, header_fill=LIGHT_GREEN)


def add_sync_and_sequence(doc, bullet_id):
    doc.add_heading("11. 跨端同步、通知与一致性", level=1)
    add_picture(doc, build_closed_loop_diagram(), "图 2  任务、提交与成绩的双端闭环", "教师发布任务、学生提交、后端评测、教师审核、学生查看的完整闭环")
    doc.add_heading("11.1 一致性分级", level=2)
    add_table(doc, ["数据", "一致性要求", "实现"], [
        ("加入关系、权限、发布状态、提交、成绩", "强一致", "主数据库事务；读取前执行对象级权限"),
        ("通知、未读数", "最终一致，目标秒级", "Outbox + 重试；页面进入时以数据库校正"),
        ("仪表盘、学情、风险预警", "最终一致，目标分钟级", "事件增量聚合 + 定时全量校准"),
        ("文件解析、AI 诊断", "异步状态一致", "job 状态机、失败重试、人工重跑"),
    ], [2800, 2400, 4160], font_size=8.7)
    doc.add_heading("11.2 推荐事件", level=2)
    events = [
        "EnrollmentJoined / EnrollmentRemoved",
        "ChapterPublished / ChapterWithdrawn",
        "MaterialVisibilityChanged",
        "TaskPublished / TaskClosed",
        "SubmissionCreated / EvaluationCompleted",
        "GradePublished / FeedbackPublished",
        "AnnouncementPublished / AnnouncementRead",
        "DiscussionPublished / DiscussionReplied / DiscussionEnded",
        "KnowledgeGraphPublished",
    ]
    for text in events:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("11.3 缓存与并发", level=2)
    add_table(doc, ["问题", "处理"], [
        ("教师刚发布，学生列表仍旧", "发布响应成功后失效目标课程/班级缓存；通知仅作提醒，不作为权限依据"),
        ("两名教师协作者同时编辑", "对象 version 自增；旧版本更新返回 409"),
        ("学生重复点击提交", "Idempotency-Key + 请求哈希，相同请求返回原结果"),
        ("定时发布任务", "服务端调度器按数据库时间执行；前端倒计时仅展示"),
        ("通知投递失败", "Outbox 保留 pending/failed，指数退避重试并告警"),
    ], [2900, 6460], font_size=8.6)


def add_testing(doc, bullet_id, number_id):
    doc.add_heading("12. 测试策略与验收用例", level=1)
    add_callout(doc, "测试口径", "本节是企业联调与验收用例基线。当前项目已有后端回归可作为现状证据；目标接口用例需在相应功能实现后执行，不得把“用例已编写”写成“功能已通过”。")
    doc.add_heading("12.1 测试层级", level=2)
    add_table(doc, ["层级", "覆盖", "门禁"], [
        ("单元测试", "状态迁移、权限谓词、可见性、分数计算、事件生成", "核心分支与异常分支必须覆盖"),
        ("API 集成", "真实数据库事务、外键、唯一约束、错误契约、幂等", "P0 接口正向与负向全部通过"),
        ("双端契约", "OpenAPI DTO、字段可见性、分页、时间和枚举", "前后端生成/校验契约无差异"),
        ("浏览器 E2E", "教师发布→学生查看→学生提交→教师评分→学生查看", "主闭环桌面与移动视口通过"),
        ("安全", "伪造身份、IDOR、越权、隐藏测试泄露、上传攻击", "P0/P1 高危为 0"),
        ("性能与恢复", "并发提交、定时发布、锁冲突、重试、备份恢复", "达到批准的 SLO/RPO/RTO"),
    ], [1600, 5000, 2760], font_size=8.6)

    doc.add_heading("12.2 跨端验收用例矩阵", level=2)
    cases = [
        ("AUTH-01", "教师正常登录", "有效账号", "返回会话；/me=teacher", "P0"),
        ("AUTH-02", "学生正常登录", "有效账号", "返回会话；/me=student", "P0"),
        ("AUTH-03", "伪造 X-User-Id", "已登录普通学生", "忽略伪造头；不能切换身份", "P0"),
        ("ENR-01", "邀请码加入", "有效邀请码、未加入", "创建 joined enrollment；教师列表出现", "P0"),
        ("ENR-02", "重复加入", "已加入", "幂等返回原关系；无重复行", "P0"),
        ("ENR-03", "无效/旧邀请码", "邀请码已轮换", "404 JOIN_CODE_INVALID", "P0"),
        ("ENR-04", "跨班级移除越权", "非课程教师", "404/403；数据不变并记录安全日志", "P0"),
        ("CNT-01", "发布章节", "草稿章节", "学生课程页出现；资料按可见性过滤", "P0"),
        ("CNT-02", "撤回章节", "已发布章节", "学生刷新后不可见；历史提交不删除", "P0"),
        ("MAT-01", "学生查看资料", "ready+students", "可读取；记录访问事件", "P1"),
        ("MAT-02", "访问教师私有资料", "visibility=teacher", "404；不泄露 URL", "P0"),
        ("TSK-01", "发布合法任务", "有测试用例和知识点", "目标班级学生可见", "P0"),
        ("TSK-02", "无测试用例发布", "test_count=0", "422；任务保持 draft", "P0"),
        ("TSK-03", "定时发布", "publish_at>now", "scheduled；到时由服务端发布", "P1"),
        ("TSK-04", "非目标班级查看", "学生不在目标班级", "列表不出现，详情 404", "P0"),
        ("SUB-01", "首次提交", "任务已发布且未截止", "version=1；创建评测/诊断", "P0"),
        ("SUB-02", "重复提交", "已有 version=1", "新业务提交 version=2；重复 HTTP 请求不增版", "P0"),
        ("SUB-03", "截止后提交", "now>due_at", "按策略拒绝或标记 late；行为稳定", "P1"),
        ("SUB-04", "隐藏测试保护", "任务含 hidden case", "学生响应不含输入、输出、权重", "P0"),
        ("GRD-01", "成绩草稿隔离", "grade=draft", "教师可见，学生结果不显示", "P0"),
        ("GRD-02", "发布成绩", "grade=draft", "学生结果显示；生成通知", "P0"),
        ("GRD-03", "查看他人成绩", "学生请求他人 submission", "404；无敏感字段泄露", "P0"),
        ("ANN-01", "公告发布", "目标课程/班级", "目标学生出现未读公告", "P1"),
        ("ANN-02", "重复标记已读", "公告可见", "两次均成功；只有一条 read 记录", "P1"),
        ("DSC-01", "讨论发布与回复", "学生已入班", "回复出现在教师端；教师收到通知", "P1"),
        ("DSC-02", "结束后回复", "discussion=ended", "404/409；不创建 reply", "P1"),
        ("GRP-01", "图谱目标班级可见", "published+target class", "目标学生可查看完整图谱", "P1"),
        ("GRP-02", "图谱跨班级越权", "学生不在 target", "404；不返回节点摘要", "P0"),
        ("NTF-01", "事件通知重试", "通知首次投递失败", "Outbox 重试且最终仅一条通知", "P1"),
        ("TXN-01", "业务事务回滚", "审计/Outbox 写失败", "主业务状态也回滚，无半成功", "P0"),
        ("CON-01", "并发编辑冲突", "旧 version 提交", "409；不覆盖新数据", "P1"),
    ]
    add_table(doc, ["编号", "场景", "前置条件", "预期结果", "级别"], cases, [950, 1800, 2200, 3680, 730], font_size=7.5, header_fill=LIGHT_GREEN)

    doc.add_heading("12.3 代表性用例详细步骤", level=2)
    detailed = [
        ("ENR-01 邀请码加入班级", [
            "教师在班级管理点击“重新生成邀请码”，记录新 join_code。",
            "学生登录后在加入班级页输入新邀请码并确认。",
            "检查响应 class_id、class_name、joined=true。",
            "教师刷新班级学生列表和加入状态。",
            "数据库检查同一 class_id+student_id 仅一条 enrollment，audit/outbox 存在。",
        ], "教师能看到该学生；班级人数只增加 1；重复请求不重复增加。"),
        ("CNT-01 发布章节并同步内容", [
            "教师创建草稿章节、知识点，并上传一份 ready 资料。",
            "将资料设为 students 可见，点击“发布章节”。",
            "学生打开对应课程学习页并刷新。",
            "检查只返回 published 章节、目标资料和已发布任务。",
            "将章节撤回后再次刷新学生页面。",
        ], "发布后可见，撤回后不可见；教师私有资料始终不下发。"),
        ("SUB-02 提交版本与幂等", [
            "学生打开已发布且未截止的任务。",
            "以 Idempotency-Key=A 提交代码，记录 submission_id/version。",
            "使用同一 key 和相同正文重试，确认返回同一 submission_id。",
            "使用 key=B 提交修改后的代码。",
            "教师查看提交列表并核对最新版本和历史版本。",
        ], "HTTP 重试不增版；真实重交生成 version+1；版本唯一约束无冲突。"),
        ("GRD-02 成绩发布闭环", [
            "教师打开学生提交，保存成绩草稿和维度评分。",
            "学生读取结果，确认看不到草稿成绩。",
            "教师点击“发布成绩”并发送学生可见反馈。",
            "学生打开通知并跳转提交详情。",
            "检查成绩、反馈、发布时间和审计记录。",
        ], "只有 published 成绩可见；通知所有者正确；他人无法读取。"),
        ("DSC-02 讨论结束门禁", [
            "教师发布讨论，学生成功回复一次。",
            "教师点击“结束讨论”。",
            "学生刷新详情，输入框和发送按钮禁用。",
            "绕过前端直接调用回复 API。",
            "检查 discussion_replies 行数未增加。",
        ], "后端拒绝 ended 状态回复，历史内容保持只读可见。"),
        ("AUTH-03 身份伪造防护", [
            "学生账号建立有效会话。",
            "请求中加入 X-User-Id: teacher-01，并调用教师课程接口。",
            "再尝试读取其他学生提交或成绩。",
            "检查响应、审计和安全告警。",
        ], "服务端只信任会话身份；所有越权请求失败且不泄露资源存在。"),
    ]
    for title, steps, expected in detailed:
        doc.add_heading(title, level=3)
        case_number_id = add_numbering_definition(doc, bullet=False)
        for step in steps:
            add_list_item(doc, step, case_number_id)
        add_callout(doc, "通过标准", expected, fill=LIGHT_GRAY)

    doc.add_heading("12.4 数据核查示例", level=2)
    add_code_block(doc, "-- 同一学生在同一班级只存在一条加入关系\nSELECT class_id, student_id, COUNT(*)\nFROM enrollments\nGROUP BY class_id, student_id\nHAVING COUNT(*) > 1;\n\n-- 已读记录不得重复\nSELECT announcement_id, user_id, COUNT(*)\nFROM announcement_reads\nGROUP BY announcement_id, user_id\nHAVING COUNT(*) > 1;")


def add_delivery_and_risk(doc, bullet_id, number_id):
    doc.add_heading("13. 实施计划与发布门禁", level=1)
    doc.add_heading("13.1 分阶段实施", level=2)
    add_table(doc, ["阶段", "范围", "完成标准"], [
        ("阶段 0：认证门禁", "统一 /auth 和 /me、Session/JWT、删除固定学生 ID、对象级权限", "AUTH/SEC 全部 P0 用例通过"),
        ("阶段 1：核心教学闭环", "入班、章节资料、任务、提交、评测、成绩学生读取", "教师→学生→教师→学生 E2E 通过"),
        ("阶段 2：互动闭环", "公告、讨论、图谱、通知、已读", "目标班级隔离和通知重试通过"),
        ("阶段 3：运营能力", "学情、预警、导出、消息实时化、性能优化", "口径、性能和恢复门槛通过"),
    ], [1800, 4400, 3160], font_size=8.6, header_fill=LIGHT_GREEN)

    doc.add_heading("13.2 发布步骤", level=2)
    release_number_id = add_numbering_definition(doc, bullet=False)
    for text in [
        "冻结 OpenAPI、数据库迁移、前端版本和本文档版本，完成评审签署。",
        "备份数据库并记录校验值；预生产执行 Alembic 升级和数据回填。",
        "部署后端认证与权限，运行负向越权测试，确认 X-User-Id 不再生效。",
        "发布教师端和学生端，执行 ENR、CNT、TSK、SUB、GRD 主闭环冒烟。",
        "检查审计、Outbox、通知、错误率、延迟和数据库事务失败。",
        "观察窗口结束后确认发布；异常时按代码、数据库和事件补偿方案回滚。",
    ]:
        add_list_item(doc, text, release_number_id)

    doc.add_heading("13.3 企业发布门禁", level=2)
    gates = [
        "P0 认证和对象级权限整改完成，无固定教师/学生身份。",
        "所有目标接口有 OpenAPI、请求响应示例、稳定业务码和负责人。",
        "数据库迁移、回填、回滚和备份恢复在预生产验证。",
        "跨端主闭环、权限负向、幂等、并发和事务回滚用例通过。",
        "隐藏测试、源代码、成绩和个人信息无越权泄露。",
        "审计、监控、告警、值班和故障升级链路可用。",
        "学生端移动端关键页面无按钮遮挡、溢出或错误状态缺失。",
    ]
    for text in gates:
        add_list_item(doc, text, bullet_id)

    doc.add_heading("14. 风险、决策与待办", level=1)
    add_table(doc, ["编号", "风险/决策", "等级", "处理", "责任角色"], [
        ("R-01", "客户端可伪造 X-User-Id", "极高", "生产前替换为 Session/JWT", "后端/安全"),
        ("R-02", "学生端接口硬编码 student-01/student-03", "极高", "统一当前会话上下文", "双端前端"),
        ("R-03", "学生读取成绩和图谱接口缺失", "高", "按第 7 章补齐并做对象权限", "后端"),
        ("R-04", "target_classes 使用 JSON", "中", "迁移 graph_targets 关系表", "后端/DBA"),
        ("R-05", "SQLite 并发写能力有限", "中", "压测；规模增长时迁移 PostgreSQL", "架构/运维"),
        ("R-06", "定时发布依赖请求触发", "中", "引入服务端调度任务", "后端/运维"),
        ("D-01", "同一业务对象双端共用表", "决策", "禁止教师端/学生端建立同义副本", "架构"),
        ("D-02", "通知采用 Outbox 最终一致", "决策", "主事务成功后可靠投递", "架构/后端"),
    ], [850, 2800, 900, 3050, 1760], font_size=8.0, header_fill=LIGHT_GREEN)

    doc.add_heading("14.1 待业务确认", level=2)
    for text in [
        "学生加入班级是自动通过还是教师审核；是否允许学生主动退出。",
        "任务截止后是否允许迟交、如何标记、是否影响评分。",
        "已发布成绩是否允许修改，是否需要学生申诉和二次发布。",
        "公告和图谱的目标范围是课程、班级还是指定学生。",
        "通知渠道只做站内信，还是同步邮件/短信/企业消息。",
        "源代码、成绩、访问日志和审计日志的保留期限。",
    ]:
        add_list_item(doc, text, bullet_id)

    doc.add_heading("15. 技术结论", level=1)
    add_callout(doc, "结论", "CodeTrack 已具备若干教师端与学生端共享后端的主链路，但当前身份模拟方式不满足真实双端上线要求。实施顺序必须先完成统一认证和对象级权限，再补齐成绩、公告、知识图谱等学生读取闭环，随后建设通知、统计和运营能力。双端应共享业务表、状态机、API 契约和审计证据，不应各自维护数据副本。", fill=LIGHT_GREEN)


def add_appendices(doc):
    heading = doc.add_heading("附录 A. 当前后端路由证据", level=1)
    heading.paragraph_format.page_break_before = True
    add_table(doc, ["模块", "已存在路由"], [
        ("班级", "POST /teacher/classes/{class_id}/join-code；POST /classes/{join_code}/join；GET /teacher/classes/{class_id}/students"),
        ("课程内容", "GET/POST /teacher/courses/{course_id}/chapters；PATCH /teacher/chapters/{chapter_id}；GET /student/courses/{course_id}/content"),
        ("任务提交", "GET/POST /teacher/tasks；POST /teacher/tasks/{task_id}/publish；GET /student/tasks；POST /student/tasks/{task_id}/submissions"),
        ("评分反馈", "GET /teacher/submissions；PUT .../grade；POST .../grade/publish；POST .../feedback"),
        ("讨论", "教师 discussions 创建/发布/结束；学生 discussions 列表/回复"),
        ("图谱", "教师 knowledge-graphs 创建、文件生成、保存、发布、删除"),
        ("通知分析", "PATCH /teacher/notifications/{id}；GET /teacher/analytics/overview"),
    ], [1800, 7560], font_size=8.4)

    doc.add_heading("附录 B. 前端联调交付清单", level=1)
    add_table(doc, ["交付物", "内容", "负责人", "状态"], [
        ("OpenAPI", "双端接口、DTO、错误码、鉴权、示例", "后端", "待补齐"),
        ("数据库迁移", "DDL、回填、回滚、索引和校验 SQL", "后端/DBA", "待补齐"),
        ("教师端映射", "按钮、接口、成功/失败状态、页面刷新", "教师端前端", "按本文执行"),
        ("学生端映射", "页面路由、接口、空/错/加载状态", "学生端前端", "待实现"),
        ("自动化测试", "单元、API、契约、E2E、安全", "测试/开发", "待补齐"),
        ("运维手册", "配置、监控、告警、备份、恢复、回滚", "运维", "待补齐"),
        ("安全评审", "认证、授权、隐私、上传、审计", "安全", "待签署"),
    ], [1900, 4060, 1800, 1600], font_size=8.5, header_fill=LIGHT_GREEN)

    doc.add_heading("附录 C. 建议错误码", level=1)
    add_table(doc, ["HTTP", "业务码", "说明", "前端动作"], [
        ("401", "AUTH_REQUIRED / TOKEN_EXPIRED", "未登录或会话过期", "刷新一次；失败跳登录"),
        ("403", "ROLE_FORBIDDEN", "角色不允许执行", "返回可访问首页"),
        ("404", "RESOURCE_NOT_FOUND", "不存在或无对象权限", "返回列表并提示资源不可用"),
        ("409", "STATE_CONFLICT", "状态机或版本冲突", "重新加载并保留用户输入"),
        ("409", "IDEMPOTENCY_CONFLICT", "同一 key 的请求正文不同", "生成新 key 后确认重试"),
        ("422", "VALIDATION_FAILED", "字段或业务前置条件不满足", "映射字段和操作提示"),
        ("429", "RATE_LIMITED", "登录/提交/下载过于频繁", "显示等待时间"),
        ("500", "INTERNAL_ERROR", "未处理异常", "显示 request_id 并允许重试"),
    ], [900, 2600, 3300, 2560], font_size=8.4)


def database_meta():
    connection = sqlite3.connect(DB_PATH)
    tables = [row[0] for row in connection.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'"
    )]
    rows = 0
    for table in tables:
        safe = table.replace('"', '""')
        rows += connection.execute(f'SELECT COUNT(*) FROM "{safe}"').fetchone()[0]
    connection.close()
    return {"tables": len(tables), "rows": rows}


def audit(doc):
    audit_document(doc)
    numbering = doc.part.numbering_part.element
    assert len(numbering.findall(qn("w:abstractNum"))) >= 2
    assert len(doc.inline_shapes) == 2
    assert any(p.style.name == "Heading 1" for p in doc.paragraphs)
    assert len(doc.tables) >= 35
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == 9360


def main():
    if not DB_PATH.exists():
        raise FileNotFoundError(DB_PATH)
    meta = database_meta()
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        set_document_header_footer(section)
    bullet_id = add_numbering_definition(doc, bullet=True)
    number_id = add_numbering_definition(doc, bullet=False)

    doc.core_properties.title = "CodeTrack 教师端与学生端协同后端接入设计说明（企业版）"
    doc.core_properties.subject = "教师端与学生端共享后端业务闭环、接口、数据、权限和测试基线"
    doc.core_properties.author = "CodeTrack 项目开发组"
    doc.core_properties.keywords = "CodeTrack, 教师端, 学生端, 后端接入, API, 数据库, 状态机, 测试"
    doc.core_properties.comments = f"{DOC_ID} {DOC_VERSION}"
    doc.core_properties.created = datetime(2026, 8, 16)
    doc.core_properties.modified = datetime(2026, 8, 16)

    add_cover(doc, meta)
    add_front_matter(doc)
    add_scope_and_principles(doc, bullet_id)
    add_baseline_and_architecture(doc, meta, bullet_id)
    add_capability_map(doc)
    add_interaction_mapping(doc)
    add_state_machines(doc)
    add_api_design(doc, bullet_id)
    add_database_design(doc, meta)
    add_security_and_frontend(doc, bullet_id)
    add_sync_and_sequence(doc, bullet_id)
    add_testing(doc, bullet_id, number_id)
    add_delivery_and_risk(doc, bullet_id, number_id)
    add_appendices(doc)
    audit(doc)
    doc.save(OUTPUT_PATH)
    print(json.dumps({
        "output": str(OUTPUT_PATH),
        "bytes": OUTPUT_PATH.stat().st_size,
        "tables": len(doc.tables),
        "paragraphs": len(doc.paragraphs),
        "database_tables": meta["tables"],
        "database_rows": meta["rows"],
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
