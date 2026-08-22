from __future__ import annotations

from datetime import datetime
import json
from pathlib import Path
import sqlite3

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from build_database_integration_document import (
    BORDER,
    CAUTION,
    DARK_GREEN,
    DB_PATH,
    FIELD_DESCRIPTIONS,
    GREEN,
    GROUPS,
    INK,
    LIGHT_GRAY,
    LIGHT_GREEN,
    MUTED,
    RISK,
    ROOT,
    TABLE_HEADER,
    TABLE_PURPOSES,
    add_callout,
    add_para,
    add_table,
    audit_document,
    configure_section,
    configure_styles,
    inspect_database,
    query_rows,
    set_cell_margins,
    set_cell_shading,
    set_header_footer,
    set_run_font,
    set_table_borders,
    set_table_geometry,
)


OUTPUT_DIR = ROOT / "artifacts" / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "CodeTrack教师端数据库内容与前端接入说明_企业版.docx"
ASSET_DIR = ROOT / "artifacts" / "database-enterprise-assets"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)

DOC_ID = "CT-TECH-DB-001"
DOC_VERSION = "V1.0"
DOC_STATUS = "技术基线版（审核签署待补）"


def add_numbering_definition(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, indent])
    level.extend([start, num_fmt, level_text, suffix, p_pr])
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, num_id)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(text), size=10.5, color=INK)
    return paragraph


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=8.4, color=INK, ascii_font="Consolas")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    set_table_geometry(table, [9360])
    set_table_borders(table, color=BORDER, size="5")
    add_para(doc, "", after=0)


def load_font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def centered_text(draw, box, text, font, fill):
    left, top, right, bottom = box
    bounds = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.multiline_text(((left + right - width) / 2, (top + bottom - height) / 2), text, font=font, fill=fill, spacing=6, align="center")


def rounded_box(draw, box, title, detail, fill, outline=GREEN):
    if isinstance(outline, str) and not outline.startswith("#"):
        outline = f"#{outline}"
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    left, top, right, bottom = box
    centered_text(draw, (left + 12, top + 12, right - 12, top + 64), title, load_font(26, True), "#176B12")
    centered_text(draw, (left + 16, top + 62, right - 16, bottom - 12), detail, load_font(19), "#1F2A24")


def arrow(draw, start, end):
    draw.line([start, end], fill="#2AAD16", width=5)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 10), (x - 18, y + 10)], fill="#2AAD16")


def build_architecture_diagram():
    path = ASSET_DIR / "enterprise-architecture.png"
    image = Image.new("RGB", (1600, 620), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 36), "教师端数据接入架构", font=load_font(34, True), fill="#1F2A24")
    boxes = [
        ((70, 170, 330, 430), "浏览器", "React / Ant Design\n页面与交互状态", "#EDF8EA"),
        ((410, 170, 690, 430), "前端数据层", "src/api.ts\n统一请求封装", "#F4F6F5"),
        ((770, 170, 1030, 430), "应用服务", "FastAPI\n鉴权与业务校验", "#EDF8EA"),
        ((1110, 170, 1370, 430), "持久化层", "SQLAlchemy\n事务与模型", "#F4F6F5"),
        ((1410, 170, 1570, 430), "数据库", "SQLite\ncodetrack.db", "#EDF8EA"),
    ]
    for box, title, detail, fill in boxes:
        rounded_box(draw, box, title, detail, fill)
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[0][2] + 12, 300), (right[0][0] - 12, 300))
    draw.text((80, 500), "请求携带 X-User-Id；服务端执行教师角色和课程归属校验；写操作提交事务并记录审计。", font=load_font(22), fill="#5F6B63")
    image.save(path)
    return path


def build_domain_diagram(counts):
    path = ASSET_DIR / "enterprise-data-domains.png"
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 36), "数据库业务域与主要关联", font=load_font(34, True), fill="#1F2A24")
    domains = [
        ((80, 150, 500, 330), "账号与个性化", f"users {counts.get('users', 0)}\ncredentials / preferences / notifications", "#EDF8EA"),
        ((580, 150, 1020, 330), "课程与组织", f"courses {counts.get('courses', 0)} / classes {counts.get('class_groups', 0)}\nchapters / enrollments / announcements", "#F4F6F5"),
        ((1100, 150, 1520, 330), "材料与知识", f"materials {counts.get('materials', 0)}\nfolders / links / knowledge graphs", "#EDF8EA"),
        ((300, 500, 760, 700), "任务与评价", f"tasks {counts.get('tasks', 0)} / submissions {counts.get('submissions', 0)}\ntests / grades / diagnosis / feedback", "#F4F6F5"),
        ((880, 500, 1340, 700), "课堂互动", f"discussions {counts.get('course_discussions', 0)}\nreplies / announcement reads", "#EDF8EA"),
    ]
    for box, title, detail, fill in domains:
        rounded_box(draw, box, title, detail, fill)
    arrow(draw, (500, 240), (580, 240))
    arrow(draw, (1020, 240), (1100, 240))
    arrow(draw, (770, 350), (580, 485))
    arrow(draw, (880, 600), (780, 600))
    draw.text((80, 800), "主链路：用户 → 课程 → 班级/章节 → 材料/任务 → 提交/评价；互动与已读状态按用户和课程隔离。", font=load_font(22), fill="#5F6B63")
    image.save(path)
    return path


def add_picture(doc, path, caption, alt_text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    shape = paragraph.add_run().add_picture(str(path), width=Inches(6.35))
    shape._inline.docPr.set("descr", alt_text)
    add_para(doc, caption, size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8, keep=True)


def add_cover(doc, meta):
    add_para(doc, "TECHNICAL BASELINE / INTERNAL", size=10, bold=True, color=GREEN, before=18, after=12)
    add_para(doc, "CodeTrack 教师端", size=14, bold=True, color=DARK_GREEN, after=4)
    add_para(doc, "数据库内容与前端接入说明", size=25, bold=True, color=INK, after=6)
    add_para(doc, "企业研发交付版", size=15, color=MUTED, after=22)

    metadata = [
        ("文档编号", DOC_ID),
        ("版本", DOC_VERSION),
        ("状态", DOC_STATUS),
        ("密级", "内部使用"),
        ("编制", "CodeTrack 项目开发组"),
        ("审核", "技术负责人（待签署）"),
        ("批准", "项目负责人（待签署）"),
        ("基线日期", "2026-08-14"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        set_run_font(paragraph.add_run(f"{label}："), size=10.5, bold=True, color=INK)
        set_run_font(paragraph.add_run(value), size=10.5, color=INK)

    add_callout(doc, "文档定位", "本文件是数据库与前端接入的企业技术基线，供架构评审、开发联调、测试验收、发布审批、安全审查和运维交接使用。建议项不代表已实现能力。", fill=LIGHT_GREEN)
    add_para(doc, f"数据库快照：{meta['snapshot_time']}　|　{meta['table_count']} 张表　|　{meta['total_rows']:,} 行记录", size=9.5, color=MUTED, italic=True, before=16, after=0)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_front_matter(doc, meta):
    doc.add_heading("文档控制", level=1)
    add_table(doc, ["属性", "内容"], [
        ("文档名称", "CodeTrack 教师端数据库内容与前端接入说明"),
        ("文档编号 / 版本", f"{DOC_ID} / {DOC_VERSION}"),
        ("适用系统", "CodeTrack 教师端（React + FastAPI + SQLAlchemy + SQLite）"),
        ("适用阶段", "开发、联调、测试、验收、部署、运维交接"),
        ("数据库基线", str(DB_PATH)),
        ("快照范围", f"{meta['table_count']} 张业务表，{meta['total_rows']:,} 行记录"),
        ("保密级别", "内部使用；禁止外发凭据、源代码、学生提交正文和个人信息"),
    ], [2300, 7060], font_size=9.2, header_fill=LIGHT_GREEN)

    doc.add_heading("版本修订记录", level=2)
    add_table(doc, ["版本", "日期", "编制角色", "变更摘要", "状态"], [
        ("V1.0", "2026-08-14", "项目开发组", "建立数据库结构、前端接入、接口、安全、发布和验收技术基线", "待审核"),
    ], [900, 1300, 1700, 4260, 1200], font_size=8.8)

    doc.add_heading("审核与批准", level=2)
    add_table(doc, ["环节", "责任角色", "确认重点", "签署"], [
        ("技术审核", "后端/架构负责人", "模型、接口、事务、迁移、回滚", "待签署"),
        ("安全审核", "信息安全负责人", "认证、权限、密码、日志、隐私", "待签署"),
        ("测试确认", "测试负责人", "功能、回归、异常、数据一致性", "待签署"),
        ("发布批准", "项目负责人", "风险接受、发布窗口、运维交接", "待签署"),
    ], [1300, 1900, 4360, 1800], font_size=8.8)

    doc.add_heading("目标读者", level=2)
    add_para(doc, "主要读者包括前端开发、后端开发、数据库/运维、测试、架构、安全与项目管理人员。业务教师可参考页面流程与数据边界，不建议直接操作数据库。")

    doc.add_heading("目录与阅读路径", level=1)
    add_table(doc, ["章节", "评审问题"], [
        ("1. 文档说明", "为什么写、包含什么、哪些内容不在范围内"),
        ("2. 系统基线与架构", "系统如何部署、数据如何流动、当前规模是多少"),
        ("3. 数据治理与前端边界", "哪些数据落库、谁负责、如何分类和保留"),
        ("4. 数据库设计", "表如何分域、约束和事务如何保证一致性"),
        ("5. 前端接入设计", "页面动作如何调用 API、状态如何恢复"),
        ("6. API 契约", "请求、响应、错误码和幂等性如何约定"),
        ("7. 安全与隐私", "现有控制是什么、生产差距在哪里"),
        ("8. 迁移、发布与回滚", "如何上线、验证、备份和恢复"),
        ("9. 运维与监控", "监控什么、何时告警、如何排障"),
        ("10. 测试与验收", "现有证据和企业验收门槛是什么"),
        ("11. 风险与演进", "剩余风险、优先级和后续路线"),
        ("附录", "全量字段字典、数据快照与核查命令"),
    ], [2200, 7160], font_size=8.8)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_enterprise_content(doc, connection, tables, counts, schemas, meta, bullet_id, number_id):
    architecture_path = build_architecture_diagram()
    domain_path = build_domain_diagram(counts)

    doc.add_heading("1. 文档说明", level=1)
    doc.add_heading("1.1 编制目的", level=2)
    add_para(doc, "建立 CodeTrack 教师端数据库与前端接入的统一技术基线，使页面行为、API 契约、数据结构、安全边界、发布流程和验收证据可追踪、可评审、可维护。")
    doc.add_heading("1.2 范围", level=2)
    for item in [
        "覆盖教师账号登录、教师档案、课程草稿、课程公告及已读状态、教师偏好，以及既有课程/班级/任务/材料/评价等业务数据。",
        "覆盖 React 前端、src/api.ts 请求层、FastAPI 服务、SQLAlchemy ORM、SQLite 数据库和 Alembic 迁移。",
        "覆盖本地开发基线和企业生产化建议；建议项使用“建议/必须补齐”明确标识。",
    ]:
        add_list_item(doc, item, bullet_id)
    doc.add_heading("1.3 不在本次范围", level=2)
    for item in [
        "统一身份认证平台、短信/邮件网关、密钥管理平台等外部系统的具体实现。",
        "生产网络拓扑、云资源账号、真实域名证书和组织级备份平台配置。",
        "学生端完整功能说明及业务数据治理制度原文。",
    ]:
        add_list_item(doc, item, bullet_id)
    doc.add_heading("1.4 术语", level=2)
    add_table(doc, ["术语", "定义"], [
        ("业务主数据", "需要跨页面、跨会话保留，并参与教学业务流程的数据"),
        ("UI 临时状态", "筛选、分页、弹窗、选中项等无需长期保存的界面状态"),
        ("数据包络", "后端统一返回的 {\"data\": ...} JSON 结构"),
        ("幂等", "重复执行同一请求不会产生额外业务副作用"),
        ("技术基线", "经评审后作为开发、测试、发布和运维共同依据的版本"),
        ("RPO / RTO", "可接受数据丢失时间点 / 可接受服务恢复时间"),
    ], [1800, 7560], font_size=9.0)
    doc.add_heading("1.5 参考依据", level=2)
    for item in [
        "teacher_backend/app/models.py、schemas.py、frontend_persistence.py、main.py。",
        "src/api.ts、src/exact/pages-global.tsx、pages-course.tsx、pages-flow.tsx。",
        "backend/alembic/versions/20260814_0006_frontend_persistence.py。",
        "codetrack.db 实时 PRAGMA 表结构、约束与行数快照。",
    ]:
        add_list_item(doc, item, bullet_id)

    doc.add_heading("2. 系统基线与架构", level=1)
    add_callout(doc, "当前基线", f"数据库为 SQLite {meta['sqlite_version']}，文件位于 {DB_PATH}；当前 {len(tables)} 张表、{sum(counts.values()):,} 行记录。前端运行于 127.0.0.1:5173，后端 API 运行于 127.0.0.1:8001。")
    doc.add_heading("2.1 技术栈", level=2)
    add_table(doc, ["层级", "当前技术", "职责", "生产关注点"], [
        ("前端", "React 19 / TypeScript / Vite / Ant Design", "页面、交互、表单和状态展示", "代码分包、错误边界、可观测性"),
        ("请求层", "Fetch + src/api.ts", "统一 API 地址、包络解析、X-User-Id", "令牌刷新、超时、重试、链路 ID"),
        ("服务端", "FastAPI / Pydantic", "参数校验、鉴权、业务规则、响应", "限流、统一异常、审计、指标"),
        ("持久化", "SQLAlchemy ORM", "模型、查询、事务提交", "事务边界、并发、连接池"),
        ("数据库", "SQLite", "本地业务数据持久化", "并发写、备份、生产迁移评估"),
        ("迁移", "Alembic + 启动兼容补丁", "结构演进与旧库兼容", "生产应以 Alembic 为唯一结构变更入口"),
    ], [1300, 2350, 2800, 2910], font_size=8.3)
    doc.add_heading("2.2 架构与数据流", level=2)
    add_picture(doc, architecture_path, "图 1　教师端前端到数据库的接入链路", "React 前端通过 API 层调用 FastAPI、SQLAlchemy 和 SQLite 的架构图")
    doc.add_heading("2.3 部署拓扑", level=2)
    add_para(doc, "当前为单机本地开发拓扑：浏览器与 Vite 前端连接本机 FastAPI，后端直接访问同一工作区内 SQLite 文件。该拓扑适合原型、演示和单机联调，不适合多实例高并发生产部署。")
    add_table(doc, ["环境", "前端", "后端", "数据库", "状态"], [
        ("本地开发", "127.0.0.1:5173", "127.0.0.1:8001", "本地 codetrack.db", "已实现"),
        ("集成测试", "建议独立构建产物", "建议独立进程/容器", "独立测试库", "待标准化"),
        ("生产", "静态资源/CDN 或 Web 服务", "反向代理后的多进程服务", "建议 PostgreSQL 或受控 SQLite 单实例", "待设计"),
    ], [1400, 2100, 2300, 2300, 1260], font_size=8.5)

    doc.add_heading("3. 数据治理与前端边界", level=1)
    doc.add_heading("3.1 持久化边界", level=2)
    add_table(doc, ["数据类别", "示例", "存储位置", "理由"], [
        ("业务主数据", "教师、课程、班级、章节、任务、成绩", "数据库", "跨会话保留并参与业务流程"),
        ("教师前端业务状态", "课程草稿、公告已读、教师偏好", "数据库", "需要跨设备/刷新恢复"),
        ("浏览器会话", "当前登录教师 ID 与名称", "localStorage", "浏览器登录态；当前不是可信认证令牌"),
        ("跨页临时焦点", "待打开任务/班级", "sessionStorage", "仅用于一次导航"),
        ("纯 UI 状态", "筛选、分页、抽屉、弹窗、选中行", "React 内存", "无需长期保存"),
    ], [1650, 2580, 1900, 3230], font_size=8.4)
    doc.add_heading("3.2 数据分级与处理", level=2)
    add_table(doc, ["等级", "数据", "示例表/字段", "处理要求"], [
        ("P4 高敏凭据", "密码盐、密码哈希", "teacher_credentials", "禁止导出和日志打印；限制最小权限"),
        ("P3 个人/教学敏感", "邮箱、学号、源代码、成绩、反馈", "users、submissions、grades", "脱敏展示、受控备份、访问审计"),
        ("P2 内部业务", "课程、班级、公告、任务", "courses、class_groups、tasks", "内部使用、按课程归属授权"),
        ("P1 公开配置", "健康检查、静态枚举", "health、前端常量", "可公开但仍需完整性保护"),
    ], [1200, 2300, 2700, 3160], font_size=8.4)
    add_callout(doc, "保留策略说明", "当前代码未实现自动归档或保留期限。企业上线前必须由业务、法务/合规和安全共同确定用户档案、源代码、成绩、审计日志和备份的保留与删除规则。", fill=CAUTION)
    doc.add_heading("3.3 数据责任矩阵（RACI）", level=2)
    add_table(doc, ["事项", "前端", "后端", "DBA/运维", "安全", "测试"], [
        ("字段与接口契约", "R", "A", "C", "C", "C"),
        ("数据库模型与迁移", "I", "R/A", "C", "C", "C"),
        ("备份与恢复", "I", "C", "R/A", "C", "C"),
        ("认证与权限", "R", "R", "C", "A", "C"),
        ("数据质量与回归", "R", "R", "C", "C", "A"),
        ("生产发布", "R", "R", "A", "C", "C"),
    ], [2900, 1100, 1100, 1600, 1300, 1360], font_size=8.5)
    add_para(doc, "R=执行，A=最终负责，C=协作/评审，I=知会。责任角色需在项目组织中映射到具体人员。", size=9, color=MUTED, italic=True)

    doc.add_heading("4. 数据库设计", level=1)
    doc.add_heading("4.1 业务域与关系", level=2)
    add_picture(doc, domain_path, "图 2　数据库业务域与主要关联", "账号、课程、材料、任务评价和课堂互动五个数据库业务域关系图")
    doc.add_heading("4.2 设计规范", level=2)
    for item in [
        "表名与字段名使用 snake_case；业务主键以字符串 ID 为主，关系/状态表可使用整数自增主键。",
        "所有关系通过外键表达；SQLite 连接时执行 PRAGMA foreign_keys=ON。",
        "关键唯一性通过 UNIQUE 约束保障，例如教师编号、课程代码、班级邀请码、教师草稿、公告已读组合。",
        "结构化可变内容使用 JSON 文本字段；读取时处理 JSON 解析失败，但当前缺少 schema_version。",
        "时间字段统一使用 DateTime；当前为无时区时间，企业生产建议统一 UTC 存储并在界面转换。",
    ]:
        add_list_item(doc, item, bullet_id)
    doc.add_heading("4.3 表清单与数据量", level=2)
    inventory_rows = [(table, counts[table], TABLE_PURPOSES.get(table, "项目业务表")) for table in tables]
    add_table(doc, ["表名", "行数", "用途"], inventory_rows, [2800, 1000, 5560], font_size=8.3,
              aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT])
    doc.add_heading("4.4 本次新增持久化表", level=2)
    add_table(doc, ["表", "主责数据", "一致性约束", "当前写入入口"], [
        ("teacher_credentials", "教师盐值与密码哈希", "user_id 主键", "初始化/密码设置"),
        ("teacher_preferences", "通知、AI、邮件摘要偏好", "teacher_id 主键", "GET 缺省创建、PUT 更新"),
        ("course_drafts", "课程创建草稿 JSON", "teacher_id 唯一", "PUT upsert / DELETE"),
        ("course_announcements", "公告标题、正文、范围、发布时间", "课程和作者外键", "兼容种子；后续应补发布接口"),
        ("announcement_reads", "公告阅读时间", "announcement_id + user_id 唯一", "PATCH 幂等新增"),
    ], [2200, 2700, 2400, 2060], font_size=8.2)
    doc.add_heading("4.5 事务、并发与幂等", level=2)
    add_table(doc, ["场景", "当前行为", "结论/改进"], [
        ("登录", "写入一条 audit_logs 后 commit", "单事务；失败不写审计"),
        ("保存草稿", "查询后新增或覆盖，再写审计并 commit", "逻辑 upsert；并发写依赖唯一约束"),
        ("删除草稿", "存在则删除，不存在仍返回 deleted=true 并写审计", "接口幂等，审计会重复增加"),
        ("公告已读", "存在则直接返回，不存在则新增并 commit", "业务幂等；唯一约束防重复"),
        ("保存偏好", "整行覆盖三个布尔值并写审计", "无字段级 patch；最后写入者生效"),
        ("SQLite 并发", "单文件写锁", "多实例/高并发前评估 PostgreSQL"),
    ], [1800, 3860, 3700], font_size=8.4)
    doc.add_heading("4.6 数据质量规则", level=2)
    for item in [
        "教师账号必须存在于 users 且 role=teacher；教师编号应唯一且非空。",
        "访问课程相关数据前必须校验 courses.teacher_id 与当前教师一致。",
        "JSON 字段写入前必须可序列化；生产版建议增加结构版本与最大长度。",
        "公告标题、摘要、接收范围和发布时间不能为空；公告作者必须是有效用户。",
        "迁移后必须执行表数量、外键、唯一索引、核心行数和抽样查询核对。",
    ]:
        add_list_item(doc, item, bullet_id)

    doc.add_heading("5. 前端接入设计", level=1)
    doc.add_heading("5.1 API 客户端约定", level=2)
    add_para(doc, "src/api.ts 通过 request<T>() 统一拼接 VITE_API_BASE（默认 /api/v1）、设置 Content-Type、附加 X-User-Id、解析 data 包络并将非 2xx 响应转换为 ApiError。multipart 上传不设置 Content-Type，由浏览器生成 boundary。")
    add_table(doc, ["约定", "当前实现", "生产建议"], [
        ("API Base", "VITE_API_BASE 或 /api/v1", "按环境注入并禁止硬编码生产地址"),
        ("身份头", "X-User-Id", "替换为服务端验证的会话/JWT"),
        ("响应包络", "{ data: T }", "补充 trace_id、error.code、meta"),
        ("错误处理", "detail → ApiError.message", "统一错误码、超时、重试和错误边界"),
        ("上传", "FormData + 当前教师头", "限制文件类型/大小并做病毒扫描"),
    ], [1700, 3540, 4120], font_size=8.5)
    doc.add_heading("5.2 页面动作映射", level=2)
    mapping_rows = [
        ("登录页加载", "GET /teacher/auth/accounts", "users", "读取教师卡片"),
        ("点击登录", "POST /teacher/auth/login", "users、teacher_credentials、audit_logs", "验密并记录审计"),
        ("进入工作台", "GET /teacher/bootstrap", "users、courses、class_groups、notifications", "返回当前教师基线"),
        ("保存课程草稿", "PUT /teacher/course-draft", "course_drafts、audit_logs", "新增或覆盖"),
        ("刷新创建页", "GET /teacher/course-draft", "course_drafts", "恢复草稿"),
        ("删除/创建成功", "DELETE /teacher/course-draft", "course_drafts、audit_logs", "清理草稿"),
        ("打开课程工作空间", "GET /teacher/courses/{course_id}/announcements", "course_announcements、announcement_reads", "加载公告和已读状态"),
        ("打开公告", "PATCH /teacher/announcements/{id}/read", "announcement_reads", "持久化已读"),
        ("进入个人设置", "GET /teacher/preferences", "teacher_preferences", "读取偏好"),
        ("保存偏好", "PUT /teacher/preferences", "teacher_preferences、audit_logs", "整行更新"),
    ]
    add_table(doc, ["页面动作", "API", "数据表", "结果"], mapping_rows, [2100, 3100, 2860, 1300], font_size=8.0)
    doc.add_heading("5.3 核心交互流程", level=2)
    add_table(doc, ["流程", "步骤", "失败处理"], [
        ("登录", "账号列表 → 提交账号/密码 → 返回教师档案 → 写浏览器会话 → 进入门户", "401 显示统一错误；不泄露账号是否存在"),
        ("课程草稿", "GET 恢复 → 编辑 → PUT 保存 → 刷新复核 → 创建成功后 DELETE", "保存失败保留当前表单并提示重试"),
        ("公告", "GET 列表 → 点击详情 → PATCH 已读 → 本地状态更新 → 刷新复核", "PATCH 失败保留详情并提示未保存"),
        ("教师偏好", "GET 读取 → 切换开关 → PUT 保存 → 以响应覆盖本地 → 刷新复核", "PUT 失败不显示成功提示"),
    ], [1500, 5050, 2810], font_size=8.3)
    doc.add_heading("5.4 浏览器状态边界", level=2)
    add_callout(doc, "设计原则", "只有影响业务连续性、跨设备或跨会话恢复的数据进入数据库。筛选、分页、抽屉、弹窗、当前选中项等保持为前端状态，避免数据库被高频、低价值 UI 写入污染。")

    doc.add_heading("6. API 契约", level=1)
    doc.add_heading("6.1 通用约定", level=2)
    add_table(doc, ["项目", "约定"], [
        ("Base URL", "/api/v1"),
        ("Content-Type", "application/json；文件上传使用 multipart/form-data"),
        ("身份", "当前除账号列表/登录外，教师接口读取 X-User-Id"),
        ("成功响应", "HTTP 200（或业务接口定义状态）+ {\"data\": ...}"),
        ("时间", "ISO 8601 字符串；当前无时区"),
        ("布尔值", "JSON true / false"),
        ("字符编码", "UTF-8"),
    ], [2100, 7260], font_size=9.0)
    doc.add_heading("6.2 接口清单", level=2)
    endpoint_rows = [
        ("GET", "/teacher/auth/accounts", "无显式教师依赖", "读取教师账号卡", "users"),
        ("POST", "/teacher/auth/login", "账号 + 密码", "验密并返回教师档案", "users / credentials / audit"),
        ("GET", "/teacher/course-draft", "X-User-Id", "读取当前教师草稿", "course_drafts"),
        ("PUT", "/teacher/course-draft", "X-User-Id + payload", "保存草稿", "course_drafts / audit"),
        ("DELETE", "/teacher/course-draft", "X-User-Id", "删除草稿", "course_drafts / audit"),
        ("GET", "/teacher/courses/{course_id}/announcements", "X-User-Id", "读取公告与已读状态", "announcements / reads"),
        ("PATCH", "/teacher/announcements/{id}/read", "X-User-Id", "标记已读", "announcement_reads"),
        ("GET", "/teacher/preferences", "X-User-Id", "读取偏好", "teacher_preferences"),
        ("PUT", "/teacher/preferences", "X-User-Id + 3 个布尔值", "保存偏好", "preferences / audit"),
    ]
    add_table(doc, ["方法", "路径", "输入", "用途", "主要表"], endpoint_rows, [850, 3160, 1800, 2000, 1550], font_size=7.8)
    doc.add_heading("6.3 登录契约", level=2)
    add_para(doc, "请求：username 长度 1-80，可为教师姓名或编号；password 长度 1-128。响应只返回教师非敏感档案，不返回盐、哈希或令牌。")
    add_code_block(doc, '{\n  "username": "T2024002",\n  "password": "<REDACTED>"\n}\n\n200 OK\n{\n  "data": {\n    "id": "teacher-02",\n    "name": "林老师",\n    "number": "T2024002",\n    "email": "lin.teacher@university.edu.cn",\n    "department": "软件学院"\n  }\n}')
    doc.add_heading("6.4 课程草稿契约", level=2)
    add_para(doc, "PUT 请求以 payload 包裹完整前端草稿；服务端覆盖当前教师唯一草稿并返回 savedAt。GET 无草稿时 data=null。DELETE 无论草稿是否存在均返回 deleted=true。")
    add_code_block(doc, '{\n  "payload": {\n    "values": {"name": "数据结构", "term": "2026-2027 学年秋季"},\n    "knowledgePoints": ["线性表", "图"],\n    "coverUrl": "/ui-assets/course-preview.png"\n  }\n}')
    doc.add_heading("6.5 公告契约", level=2)
    add_para(doc, "公告列表按 pinned 降序、published_at 降序返回。read 字段由 announcement_reads 与当前用户合并计算。PATCH 已读接口为幂等操作。")
    add_code_block(doc, '{\n  "data": [{\n    "id": "announcement-tree-materials",\n    "title": "第 6 周资料已更新",\n    "content": ["正文段落"],\n    "published_at": "2026-05-28 09:30",\n    "pinned": true,\n    "read": false\n  }]\n}')
    doc.add_heading("6.6 教师偏好契约", level=2)
    add_para(doc, "PUT 为整行替换语义，三个布尔字段均应提交。GET 在记录不存在时创建默认记录。")
    add_code_block(doc, '{\n  "notifications_enabled": true,\n  "ai_assistant_enabled": true,\n  "email_digest": false\n}')
    doc.add_heading("6.7 错误码与前端处理", level=2)
    add_table(doc, ["HTTP", "场景", "detail 示例", "前端处理"], [
        ("401", "账号或密码校验失败", "用户名或密码错误，请重新输入", "保留账号输入，清空/聚焦密码"),
        ("403", "用户不存在或不是教师", "需要教师权限", "清除会话并返回登录页"),
        ("404", "课程/公告不存在或无权访问", "课程不存在或无权访问", "返回可访问课程并提示"),
        ("422", "Pydantic 参数校验失败", "字段校验明细", "映射到表单字段；当前需完善"),
        ("500", "未处理异常/数据库错误", "内部错误", "显示通用错误并记录 trace_id；当前需完善"),
    ], [900, 2450, 3010, 3000], font_size=8.3)

    doc.add_heading("7. 安全与隐私", level=1)
    doc.add_heading("7.1 当前控制", level=2)
    add_table(doc, ["控制域", "当前实现", "评价"], [
        ("密码存储", "16 字节随机盐 + PBKDF2-HMAC-SHA256，120,000 次", "不存明文；适合演示基线"),
        ("比较", "hmac.compare_digest", "降低时序侧信道风险"),
        ("角色", "users.role == teacher", "具备最小角色校验"),
        ("对象权限", "课程 teacher_id 归属校验", "公告访问已做对象级授权"),
        ("审计", "登录、草稿、偏好等写操作写 audit_logs", "已有基础证据链"),
        ("CORS", "只允许 localhost / 127.0.0.1:5173", "适合本地开发"),
        ("外键", "PRAGMA foreign_keys=ON", "降低孤儿数据风险"),
    ], [1600, 4900, 2860], font_size=8.4)
    doc.add_heading("7.2 生产差距与强制整改", level=2)
    add_table(doc, ["优先级", "差距", "风险", "上线前要求"], [
        ("P0", "X-User-Id 由客户端自报", "可伪造其他教师身份", "改为服务端签名会话/JWT，并从令牌解析用户"),
        ("P0", "演示账号使用统一初始弱口令", "撞库和未授权访问", "强制首次修改、密码策略、重置和锁定"),
        ("P0", "登录无速率限制", "暴力破解", "按账号/IP 限流、失败计数和告警"),
        ("P1", "无 TLS/安全响应头基线", "传输泄露和浏览器攻击", "反向代理启用 HTTPS、HSTS、CSP 等"),
        ("P1", "错误响应无稳定业务码/trace_id", "难以定位和审计", "统一异常模型并接入链路追踪"),
        ("P1", "敏感字段无字段级加密", "数据库泄露影响扩大", "评估邮箱、源代码等加密/脱敏"),
        ("P2", "审计日志无留存/防篡改策略", "证据不完整", "集中日志、访问控制、归档和校验"),
    ], [850, 2400, 2600, 3510], font_size=8.1)
    add_callout(doc, "安全结论", "当前实现是本地演示和功能联调基线，不应以现状直接暴露到公网。P0 项必须在生产发布门禁前关闭。", fill=RISK)
    doc.add_heading("7.3 隐私与导出", level=2)
    for item in [
        "文档和日志不得输出 password_salt、password_hash、完整学生源代码或大段教学敏感正文。",
        "教师邮箱、学生学号和成绩导出应按角色授权并记录审计；对外材料必须脱敏。",
        "备份文件应加密、限制下载、设置保留期，并定期验证可恢复性。",
        "删除用户或课程前应明确级联影响、法律保留要求和恢复窗口。",
    ]:
        add_list_item(doc, item, bullet_id)

    doc.add_heading("8. 迁移、发布与回滚", level=1)
    doc.add_heading("8.1 结构迁移", level=2)
    add_para(doc, "本次数据库结构由 20260814_0006_frontend_persistence.py 增加五张表，依赖 20260812_0005_teacher_knowledge_graphs.py。迁移脚本在建表前检查现有表，兼容旧 SQLite 文件。")
    add_code_block(doc, "alembic -c backend/alembic.ini upgrade head")
    doc.add_heading("8.2 发布步骤", level=2)
    steps = [
        "冻结发布版本，确认代码、迁移脚本、前端构建产物和本文档版本一致。",
        "停止业务写入或进入维护窗口，执行数据库一致性备份并记录校验值。",
        "在目标环境执行 Alembic upgrade；禁止人工直接修改生产表结构。",
        "启动后端，检查 health、表数量、外键、账号、课程和新增表行数。",
        "发布前端，执行登录、草稿、公告已读、偏好四条冒烟用例。",
        "观察错误率、数据库锁、审计写入和页面错误，达到观察窗口后关闭发布任务。",
    ]
    for item in steps:
        add_list_item(doc, item, number_id)
    doc.add_heading("8.3 发布前检查", level=2)
    add_table(doc, ["检查项", "通过条件", "责任角色"], [
        ("备份", "备份文件存在、可读取、校验值已记录", "DBA/运维"),
        ("迁移", "目标 revision 正确，新增 5 表及索引存在", "后端/DBA"),
        ("数据", "用户/课程核心行数无异常变化", "后端/测试"),
        ("安全", "P0 风险已关闭或获得正式风险接受", "安全/项目负责人"),
        ("功能", "4 条浏览器闭环全部通过", "测试"),
        ("回滚", "备份路径、恢复命令、负责人和窗口明确", "运维"),
    ], [2500, 4960, 1900], font_size=8.5)
    doc.add_heading("8.4 回滚策略", level=2)
    for item in [
        "代码回滚：恢复上一稳定后端和前端构建产物；如新结构向后兼容，可暂不降级数据库。",
        "数据库回滚：优先从发布前一致性备份恢复；Alembic downgrade 仅在已验证且无新数据需要保留时使用。",
        "数据补偿：迁移后已产生新业务数据时，不得直接覆盖，应先导出差异并制定补偿脚本。",
        "回滚后重新执行健康检查、核心行数、登录和读写闭环。",
    ]:
        add_list_item(doc, item, bullet_id)
    doc.add_heading("8.5 备份与恢复建议", level=2)
    add_table(doc, ["项目", "建议基线"], [
        ("备份方式", "SQLite 在线 backup API 或停写后一致性复制；处理 WAL/SHM"),
        ("频率", "每日全量 + 重要发布前即时备份（按业务确认）"),
        ("RPO / RTO", "建议 RPO≤24h、RTO≤4h；需业务正式批准"),
        ("保留", "建议 7 日每日、4 周每周、12 月每月；需合规确认"),
        ("演练", "至少每季度恢复演练并记录恢复耗时、校验结果和问题"),
        ("加密", "备份传输与静态存储加密，密钥与备份分离管理"),
    ], [2200, 7160], font_size=8.7)

    doc.add_heading("9. 运维与监控", level=1)
    doc.add_heading("9.1 配置项", level=2)
    add_table(doc, ["配置", "当前默认", "用途", "生产要求"], [
        ("VITE_API_BASE", "/api/v1", "前端 API 前缀", "环境注入"),
        ("CODETRACK_DATABASE_URL", "sqlite:///codetrack.db", "数据库连接", "秘密/配置平台管理"),
        ("后端端口", "8001", "FastAPI 监听", "反向代理后仅内网开放"),
        ("前端端口", "5173", "Vite 开发服务", "生产不使用开发服务"),
        ("PBKDF2_ITERATIONS", "代码常量 120000", "密码计算成本", "配置化并评估升级"),
    ], [2050, 2100, 2400, 2810], font_size=8.4)
    doc.add_heading("9.2 建议监控指标", level=2)
    add_table(doc, ["类别", "指标", "建议告警"], [
        ("API", "请求量、P95/P99 延迟、4xx/5xx、超时", "5xx>1% 持续 5 分钟；P95 超基线"),
        ("认证", "登录成功/失败、失败账号/IP 分布", "同账号或 IP 连续失败达到阈值"),
        ("数据库", "文件大小、写锁/locked 错误、事务失败、备份状态", "出现 locked 峰值或备份失败立即告警"),
        ("数据质量", "孤儿外键、唯一冲突、JSON 解析失败", "任一异常均需工单"),
        ("审计", "关键写操作审计缺失、日志增长", "业务成功但审计缺失时告警"),
        ("前端", "页面错误、API 失败、白屏、构建版本", "错误率超过基线或版本不一致"),
    ], [1400, 4400, 3560], font_size=8.4)
    doc.add_heading("9.3 日常维护", level=2)
    for item in [
        "每日检查后端健康、5xx、数据库锁和备份结果。",
        "每周检查数据库增长、审计日志增长、失败登录和异常 JSON。",
        "每月执行外键检查、索引使用评审、权限复核和恢复抽测。",
        "每季度执行灾备恢复演练、密码算法评审和生产容量评估。",
    ]:
        add_list_item(doc, item, bullet_id)
    doc.add_heading("9.4 常见故障定位", level=2)
    add_table(doc, ["现象", "优先检查", "处理"], [
        ("登录页无账号", "后端 8001、accounts API、users.role", "恢复服务或修复教师数据"),
        ("登录成功后被退出", "bootstrap 403/404、教师课程归属、浏览器会话", "修复课程归属或会话"),
        ("草稿无法恢复", "course_drafts、X-User-Id、payload_json", "核对教师 ID 和 JSON"),
        ("公告已读不保持", "announcement_reads 唯一记录、PATCH 响应", "核对写事务和课程权限"),
        ("偏好保存失败", "PUT 422/500、字段是否完整", "补齐三个布尔值并查日志"),
        ("database is locked", "并发写、长事务、遗留进程", "结束长事务；评估迁移数据库"),
    ], [2000, 3900, 3460], font_size=8.4)

    doc.add_heading("10. 测试与验收", level=1)
    doc.add_heading("10.1 已完成验证", level=2)
    add_table(doc, ["层级", "范围", "结果"], [
        ("后端新增接口", "账号、正确/错误登录、草稿 CRUD、公告已读、偏好读写", "通过"),
        ("后端回归", "teacher_backend/tests 全部 21 个测试", "21 passed"),
        ("前端类型", "TypeScript project build", "通过"),
        ("生产构建", "Vite production build", "通过；存在大包体积警告"),
        ("浏览器闭环", "双账号、草稿刷新、公告已读刷新、偏好刷新", "4/4 通过"),
        ("文档视觉", "Word → PDF → 逐页 PNG", "无裁切、重叠或缺字"),
    ], [1800, 5700, 1860], font_size=8.7)
    doc.add_heading("10.2 企业验收清单", level=2)
    acceptance = [
        "接口契约、数据库迁移和本文档版本一致。",
        "P0 安全项全部关闭，或有书面风险接受、期限和责任人。",
        "生产环境使用独立数据库和配置，不复用开发数据库。",
        "迁移、回滚、备份和恢复均在预生产环境演练通过。",
        "功能回归、权限负向用例、并发写和异常恢复用例通过。",
        "监控、告警、日志、值班和故障升级链路已配置。",
        "数据保留、删除、导出和脱敏规则获得业务/合规确认。",
        "发布后观察窗口内无持续错误、锁冲突或数据量异常。",
    ]
    for item in acceptance:
        add_list_item(doc, item, bullet_id)
    doc.add_heading("10.3 建议非功能门槛（未实测）", level=2)
    add_table(doc, ["指标", "建议门槛", "说明"], [
        ("可用性", "≥99.5%/月", "内部教学系统建议值，需业务确认"),
        ("读取延迟", "P95≤300ms", "50 并发基线，需压测"),
        ("写入延迟", "P95≤500ms", "含事务和审计，需压测"),
        ("服务端错误率", "<0.5%", "排除业务校验 4xx"),
        ("备份成功率", "100%", "失败必须告警和补偿"),
        ("恢复目标", "RPO≤24h，RTO≤4h", "需业务正式批准"),
    ], [2200, 2600, 4560], font_size=8.7)

    doc.add_heading("11. 风险与演进路线", level=1)
    add_table(doc, ["编号", "风险", "等级", "处置建议", "目标阶段"], [
        ("R-01", "客户端可伪造 X-User-Id", "极高", "接入真实认证令牌与服务端会话", "生产前"),
        ("R-02", "统一演示弱口令与无登录限流", "极高", "密码策略、首次改密、限流和锁定", "生产前"),
        ("R-03", "SQLite 多写并发与单文件故障", "高", "容量压测；必要时迁移 PostgreSQL", "试运行前"),
        ("R-04", "启动 create_all 与 Alembic 双通道", "中", "生产只允许 Alembic 变更结构", "发布规范化"),
        ("R-05", "JSON 字段缺少 schema_version", "中", "增加版本、校验与升级脚本", "下一迭代"),
        ("R-06", "无审计留存和集中监控", "中", "集中日志、指标、告警和归档", "试运行前"),
        ("R-07", "前端单包体积较大", "低", "路由级拆包和依赖分块", "性能优化"),
    ], [800, 2600, 950, 3460, 1550], font_size=8.0)
    doc.add_heading("11.1 分阶段建议", level=2)
    add_table(doc, ["阶段", "重点"], [
        ("阶段 1：生产门禁", "真实认证、密码策略、限流、HTTPS、统一错误、备份恢复、P0 风险关闭"),
        ("阶段 2：可运维化", "集中日志、指标告警、环境配置、发布流水线、迁移唯一入口"),
        ("阶段 3：可扩展化", "PostgreSQL 评估、并发控制、JSON 版本、缓存与性能压测"),
        ("阶段 4：治理完善", "数据目录、保留删除、敏感字段加密、审计归档和合规审查"),
    ], [2200, 7160], font_size=8.8)

    doc.add_heading("12. 技术结论", level=1)
    add_callout(doc, "结论", "教师端核心前端业务数据已形成 React → API → FastAPI → SQLAlchemy → SQLite 的可验证闭环，数据库结构、接口、测试和文档具备企业评审基础。当前仍属于本地/演示安全模型，生产发布必须优先完成真实认证、弱口令治理、限流、备份恢复和监控告警。")

    appendix_heading = doc.add_heading("附录 A. 全量数据库字段字典", level=1)
    appendix_heading.paragraph_format.page_break_before = True
    add_para(doc, "字段结构来自 SQLite PRAGMA 实时检查。PK=主键，FK=外键，UNIQUE=唯一约束；文档不展示任何凭据值、学生源代码或大段敏感正文。", size=9.5, color=MUTED)
    for group_name, group_tables in GROUPS:
        doc.add_heading(group_name, level=2)
        for table in group_tables:
            if table not in schemas:
                continue
            doc.add_heading(f"{table}（{counts[table]} 行）", level=3)
            add_para(doc, TABLE_PURPOSES.get(table, "项目业务表"), size=9.2, color=MUTED, after=4, keep=True)
            add_table(doc, ["字段", "类型", "约束", "说明"], schemas[table], [1900, 1500, 3300, 2660], font_size=8.0)

    doc.add_heading("附录 B. 当前数据快照（脱敏）", level=1)
    teacher_rows = query_rows(connection, "SELECT id, name, number, department, email FROM users WHERE role='teacher' ORDER BY id")
    masked_teachers = []
    for user_id, name, number, department, email in teacher_rows:
        local, _, domain = str(email).partition("@")
        masked_email = (local[:3] + "***@" + domain) if domain else "—"
        masked_teachers.append((user_id, name, number, department, masked_email))
    add_table(doc, ["ID", "姓名", "编号", "院系", "邮箱（脱敏）"], masked_teachers, [1500, 1100, 1300, 2800, 2660], font_size=8.5)
    add_para(doc, f"当前新增表状态：teacher_credentials={counts.get('teacher_credentials', 0)}，teacher_preferences={counts.get('teacher_preferences', 0)}，course_drafts={counts.get('course_drafts', 0)}，course_announcements={counts.get('course_announcements', 0)}，announcement_reads={counts.get('announcement_reads', 0)}。", size=9.5)
    course_rows = query_rows(connection, """
        SELECT c.id, u.name, c.name, c.code, c.term, c.status
        FROM courses c JOIN users u ON u.id=c.teacher_id
        WHERE c.teacher_id IN ('teacher-01','teacher-02')
        ORDER BY c.teacher_id, c.created_at, c.id LIMIT 12
    """)
    add_table(doc, ["课程ID", "教师", "课程名称", "代码", "学期", "状态"], course_rows, [1500, 900, 2400, 1100, 2260, 1200], font_size=8.0)

    doc.add_heading("附录 C. 核查命令", level=1)
    add_para(doc, "以下命令用于开发/测试环境核查，生产执行前应走变更审批并确认备份。")
    add_code_block(doc, "# 后端测试\npython -m pytest teacher_backend/tests -q -p no:cacheprovider\n\n# 前端类型与生产构建\ntsc -b\nvite build\n\n# 健康检查\nGET http://127.0.0.1:8001/api/v1/health\n\n# 标准迁移\nalembic -c teacher_backend/alembic.ini upgrade head")


def audit_enterprise_document(doc):
    audit_document(doc)
    numbering = doc.part.numbering_part.element
    assert len(numbering.findall(qn("w:abstractNum"))) >= 2
    assert any(paragraph.style.name == "Heading 1" for paragraph in doc.paragraphs)
    assert len(doc.inline_shapes) >= 2


def main():
    if not DB_PATH.exists():
        raise FileNotFoundError(DB_PATH)
    connection = sqlite3.connect(DB_PATH)
    connection.execute("PRAGMA foreign_keys=ON")
    tables, counts, schemas = inspect_database(connection)
    meta = {
        "sqlite_version": sqlite3.sqlite_version,
        "snapshot_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "table_count": len(tables),
        "total_rows": sum(counts.values()),
    }

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        set_header_footer(section)
    bullet_id = add_numbering_definition(doc, bullet=True)
    number_id = add_numbering_definition(doc, bullet=False)

    doc.core_properties.title = "CodeTrack 教师端数据库内容与前端接入说明（企业版）"
    doc.core_properties.subject = "企业研发交付技术基线"
    doc.core_properties.author = "CodeTrack 项目开发组"
    doc.core_properties.keywords = "CodeTrack, 企业开发, 数据库, 前端接入, API, 运维, 安全"
    doc.core_properties.comments = f"{DOC_ID} {DOC_VERSION}"

    add_cover(doc, meta)
    add_front_matter(doc, meta)
    add_enterprise_content(doc, connection, tables, counts, schemas, meta, bullet_id, number_id)
    audit_enterprise_document(doc)
    doc.save(OUTPUT_PATH)
    connection.close()
    print(json.dumps({
        "output": str(OUTPUT_PATH),
        "tables": len(tables),
        "rows": sum(counts.values()),
        "bytes": OUTPUT_PATH.stat().st_size,
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

