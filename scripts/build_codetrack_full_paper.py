from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import build_codetrack_paper as base


ROOT = Path(r"E:\teacher")
OUT_DIR = ROOT / "artifacts"
ASSET_DIR = OUT_DIR / "codetrack_full_paper_assets"
OUT_PATH = OUT_DIR / "CodeTrack_面向学科生态的垂直大模型与自主化学科智能体集成方案研究_完整论文稿.docx"


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(0)
    for name, east, size, bold in (("Heading 1", "宋体", 14, True),
                                   ("Heading 2", "仿宋_GB2312", 12, True),
                                   ("Heading 3", "宋体", 12, False)):
        s = doc.styles[name]
        s.font.name = "Times New Roman"; s.font.size = Pt(size); s.font.bold = bold
        s._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        s.font.color.rgb = RGBColor.from_string("1E2D38")
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        s.paragraph_format.space_before = Pt(8); s.paragraph_format.space_after = Pt(8)
        s.paragraph_format.keep_with_next = True


def set_page_number(section, start=1):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end]); base.set_run_font(run, size=9)
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType"); sect_pr.append(pg)
    pg.set(qn("w:start"), str(start))


def title_line(doc, text, size, bold=False, color="1E2D38", before=0, after=0):
    p = doc.add_paragraph(); base.set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=before, after=after)
    r = p.add_run(text); base.set_run_font(r, east_asia="宋体", size=size, bold=bold, color=color)
    return p


def add_h(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    size = 14 if level == 1 else 12
    base.set_run_font(r, east_asia="宋体" if level != 2 else "仿宋_GB2312", size=size,
                      bold=(level <= 2), color="1E2D38")
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph(); base.set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=12)
    r = p.add_run(text); base.set_run_font(r, east_asia="宋体", size=14, bold=True)
    return p


def add_toc(doc):
    add_section_title(doc, "目  录")
    p = doc.add_paragraph(); base.set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "目录将在Word中自动更新"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, placeholder, end])


def make_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    base.ASSET_DIR = ASSET_DIR
    base.build_diagrams()

    im = Image.new("RGB", (2100, 820), "white"); d = ImageDraw.Draw(im)
    phases = [
        ("真实需求", "课堂观察\n教师访谈\n答疑记录"),
        ("学科数据", "四元组标注\n质量复核\n版本管理"),
        ("模型适配", "Qwen2.5-7B\nLoRA微调\n独立评测"),
        ("知识增强", "课程知识库\n混合检索\n证据核验"),
        ("智能协同", "分类·检索\n拆解·质检\n有限回溯"),
        ("产品验证", "三端接入\n教师复核\n用户试用"),
    ]
    for i, (title, detail) in enumerate(phases):
        x = 65 + i * 340
        base.draw_box(d, (x, 220, x + 275, 500), f"{title}\n{detail}", "#F4F9FB" if i % 2 == 0 else "#EEF7F3",
                      "#2F6F9F" if i % 2 == 0 else "#3A8D74", 25)
        if i < len(phases) - 1:
            base.arrow(d, (x + 275, 360), (x + 330, 360))
    d.text((65, 80), "真实问题驱动、可验证方法与持续反馈闭环", font=base.img_font(38, True), fill="#1E2D38")
    im.save(ASSET_DIR / "technical_route.png")

    im = Image.new("RGB", (1800, 900), "white"); d = ImageDraw.Draw(im)
    boxes = [
        (640, 70, 1160, 210, "使用与运行日志", "#EAF3F8", "#2F6F9F"),
        (1190, 320, 1710, 460, "学生反馈与教师复核", "#EAF5F1", "#3A8D74"),
        (1030, 650, 1550, 790, "离线评测与安全检查", "#FFF5E6", "#C8892E"),
        (250, 650, 770, 790, "知识库/提示/数据更新", "#FDECEC", "#C65B5B"),
        (90, 320, 610, 460, "问题归类与版本审批", "#F1F1F1", "#7D8790"),
    ]
    for x1, y1, x2, y2, text, fill, edge in boxes:
        base.draw_box(d, (x1, y1, x2, y2), text, fill, edge, 28)
    base.arrow(d, (1160, 140), (1450, 320)); base.arrow(d, (1450, 460), (1290, 650))
    base.arrow(d, (1030, 720), (770, 720)); base.arrow(d, (510, 650), (350, 460)); base.arrow(d, (610, 390), (900, 210))
    d.text((540, 455), "所有更新先离线验证，后小范围试用", font=base.img_font(26, True), fill="#5B6570")
    im.save(ASSET_DIR / "feedback_loop.png")

    im = Image.new("RGB", (2100, 1000), "white"); d = ImageDraw.Draw(im)
    for i, label in enumerate(["学生端", "教师端", "管理员端", "科研入口"]):
        base.draw_box(d, (100 + i * 490, 70, 480 + i * 490, 205), label, "#F4F9FB", "#2F6F9F", 28)
        base.arrow(d, (290 + i * 490, 205), (1050, 330))
    base.draw_box(d, (670, 330, 1430, 470), "统一API网关：鉴权、限流、追踪、版本路由", "#EAF3F8", "#2F6F9F", 28)
    services = ["模型服务", "知识服务", "协同服务", "安全服务", "分析服务"]
    for i, label in enumerate(services):
        x = 100 + i * 400
        base.draw_box(d, (x, 610, x + 320, 735), label, "#EAF5F1", "#3A8D74", 27)
        base.arrow(d, (1050, 470), (x + 160, 610))
    base.draw_box(d, (340, 840, 1760, 950), "统一数据对象 · Chroma向量索引 · LoRA适配器 · 代码沙箱 · 日志审计", "#F7F7F7", "#7D8790", 25)
    for i in range(5): base.arrow(d, (260 + i * 400, 735), (650 + i * 200, 840))
    im.save(ASSET_DIR / "deployment.png")


def body(doc, text, bold_prefix=None):
    return base.add_body(doc, text, bold_prefix=bold_prefix)


def table(doc, caption, headers, rows, widths=None, page_break=False):
    return base.add_table(doc, caption, headers, rows, widths, page_break_before=page_break)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_diagrams()
    doc = Document(); configure(doc)

    # Cover
    title_line(doc, "信阳学院大学生学术科技作品", 12, bold=True, before=30, after=56)
    title_line(doc, "面向学科生态的垂直大模型构建与", 20, bold=True, after=8)
    title_line(doc, "自主化学科智能体集成方案研究", 20, bold=True, after=22)
    title_line(doc, "——以 CodeTrack 人工智能专业课程平台为例", 14, color="5B6570", after=70)
    table(doc, "", ["项  目", "内  容"], [
        ["作品名称", "CodeTrack"],
        ["团队名称", "信赛创翼队"],
        ["项目类别", "自然学术类"],
        ["所属学院", "计算机与人工智能学院"],
        ["完成时间", "2026年8月"],
    ], [4.2, 10.8])
    title_line(doc, "本稿为完整论文初稿，实验结果须以实际训练与测试数据补充", 10.5, color="7A5555", before=65)
    doc.add_page_break()

    # Chinese abstract
    add_section_title(doc, "摘  要")
    body(doc, "通用大语言模型在人工智能专业课程中具备较强的语言理解与生成能力，但面对课程专有术语、公式符号、代码语义和多步骤推理任务时，仍可能出现知识偏差、推导跳跃、教学方法不一致和引用依据不清等问题。现有教学辅助工具往往将问答、题库、代码运行、学习分析和资源管理分散实现，难以形成数据可追溯、模型可更新、过程可检查的统一学科智能服务。基于高校计算机与人工智能方向课程的需求，本研究提出面向学科生态的垂直大模型与自主化学科智能体集成方案，并以CodeTrack教师—学生—管理员一体化平台作为应用载体。")
    body(doc, "研究以Qwen2.5-7B-Instruct为基座模型，计划构建不少于350条“问题—题型—分步推导—答案”四元组结构化推理数据，采用LoRA完成参数高效适配；在外部知识层，将课程文档、结构化题库和知识点关系组织为具有来源、版本、权限和审核状态的课程知识库，通过关键词检索、向量检索、融合排序和证据核验形成可追溯RAG链路；在任务层设置问题分类、检索增强、步骤拆解和答案质检四类专用智能体，以结构化状态、有限回溯和人工复核实现复杂问题的受控求解。系统进一步将模型、知识、协同、安全和分析能力通过统一API接入三端业务，并保留科研辅助入口。")
    body(doc, "本文给出需求分析、数据规范、模型方法、知识库与RAG、多智能体协同、平台集成、安全边界以及对比实验和用户试用方案。现阶段网站原型与三端业务设计已经形成，垂类模型训练、端到端智能体联调和真实用户效果验证仍在推进。因此，本文不使用预期指标替代实验结论，而将相关参数和指标作为后续实现、评测与验收基线。研究旨在形成一套算力成本可控、数据责任清晰、结果可核验且可向同类课程迁移的学科智能平台建设方法。")
    p = doc.add_paragraph(); base.set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=10)
    r = p.add_run("关键词："); base.set_run_font(r, east_asia="黑体", size=12, bold=True)
    r = p.add_run("学科垂类大模型；LoRA；检索增强生成；多智能体协同；智能教学；CodeTrack"); base.set_run_font(r)
    doc.add_page_break()

    # English abstract
    add_section_title(doc, "ABSTRACT")
    body(doc, "General-purpose large language models show strong language understanding and generation capabilities, yet their direct use in artificial intelligence courses may lead to terminology confusion, skipped reasoning steps, inconsistent instructional methods, and unsupported answers. This study proposes an integrated approach that combines a discipline-specific language model, a course knowledge base, retrieval-augmented generation, and autonomous academic agents. CodeTrack, a unified platform for students, teachers, and administrators, is used as the application carrier.")
    body(doc, "The proposed method adopts Qwen2.5-7B-Instruct as the base model and uses LoRA for parameter-efficient adaptation. A structured dataset of at least 350 quadruples consisting of question, problem type, step-by-step reasoning, and answer is planned. Course documents, structured exercises, and knowledge-point relations are organized into a versioned and permission-aware knowledge base. A verifiable RAG pipeline combines lexical retrieval, dense retrieval, fusion ranking, and evidence checking. Four specialized agents are defined for problem classification, evidence retrieval, step decomposition, and answer verification. Their collaboration is governed by explicit state, bounded retry, tool permissions, and human review.")
    body(doc, "This paper presents the research design, data specification, model adaptation method, knowledge and agent architecture, system integration, safety constraints, and experimental protocol. The website prototype and multi-role product design have been completed, while model training, end-to-end agent evaluation, and user studies remain in progress. Therefore, the paper distinguishes planned indicators from verified results and does not report unobserved performance claims.")
    p = doc.add_paragraph(); base.set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=10)
    r = p.add_run("KEY WORDS: "); base.set_run_font(r, east_asia="黑体", size=12, bold=True)
    r = p.add_run("discipline-specific large language model; LoRA; retrieval-augmented generation; multi-agent collaboration; intelligent education"); base.set_run_font(r)
    doc.add_page_break(); add_toc(doc)

    # Body section, page numbering starts at 1
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54); sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)
    set_page_number(sec, 1)

    # Chapter 1
    add_h(doc, "第一章 绪论", 1)
    add_h(doc, "1.1 研究背景", 2)
    body(doc, "生成式人工智能正在由开放域问答逐步进入教学、科研和专业服务场景。高校计算机与人工智能方向课程具有概念更新快、公式与代码并存、知识依赖关系复杂等特点，学生不仅需要得到最终答案，更需要理解每一步推导依据、代码执行结果与课程方法之间的联系。通用模型的训练目标并不针对具体课程，其回答可能在语言上连贯，却未必符合课程大纲、教材版本和任课教师的讲解规范。")
    body(doc, "已有研究表明，领域适配、外部知识增强和任务分解能够分别改善模型的专业表达、知识时效性和复杂任务执行能力[1-6]。然而，将三类方法应用于真实课程时仍存在四项困难：一是课程数据来源分散，公式、代码和解题步骤缺少统一结构；二是模型参数知识与可更新课程知识边界不清；三是智能体协作容易产生错误传递和职责重叠；四是学生隐私、试题权限、代码执行和成绩评价具有明显的教育安全边界。")
    add_h(doc, "1.2 问题来源与前期调研", 2)
    body(doc, "根据项目申报阶段记录，团队通过课堂观察、教师访谈和历史答疑数据整理识别教学痛点。课堂观察连续覆盖3次习题课并形成36个时间切片；教师访谈涉及3位授课教师和11个典型答疑事件；历史答疑记录用于归纳概念不清、计算错误、推导中断和结果验算等问题类型。调研材料显示，综合题的困难主要集中在推导链中后段，教师最关注的不是模型是否给出另一个可行答案，而是其方法是否与课程要求一致、学生是否能够跟随。上述数据来自申报材料，正式发表前应保留原始记录、编码表和知情说明以便复核。")
    table(doc, "表1-1 前期需求调研材料", ["方法", "申报阶段记录", "主要用途", "正式论文补充材料"], [
        ["课堂观察", "3次习题课、36个时间切片", "识别卡点与行为模式", "观察表、课程信息、编码规则"],
        ["教师访谈", "3名教师、11个关键事件", "提炼课堂方法一致性需求", "提纲、匿名转录与主题编码"],
        ["答疑记录", "课程历史问题记录", "归纳高频问题与题型", "来源范围、脱敏和统计过程"],
    ], [3.0, 4.0, 4.2, 4.8], page_break=True)
    add_h(doc, "1.3 研究目标与研究问题", 2)
    body(doc, "本研究的总体目标是构建面向人工智能专业课程的垂类模型与自主化智能体集成方案，使系统能够在有限算力和明确数据边界下提供有依据、可检查、可反馈的教学辅助。围绕该目标提出四个研究问题。")
    table(doc, "表1-2 研究问题与验证证据", ["编号", "研究问题", "主要证据"], [
        ["RQ1", "结构化推理数据与LoRA能否改善学科表达和分步推导？", "基座/LoRA对比、错误类型分析"],
        ["RQ2", "课程知识库与RAG能否提高证据命中和回答忠实度？", "Recall@k、MRR、引用准确率"],
        ["RQ3", "四智能体协同能否提高复杂任务完成率并控制错误传播？", "消融实验、回溯率、人工转交率"],
        ["RQ4", "三端集成能否在权限、安全和教学责任边界内形成闭环？", "安全测试、用户试用、运行日志"],
    ], [2.0, 8.0, 5.0])
    add_h(doc, "1.4 研究内容与方法", 2)
    body(doc, "研究内容包括七个相互关联的部分：教学需求与范围界定、结构化推理数据集构建、LoRA垂类适配、课程知识库与RAG、四智能体协同、三端平台集成，以及技术—教学—安全三维效果验证。研究采用文献分析、需求调研、系统设计、对比实验、消融实验、教师盲评和用户试用等方法。")
    add_h(doc, "1.5 技术路线", 2)
    body(doc, "技术路线遵循“真实需求—学科数据—模型适配—知识增强—智能协同—产品验证”的顺序。模型微调用于学习相对稳定的术语表达和分步推理形式，知识库用于保存可更新、可授权的课程资料，多智能体用于组织任务与质检，平台负责权限、日志、代码执行和用户反馈。")
    base.add_picture(doc, ASSET_DIR / "technical_route.png", "图1-1 项目研究与技术路线", width=15)
    add_h(doc, "1.6 研究创新与论文结构", 2)
    body(doc, "本研究的创新性主要体现在三方面。第一，以“问题—题型—分步推导—答案”四元组描述课程推理样本，使解题过程本身成为训练和评价对象。第二，将课程证据检索、题型策略、子任务拆解与答案质检组织为可回溯的四智能体链路，避免将所有能力压缩到单次提示中。第三，将模型输出与课程版本、引用来源、代码结果、权限和人工复核绑定，使教学辅助从单次生成转向可持续治理。后续章节依次介绍需求、理论基础、垂类模型、智能体、系统与实验方案。")

    # Chapter 2
    add_h(doc, "第二章 研究设计与需求分析", 1)
    add_h(doc, "2.1 研究对象与应用边界", 2)
    body(doc, "研究对象为本科计算机与人工智能方向核心课程学习者及相关教师，首阶段聚焦程序设计、数据结构、数据库和机器学习等课程。学生场景包括课程问答、步骤提示、错题复习和学习路径建议；教师场景包括资料审核、练习生成草稿、批改辅助和学情解释；管理员场景包括用户组织、权限、知识版本、模型状态和审计日志。科研入口保留文献理解、代码说明和研究资料检索，但不与教学问答共用同一组效果结论。")
    add_h(doc, "2.2 用户需求", 2)
    body(doc, "学生需要系统识别口语化问题、保留上下文、解释中间步骤并指出依据，而不是直接展示考试或受限作业的最终答案。教师需要控制资料、方法和评分边界，查看AI生成内容的来源与修改记录。管理员需要确保不同用户只能访问所属课程和被授权资源，并能追踪模型、知识库和策略版本。")
    table(doc, "表2-1 三类用户需求与责任边界", ["用户", "主要需求", "AI可执行内容", "必须人工确认"], [
        ["学生", "答疑、提示、复习、反馈", "解释、示例、个性化建议", "考试最终答案、成绩改动"],
        ["教师", "备课、练习、批改、学情", "草稿、归纳、风险提示", "发布、评分、正式知识修改"],
        ["管理员", "权限、版本、运行、安全", "监测、归类、版本比较", "权限变更、模型上线与回滚"],
    ], [2.3, 4.2, 4.4, 4.1])
    add_h(doc, "2.3 功能与非功能需求", 2)
    table(doc, "表2-2 系统核心需求", ["类别", "需求", "验收方式"], [
        ["场景适配", "回答符合课程术语、符号和课堂方法", "典型问题与教师盲评"],
        ["自然交互", "支持上下文、歧义澄清和多轮追问", "多轮任务脚本"],
        ["知识可信", "资料有来源、版本和审核状态", "引用核对与版本测试"],
        ["过程可追溯", "保留检索、拆解、工具和质检记录", "结构化日志检查"],
        ["个性化", "基于学习记录调整提示和复习建议", "规则核验与用户反馈"],
        ["安全稳定", "权限最小化、代码隔离、失败可回退", "安全测试与压力测试"],
    ], [2.8, 7.3, 4.9])
    add_h(doc, "2.4 数据、伦理与知识产权要求", 2)
    body(doc, "教材、试卷、作业、课堂录音、学生代码和学习记录在使用前须明确授权范围。个人信息不直接用于模型训练，需删除姓名、学号、联系方式和其他直接标识。考试题和教师自建资源不因进入知识库而改变访问权限。平台应明确标注“AI生成内容”，不得生成虚假实验数据、伪造文献或自动形成成绩和处分。公开数据集、模型和代码时，应分别核对版权、许可证和隐私条件。")
    add_h(doc, "2.5 评价框架", 2)
    body(doc, "项目评价分为技术效果、教学使用效果和安全性三个维度。技术层关注答案、推理、检索和系统性能；教学层关注方法一致性、可跟随性和教师修改成本；安全层关注越权、隐私、提示注入、危险代码和高影响操作。三个维度共同决定版本是否可以进入小范围试用，不能用单一正确率替代综合评价。")

    # Chapter 3
    add_h(doc, "第三章 文献综述与理论基础", 1)
    add_h(doc, "3.1 垂直领域大模型", 2)
    body(doc, "大语言模型的开放域预训练使其具备较广泛的语言能力，但垂直领域应用仍面临专业知识稀疏、术语歧义、知识更新滞后和生成结果难以解释等问题。领域适配通常包括继续预训练、监督微调、参数高效微调、提示工程和外部知识增强。对于高校课程项目，数据规模和算力有限，全参数微调成本较高，且课程资料更新频繁，因此需要将稳定能力与动态知识分开处理。")
    add_h(doc, "3.2 参数高效微调与LoRA", 2)
    body(doc, "LoRA通过冻结基座模型参数并在部分线性层引入低秩矩阵，以较少可训练参数完成任务适配[2]。其优势是训练成本较低、适配器便于保存和切换，适合按课程管理；局限是小样本容易过拟合，数据质量和目标层选择会显著影响结果，且微调不能自动保证知识事实正确。因而LoRA应与独立测试、通用能力检查和知识库结合。")
    add_h(doc, "3.3 检索增强生成", 2)
    body(doc, "RAG通过检索外部资料并将结果作为生成上下文，缓解模型参数知识滞后和来源不可见问题[3-4]。教学场景中的检索对象不仅是普通文本，还包括公式、代码、例题、评分点和知识依赖。固定长度切块可能破坏这些结构；单纯向量相似度也可能召回主题相近但条件不同的题目。因此需要结合关键词、向量、元数据、课程版本和权限过滤，并单独评价检索与生成。")
    add_h(doc, "3.4 大模型智能体与多智能体协同", 2)
    body(doc, "智能体通常包含目标、状态、推理、工具和反馈等要素，能够根据任务状态选择下一步行动。ReAct等方法将推理与行动结合[5]，多智能体方法则通过角色分工降低单一模型在复杂任务中的认知负担[6]。但智能体数量增加并不必然提高效果，职责重叠、状态不一致和错误传播可能降低稳定性。学科问答需要明确每个智能体的输入输出、工具权限、终止条件和人工接管条件。")
    add_h(doc, "3.5 教育大模型与人机协同", 2)
    body(doc, "教育大模型应用正在从通用问答扩展到备课、评测、学习分析和个性化辅导[8-10]。教育场景与一般聊天不同，输出会影响学习路径和教师判断，因此必须保留教师责任、学生申诉和证据解释。个性化也不应简化为固定学生标签，而应基于可解释的课程行为和当前任务动态调整提示强度。")
    add_h(doc, "3.6 现有研究不足与理论框架", 2)
    table(doc, "表3-1 相关技术的作用与研究缺口", ["技术", "主要作用", "在本研究中的缺口"], [
        ["LoRA", "低成本领域适配", "需验证小样本、通用能力与跨课程迁移"],
        ["RAG", "引入动态知识与来源", "需处理教学语义切块、版本与权限"],
        ["多智能体", "任务分工与工具协同", "需控制状态一致性和错误传播"],
        ["学习分析", "识别知识短板与行为模式", "需防止标签化和无依据评价"],
        ["教育安全", "约束隐私、考试与高影响操作", "需与模型、知识和工具同步实施"],
    ], [2.6, 5.0, 7.4])
    body(doc, "基于上述分析，本文形成“领域适配—外部知识—任务分解—人机协同”的理论框架：LoRA学习稳定的学科表达，RAG提供可更新证据，智能体组织复杂步骤，人机协同确保教学责任和安全边界。四者通过版本、日志和反馈机制连接，构成后续方法与系统设计的依据。")

    # Chapter 4
    add_h(doc, "第四章 学科垂类大模型构建与知识增强", 1)
    add_h(doc, "4.1 学科数据来源与四元组结构", 2)
    body(doc, "首阶段数据覆盖人工智能专业代表性课程，来源包括课程标准、经授权教材与课件、实验指导、公开习题、教师确认解答和脱敏后的典型错误。项目计划构建不少于350条“问题—题型—分步推导—答案”四元组，题型分为概念辨析、数值计算、证明推理和实际应用，并按照8:1:1划分训练、验证和测试集。相似题、数值替换题和同源改写题必须按题源分组后再划分，以防止测试泄漏。")
    table(doc, "表4-1 学科推理数据字段", ["字段", "内容", "作用", "质量要求"], [
        ["问题", "题干、条件、课程、知识点", "确定任务范围", "表述完整，无条件缺失"],
        ["题型", "概念、计算、证明、应用", "选择策略与检索范围", "双人标注并复核"],
        ["分步推导", "步骤、依据、中间结果", "训练和评价推理过程", "每步可检查、前后可导出"],
        ["答案", "结论、公式或代码", "核验最终结果", "教师或标准答案确认"],
        ["证据", "文档、章节、页码、题号", "引用与追溯", "来源有效且权限清楚"],
        ["版本", "来源时间、审核人、状态", "更新与回退", "变更过程可查询"],
    ], [2.3, 4.6, 3.7, 4.4])
    add_h(doc, "4.2 数据处理与质量控制", 2)
    body(doc, "数据处理采用“授权登记—结构解析—清理去重—语义标注—交叉复核—数据分集—版本发布”的流程。解析阶段保留公式、代码、表格和题目层级；清理阶段统一符号、编码和LaTeX格式；标注阶段补充题型、知识点、推导步骤和评分要点；复核阶段由两名标注者独立检查并由任课教师处理分歧。正式实验应报告清洗前后数量、重复率、字段完整率和标注一致性。")
    table(doc, "表4-2 数据质量指标", ["指标", "计算对象", "用途"], [
        ["字段完整率", "问题、题型、步骤、答案、证据", "判断样本是否可训练与核验"],
        ["重复率", "文本与语义近重复样本", "防止数据冗余和分集泄漏"],
        ["标注一致性", "题型、知识点和步骤边界", "评价标注规范可靠性"],
        ["教师通过率", "答案与推导链", "控制正式数据质量"],
        ["来源可追溯率", "资料编号与权限信息", "支持引用、下线与审计"],
    ], [3.1, 6.0, 5.9])
    add_h(doc, "4.3 基座模型与LoRA微调", 2)
    body(doc, "依据申报方案，首版采用Qwen2.5-7B-Instruct[7]作为基座模型。LoRA冻结原始权重W，并通过低秩矩阵B和A学习增量ΔW：")
    base.add_equation(doc, "W′ = W + ΔW，ΔW = BA，r ≪ min(d, k)", "4-1")
    body(doc, "其中r为低秩维度，方案初值设为16。训练目标不是记忆特定题目，而是学习课程术语、公式符号、分步表达、代码说明和教师认可的解题风格。数据版本、模型配置、随机种子、训练日志和评测结果均随适配器保存。")
    table(doc, "表4-3 首版LoRA训练方案", ["项目", "方案值", "正式报告要求"], [
        ["基座模型", "Qwen2.5-7B-Instruct", "精确版本、许可证与校验值"],
        ["数据规模", "不少于350条四元组", "清洗前后数量与课程分布"],
        ["数据划分", "8:1:1", "按题源分组，防止相似题泄漏"],
        ["LoRA秩", "r=16", "与其他秩设置进行消融"],
        ["初始学习率", "1×10⁻⁴", "调度器、轮数、早停与优化器"],
        ["硬件", "单张RTX 3090，24 GB", "驱动、框架、显存、耗时与能耗"],
    ], [3.2, 4.6, 7.2])
    add_h(doc, "4.4 课程知识库构建", 2)
    body(doc, "课程知识库由文档资料、结构化题库和知识点关系组成。文档依据章节、概念、例题、函数、实验任务等教学语义单元分块，避免破坏公式、代码和题目结构。每个知识单元记录课程、章节、知识点、资料来源、版本状态、适用对象、审核人和变更说明。只有现行、已审核且用户有权访问的内容可以进入生成上下文。")
    table(doc, "表4-4 知识库元数据", ["元数据", "示例", "功能"], [
        ["课程范围", "课程、章节、知识点", "限制检索空间"],
        ["资料来源", "教材、课件、教案、题库", "展示引用依据"],
        ["版本状态", "待审核、现行、已下线", "避免使用过期资料"],
        ["适用对象", "教师、学生、管理员", "执行权限过滤"],
        ["内容类型", "概念、例题、代码、实验", "选择切块与检索策略"],
        ["管理记录", "责任教师、时间、变更说明", "支持复核与回退"],
    ], [3.0, 5.2, 6.8])
    add_h(doc, "4.5 混合检索与可验证RAG", 2)
    body(doc, "检索链路先识别用户角色、课程、题型和对话上下文，再将原问题与规范化查询同时用于关键词检索和向量检索。候选资料根据相关性、来源可靠性、课程版本和权限重排，融合得分定义为：")
    base.add_equation(doc, "S(d,q)=αS_keyword(d,q)+βS_dense(d,q)+γS_authority(d)", "4-2")
    body(doc, "权重α、β和γ以及Top-k、相似度阈值和分块参数由验证集确定。生成阶段要求关键结论标注资料编号；证据不足、资料冲突或引用与结论不一致时，系统重新检索、降低结论强度或转教师复核。RAG只能降低无依据生成风险，不能被表述为自动消除模型幻觉。")
    base.add_picture(doc, ASSET_DIR / "rag_flow.png", "图4-1 面向课程问答的可验证RAG链路", width=15)
    add_h(doc, "4.6 模型与知识版本管理", 2)
    body(doc, "模型适配器、知识索引和提示策略分别版本化，并通过发布清单绑定。知识资料更新不必立即重新训练模型，而是先重建受影响的索引并执行回归测试；只有稳定表达或推理行为需要改变时才考虑增量微调。上线前依次通过离线评测、教师盲评、安全检查和小范围试用，出现明显退步时恢复上一版本。")

    # Chapter 5
    add_h(doc, "第五章 自主化学科智能体集成方案", 1)
    add_h(doc, "5.1 智能体定义与设计原则", 2)
    body(doc, "本文将智能体定义为具有明确目标、结构化状态、可调用工具、决策条件和终止规则的任务单元。普通数据库查询、接口封装和页面组件不因使用大模型而自动成为智能体。系统遵循职责单一、输入输出明确、权限最小、失败可见和人工可接管等原则。")
    p = add_h(doc, "5.2 四类核心智能体", 2)
    p.paragraph_format.page_break_before = True
    table(doc, "表5-1 核心智能体职责", ["智能体", "输入", "主要任务", "输出"], [
        ["问题分类", "问题、课程、上下文", "识别题型、意图和风险", "题型标签、约束条件"],
        ["检索增强", "标签、查询、权限", "召回并核验课程证据", "证据集合、来源、得分"],
        ["步骤拆解", "问题、题型、证据", "生成可检查子任务链", "步骤、依赖、工具请求"],
        ["答案质检", "步骤结果、证据、规则", "检查逻辑、引用和安全", "通过/不通过、问题清单"],
    ], [2.5, 3.8, 5.0, 3.7])
    body(doc, "主模型求解被视为受控工具调用，不单独计为第五个核心智能体。练习生成、批改辅助和学情分析属于教学业务模块，可复用核心问答链路，但不与四个核心智能体处于同一层级。该区分能够保持申报方案与系统实现的一致性。")
    add_h(doc, "5.3 协同流程与状态控制", 2)
    body(doc, "调度器维护用户角色、课程、原问题、题型、证据集合、子任务、工具结果、质检状态、重试次数和人工复核标记。只有上一步输出满足模式约束、权限检查通过且风险可控时，流程才进入下一阶段。质检不通过时最多回溯2次；连续失败、资料冲突、代码工具异常或涉及成绩发布等高影响操作时停止自动执行。")
    base.add_picture(doc, ASSET_DIR / "agent_flow.png", "图5-1 四智能体协同与人工复核机制", width=15)
    add_h(doc, "5.4 教学业务模块与个人画像", 2)
    body(doc, "练习生成模块根据课程目标、知识点和难度生成题目草稿；批改辅助模块结合代码运行结果、常见错误和评分点给出建议；学情分析模块按照班级、知识点和错误类型汇总统计；教学改进模块将诊断结果转化为下一轮教学建议。所有内容由教师最终确认。")
    body(doc, "个人画像用于描述学习者在课程范围内的知识掌握、错误类型、任务完成和反馈偏好，不使用家庭背景、性别、地域等与学习无关属性作价值判断。画像结论必须关联到可核验的学习记录，并允许学生查看主要依据、提出纠错或申诉。")
    add_h(doc, "5.5 科研入口", 2)
    body(doc, "科研入口保留文献检索与归纳、代码解释、数据处理辅助和格式规范检查等功能。该入口与教学端共享知识版本、引用展示、权限和安全审计，但不得生成虚假数据或伪造文献。涉及研究结论时，应展示数据来源、分析方法和不确定性；涉及文献时，应提供可核查的题名、作者和链接。科研场景的效果指标应独立设置，不直接套用教学问答正确率。")
    add_h(doc, "5.6 反馈闭环", 2)
    body(doc, "学生可以标记“看不懂、依据不足、答案有误”，教师可以确认方法不一致、评分点不合理或知识过期，管理员将问题归类为数据补充、知识库修订、提示与流程调整、模型更新或安全策略调整。更新内容先通过离线评测与安全检查，再进入小范围试用。")
    base.add_picture(doc, ASSET_DIR / "feedback_loop.png", "图5-2 使用反馈与版本迭代闭环", width=14.5)

    # Chapter 6
    add_h(doc, "第六章 系统设计、实现与实验方案", 1)
    add_h(doc, "6.1 系统总体架构", 2)
    body(doc, "CodeTrack采用应用层、智能层、知识层以及模型与平台层的分层架构。应用层承载学生端、教师端、管理员端和科研入口；智能层组织问题分类、混合检索、步骤拆解、主模型求解和答案质检；知识层维护课程资料、题库、知识点关系和向量索引；模型与平台层提供基座模型、LoRA适配器、统一数据对象、权限、日志和代码沙箱。")
    base.add_picture(doc, ASSET_DIR / "architecture.png", "图6-1 CodeTrack垂类模型与智能体总体架构", width=15)
    add_h(doc, "6.2 服务部署与接口组织", 2)
    body(doc, "模型服务负责版本加载和推理请求，知识服务负责资料入库、索引、检索和来源定位，协同服务维护智能体状态和人工复核，安全服务执行鉴权、敏感信息检测和审计，分析服务归纳错误类型、知识掌握和用户反馈。统一API网关完成身份验证、限流、链路追踪和版本路由，避免三端分别实现AI逻辑。")
    base.add_picture(doc, ASSET_DIR / "deployment.png", "图6-2 CodeTrack服务部署与统一接入关系", width=15)
    add_h(doc, "6.3 三端数据与接口边界", 2)
    table(doc, "表6-1 三端共享数据对象", ["对象", "关键字段", "主要使用方", "控制要求"], [
        ["课程与班级", "编号、版本、成员、学期", "三端", "按角色与班级授权"],
        ["资源与知识", "来源、课程、审核、版本", "教师、学生、管理端", "现行版本和对象权限"],
        ["题目与任务", "题型、知识点、截止时间", "教师、学生", "受限任务控制答案"],
        ["学习记录", "提交、错误、反馈、引用", "学生、教师", "最小化保存与可申诉"],
        ["模型运行", "版本、时延、状态、风险", "管理员", "脱敏、审计、告警"],
    ], [2.8, 5.0, 3.8, 3.4])
    add_h(doc, "6.4 代码运行与工具安全", 2)
    body(doc, "代码运行沿用平台统一隔离流程。每次任务在临时目录中执行，限制时间、内存、进程数、文件类型、输出长度和网络访问。模型只能提交结构化工具请求，后端完成参数、权限和资源检查后执行；模型不直接获得系统命令、数据库连接或任意文件访问权限。编译、输入、输出、异常和资源使用记录用于批改与复核。")
    p = add_h(doc, "6.5 六层安全控制", 2)
    p.paragraph_format.page_break_before = True
    table(doc, "表6-2 安全控制与测试重点", ["层次", "主要措施", "测试重点"], [
        ["数据", "授权、脱敏、最小化保存", "隐私泄露、未授权训练"],
        ["身份", "角色、课程、班级权限", "水平与垂直越权"],
        ["知识", "审核、版本、恶意文档检测", "提示注入、过期资料"],
        ["生成", "敏感信息、证据与受限答案", "无依据结论、考试答案"],
        ["工具", "结构化调用和代码沙箱", "危险代码、资源耗尽"],
        ["上线", "门控、灰度、告警、回滚", "异常扩散和不可追踪"],
    ], [2.4, 7.0, 5.6])
    add_h(doc, "6.6 当前实现状态", 2)
    body(doc, "为避免将设计方案写成已完成成果，本文依据当前项目状态区分已形成基础、已完成设计和待实现验证。后续每次论文更新应由代码仓库、模型文件、日志或用户记录支持状态变更。")
    table(doc, "表6-3 项目当前状态说明", ["模块", "当前状态", "论文可使用表述"], [
        ["网站原型与三端界面", "已形成产品原型", "可描述页面、流程和数据对象"],
        ["课程/班级/资源等平台能力", "已有开发方案与部分实现", "按真实接口和截图说明"],
        ["结构化数据集", "规范已设计，数据待完成核验", "不得宣称已完成350条"],
        ["Qwen2.5 LoRA模型", "方案已确定，真实训练结果暂无", "只能写训练方案与参数基线"],
        ["RAG与四智能体", "流程和接口已设计，端到端代码待完成", "不得宣称稳定运行"],
        ["用户试用与效果数据", "待开展", "不得填写预期正确率和满意度"],
    ], [4.0, 5.0, 6.0], page_break=True)
    add_h(doc, "6.7 对比与消融实验", 2)
    body(doc, "实验应在相同测试集、可见资料、解码参数和评分规则下逐步增加组件，区分模型适配、外部知识和智能体流程的贡献。DeepSeek-V3等通用模型可以作为外部参照，但其服务版本、访问时间和参数设置必须记录。")
    table(doc, "表6-4 对比实验分组", ["组别", "模型与组件", "研究目的"], [
        ["G0", "原始Qwen2.5-7B-Instruct", "获得基座表现"],
        ["G1", "基座模型+RAG", "评估外部知识增益"],
        ["G2", "LoRA模型", "评估参数适配增益"],
        ["G3", "LoRA模型+RAG", "评估两类知识注入组合效果"],
        ["G4", "LoRA+RAG+四智能体", "评估拆解、质检和回溯贡献"],
        ["G5", "DeepSeek-V3等通用模型", "提供外部通用模型参照"],
    ], [2.0, 7.0, 6.0])
    p = add_h(doc, "6.8 评价指标与统计方法", 2)
    p.paragraph_format.page_break_before = True
    table(doc, "表6-5 技术、教学与安全评价指标", ["维度", "指标", "评价方式"], [
        ["答案质量", "正确率、公式/代码可执行率", "标准答案与教师核验"],
        ["推理过程", "完整性、步骤可导出性", "教师5级量表盲评"],
        ["教学一致性", "课堂方法一致性、修改成本", "教师独立评价与修改记录"],
        ["检索质量", "Recall@k、MRR、证据命中率", "带相关性标注的问题集"],
        ["可信生成", "引用准确率、答案忠实度", "逐结论核对证据"],
        ["协同稳定性", "完成率、回溯率、人工转交率", "结构化运行日志"],
        ["系统性能", "首字延迟、总时延、显存占用", "统一硬件重复测量"],
        ["安全性", "越权、泄露、提示注入通过率", "专项安全测试集"],
    ], [2.7, 6.0, 6.3])
    body(doc, "推导完整性和方法一致性拟由2—3名教师盲评，并以Kendall's W报告评价者一致性。除均值外报告样本量、标准差或置信区间；多组比较根据数据分布选择适当显著性检验。模型生成存在随机性时，应固定随机种子或重复运行并报告波动。")
    add_h(doc, "6.9 典型问题与用户验证", 2)
    table(doc, "表6-6 三类典型测试案例设计", ["案例", "任务内容", "主要验证点"], [
        ["案例A：概念辨析", "比较过拟合与欠拟合并给出判断依据", "术语准确、证据引用、歧义澄清"],
        ["案例B：代码推理", "定位遍历程序错误并解释复杂度", "代码执行、步骤拆解、错误定位"],
        ["案例C：综合应用", "完成机器学习参数更新与结果解释", "多知识点依赖、公式和质检回溯"],
    ], [3.2, 7.0, 4.8])
    body(doc, "用户验证至少包含2名真实目标用户的结构化记录。项目计划邀请15—20名学生和2—3名教师开展两周小范围试用，记录使用场景、频次、完成时间、教师修改量、错误类型和主观评价。若累积50条经教师确认的有效纠错样本，可进入增量训练候选集，但是否更新仍由离线评测和版本审批决定。")
    add_h(doc, "6.10 结果报告规范", 2)
    body(doc, "由于真实模型训练与用户试验尚未完成，本稿不填写虚构结果。正式论文应至少补充五类结果：数据集统计与标注一致性；各实验组的答案、推理和检索指标；典型成功与失败案例；系统时延、显存与稳定性；用户试用和安全测试。结果分析应同时报告提升与退化，区分模型、检索和流程原因，并将原始输出、评分表和日志作为可复核材料保存。")

    # Chapter 7
    add_h(doc, "第七章 结论与展望", 1)
    add_h(doc, "7.1 研究结论", 2)
    body(doc, "本文围绕人工智能专业课程中的术语偏差、推导跳跃、证据不足和教学方法不一致问题，提出面向学科生态的垂类模型与自主化智能体集成方案。方案以结构化推理数据为基础，以Qwen2.5-7B-Instruct和LoRA完成低成本适配，以课程知识库和可验证RAG提供动态证据，以四智能体链路组织分类、检索、拆解和质检，并通过CodeTrack三端平台实现权限、日志、代码工具和用户反馈的统一管理。")
    body(doc, "本研究的核心价值不在于堆叠平台功能，而在于将稳定表达、动态知识、复杂任务和教学责任分别交给适合的机制处理，并通过版本与证据连接。该思路为有限算力高校建设可复制、可维护的学科智能平台提供了方法基础。由于模型和用户实验仍在推进，本文结论仅限于方案、规范和实验协议的形成，不能替代后续效果验证。")
    add_h(doc, "7.2 研究不足", 2)
    body(doc, "第一，结构化数据规模较小且集中于计算机与人工智能方向课程，跨课程代表性有限。第二，LoRA参数、检索权重和智能体策略尚需真实实验确定。第三，小规模短期试用难以证明长期学习效果，用户自选择也可能带来偏差。第四，多智能体增加了链路长度与系统成本，是否优于更简单流程需要消融实验判断。第五，科研入口的任务类型和评价标准尚未充分展开。")
    add_h(doc, "7.3 未来工作", 2)
    body(doc, "后续工作首先完成数据授权、标注和独立测试集建设，按本论文配置开展LoRA训练与消融实验；其次实现RAG、四智能体和代码工具的端到端联调，形成完整日志和安全测试；再次开展教师盲评与真实用户试用，补充典型案例和量化结果；最后从单课程扩展至课程群，并在数学、物理等理工科场景验证迁移成本与方法适用性。所有迭代均应保留数据、模型、知识库和代码版本，确保结论可复核。")

    # References
    add_h(doc, "参考文献", 1)
    refs = [
        "[1] 籍欣萌, 昝红英, 崔婷婷, 张坤丽. 大模型在垂直领域应用的现状与挑战[J]. 计算机工程与应用, 2025, 61(12): 1-11.",
        "[2] HU E J, SHEN Y, WALLIS P, et al. LoRA: Low-Rank Adaptation of Large Language Models[C]//International Conference on Learning Representations. 2022.",
        "[3] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems. 2020, 33: 9459-9474.",
        "[4] GAO Y, XIONG Y, GAO X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey[EB/OL]. arXiv:2312.10997, 2023.",
        "[5] YAO S, ZHAO J, YU D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]//International Conference on Learning Representations. 2023.",
        "[6] WANG L, MA C, FENG X, et al. A Survey on Large Language Model Based Autonomous Agents[J]. Frontiers of Computer Science, 2024, 18(6): 186345.",
        "[7] QWEN TEAM. Qwen2.5 Technical Report[EB/OL]. arXiv:2412.15115, 2024.",
        "[8] 刘三女牙, 郝晓丽, 李卿, 等. 教育大语言模型的内涵、构建和挑战[J]. 现代远程教育研究, 2024(5).",
        "[9] 朱孟潇, 陆文灏, 王禧玥, 等. 智能教育中的大模型智能体：基础、模式和挑战[J]. 教师教育论坛, 2025, 38(2): 49-60.",
        "[10] 孙茂松, 刘洋, 李涓子, 等. 生成式人工智能在教育中的应用与反思[J]. 教育研究, 2024, 45(12): 78-90.",
        "[11] CHROMA. Chroma Documentation[EB/OL]. https://docs.trychroma.com/.",
        "[12] LANGCHAIN. LangChain Documentation[EB/OL]. https://python.langchain.com/.",
    ]
    for ref in refs:
        p = doc.add_paragraph(); base.set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=1)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0.85); p.paragraph_format.first_line_indent = Cm(-0.85)
        r = p.add_run(ref); base.set_run_font(r, size=9.5)

    # Appendices
    p = add_h(doc, "附录A 教师盲评量表（建议稿）", 1)
    p.paragraph_format.page_break_before = True
    table(doc, "表A-1 教师盲评评分说明", ["维度", "1分", "3分", "5分"], [
        ["答案正确性", "结论错误或无法使用", "主要结论正确但有局部错误", "结论、公式或代码均正确"],
        ["推导完整性", "关键步骤缺失", "步骤基本完整但解释不足", "步骤完整且前后可导出"],
        ["方法一致性", "明显偏离课程要求", "基本符合但需修改", "与课堂方法和符号一致"],
        ["证据可信度", "无来源或引用错误", "有来源但支撑不充分", "关键结论均有有效证据"],
        ["可理解性", "学生难以跟随", "部分说明需要教师补充", "层次清楚，可直接用于辅导"],
    ], [3.2, 4.0, 4.1, 3.7])
    p = add_h(doc, "附录B 用户试用记录字段", 1)
    p.paragraph_format.page_break_before = True
    table(doc, "表B-1 用户试用记录模板", ["字段", "记录内容"], [
        ["用户身份", "匿名编号、教师/学生、课程与年级"],
        ["使用任务", "问题类型、功能入口、是否使用代码工具"],
        ["过程数据", "开始时间、总时长、追问次数、引用展开"],
        ["结果评价", "正确性、可跟随性、方法一致性、满意度"],
        ["人工修改", "教师修改内容、修改时间和原因"],
        ["问题反馈", "答案有误、依据不足、看不懂、权限或安全问题"],
        ["授权说明", "知情同意、匿名化和数据使用范围"],
    ], [4.0, 11.0])
    p = add_h(doc, "附录C 实验结果填报清单", 1)
    p.paragraph_format.page_break_before = True
    body(doc, "正式提交前应补齐：数据集统计表；训练损失与验证曲线；各实验组完整指标；教师一致性统计；至少3个典型问题的完整输出；至少2名真实目标用户的结构化记录；模型、知识库和代码版本；安全测试结果；运行环境、时延、显存和异常日志。任何尚未获得的数据均应保留为空或明确标注“待实验”，不得用预期值代替。")

    doc.core_properties.title = "面向学科生态的垂直大模型构建与自主化学科智能体集成方案研究"
    doc.core_properties.subject = "CodeTrack完整学术论文初稿"
    doc.core_properties.author = "信赛创翼队"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
