from pathlib import Path
from datetime import date

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\teacher")
SOURCE_IMAGES = ROOT / "artifacts" / "operation-manual"
WORK_DIR = ROOT / "artifacts" / "operation-manual-docx"
OUTPUT_DIR = ROOT / "artifacts" / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "CodeTrack教师端详细操作手册_实测版.docx"

WORK_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# compact_reference_guide token map with two named overrides:
# 1. CJK font fallback: Microsoft YaHei for East Asian glyphs.
# 2. Product accent: CodeTrack green replaces preset heading blue consistently.
GREEN = "2AAD16"
DARK_GREEN = "176B12"
INK = "1F2A24"
MUTED = "5F6B63"
LIGHT_GREEN = "EDF8EA"
LIGHT_GRAY = "F4F6F5"
TABLE_HEADER = "E8EEF5"
BORDER = "D8E0DA"
CAUTION = "FFF5D9"
RISK = "FCE8E6"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, italic=None, ascii_font="Calibri"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, bold=False, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths_dxa, header_fill=TABLE_HEADER, font_size=9.2, aligns=None, keep_together=False):
    table = doc.add_table(rows=1, cols=len(headers))
    set_repeat_table_header(table.rows[0])
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = str(text)
        set_cell_shading(table.rows[0].cells[i], header_fill)
        style_cell_text(table.rows[0].cells[i], bold=True, color=INK, size=font_size,
                        align=(aligns[i] if aligns else WD_ALIGN_PARAGRAPH.LEFT))
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
            style_cell_text(cells[i], size=font_size,
                            align=(aligns[i] if aligns else WD_ALIGN_PARAGRAPH.LEFT))
    for row_index, row in enumerate(table.rows):
        set_row_cant_split(row)
        if keep_together and row_index < len(table.rows) - 1:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc, marker="%1.", bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, indent])
    level.extend([start, fmt, lvl_text, suff, p_pr])
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_steps(doc, steps):
    num_id = add_numbering_definition(doc)
    for text in steps:
        p = doc.add_paragraph(style="Procedure Step")
        apply_num(p, num_id)
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return num_id


def add_bullets(doc, items, num_id=None):
    if num_id is None:
        num_id = add_numbering_definition(doc, marker="●", bullet=True)
    for text in items:
        p = doc.add_paragraph(style="Compact Bullet")
        apply_num(p, num_id)
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return num_id


def add_callout(doc, label, text, kind="info"):
    fill = LIGHT_GREEN if kind == "info" else CAUTION if kind == "caution" else RISK
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, bold=True, color=DARK_GREEN if kind == "info" else INK)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def crop_image(source_name, output_name, crop_box):
    source = SOURCE_IMAGES / source_name
    output = WORK_DIR / output_name
    with Image.open(source) as image:
        image.crop(crop_box).save(output, quality=92)
    return output


def add_figure(doc, image_path, caption, width=6.35, alt_text=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text or caption)
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.keep_together = True
    c.add_run(caption)


def page_break(doc):
    # Heading 1 owns page starts. Unlike a hard break paragraph, its
    # page-break-before property cannot create a blank page at a boundary.
    return None


def add_h1(doc, text):
    return doc.add_paragraph(text, style="Heading 1")


def add_h2(doc, text):
    return doc.add_paragraph(text, style="Heading 2")


def add_h3(doc, text):
    return doc.add_paragraph(text, style="Heading 3")


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, GREEN),
        ("Heading 2", 13, 14, 7, DARK_GREEN),
        ("Heading 3", 12, 10, 5, INK),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = name == "Heading 1"

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)

    step_style = styles.add_style("Procedure Step", 1)
    step_style.base_style = normal
    step_style.paragraph_format.left_indent = Inches(0.375)
    step_style.paragraph_format.first_line_indent = Inches(-0.188)
    step_style.paragraph_format.space_after = Pt(4)
    step_style.paragraph_format.line_spacing = 1.25

    bullet_style = styles.add_style("Compact Bullet", 1)
    bullet_style.base_style = normal
    bullet_style.paragraph_format.left_indent = Inches(0.375)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.188)
    bullet_style.paragraph_format.space_after = Pt(4)
    bullet_style.paragraph_format.line_spacing = 1.25


def configure_document(doc):
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    # The default Word template can enable separate even-page header/footer
    # stories. This guide intentionally uses one quiet running header/footer.
    settings = doc.settings.element
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is not None:
        settings.remove(even_odd)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("CodeTrack 教师端  |  操作手册（实测版）")
    set_run_font(r, size=9, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def build():
    doc = Document()
    configure_document(doc)

    dash_crop = crop_image("01-dashboard-desktop.png", "dashboard-crop.png", (0, 0, 1440, 900))
    content_crop = crop_image("03-content-desktop.png", "content-crop.png", (0, 0, 1440, 1000))
    task_crop = crop_image("04-tasks-desktop.png", "tasks-crop.png", (0, 0, 1440, 1050))
    material_crop = crop_image("06-materials-desktop.png", "materials-crop.png", (0, 0, 1440, 970))
    graph_current = ROOT / "artifacts" / "knowledge-graph-check-desktop.png"
    analytics_crop = crop_image("08-analytics-desktop.png", "analytics-crop.png", (0, 0, 1440, 1050))

    # Cover: editorial_cover pattern.
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CODETRACK TEACHER")
    set_run_font(r, size=11, bold=True, color=GREEN)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph("CodeTrack 教师端详细操作手册", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("基于多轮自动化与浏览器实测的交付版说明", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("适用对象：任课教师、教务验收人员、系统管理员")
    set_run_font(r, size=11, color=MUTED)
    add_table(doc, ["文档属性", "内容"], [
        ["文档版本", "V1.1 · 实测版"],
        ["测试日期", "2026-08-14（Asia/Shanghai）"],
        ["测试基线", "Git 66706fb；包含当前工作区待提交功能"],
        ["验证结论", "新版知识图谱接口、前端构建与浏览器交互复核通过"],
    ], [1875, 7485], header_fill=LIGHT_GREEN, font_size=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("编制日期：2026 年 8 月 14 日")
    set_run_font(r, size=10.5, color=MUTED)

    page_break(doc)
    add_h1(doc, "文档说明")
    add_callout(doc, "结论", "本手册中的主要操作路径均已结合接口自动化、生产构建或真实浏览器检查。教师可从“快速上手”开始；验收人员可直接查看“测试用例与实测结果”。")
    add_h2(doc, "适用范围")
    add_bullets(doc, [
        "教师端：登录、工作台、课程、章节、班级、邀请、任务、批改、资料、知识图谱、课堂讨论、学情分析、通知与设置。",
        "部署与启动：给出最短启动和健康检查步骤；更完整的运维说明继续以项目 RUNBOOK.md 为准。",
        "学生端：仅覆盖教师可见的“学生视角预览”和接口闭环，不作为完整学生端用户手册。",
    ])
    add_h2(doc, "功能状态标识")
    add_table(doc, ["标识", "含义", "使用建议"], [
        ["浏览器已验收", "真实页面已打开并检查可见性、溢出、接口和控制台错误", "可按本手册直接操作"],
        ["接口已验收", "后端业务闭环已由自动化用例验证", "页面入口未联调时可由管理员/API 使用"],
        ["展示入口", "当前页面有按钮或提示，但业务动作尚未形成持久化闭环", "仅演示，不作为正式数据操作"],
    ], [1800, 3960, 3600], font_size=9.5)
    add_h2(doc, "目录")
    add_body(doc, "第 1-7 章覆盖快速上手、登录导航、课程、章节、班级、任务与批改；第 8-16 章覆盖资料、图谱、讨论、学情、设置、故障处理及测试结果；附录提供日常检查清单、数据影响与已知限制。")

    page_break(doc)
    add_h1(doc, "1. 快速上手")
    add_h2(doc, "1.1 启动系统")
    add_body(doc, "开发环境默认使用前端 5173 端口和后端 8001 端口。所有命令在项目根目录 E:\\teacher 执行。")
    add_steps(doc, [
        "首次使用时安装 Python 依赖：python -m pip install -r backend\\requirements.txt。",
        "安装前端依赖：pnpm install。若已有可用 node_modules，不要使用不同版本的 pnpm 强制重建。",
        "设置 Python 解释器（推荐）：$env:CODETRACK_PYTHON = 'D:\\Anaconda\\python.exe'。",
        "启动前后端：pnpm run dev。",
        "浏览器打开 http://127.0.0.1:5173/。",
    ])
    add_callout(doc, "启动成功判定", "健康接口 http://127.0.0.1:8001/api/v1/health 返回 status=ok；前端地址返回 HTTP 200。")
    add_h2(doc, "1.2 最短教学闭环")
    add_steps(doc, [
        "使用王老师登录并进入“教学工作台”。",
        "在“我的课程”选择课程，点击“进入课程”。",
        "在“课程章节内容”创建或打开章节，添加知识点、课件和练习，并将章节发布。",
        "在“班级管理”建立班级，在“邀请学生”复制链接、邀请码或下载二维码。",
        "在“任务管理”创建任务、配置测试用例、选择班级并发布。",
        "学生提交后进入“查看成绩”，确认评分依据、反馈和最终得分，再发布成绩。",
        "在“学情分析”查看班级总览、个体诊断和预警。",
    ])
    add_callout(doc, "关键原则", "草稿不会自动对学生可见；章节、资料、任务和成绩均需确认可见性或执行发布动作。", kind="caution")

    page_break(doc)
    add_h1(doc, "2. 登录、页面结构与导航")
    add_h2(doc, "2.1 演示账号登录")
    add_table(doc, ["教师", "账号/编号", "演示密码", "说明"], [
        ["王老师", "王老师 / T2024001", "123456", "浏览器验收默认账号"],
        ["林老师", "林老师 / T2024002", "123456", "第二演示教师"],
    ], [1600, 2700, 1600, 3460], font_size=9.5)
    add_steps(doc, [
        "在“教师登录”页输入教师姓名和密码，或选择演示账号快速登录。",
        "点击“登录教师端”。登录后点击“进入工作台”。",
        "确认页面右上角显示当前教师姓名；需要退出时点击姓名右侧下拉箭头，再点“退出登录”。",
    ])
    add_callout(doc, "安全说明", "123456 和 X-User-Id 仅用于演示。接入真实学校数据前必须替换为 SSO、JWT 或其他正式身份系统。", kind="caution")
    add_h2(doc, "2.2 导航结构")
    add_table(doc, ["区域", "入口", "作用"], [
        ["全局导航", "工作台首页、我的课程、个人设置", "跨课程管理与教师账号偏好"],
        ["课程导航", "工作空间首页、章节、班级、任务、资料、图谱、学情、讨论", "只作用于当前选中的课程"],
        ["顶部栏", "通知、回到工作台首页、教师菜单", "处理全局提醒与快速返回"],
        ["左下角", "返回我的课程、AI 助教", "切换课程或进入 AI 审核中心"],
    ], [1600, 2800, 4960], font_size=9.5)

    page_break(doc)
    add_h1(doc, "3. 工作台与课程管理")
    add_h2(doc, "3.1 使用工作台")
    add_bullets(doc, [
        "顶部指标：我的课程、待发布任务、待批改提交、学情提醒。点击对应区域可进入后续处理。",
        "我的课程卡片：显示课程状态、进度和班级信息；点击“进入课程”进入课程工作空间。",
        "今日待办与 AI 教学建议：用于定位批改、风险学生和教学建议；动作完成后可刷新数据。",
        "通知中心：点击顶部铃铛查看任务、预警和 AI 类型通知，点击通知即可标记为已读。",
    ])
    add_h2(doc, "3.2 创建课程")
    add_steps(doc, [
        "在工作台或“我的课程”点击“创建课程”。",
        "基础信息：填写课程名称、课程代码、学期、专业/方向、课程说明和封面。",
        "教学设置：设置开课周、结课周、总学时、每周课时、授课模式，并维护章节概览。",
        "完成创建：检查右侧学生端预览和“创建进度检查”，确认无误后完成创建。",
        "创建成功后系统进入邀请学生页面；继续建立或选择班级并发出邀请。",
    ])
    add_callout(doc, "成功判定", "新课程出现在“我的课程”列表；课程状态为筹备中或进行中；进入课程后左侧出现课程工作空间导航。")
    add_h2(doc, "3.3 归档、恢复与删除")
    add_steps(doc, [
        "在课程卡片的更多菜单中选择“归档课程”，归档项会与活动课程分开显示。",
        "对已归档课程选择“移出归档”，课程恢复到原有状态。",
        "删除前确认课程不再需要；删除动作会级联清理其课程业务数据，属于不可恢复操作。",
    ])
    add_callout(doc, "风险提示", "删除课程不是回收站操作。正式环境应先备份数据库，并由具备权限的管理员执行。", kind="risk")

    page_break(doc)
    add_h1(doc, "4. 课程章节内容")
    add_body(doc, "入口：进入课程后，在左侧选择“课程章节内容”。该模块统一组织知识点、课件和课后练习，并控制学生端可见性。")
    add_figure(doc, content_crop, "图 1  课程章节目录（桌面视口实测）", alt_text="课程章节内容页面，显示六个章节、发布状态、知识点、课件和练习数量")
    add_h2(doc, "4.1 新建章节")
    add_steps(doc, [
        "点击右上角“新建章节”或目录底部“添加新章节”。",
        "填写章节名称和说明，选择教学方式：理论讲授、翻转课堂、案例教学、项目制教学、实验实训或混合式教学。",
        "点击“创建章节”。新章节默认以草稿状态保存。",
    ])
    add_h2(doc, "4.2 维护章节内容")
    add_steps(doc, [
        "点击章节行打开右侧章节详情。",
        "在“知识点”页签点击“添加知识点”，填写名称、说明和难度。",
        "在“教学课件”页签点击“上传课件”；用“学生可见”开关控制单份资料。",
        "在“课后练习”页签点击“创建练习”，填写名称、题型、难度和截止时间；选择保存草稿或发布给学生。",
        "需要更改课堂组织方式时，直接修改“教学方式”下拉框。",
    ])
    add_h2(doc, "4.3 发布与学生视角")
    add_steps(doc, [
        "在章节详情中打开“学生端状态”开关。发布后，与该章节同名且状态正常的资料会同步调整为学生可见。",
        "关闭开关会撤回章节，相关资料恢复为仅教师可见。",
        "点击页面右上角“查看学生视角”，检查学生实际能看到的已发布章节、知识点、资料和练习。",
    ])
    add_callout(doc, "验收要点", "学生预览只显示已发布章节；草稿章节不出现。该可见性闭环已由自动化用例验证。")

    page_break(doc)
    add_h1(doc, "5. 班级管理与邀请学生")
    add_h2(doc, "5.1 新建与查看班级")
    add_steps(doc, [
        "进入“班级管理”，点击“新建班级”。",
        "填写班级名称、年级、专业方向、班级状态、上课安排和任课教师，点击“创建班级”。",
        "点击班级卡片或“查看详情”，检查班级规模、邀请码、上课安排和最近加入状态。",
        "点击“查看全部”打开学生加入状态，可按状态筛选并搜索姓名或学号。",
    ])
    add_h2(doc, "5.2 发出邀请")
    add_steps(doc, [
        "从班级卡片点击“邀请学生”，或在课程工作空间进入“邀请学生加入”。",
        "页面右上角选择目标班级，避免把邀请信息发给错误班级。",
        "分享链接：点击“复制”，将链接发给学生。",
        "二维码：点击“下载二维码”，可投屏、打印或发到班群。",
        "班级邀请码：点击“复制”，学生在加入页面输入邀请码。",
        "需要让旧邀请失效时点击“重新生成链接/邀请码”；系统会生成新的 8 位代码。",
    ])
    add_callout(doc, "注意", "重新生成后，旧链接、旧二维码和旧邀请码不应继续使用；请立即把新邀请信息发给学生。", kind="caution")
    add_h2(doc, "5.3 批量导入状态")
    add_callout(doc, "当前版本", "批量导入后端接口已验证：重复学号会跳过，导入结果会返回创建、加入和跳过数量。但当前页面的 Excel/CSV 选择与模板下载仍是展示入口，尚未完成 UI 提交联调。需批量导入时由管理员调用接口，或等待页面联调完成。", kind="caution")

    page_break(doc)
    add_h1(doc, "6. 任务创建、测试用例与发布")
    add_figure(doc, task_crop, "图 2  任务管理列表（桌面视口实测）", alt_text="任务管理页面，包含筛选、任务列表、提交概览和查看成绩入口")
    add_h2(doc, "6.1 创建任务")
    add_steps(doc, [
        "进入“任务管理”，点击“新建作业”。",
        "基本信息：填写任务标题、题目说明、知识点、任务类型和难度。",
        "测试用例：为编程题配置公开和隐藏测试，并为各用例设置权重。公开测试对学生可见，隐藏测试只对教师可见。",
        "发布设置：填写截止时间，选择下发班级和提示级别。",
        "预览发布：复核题目与范围，点击“确认发布”；暂不发布时在第一步选择“保存为草稿”。",
    ])
    add_callout(doc, "校验规则", "编程任务必须至少包含一个测试用例；空测试用例任务会返回 422，已由业务校验用例验证。")
    add_h2(doc, "6.2 使用 AI 生成草稿")
    add_steps(doc, [
        "在任务创建面板点击“与 AI 对话生成练习草稿”。",
        "描述知识点、题型、难度和班级薄弱项，或选择预设提示。",
        "点击“生成到草稿”，逐项检查题干、答案/代码、测试用例、权重和截止时间。",
        "教师确认后再发布。AI 只提供草稿，不应绕过教师审核。",
    ])
    add_h2(doc, "6.3 任务状态与后续操作")
    add_table(doc, ["状态", "学生可见", "教师可执行"], [
        ["草稿", "否", "编辑、继续发布"],
        ["待发布", "按发布时间", "检查发布设置"],
        ["已发布", "是", "预览学生视角、查看成绩、监控提交"],
        ["已结束", "按课程规则", "查看成绩与历史数据"],
    ], [1600, 1800, 5960], font_size=9.5)

    page_break(doc)
    add_h1(doc, "7. 任务监控、批改与成绩发布")
    add_h2(doc, "7.1 查看提交与测试结果")
    add_steps(doc, [
        "在任务管理中对已发布任务点击“预览学生视角”进入任务监控，或通过任务导航进入“任务监控”。",
        "选择任务，按学生姓名/学号和状态筛选提交记录。",
        "点击一条提交，检查版本、源代码、通过用例数、运行时间、用例明细和 AI 诊断。",
        "隐藏测试明细只对教师显示；学生端只获得公开测试信息。",
    ])
    add_h2(doc, "7.2 批改与发布成绩")
    add_steps(doc, [
        "在任务卡片点击“查看成绩”，或从监控详情点击“进入批改”。",
        "选择目标学生，检查提交版本、代码、每个测试用例和当前状态。",
        "调整四项评分依据：自动测试 40 分、代码质量 30 分、实验报告 20 分、课堂表现 10 分。",
        "四项之和必须等于最终得分；每项不能超过对应上限。",
        "填写具体教师反馈。点击“保存草稿”只保存批改；点击“发布成绩”才对学生发布。",
    ])
    add_figure(doc, SOURCE_IMAGES / "05-grading-desktop.png", "图 3  学生成绩与四维评分依据（桌面视口实测）", alt_text="学生成绩页面，左侧学生列表，中间代码与测试结果，右侧四维评分和教师反馈")
    add_callout(doc, "验收要点", "已验证教师创建任务、发布到班级、学生提交、自动评测、教师保存四维评分并发布成绩的完整闭环。")

    page_break(doc)
    add_h1(doc, "8. 资料管理与回收站")
    add_figure(doc, material_crop, "图 4  教学材料管理（桌面视口实测）", alt_text="教学材料页面，包含文件夹、类型筛选、资料列表、资料预览和 AI 大纲入口")
    add_h2(doc, "8.1 上传与整理资料")
    add_steps(doc, [
        "进入“资料管理”，选择目标文件夹或新建文件夹。",
        "点击“上传课件”或“上传讲义”选择本地文件；也可点击“添加链接”添加外部资料。",
        "上传后在右侧预览区检查标题、章节、文件大小、格式、状态和内容链接。",
        "需要按章节管理时，进入章节详情的“教学课件”上传，资料会自动归入当前章节。",
        "只有 visibility=students 的资料可由学生下载；教师专用资料对学生请求返回 403。",
    ])
    add_h2(doc, "8.2 删除、恢复和移动")
    add_steps(doc, [
        "删除单份资料后，资料进入回收站，正常列表和下载入口不再显示。",
        "进入左侧回收站，找到资料后执行恢复。恢复后资料回到正常状态并可再次下载。",
        "删除空文件夹会永久删除该空目录；删除包含资料的文件夹会将目录及其资料放入回收站。",
        "在已删除目录中，可将资料单独保留到“未分类资料”，或移动到现有文件夹。最后一份资料移出后，空的已删除目录会被清理。",
    ])
    add_callout(doc, "数据安全", "永久删除前先确认目录为空；对重要资料，建议先保留原文件和数据库备份。", kind="risk")
    add_h2(doc, "8.3 AI 生成材料目录")
    add_steps(doc, [
        "点击页面底部“生成大纲”。",
        "在候选目录中勾选需要的项，点击“确认并生成目录”。",
        "确认后目录才写入课程资料库；未确认的 AI 候选不会保存。",
    ])

    page_break(doc)
    add_h1(doc, "9. 课程知识图谱")
    add_figure(doc, graph_current, "图 5  新版课程知识图谱编辑器（桌面视口实测）", alt_text="新版课程知识图谱页面，左侧包含资料生成和图谱列表，中间是节点关系画布，右侧显示属性面板和来源资料")
    add_h2(doc, "9.1 新建与打开图谱")
    add_steps(doc, [
        "进入“课程知识图谱”。顶部统计卡显示图谱总数、已发布数量、当前节点数和当前关系数。",
        "需要从零开始时，点击左侧“新建空白图谱”。系统会创建一张含默认核心节点的草稿图谱。",
        "左侧“图谱列表”按更新时间显示当前课程的全部图谱；点击列表项加载对应图谱，当前项以蓝色边框高亮。",
        "点击顶部“刷新”可重新加载图谱列表和当前图谱。刷新会覆盖尚未保存的画布修改。",
    ])
    add_h2(doc, "9.2 上传资料生成图谱")
    add_steps(doc, [
        "在左侧“资料生成”区域填写图谱名称、发布班级和说明。发布班级可使用中文逗号、英文逗号或换行分隔。",
        "点击上传区域选择 PDF、Word DOCX、Markdown 或 TXT 文件；选中后检查文件名和大小。",
        "点击“分析并生成图谱”。系统上传资料、抽取知识点和关系，并创建一张草稿图谱。",
        "生成完成后，在右侧“来源资料”检查文件名、大小和资料摘要，再继续人工校正节点与关系。",
    ])
    add_h2(doc, "9.3 编辑节点与关系")
    add_steps(doc, [
        "工具栏选择“选择”模式，点击节点或关系，在右侧属性面板编辑名称、类型、难度、说明或关系类型。",
        "点击“节点”新增自定义节点。节点默认只存在于当前前端编辑状态，保存草稿后才会持久化。",
        "选择“连接”模式和“前驱/后继/相关”关系类型，依次点击起始节点和目标节点创建关系。",
        "拖拽节点可调整位置；点击“布局”可将核心节点居中并把其余节点按椭圆分布。",
        "选中节点或关系后点击“删除”。删除节点时，其关联关系会同步从当前画布移除。",
    ])
    add_h2(doc, "9.4 保存、发布与删除")
    add_steps(doc, [
        "完成编辑后点击顶部“保存草稿”，保存图谱信息、节点、关系和坐标。刷新页面后应能恢复保存结果。",
        "点击“发布”时，系统会先保存当前编辑，再将图谱状态更新为已发布。",
        "右侧“关联节点”用于查看当前节点的入边和出边；点击关联项可切换到对应关系。",
        "需要删除整张图谱时，点击右侧“删除当前图谱”并确认。该操作会从服务端永久删除当前图谱。",
    ])
    add_callout(doc, "保存边界", "新增节点、属性修改、连边、拖拽、布局和画布删除会先进入前端状态。切换图谱或刷新页面前必须点击“保存草稿”，否则未保存修改可能丢失。", kind="caution")

    page_break(doc)
    add_h1(doc, "10. 课堂讨论")
    add_h2(doc, "10.1 创建、发布与结束")
    add_steps(doc, [
        "在课程工作空间点击快捷入口“课堂讨论”，或从课程左侧导航进入。",
        "选择班级，填写讨论标题和讨论内容。",
        "选择直接发布或保存为草稿。草稿不会出现在学生讨论列表。",
        "发布后进入实时讨论视图，系统定时刷新回复；学生可提交回复。",
        "讨论结束时点击“结束讨论”。结束后的讨论进入历史记录，并从学生可参与列表移除。",
    ])
    add_h2(doc, "10.2 状态与行为")
    add_table(doc, ["状态", "学生端", "教师端"], [
        ["草稿", "不可见", "可发布"],
        ["进行中/已发布", "可见并可回复", "实时查看回复、结束讨论"],
        ["已结束", "不再出现在可参与列表", "保留历史记录"],
    ], [1800, 3200, 4360], font_size=9.5)
    add_callout(doc, "接口实测", "草稿隐藏、发布可见、学生回复计数、教师读取回复、结束后禁止继续回复均已通过自动化用例。")

    page_break(doc)
    add_h1(doc, "11. 学情分析与 AI 审核")
    add_figure(doc, analytics_crop, "图 6  学情分析班级总览（桌面视口实测）", alt_text="学情分析页面，显示平均完成率、平均得分、逾期率、提示等级、风险学生和知识点掌握度")
    add_h2(doc, "11.1 班级总览")
    add_bullets(doc, [
        "摘要指标：平均完成率、平均得分、逾期率、平均提示等级、风险学生、薄弱知识点。",
        "图表：成绩分布、高频错误和知识点掌握度。",
        "切换班级或任务后重新检查数据口径；导出报告按钮当前提供页面成功反馈，导出文件格式需后续联调确认。",
    ])
    add_h2(doc, "11.2 个体诊断与预警")
    add_steps(doc, [
        "点击“个体诊断”，选择班级，再按姓名或学号搜索学生。",
        "查看课程进度、平均得分、提交次数、最高提示等级和分知识点掌握度。",
        "点击“预警中心”，查看风险等级、命中规则和最近活跃时间。",
        "结合原始提交、测试结果和课堂记录进行人工判断，不要只依据 AI 解释。",
    ])
    add_h2(doc, "11.3 AI 审核中心")
    add_steps(doc, [
        "从左下角“AI 助教”进入审核中心。",
        "选择待审核记录，检查学生、任务、诊断类型、置信度、来源和是否为规则兜底。",
        "选择“接受诊断”“修改后接受”或“驳回”。",
        "修改后接受会单独保存教师修订版本；原始 AI 诊断不会被覆盖。",
    ])
    add_callout(doc, "责任边界", "风险等级由后端规则计算，AI 只负责解释证据；教师应对最终干预和评价负责。", kind="caution")

    page_break(doc)
    add_h1(doc, "12. 通知、个人设置与移动端")
    add_h2(doc, "12.1 通知中心")
    add_steps(doc, [
        "点击顶部铃铛，打开通知中心。",
        "通知按 AI、预警、任务等类型显示；未读数量显示在角标。",
        "点击通知标记为已读；角标会随未读数量变化。",
    ])
    add_h2(doc, "12.2 个人设置")
    add_bullets(doc, [
        "个人资料页展示姓名、教师编号、院系和邮箱，当前由统一认证服务维护，页面为只读。",
        "通知设置、AI 助教偏好、数据与隐私为设置分区；当前“保存偏好”主要提供页面反馈，持久化范围应在部署前再次确认。",
    ])
    add_h2(doc, "12.3 移动端使用")
    add_body(doc, "现版本已在 390 × 844 手机视口检查工作台和课程章节内容，无页面横向溢出。移动端适合浏览、检查章节和完成轻量操作；复杂批改、图谱编辑和多列表格仍建议使用桌面端。")
    add_figure(doc, SOURCE_IMAGES / "10-content-mobile.png", "图 7  课程章节内容（390 px 手机视口实测）", width=2.35, alt_text="手机端课程章节内容页面，无横向溢出")

    page_break(doc)
    add_h1(doc, "13. 常见故障与处理")
    add_table(doc, ["现象", "可能原因", "处理方法"], [
        ["无法连接教学后端", "8001 未启动或端口被占用", "访问健康接口；检查监听端口；重启后端"],
        ["前端能打开但无数据", "Vite 代理目标错误或后端异常", "确认 vite.config.ts 指向 8001；查看后端日志"],
        ["403 无权限", "用户角色不匹配或学生未加入课程", "确认演示账号；检查 X-User-Id 与课程归属"],
        ["422 校验失败", "必填字段、测试用例、评分维度不完整", "按错误提示补齐；评分维度之和须等于总分"],
        ["资料下载 404", "资料在回收站或文件不存在", "检查回收站并恢复；核对 backend/uploads"],
        ["深层地址刷新 404", "生产静态托管没有 SPA 回退", "从根地址进入；正式部署配置 Nginx try_files"],
        ["pnpm 要求删除 node_modules", "pnpm/Node 版本与现有依赖目录不同", "使用锁文件对应版本；不要在无确认时强制重建"],
        ["旧浏览器脚本超时", "按钮或完成提示文案已变化", "使用 operation-manual-check.mjs；同步更新旧脚本定位器"],
    ], [2200, 2800, 4360], font_size=8.8)
    add_h2(doc, "13.1 健康检查命令")
    add_bullets(doc, [
        "后端：Invoke-RestMethod http://127.0.0.1:8001/api/v1/health",
        "前端：(Invoke-WebRequest http://127.0.0.1:5173/ -UseBasicParsing).StatusCode",
        "后端测试：D:\\Anaconda\\python.exe -m pytest backend\\tests -q -p no:cacheprovider",
        "前端构建：node_modules\\.bin\\tsc.cmd -b；随后 node_modules\\.bin\\vite.cmd build",
        "现版本浏览器验收：node scripts\\operation-manual-check.mjs",
    ])

    page_break(doc)
    add_h1(doc, "14. 测试环境与总体结果")
    add_h2(doc, "14.1 实测环境")
    add_table(doc, ["项目", "实测值"], [
        ["日期/时区", "2026-08-12 / Asia/Shanghai"],
        ["工作区", "E:\\teacher；分支 teacher-develop；Git 66706fb + 当前待提交改动"],
        ["操作系统", "Windows PowerShell"],
        ["Node.js / pnpm", "v24.15.0 / 11.5.0"],
        ["Python", "3.13.9（D:\\Anaconda\\python.exe）"],
        ["浏览器", "Google Chrome，Playwright 无头模式"],
        ["视口", "桌面 1440×900；手机 390×844；另有班级页 1280×720"],
        ["服务", "UI 5173；API 8001；SQLite"],
    ], [2400, 6960], font_size=9.5)
    add_h2(doc, "14.2 总体结果")
    add_table(doc, ["验证层", "结果", "证据"], [
        ["后端自动化", "通过", "13 passed，1 条依赖弃用警告"],
        ["前端类型检查与构建", "通过", "3808 modules transformed；构建完成"],
        ["浏览器页面验收", "通过", "10 个主要视图 + 2 个视口；API 错误 0；控制台错误 0"],
        ["响应式检查", "通过", "桌面与手机检查页面横向溢出均为 false"],
        ["旧视觉脚本", "需更新", "旧按钮/提示文案导致定位器超时，不属于产品失败"],
    ], [2500, 1500, 5360], font_size=9.2)
    add_callout(doc, "构建提醒", "生产 JS 包约 1,958.15 kB（gzip 613.06 kB），超过 Vite 500 kB 提醒阈值。当前不影响功能，但建议后续使用路由懒加载或 manualChunks 优化。", kind="caution")

    page_break(doc)
    add_h1(doc, "15. 后端自动化测试用例明细")
    backend_cases = [
        ["B01", "健康与教师启动数据", "health=ok；王老师与课程返回", "通过"],
        ["B02", "任务-提交-批改-发布闭环", "任务发布、学生提交、评测、四维评分、成绩发布", "通过"],
        ["B03", "权限与业务校验", "学生访问教师接口 403；空测试用例任务 422", "通过"],
        ["B04", "课程归档、恢复与删除", "状态转换正确；删除返回 deleted=true", "通过"],
        ["B05", "章节发布与学生可见性", "仅已发布章节出现；撤回后隐藏", "通过"],
        ["B06", "邀请码刷新与批量导入", "8 位新码；重复学号跳过；加入状态正确", "通过"],
        ["B07", "课堂讨论闭环", "草稿隐藏、发布、回复、结束、禁止迟到回复", "通过"],
        ["B08", "空文件夹删除", "空目录永久删除且不进入回收站", "通过"],
        ["B09", "非空目录恢复", "目录与资料进入回收站并可整体恢复", "通过"],
        ["B10", "已删目录资料处置", "可保留到未分类或移动到现有目录", "通过"],
        ["B11", "资料关联知识图谱", "显式关联返回并出现在图谱节点详情", "通过"],
        ["B12", "评分演示数据幂等", "重复初始化不产生重复任务、提交、评测和成绩", "通过"],
        ["B13", "真实文件上传与权限", "上传/下载、403、删除、404、恢复完整通过", "通过"],
    ]
    add_table(doc, ["编号", "用例", "关键断言", "结果"], backend_cases,
              [720, 2460, 4980, 1200], font_size=8.2,
              aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    add_body(doc, "执行结果：13 passed, 1 warning in 3.87s。警告来自 Starlette TestClient 对 httpx 的弃用提示，不影响本轮业务断言。")

    page_break(doc)
    add_h1(doc, "16. 浏览器验收用例明细")
    browser_cases = [
        ["U01", "教师工作台", "/teacher/dashboard", "标题可见；无溢出"],
        ["U02", "我的课程", "/teacher/courses", "标题可见；无溢出"],
        ["U03", "课程章节内容", "/content", "标题可见；学生预览可打开"],
        ["U04", "班级管理", "/classes", "7 张班级卡片；无溢出"],
        ["U05", "邀请学生", "/invite", "邀请页可见；无溢出"],
        ["U06", "任务管理", "/tasks", "9 条任务；无溢出"],
        ["U07", "学生成绩", "/grading", "批改页可见；无溢出"],
        ["U08", "教学材料", "/materials", "列表与预览可见；无溢出"],
        ["U09", "课程知识图谱", "/graph", "交互画布可见；无溢出"],
        ["U10", "学情分析", "/analytics", "总览与个体诊断可达"],
        ["U11", "手机端工作台", "390×844", "scrollWidth=390；无溢出"],
        ["U12", "手机端章节内容", "390×844", "scrollWidth=390；无溢出"],
    ]
    add_table(doc, ["编号", "页面/场景", "地址/视口", "结果"], browser_cases,
              [720, 2400, 2580, 3660], font_size=8.7,
              aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    add_h2(doc, "16.1 浏览器验收汇总")
    add_bullets(doc, [
        "最终脚本结论 passed=true。",
        "所有检查页面 apiErrors=[]；consoleErrors=[]。",
        "桌面和手机视口均未发现文档宽度大于视口宽度。",
        "任务列表断言第一次因测试脚本选择器写成 .task-v2-card 产生假阴性；修正为 .task-v2-list > article 后重跑通过，页面实际返回 9 条任务。",
    ])
    add_h2(doc, "16.2 旧脚本失配记录")
    add_bullets(doc, [
        "fullstack-check.mjs 等待旧按钮“邀请学生加入”；当前班级卡片按钮为“邀请学生”。应更新定位器，不判定为产品失败。",
        "exact-visual-check.mjs 等待旧提示“课程创建准备完成”；当前完成页为“课程创建完成”。应更新完成提示，不判定为产品失败。",
    ])

    page_break(doc)
    add_h1(doc, "附录 A. 日常操作检查清单")
    add_h2(doc, "课前")
    add_bullets(doc, [
        "确认当前课程和班级选择正确。",
        "确认本周章节已发布，学生视角可以看到知识点、课件和练习。",
        "确认任务截止时间、下发班级、公开/隐藏测试和提示级别。",
        "确认邀请链接/邀请码仍有效，班级容量和加入状态正常。",
        "确认资料权限：学生资料为 students，教师资料为 teacher。",
    ])
    add_h2(doc, "课中")
    add_bullets(doc, [
        "课堂讨论发布到正确班级；结束后保存历史记录。",
        "任务监控中关注未提交、连续失败和高提示等级学生。",
        "需要解释 AI 诊断时回到原始代码、测试和课堂证据。",
    ])
    add_h2(doc, "课后")
    add_bullets(doc, [
        "批改时四维评分不越上限且总和等于最终得分。",
        "填写具体反馈后再发布成绩，避免只给分不说明。",
        "在学情分析中复核薄弱知识点与风险学生，并记录干预动作。",
        "定期清理回收站前备份数据库和重要资料。",
    ])
    add_h2(doc, "上线前")
    add_bullets(doc, [
        "替换演示登录为正式 SSO/JWT，并启用 HTTPS、限制 CORS、设置上传类型和大小限制。",
        "配置数据库与上传目录备份并验证恢复；配置 SPA 深层路由回退、进程守护和反向代理。",
        "完成批量导入、模板下载、导出报告、设置持久化等展示入口的联调。",
    ])

    page_break(doc)
    add_h1(doc, "附录 B. 数据影响与已知限制")
    add_table(doc, ["操作", "数据影响", "可恢复性"], [
        ["发布/撤回章节", "切换章节状态，并同步同章节资料的学生可见性", "可反向切换"],
        ["重新生成邀请码", "替换班级 join_code", "旧码不可继续使用"],
        ["发布任务", "任务绑定班级并对学生可见", "当前无独立撤回说明"],
        ["发布成绩", "成绩状态变为 grade_published", "应通过后续更正流程处理"],
        ["删除资料", "进入回收站，下载返回 404", "可恢复"],
        ["删除空文件夹", "目录永久删除", "不可恢复"],
        ["删除课程", "级联删除课程业务数据", "不可恢复，需备份"],
    ], [2200, 4800, 2360], font_size=9)
    add_h2(doc, "已知限制")
    add_bullets(doc, [
        "演示认证不是生产级安全方案。",
        "批量导入页面上传和 Excel 模板下载尚未完成 UI 联调；后端接口已通过。",
        "学情导出和部分设置保存按钮当前以页面反馈为主，交付前应确认最终文件或持久化行为。",
        "前端生产包较大，建议后续拆包优化首屏性能。",
        "旧视觉检查脚本需要随当前文案和按钮结构更新。",
        "FastAPI 直接托管 dist 时，深层 SPA 地址刷新可能 404；生产环境需配置回退。",
    ])
    add_callout(doc, "最终结论", "当前版本已具备课程、章节、班级、任务、提交、批改、资料、图谱、讨论与学情的主要教师端闭环；核心后端与现版本页面验收通过。正式上线前需优先完成正式认证、备份恢复、展示入口联调和性能拆包。")

    # Core properties and final save.
    props = doc.core_properties
    props.title = "CodeTrack 教师端详细操作手册（实测版）"
    props.subject = "教师端操作、测试用例与验收结果"
    props.author = "CodeTrack 项目组"
    props.keywords = "CodeTrack, 教师端, 操作手册, 测试用例, 验收"
    props.comments = "基于 2026-08-14 当前代码、浏览器复核与新版知识图谱界面更新"
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
