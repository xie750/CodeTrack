from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"E:\teacher")
OUT_DIR = ROOT / "artifacts"
ASSET_DIR = OUT_DIR / "codetrack_paper_assets"
OUT_PATH = OUT_DIR / "CodeTrack_垂类大模型与核心方法_学术论文修改版.docx"

BLUE = "2F6F9F"
GREEN = "3A8D74"
LIGHT_BLUE = "EAF3F8"
LIGHT_GREEN = "EAF5F1"
GRAY = "5B6570"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_three_line_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("left", "right", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    for edge, size in (("top", "12"), ("bottom", "12"), ("insideH", "8")):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), "333333")
        borders.append(node)


def set_run_font(run, east_asia="宋体", western="Times New Roman", size=12, bold=None, color=None):
    run.font.name = western
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_format(p, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Cm(0.85) if indent else None


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_para_format(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, east_asia="黑体", size=12, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=8)
    size = 14 if level == 1 else 12
    east = "宋体" if level == 1 else "仿宋_GB2312"
    r = p.add_run(text)
    set_run_font(r, east_asia=east, size=size, bold=(level <= 2), color="1E2D38")
    p.paragraph_format.keep_with_next = True
    return p


def add_caption(doc, text, is_table=False):
    p = doc.add_paragraph()
    set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=4)
    r = p.add_run(text)
    set_run_font(r, size=12)
    p.paragraph_format.keep_with_next = is_table
    return p


def add_table(doc, caption, headers, rows, widths=None, page_break_before=False):
    cap = add_caption(doc, caption, is_table=True)
    if page_break_before:
        cap.paragraph_format.page_break_before = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            cell.width = Cm(widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(header)
        set_run_font(r, east_asia="宋体", size=10.5, bold=True)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            if widths:
                cell.width = Cm(widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            r = p.add_run(str(value))
            set_run_font(r, east_asia="宋体", size=10.5)
    set_three_line_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_equation(doc, equation, number):
    p = doc.add_paragraph()
    set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4)
    r = p.add_run(f"{equation}    （{number}）")
    set_run_font(r, east_asia="宋体", size=12)


def add_picture(doc, path, caption, width=15.0):
    p = doc.add_paragraph()
    set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(path), width=Cm(width))
    p.paragraph_format.keep_with_next = True
    add_caption(doc, caption)


def img_font(size, bold=False):
    path = r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"
    return ImageFont.truetype(path, size)


def draw_box(draw, box, text, face, edge, size=28):
    draw.rounded_rectangle(box, radius=16, fill=face, outline=edge, width=3)
    font = img_font(size)
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=7, align="center")
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    x1, y1, x2, y2 = box
    draw.multiline_text(((x1 + x2 - tw) / 2, (y1 + y2 - th) / 2 - 2), text,
                        fill="#1E2D38", font=font, spacing=7, align="center")


def arrow(draw, start, end, color="#64737F", width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        pts = [(ex, ey), (ex - 14 if ex > sx else ex + 14, ey - 9),
               (ex - 14 if ex > sx else ex + 14, ey + 9)]
    else:
        pts = [(ex, ey), (ex - 9, ey - 14 if ey > sy else ey + 14),
               (ex + 9, ey - 14 if ey > sy else ey + 14)]
    draw.polygon(pts, fill=color)


def build_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    im = Image.new("RGB", (2100, 1200), "white"); d = ImageDraw.Draw(im)
    labels = ["应用层", "智能层", "知识层", "模型与平台层"]
    for y, label in zip((90, 370, 650, 930), labels):
        d.text((45, y), label, font=img_font(34, True), fill="#1E2D38")
    for i, label in enumerate(["学生端\n分步答疑·错题复习", "教师端\n资源审核·批改辅助", "管理端\n版本管理·安全审计", "科研入口\n文献与代码辅助"]):
        draw_box(d, (300+i*430, 55, 650+i*430, 205), label, "#F6FAFC", "#2F6F9F", 26)
    for i, label in enumerate(["问题分类", "混合检索", "步骤拆解", "主模型求解", "答案质检"]):
        x = 245+i*365; draw_box(d, (x, 340, x+285, 480), label, "#EAF5F1", "#3A8D74", 26)
        if i < 4: arrow(d, (x+285, 410), (x+355, 410))
    for i, label in enumerate(["课程文档库", "结构化题库", "知识点关系", "向量索引与元数据"]):
        draw_box(d, (300+i*430, 625, 650+i*430, 760), label, "#EAF3F8", "#2F6F9F", 25)
    for i, label in enumerate(["Qwen2.5-7B-Instruct\n+ LoRA适配器", "CodeTrack统一数据对象", "权限·日志·代码沙箱"]):
        draw_box(d, (315+i*560, 900, 790+i*560, 1065), label, "#F7F7F7", "#7D8790", 25)
    for y1, y2 in ((205, 340), (480, 625), (760, 900)): arrow(d, (1050, y1+8), (1050, y2-8), "#7D8790")
    im.save(ASSET_DIR / "architecture.png")

    im = Image.new("RGB", (2100, 820), "white"); d = ImageDraw.Draw(im)
    labels = ["问题与角色识别", "查询改写", "关键词+向量检索", "融合排序与重排", "上下文组装", "受约束生成", "证据一致性检查"]
    colors = ["#EAF3F8", "#EAF3F8", "#EAF5F1", "#EAF5F1", "#FFF5E6", "#FFF5E6", "#FDECEC"]
    edges = ["#2F6F9F", "#2F6F9F", "#3A8D74", "#3A8D74", "#C8892E", "#C8892E", "#C65B5B"]
    for i, label in enumerate(labels):
        x = 55+i*290; draw_box(d, (x, 200, x+240, 345), label, colors[i], edges[i], 23)
        if i < 6: arrow(d, (x+240, 272), (x+280, 272))
    draw_box(d, (680, 535, 1420, 690), "资料不足 / 引用冲突 / 权限不符\n重新检索或转教师复核", "#FDECEC", "#C65B5B", 26)
    arrow(d, (1860, 345), (1420, 535), "#C65B5B"); arrow(d, (680, 615), (250, 345), "#C65B5B")
    im.save(ASSET_DIR / "rag_flow.png")

    im = Image.new("RGB", (2100, 940), "white"); d = ImageDraw.Draw(im)
    draw_box(d, (730, 70, 1370, 205), "调度器：状态、权限与重试控制", "#EAF3F8", "#2F6F9F", 28)
    agents = [(90, "问题分类智能体\n输出题型标签"), (590, "检索增强智能体\n输出证据集合"), (1090, "步骤拆解智能体\n输出子任务链"), (1590, "答案质检智能体\n输出通过/问题清单")]
    for x, label in agents:
        draw_box(d, (x, 345, x+420, 515), label, "#EAF5F1", "#3A8D74", 25)
        arrow(d, (1050, 205), (x+210, 345))
    draw_box(d, (470, 650, 900, 795), "主模型与代码工具\n仅接受受控调用", "#FFF5E6", "#C8892E", 25)
    draw_box(d, (1200, 650, 1630, 795), "教师/管理员复核\n高影响操作最终确认", "#FDECEC", "#C65B5B", 25)
    arrow(d, (1300, 515), (900, 650)); arrow(d, (1800, 515), (1415, 650), "#C65B5B")
    note = "质检未通过时最多回溯2次；连续失败、资料冲突或权限风险时转人工"
    bbox = d.textbbox((0, 0), note, font=img_font(24)); d.text(((2100-(bbox[2]-bbox[0]))/2, 865), note, font=img_font(24), fill="#5B6570")
    im.save(ASSET_DIR / "agent_flow.png")


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(0)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run = footer.add_run(); run._r.extend([fld_begin, instr, fld_end]); set_run_font(run, size=9)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_diagrams()
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph(); set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=54, after=14)
    r = p.add_run("垂类大模型与核心方法"); set_run_font(r, east_asia="宋体", size=18, bold=True, color="1E2D38")
    p = doc.add_paragraph(); set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    r = p.add_run("面向人工智能专业课程的 CodeTrack 一体化智能教学平台"); set_run_font(r, size=13, color=GRAY)

    p = doc.add_paragraph(); set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=5)
    r = p.add_run("摘  要："); set_run_font(r, east_asia="黑体", size=12, bold=True)
    r = p.add_run("针对通用大语言模型在人工智能专业课程中容易出现术语混淆、推导跳跃、方法与课堂教学不一致以及回答依据难以核验等问题，本文提出面向课程生态的垂类模型与自主化学科智能体集成方案。该方案以Qwen2.5-7B-Instruct为基座模型，采用LoRA进行参数高效适配，以“问题—题型—分步推导—答案”四元组组织学科推理数据；在外部知识层构建具有课程、章节、版本、权限和来源信息的知识库，并通过关键词检索、向量检索、融合排序和证据检查形成可追溯RAG链路；在任务层设置问题分类、检索增强、步骤拆解和答案质检四类智能体，通过受控调度、有限回溯和人工复核完成复杂问题求解。本文进一步给出数据质量控制、三端集成、安全边界和对比实验方案。鉴于垂类模型训练与真实用户试验尚在推进，本文不报告未经验证的性能结论，而将所述指标作为后续实验与验收依据。"); set_run_font(r)
    p = doc.add_paragraph(); set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=28)
    r = p.add_run("关键词："); set_run_font(r, east_asia="黑体", size=12, bold=True)
    r = p.add_run("学科垂类大模型；参数高效微调；检索增强生成；多智能体协同；智能教学"); set_run_font(r)

    add_heading(doc, "1 引言", 1)
    add_body(doc, "高校人工智能专业课程同时包含概念、公式、代码和复杂推理任务。通用大语言模型虽然具备较强的语言生成能力，但其开放域训练分布与具体课程的术语体系、解题规范和教学方法并不完全一致，容易出现最终答案看似合理而中间步骤无法核验的情况。参数高效微调能够以较低资源完成领域适配[1-2]，RAG能够引入可更新的外部知识并提供来源[3-4]，智能体方法则可将复杂任务拆分为多个受控环节[5-6]。三者结合为课程问答中的“知识准确、过程可查、方法一致”提供了可行技术路径。")
    add_body(doc, "CodeTrack面向本科计算机与人工智能方向核心课程，服务学生、教师和管理员三类用户。研究重点不是建设一个功能齐全的通用教务平台，而是回答三个具体问题：第一，低资源条件下如何利用结构化推理数据增强模型的学科表达与分步推导能力；第二，如何将课程知识检索、问题拆解和结果质检组织为可追溯的协同链路；第三，如何在教育场景中约束模型的权限、数据和高影响操作。")
    add_body(doc, "本文属于方法设计与实验协议稿。现阶段已形成网站原型及三端业务设计，但垂类模型训练、智能体联调和用户效果验证尚未全部完成。因此，文中涉及的模型版本、数据规模、参数配置和评测指标均为依据项目申报方案确定的实施基线，后续正式论文应以真实代码、日志和实验结果替换计划性表述。")

    add_heading(doc, "2 研究对象与总体架构", 1)
    add_heading(doc, "2.1 研究对象与范围", 2)
    add_body(doc, "研究对象为本科计算机与人工智能方向核心课程中的课后习题求解与课程答疑，优先选择程序设计、数据结构、数据库和机器学习等课程开展试点。学生端提供课程问答、分步提示、错题复习和学习建议；教师端承担资料审核、练习确认、批改复核与学情解释；管理员端负责权限、知识库版本、模型运行和安全日志。科研入口作为平台扩展场景保留，用于文献理解、代码说明和研究资料检索，但不纳入本阶段教学问答实验的效果结论。")
    add_heading(doc, "2.2 总体技术架构", 2)
    add_body(doc, "系统采用应用层、智能层、知识层以及模型与平台层的分层架构。应用层承载三端业务和科研入口；智能层组织分类、检索、拆解、求解与质检；知识层维护课程资料、结构化题库、知识点关系和向量索引；模型与平台层提供基座模型、LoRA适配器、统一数据对象、权限控制、日志审计和代码沙箱。各层通过明确接口衔接，使课程数据更新、模型版本切换和业务功能扩展相互解耦。")
    add_picture(doc, ASSET_DIR / "architecture.png", "图2-1 CodeTrack垂类模型与智能体总体架构")

    add_heading(doc, "3 学科结构化数据集构建", 1)
    add_heading(doc, "3.1 数据来源与样本结构", 2)
    add_body(doc, "数据拟来源于课程标准、经授权的教材与课件、实验指导、公开习题、教师确认的解答以及经过脱敏处理的典型错误。项目申报方案计划构建不少于350条“问题—题型—分步推导—答案”四元组样本，题型覆盖概念辨析、数值计算、证明推理和实际应用，并按8:1:1划分训练集、验证集和测试集。正式实施时应按课程、来源和题型分别统计数量，避免同一道题及其改写版本同时出现在训练集和测试集中。")
    add_table(doc, "表3-1 学科推理数据字段设计", ["字段", "主要内容", "质量要求"], [
        ["问题", "题干、条件、课程与知识点", "语义完整，条件无缺失"],
        ["题型", "概念、计算、证明、应用", "由双人独立标注并复核"],
        ["分步推导", "步骤、依据、中间结果", "每步可独立检查，前后可导出"],
        ["答案", "结论、公式或可执行代码", "经教师或标准答案核验"],
        ["证据与版本", "资料编号、章节、权限、版本", "来源有效，可追溯、可回退"],
    ], [3.0, 6.0, 6.0])
    add_heading(doc, "3.2 数据处理与质量控制", 2)
    add_body(doc, "数据处理遵循“登记—解析—清理—标注—复核—分集—版本化”的流程。登记阶段记录资料权属、适用课程和使用目的；解析阶段保留公式、代码、表格和题目层级；清理阶段统一符号与LaTeX格式并进行语义去重；标注阶段补充题型、知识点、推导步骤和评分要点；复核阶段由两名标注者交叉检查并由任课教师处理分歧。正式论文应报告标注一致性指标，并公开不含隐私和版权限制的样本说明。")
    add_body(doc, "学生姓名、学号、原始作业、对话记录和密钥等信息不直接进入训练集。考试题、教材内容和校内资料须在授权范围内使用；不能仅以“公开可获得”推定可以训练或公开发布。数据集版本应与模型和知识库版本绑定，以支持实验复现与问题追踪。")

    add_heading(doc, "4 垂类模型构建", 1)
    add_heading(doc, "4.1 基座模型与LoRA适配", 2)
    add_body(doc, "依据项目申报方案，首版基座模型确定为Qwen2.5-7B-Instruct[7]。选择该模型主要考虑其中文理解、代码能力、开放许可和消费级GPU部署条件。为降低训练成本，采用LoRA在目标线性层中引入低秩增量矩阵，并冻结基座模型原始权重。其参数更新可表示为：")
    add_equation(doc, "W′ = W + ΔW，ΔW = BA，r ≪ min(d, k)", "1")
    add_body(doc, "式中，W为冻结的原始权重，B与A为可训练低秩矩阵，r为低秩维度。申报方案将r设为16，并以验证集损失和答案质量共同选择检查点。该设置属于预设方案，正式论文必须依据真实训练日志报告目标层、缩放系数、Dropout、训练轮数、批大小、最大序列长度、优化器、精度类型和随机种子。")
    add_table(doc, "表4-1 首版模型训练方案", ["项目", "方案值", "论文报告要求"], [
        ["基座模型", "Qwen2.5-7B-Instruct", "记录精确版本与许可证"],
        ["数据规模", "不少于350条四元组样本", "报告清洗前后数量与分布"],
        ["数据划分", "8:1:1", "按题源分组，防止相似题泄漏"],
        ["LoRA秩", "r=16", "结合验证集与消融实验确认"],
        ["初始学习率", "1×10⁻⁴", "报告调度器、轮数与早停条件"],
        ["训练环境", "单张RTX 3090，24 GB", "报告驱动、框架、显存与耗时"],
    ], [3.4, 4.5, 7.1], page_break_before=True)
    add_heading(doc, "4.2 模型发布与版本控制", 2)
    add_body(doc, "模型发布采用“离线评测—教师盲评—安全测试—小范围试用”的门控流程。每个适配器保存数据版本、训练配置、代码提交标识和评测结果；若专业能力提升伴随通用能力明显退化，或出现敏感信息泄露、越权建议等风险，则不得上线。平台按课程加载对应LoRA适配器和知识库，基础模型共享，但这一设计是否能够减少课程间干扰仍需通过跨课程测试验证。")

    add_heading(doc, "5 课程知识库与RAG", 1)
    add_heading(doc, "5.1 知识组织与索引", 2)
    add_body(doc, "课程知识库由文档资料、结构化题库和知识点关系组成。文档不按固定字符数机械截断，而是依据章节、概念、例题、函数、实验任务等教学语义单元分块，并保留标题、页码、题号、课程版本、适用角色和审核状态。原型方案使用Chroma管理向量索引，具体嵌入模型、分块长度和重叠比例须在实现后依据检索实验确定。")
    add_heading(doc, "5.2 混合检索与可验证生成", 2)
    add_body(doc, "检索链路先识别用户角色、课程、问题类型和对话上下文，再将原问题与规范化查询同时用于关键词检索和向量检索。候选结果经过融合排序，并按照相关性、来源可靠性、课程版本和访问权限进行重排。融合得分可表示为：")
    add_equation(doc, "S(d,q) = αS_keyword(d,q) + βS_dense(d,q) + γS_authority(d)", "2")
    add_body(doc, "其中，S_keyword为关键词相关度，S_dense为向量相似度，S_authority表示资料审核等级；α、β和γ由验证集确定。生成阶段只使用通过权限过滤的上下文，并要求输出关键结论对应的资料编号。证据检查发现资料不足、引用冲突或回答超出证据范围时，系统重新检索、降低结论强度或转教师复核，而不是继续生成确定性答案。")
    add_picture(doc, ASSET_DIR / "rag_flow.png", "图5-1 面向课程问答的可验证RAG链路")

    add_heading(doc, "6 自主化多智能体协同", 1)
    add_heading(doc, "6.1 智能体角色与边界", 2)
    add_body(doc, "核心问答链保留申报表确定的四个专用智能体，主模型求解作为受控工具调用，不额外计为独立智能体。练习生成、批改辅助和学情分析属于教学业务模块，可复用核心问答能力，但不与四个核心智能体混为同一层级。")
    add_table(doc, "表6-1 核心智能体职责与输出", ["智能体", "主要职责", "结构化输出"], [
        ["问题分类", "识别课程、题型、意图与风险", "题型标签、约束条件"],
        ["检索增强", "查询知识库并校验证据权限", "证据集合、来源与得分"],
        ["步骤拆解", "按题型策略生成可检查子任务", "有序步骤链、依赖关系"],
        ["答案质检", "检查逻辑、结论、引用和安全", "通过/不通过、问题清单"],
    ], [3.2, 6.4, 5.4], page_break_before=True)
    add_heading(doc, "6.2 状态控制与失败处理", 2)
    add_body(doc, "调度器维护用户角色、课程、原问题、题型、证据集合、子任务、工具结果、质检状态和人工复核标记。只有当上一步输出符合模式约束、权限检查通过且风险可控时，流程才进入下一阶段。质检未通过时最多回溯2次；连续失败、资料冲突、代码工具异常或涉及成绩发布等高影响操作时，流程停止自动执行并转交教师或管理员。")
    add_picture(doc, ASSET_DIR / "agent_flow.png", "图6-1 四智能体协同与人工复核机制")

    add_heading(doc, "7 安全、伦理与教育边界", 1)
    add_body(doc, "系统按照数据、身份、知识、生成、工具和上线六个层次设置控制。数据层执行授权登记、脱敏和最小化保存；身份层依据角色、课程和班级过滤内容；知识层防止未审核资料和提示注入文档进入正式索引；生成层检测敏感信息、无依据结论和受限答案；工具层在隔离环境中限制运行时间、内存、进程、文件和网络；上线层通过安全测试、灰度试用、告警和回滚控制风险。")
    add_body(doc, "教育场景中，系统不得依据与学习无关的个人属性评价学生，不自动形成成绩、处分或正式教学结论；练习发布、评分确认和知识库修改由教师负责。界面应明确标注“AI生成内容”，并提供引用展开、问题反馈和人工申诉入口。科研入口不得生成虚假实验数据或伪造文献，文献和数据分析结果必须保留来源及不确定性说明。")

    add_heading(doc, "8 实验设计与效果验证", 1)
    add_heading(doc, "8.1 对比与消融实验", 2)
    add_body(doc, "为区分各模块的实际贡献，实验不应只比较完整系统与通用模型，而应设置逐步增加组件的对照组。在相同测试集、可见资料、解码配置和评分规则下比较各组，并记录模型、数据和知识库版本。")
    add_table(doc, "表8-1 对比实验分组", ["组别", "模型与组件", "研究目的"], [
        ["G0", "原始Qwen2.5-7B-Instruct", "获得基座模型表现"],
        ["G1", "基座模型+RAG", "评估外部知识增益"],
        ["G2", "LoRA模型", "评估参数适配增益"],
        ["G3", "LoRA模型+RAG", "观察两类知识注入的组合效果"],
        ["G4", "LoRA+RAG+四智能体", "评估拆解、质检与回溯贡献"],
        ["G5", "DeepSeek-V3等通用模型", "提供外部通用模型参照"],
    ], [2.0, 7.0, 6.0])
    add_heading(doc, "8.2 评价指标", 2)
    add_table(doc, "表8-2 评价指标与计算方式", ["维度", "指标或方法", "说明"], [
        ["答案质量", "正确率、公式/代码可执行率", "教师或标准答案核验"],
        ["推理过程", "完整性、步骤可导出性", "教师5级量表盲评"],
        ["教学一致性", "课堂方法一致性评分", "由任课教师独立评价"],
        ["检索质量", "Recall@k、MRR、证据命中率", "使用带相关性标注的问题集"],
        ["可信生成", "引用准确率、答案忠实度", "核对结论是否受证据支持"],
        ["协同稳定性", "任务完成率、回溯率、人工转交率", "依据结构化运行日志统计"],
        ["系统性能", "首字延迟、总时延、显存占用", "统一硬件环境重复测量"],
        ["安全性", "越权、泄露、提示注入通过率", "使用专项安全测试集"],
    ], [2.5, 5.6, 6.9], page_break_before=True)
    add_body(doc, "推导完整性和方法一致性拟由2—3名教师进行盲评，并使用Kendall's W报告评价者一致性。除均值外，应报告样本量、标准差或置信区间；涉及多组比较时采用适合数据分布的显著性检验。至少选择3个典型问题展示完整输入、检索证据、推导过程、最终答案及与标准答案的比较。")
    add_heading(doc, "8.3 用户试用与当前状态", 2)
    add_body(doc, "比赛效果验证要求至少包含2名真实目标用户的结构化试用记录。项目计划邀请15—20名学生和2—3名教师开展两周小范围试用，记录使用场景、频次、完成时间、修改量、错误类型和主观评价。若累积50条经教师确认的有效纠错样本，可进入增量训练候选集，但是否触发更新仍应由离线评测和版本审批共同决定。")
    add_body(doc, "截至本稿生成时，网站原型与三端界面属于已形成的产品基础，垂类模型代码、LoRA训练结果、四智能体端到端运行数据和用户试验结果尚不能作为已完成成果报告。因此，本节给出的是实验协议而非实验结论，后续不得用预期指标替代真实结果。")

    add_heading(doc, "9 平台集成与持续改进", 1)
    add_body(doc, "模型服务、知识服务、协同服务、安全服务和分析服务通过统一API接入CodeTrack现有平台，并复用课程、班级、用户、资源、题目、知识点和代码运行记录等数据对象。学生可标记“看不懂、依据不足、答案有误”，教师可确认方法不一致、评分点不合理或知识过期，管理员将问题分为数据补充、知识库修订、提示与流程调整以及模型更新。所有更新先通过离线评测和安全检查，再进行小范围试用。")
    add_body(doc, "科研入口与教学入口共享权限、知识版本、引用展示和安全审计能力，但科研任务的评价标准与教学问答不同。前者更关注文献覆盖、代码可执行性、数据来源和学术诚信，后者更关注答案正确、推导可跟随和课堂方法一致。正式论文应分别说明两类场景，不应使用同一组指标笼统评价。")

    add_heading(doc, "10 结论与研究限制", 1)
    add_body(doc, "本文提出面向人工智能专业课程的垂类模型与自主化学科智能体集成方案，将结构化推理数据、LoRA适配、课程知识库、可验证RAG、四智能体协同和教育安全机制组织为统一技术链路。相较于单纯增加平台功能，该方案把研究重点放在学科知识如何注入、推理过程如何拆解、生成结论如何核验以及高影响操作如何受控。")
    add_body(doc, "当前工作的主要限制是数据规模较小、模型训练和端到端联调尚未完成、真实用户试用时间有限，且跨课程和科研场景的迁移能力尚未验证。因此，现阶段不能宣称系统已经提升正确率或教学效果。后续工作应依据本章实验协议完成模型训练、消融实验、安全测试和用户研究，并以真实结果修订模型参数、检索策略和智能体流程。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] 籍欣萌, 昝红英, 崔婷婷, 等. 大模型在垂直领域应用的现状与挑战[J]. 计算机工程与应用, 2025, 61(12): 1-11.",
        "[2] HU E J, SHEN Y, WALLIS P, et al. LoRA: Low-Rank Adaptation of Large Language Models[C]//International Conference on Learning Representations. 2022.",
        "[3] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems. 2020, 33: 9459-9474.",
        "[4] GAO Y, XIONG Y, GAO X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey[EB/OL]. arXiv:2312.10997, 2023.",
        "[5] YAO S, ZHAO J, YU D, et al. ReAct: Synergizing Reasoning and Acting in Language Models[C]//International Conference on Learning Representations. 2023.",
        "[6] WANG L, MA C, FENG X, et al. A Survey on Large Language Model Based Autonomous Agents[J]. Frontiers of Computer Science, 2024, 18(6): 186345.",
        "[7] QWEN TEAM. Qwen2.5 Technical Report[EB/OL]. arXiv:2412.15115, 2024.",
        "[8] 刘三女牙, 郝晓丽, 李卿, 等. 教育大语言模型的内涵、构建和挑战[J]. 现代远程教育研究, 2024(5).",
    ]
    for ref in refs:
        p = doc.add_paragraph(); set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3)
        p.paragraph_format.left_indent = Cm(0.85); p.paragraph_format.first_line_indent = Cm(-0.85)
        r = p.add_run(ref); set_run_font(r, size=10.5)

    doc.core_properties.title = "垂类大模型与核心方法"
    doc.core_properties.subject = "CodeTrack学术论文方法章节"
    doc.core_properties.author = "信赛创翼队"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
