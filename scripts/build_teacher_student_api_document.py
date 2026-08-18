from __future__ import annotations

from datetime import datetime
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_database_integration_document import (
    BORDER,
    GREEN,
    INK,
    LIGHT_GRAY,
    LIGHT_GREEN,
    MUTED,
    RISK,
    ROOT,
    add_callout,
    add_field,
    add_para,
    add_table,
    audit_document,
    configure_section,
    configure_styles,
    set_run_font,
    set_table_geometry,
    style_cell,
)
from build_database_enterprise_document import add_code_block, add_list_item
from build_teacher_student_backend_document import (
    add_numbering_definition,
    build_closed_loop_diagram,
    add_picture,
)


OUTPUT_DIR = ROOT / "artifacts" / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "CodeTrack教师端与学生端协同API接口文档_企业版.docx"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

DOC_ID = "CT-API-TS-001"
DOC_VERSION = "V1.0"
BASELINE_DATE = "2026-08-16"


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def set_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = "CodeTrack 双端协同 API"
    table.cell(0, 1).text = "教师端—学生端接口契约"
    set_table_geometry(table, [3600, 5760], indent_dxa=0)
    style_cell(table.cell(0, 0), bold=True, color=GREEN, size=8.5)
    style_cell(table.cell(0, 1), color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = f"{DOC_ID}  |  {DOC_VERSION}  |  内部开发使用"
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.add_run("第 "), size=8.5, color=MUTED)
    add_field(right, "PAGE")
    set_run_font(right.add_run(" 页 / 共 "), size=8.5, color=MUTED)
    add_field(right, "NUMPAGES")
    set_run_font(right.add_run(" 页"), size=8.5, color=MUTED)
    set_table_geometry(table, [5000, 4360], indent_dxa=0)
    style_cell(table.cell(0, 0), size=8.5, color=MUTED)


def add_cover(doc):
    add_para(doc, "CODETRACK  /  API CONTRACT", size=10, bold=True, color=GREEN, after=16)
    add_para(doc, "教师端—学生端协同", size=28, bold=True, color=INK, after=2)
    add_para(doc, "API 接口文档", size=28, bold=True, color=INK, after=12)
    add_para(doc, "面向双端前端、后端与测试的可执行接口契约", size=14, color=MUTED, after=24)
    add_callout(
        doc,
        "文档定位",
        "本文件只描述教师端与学生端需要通过共享后端互通的 API。每个接口均标注当前实现状态、调用方、数据可见性、状态变化和联调要求；它不是普通操作手册，也不是泛化的后端架构说明。",
        fill=LIGHT_GREEN,
    )
    add_table(
        doc,
        ["文档属性", "内容"],
        [
            ("文档编号", DOC_ID),
            ("版本 / 基线", f"{DOC_VERSION} / {BASELINE_DATE}"),
            ("适用项目", "CodeTrack 教师端、学生端、FastAPI 共享后端"),
            ("基础路径", "/api/v1"),
            ("主要读者", "教师端前端、学生端前端、后端、测试、架构与产品"),
            ("事实来源", "当前代码路由、src/api.ts、SQLAlchemy 模型及后端自动化测试"),
        ],
        [1900, 7460],
        font_size=9.2,
    )
    add_para(doc, "结论先行", size=10.5, bold=True, color=GREEN, before=18, after=4)
    add_para(
        doc,
        "当前已形成入班、课程内容发布、任务提交批改、课堂讨论 4 条跨端链路；成绩回读、学生公告、学生知识图谱、统一通知和真实身份认证仍需补齐。生产接入前必须移除 X-User-Id 与固定 student-01/student-03 身份模拟。",
        size=11,
        after=0,
    )


def add_control_and_toc(doc):
    heading = doc.add_heading("文档控制", level=1)
    heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ["角色", "评审重点", "输出"],
        [
            ("双端前端负责人", "页面动作、接口时机、加载/空/错状态、字段兼容", "接口调用映射确认"),
            ("后端负责人", "鉴权、对象权限、事务、幂等、错误码和状态机", "OpenAPI 与实现确认"),
            ("测试负责人", "正向、负向、越权、并发与跨端可见性", "联调和验收报告"),
            ("产品负责人", "跨端业务闭环、发布口径和状态含义", "范围与优先级确认"),
        ],
        [1850, 5100, 2410],
        font_size=8.8,
    )
    doc.add_heading("阅读导航", level=1)
    add_table(
        doc,
        ["章节", "用途"],
        [
            ("1-3", "确认范围、现状、统一协议和身份安全要求"),
            ("4", "按教师按钮、接口、学生页面快速查找"),
            ("5-8", "查看 4 条已实现跨端链路的详细契约"),
            ("9", "查看必须新增的学生端 API 契约"),
            ("10-12", "字段字典、错误码、数据库与权限规则"),
            ("13-15", "前端接入、联调用例、发布门禁"),
        ],
        [1750, 7610],
        font_size=9.2,
    )
    add_callout(doc, "状态说明", "“已实现”表示当前后端已有路由且代码链路存在；“部分实现”表示教师侧写入存在但学生侧读取或生产能力不完整；“待实现”是本文件给出的目标契约，不能当作当前可调用接口。")


def add_scope(doc, bullet_id):
    heading = doc.add_heading("1. 范围与边界", level=1)
    heading.paragraph_format.page_break_before = True
    add_para(doc, "本文件覆盖同一业务对象在教师端写入、学生端读取或回写，再由教师端查看结果的接口。仅教师本人使用且不会影响学生端的草稿偏好、教师账号列表等接口不在详细契约范围内。")
    doc.add_heading("1.1 纳入范围", level=2)
    for text in [
        "班级邀请码、学生入班和教师查看入班结果。",
        "章节、资料与任务发布后在学生课程页的可见性。",
        "学生任务提交、自动评测、教师批改、成绩与反馈发布。",
        "教师发起讨论、学生回复、教师端实时回流。",
        "教师发布知识图谱后学生端读取目标班级图谱。",
        "公告、通知、身份认证、对象级权限和跨端状态一致性。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("1.2 明确不等于", level=2)
    add_table(doc, ["文档类型", "本文件是否承担", "说明"], [
        ("前端操作手册", "否", "操作手册说明用户怎么点；本文件说明点击后调用哪个 API、改哪些数据。"),
        ("数据库设计说明", "部分", "仅保留接口所需表映射和事务边界，不展开全部表结构。"),
        ("后端总体架构", "否", "仅描述支撑跨端 API 的认证、权限、幂等和事件要求。"),
        ("OpenAPI 机器文件", "否", "本文是人读契约；发布时仍需生成并冻结 openapi.json。"),
    ], [2100, 1500, 5760], font_size=8.8)


def add_current_state(doc):
    doc.add_heading("2. 当前实现结论", level=1)
    add_table(doc, ["跨端能力", "状态", "教师端", "学生端", "关键缺口"], [
        ("邀请码入班", "已实现", "生成/刷新邀请码、查看成员", "按邀请码加入", "无申请审核、移除和容量强校验"),
        ("课程内容发布", "已实现", "发布/撤回章节，资料自动切换可见性", "读取已发布章节、资料、任务", "缺学生课程列表与版本/缓存策略"),
        ("任务与提交", "已实现", "建任务、发布、看提交与评测", "任务列表、提交代码", "定时发布无调度器；学生结果详情缺失"),
        ("批改与反馈", "部分实现", "保存/发布成绩、发布反馈", "无本人结果读取 API", "必须补学生成绩与反馈读取"),
        ("课堂讨论", "已实现", "创建、发布、结束、查看回复", "列表、回复", "无分页、撤回、敏感内容治理"),
        ("知识图谱", "部分实现", "创建、保存、发布", "无读取 API", "目标班级目前是文本数组，需绑定 class_id"),
        ("公告与通知", "部分实现", "教师公告读取、教师通知", "无标准公告/通知入口", "缺学生列表、详情、已读和统一通知"),
        ("统一认证", "待实现", "X-User-Id 模拟", "固定 student-01/student-03", "生产必须使用 Session/JWT"),
    ], [1600, 1050, 2200, 1900, 2610], font_size=7.7)
    add_picture(doc, build_closed_loop_diagram(), "图 1  教师端—学生端 API 数据闭环", "教师发布、学生参与、后端处理、教师查看结果的跨端数据闭环")


def add_protocol(doc, bullet_id):
    doc.add_heading("3. 统一协议约定", level=1)
    doc.add_heading("3.1 环境与请求", level=2)
    add_table(doc, ["项目", "约定"], [
        ("开发基础地址", "http://127.0.0.1:8001/api/v1"),
        ("生产基础地址", "https://<api-domain>/api/v1"),
        ("数据格式", "application/json; charset=utf-8；文件上传使用 multipart/form-data"),
        ("时间格式", "ISO 8601，服务端存 UTC，响应带时区；前端按用户时区显示"),
        ("标识符", "业务对象 id 使用字符串；教师知识图谱当前 id 为整数"),
        ("幂等", "GET/PUT/PATCH/DELETE 按语义幂等；创建与发布接口建议支持 Idempotency-Key"),
        ("分页", "目标列表接口统一 page、page_size，响应包含 total/page/page_size"),
    ], [2200, 7160], font_size=9.0)
    doc.add_heading("3.2 当前响应包", level=2)
    add_code_block(doc, '{\n  "data": { "id": "..." }\n}')
    add_para(doc, "当前成功响应统一包裹在 data 字段中；FastAPI 校验错误使用 detail。目标生产契约建议统一 trace_id 和稳定业务 code，便于双端定位同一次请求。", size=9.6, color=MUTED)
    add_code_block(doc, '{\n  "code": "TASK_NOT_SUBMITTABLE",\n  "message": "任务不可提交",\n  "trace_id": "01J...",\n  "details": {}\n}')
    doc.add_heading("3.3 身份与权限", level=2)
    add_callout(doc, "上线阻断项", "当前请求通过 X-User-Id 自报身份，且 src/api.ts 内默认 teacher-01、学生调用固定 student-01/student-03。该方式只适合本地原型，任何调用方都可伪装其他用户。生产上线前必须改为服务端验证的 Session 或 JWT，并从会话解析 user_id、role 和 tenant_id。", fill="FDECEC")
    for text in [
        "教师接口：必须校验课程 teacher_id 等于当前教师，不能只校验 role=teacher。",
        "学生接口：必须校验 Enrollment(class_id, student_id) 存在且状态有效。",
        "提交、成绩、反馈：学生只能读取本人数据；教师只能读取本人课程下的数据。",
        "隐藏测试用例：学生任务响应不得包含 hidden=true 用例、输入、期望输出或判题细节。",
    ]:
        add_list_item(doc, text, bullet_id)


def add_matrix(doc):
    doc.add_heading("4. 教师按钮—API—学生页面总表", level=1)
    rows = [
        ("班级管理 / 刷新邀请码", "POST /teacher/classes/{class_id}/join-code", "学生加入班级页", "新邀请码生效，旧码失效", "已实现"),
        ("学生加入 / 输入邀请码", "POST /classes/{join_code}/join", "教师班级成员页", "新增 enrollment，教师刷新可见", "已实现"),
        ("课程内容 / 发布章节", "PATCH /teacher/chapters/{chapter_id}", "学生课程内容页", "章节及 ready 资料可见", "已实现"),
        ("课程内容 / 撤回章节", "PATCH /teacher/chapters/{chapter_id}", "学生课程内容页", "章节与资料隐藏", "已实现"),
        ("任务 / 发布", "POST /teacher/tasks/{task_id}/publish", "学生任务列表", "目标班级出现任务", "已实现"),
        ("学生任务 / 提交", "POST /student/tasks/{task_id}/submissions", "教师提交监控", "生成提交、评测和诊断", "已实现"),
        ("批改 / 发布成绩", "POST /teacher/submissions/{id}/grade/publish", "学生任务结果页", "学生可读成绩与反馈", "学生读取待实现"),
        ("课堂讨论 / 发布", "POST /teacher/discussions/{id}/publish", "学生课堂讨论页", "讨论可见并生成通知", "已实现"),
        ("学生讨论 / 回复", "POST /student/discussions/{id}/replies", "教师课堂讨论页", "回复与参与人数更新", "已实现"),
        ("知识图谱 / 发布", "POST /teacher/knowledge-graphs/{id}/publish", "学生知识图谱页", "目标班级读取发布版本", "学生读取待实现"),
        ("公告 / 发布", "POST /teacher/announcements", "学生消息/公告页", "目标范围可见并通知", "创建与学生读取待实现"),
    ]
    add_table(doc, ["教师/学生动作", "API", "另一端页面", "跨端结果", "状态"], rows, [1900, 2900, 1700, 1900, 960], font_size=7.5)
    add_callout(doc, "前端刷新原则", "写接口成功后，以响应数据更新当前页面；另一端通过重新拉取、轮询、SSE/WebSocket 或通知中心获知变化。不要让教师端直接写学生端本地缓存，也不要在双端分别维护同一业务状态。")


def endpoint(doc, code, title, status, method, path, caller, trigger, permission, request_rows, response_rows, effects, errors, request_json=None, response_json=None):
    heading = doc.add_heading(f"{code} {title}", level=2)
    if code in {"API-TA-06", "API-GR-01"}:
        heading.paragraph_format.page_break_before = True
    fill = LIGHT_GREEN if status == "已实现" else "FFF4D6" if status == "部分实现" else "FDECEC"
    add_callout(doc, status, f"{method} {path}  |  调用方：{caller}  |  页面触发：{trigger}", fill=fill)
    add_para(doc, f"权限：{permission}", size=9.4, bold=True, color=INK, after=5)
    if request_rows:
        add_para(doc, "请求字段", size=10.2, bold=True, color=GREEN, before=4, after=4, keep=True)
        add_table(doc, ["字段", "位置/类型", "必填", "规则与说明"], request_rows, [2200, 1900, 900, 4360], font_size=8.2, keep_together=True)
    if request_json:
        add_para(doc, "请求示例", size=9.5, bold=True, color=INK, before=3, after=2, keep=True)
        add_code_block(doc, request_json)
    if response_rows:
        add_para(doc, "响应字段", size=10.2, bold=True, color=GREEN, before=4, after=4, keep=True)
        add_table(doc, ["字段", "类型", "说明"], response_rows, [2500, 1600, 5260], font_size=8.2, keep_together=True)
    if response_json:
        add_para(doc, "响应示例", size=9.5, bold=True, color=INK, before=3, after=2, keep=True)
        add_code_block(doc, response_json)
    add_para(doc, "跨端数据结果", size=10.2, bold=True, color=GREEN, before=4, after=3, keep=True)
    add_para(doc, effects, size=9.4, after=5)
    add_para(doc, "主要错误", size=10.2, bold=True, color=GREEN, before=2, after=3, keep=True)
    add_para(doc, errors, size=9.4, color=MUTED, after=9)


def add_join_apis(doc):
    doc.add_heading("5. 班级加入链路 API", level=1)
    add_para(doc, "链路：教师生成邀请码 → 学生提交邀请码 → 后端创建 enrollment → 教师查看成员。当前加入操作为直接通过，不存在待审核状态。")
    endpoint(doc, "API-CL-01", "刷新班级邀请码", "已实现", "POST", "/api/v1/teacher/classes/{class_id}/join-code", "教师端", "班级管理 > 刷新邀请码", "当前教师拥有该班级所属课程", [
        ("class_id", "path / string", "是", "教学班 ID"),
    ], [("data.class_id", "string", "教学班 ID"), ("data.join_code", "string", "新邀请码；旧邀请码立即失效")], "更新 class_groups.join_code，写入审计日志；学生随后只能使用新码加入。", "403 非教师；404 班级不存在或不属于当前教师。", response_json='{\n  "data": {"class_id": "class-se1", "join_code": "A8K2M9QX"}\n}')
    endpoint(doc, "API-CL-02", "学生按邀请码加入班级", "已实现", "POST", "/api/v1/classes/{join_code}/join", "学生端", "加入班级 > 确认", "当前会话必须是 student；邀请码对应有效班级", [
        ("join_code", "path / string", "是", "教师端展示的邀请码，URL 编码"),
    ], [("data.class_id", "string", "加入的班级 ID"), ("data.class_name", "string", "班级名称"), ("data.joined", "boolean", "成功或已加入均为 true")], "首次请求新增 enrollments(class_id, student_id)；重复请求不重复插入，返回相同成功语义。教师端成员列表刷新后可见。", "403 非学生；404 邀请码无效。当前实现未校验班级 closed、容量上限或申请审核。", response_json='{\n  "data": {"class_id": "class-se1", "class_name": "软件工程 1 班", "joined": true}\n}')
    endpoint(doc, "API-CL-03", "教师查看入班状态", "已实现", "GET", "/api/v1/teacher/classes/{class_id}/join-status", "教师端", "班级管理 > 加入进度", "当前教师拥有班级所属课程", [("class_id", "path / string", "是", "教学班 ID")], [
        ("data.summary.joined", "integer", "已加入人数"), ("data.summary.pending", "integer", "当前固定为 0"), ("data.rows[]", "array", "学生、学号、加入时间和方式"),
    ], "读取 enrollments 与 users。当前 join_method 和 last_active 含原型展示值，不应作为生产统计证据。", "403 非教师；404 班级不存在。")


def add_content_apis(doc):
    doc.add_heading("6. 课程内容发布链路 API", level=1)
    add_para(doc, "链路：教师将章节状态改为 published → 同章节 ready 资料 visibility 自动改为 students → 已入班学生读取课程内容。撤回时执行相反变化。")
    endpoint(doc, "API-CO-01", "发布或撤回章节", "已实现", "PATCH", "/api/v1/teacher/chapters/{chapter_id}", "教师端", "课程内容 > 发布/撤回", "章节所属课程必须归当前教师", [
        ("chapter_id", "path / string", "是", "章节 ID"),
        ("status", "body / enum", "否", "draft | published"),
        ("title", "body / string", "否", "2-160 字符"),
        ("teaching_mode", "body / string", "否", "理论讲授、翻转课堂、案例教学、项目制教学、实验实训、混合式教学"),
    ], [("data.status", "enum", "最终章节状态"), ("data.title", "string", "章节标题"), ("data.teaching_mode", "string", "教学方式")], "status=published 时，将同课程、chapter_label 等于章节标题且 status=ready 的资料改为 students；status=draft 时改回 teacher。写审计日志。", "404 章节不存在；422 教学方式或状态非法。", request_json='{\n  "status": "published",\n  "teaching_mode": "案例教学"\n}')
    endpoint(doc, "API-CO-02", "学生读取已发布课程内容", "已实现", "GET", "/api/v1/student/courses/{course_id}/content", "学生端", "课程详情 > 内容", "学生必须通过 enrollment 加入该课程下至少一个班级", [("course_id", "path / string", "是", "课程 ID")], [
        ("data[]", "array", "仅 published 章节"),
        ("data[].knowledge_points", "array", "章节知识点"),
        ("data[].materials", "array", "仅 visibility=students 且 status=ready 的同章节资料"),
        ("data[].tasks", "array", "仅学生所在班级已发布任务"),
    ], "按当前学生班级过滤内容。教师撤回章节后，下一次请求中该章节整体消失；隐藏资料不得返回 content_url。", "403 未加入该课程。当前实现无 404 课程不存在的独立语义、无分页和版本号。", response_json='{\n  "data": [{\n    "id": "chapter-01", "title": "第 1 章 绪论", "status": "published",\n    "materials": [{"id": "material-01", "title": "课程讲义", "type": "pdf", "content_url": "/api/v1/material-files/..."}],\n    "tasks": [{"id": "task-01", "title": "线性表练习", "due_at": "2026-08-30T23:59:00"}]\n  }]\n}')


def add_task_apis(doc):
    doc.add_heading("7. 任务、提交与批改链路 API", level=1)
    add_para(doc, "链路：教师创建草稿 → 发布到班级 → 学生读取任务 → 学生提交 → 后端评测与诊断 → 教师查看、评分和发布结果。")
    endpoint(doc, "API-TA-01", "创建任务草稿", "已实现", "POST", "/api/v1/teacher/tasks", "教师端", "任务管理 > 新建任务 > 保存", "课程归当前教师", [
        ("course_id", "body / string", "是", "课程 ID"), ("class_id", "body / string|null", "否", "草稿阶段可为空"),
        ("title", "body / string", "是", "2-200 字符"), ("chapter_label", "body / string", "是", "必须与章节标题保持稳定关联"),
        ("due_at", "body / datetime", "是", "ISO 8601"), ("test_cases", "body / array", "是", "至少 1 条；weight 0-100"),
    ], [("data.id", "string", "任务 ID"), ("data.status", "string", "draft"), ("data.test_cases", "array", "教师端完整测试用例")], "写 tasks、test_cases、audit_logs；学生端不可见。", "403 非课程教师；422 无测试用例、字段或时间格式非法。")
    endpoint(doc, "API-TA-02", "发布任务到班级", "已实现", "POST", "/api/v1/teacher/tasks/{task_id}/publish", "教师端", "任务管理 > 发布", "任务归当前教师且目标班级属于课程", [
        ("task_id", "path / string", "是", "任务 ID"), ("class_id", "body / string", "是", "目标班级 ID"),
        ("publish_at", "body / datetime|null", "否", "为空表示立即发布"), ("due_at", "body / datetime", "是", "必须晚于 publish_at"),
    ], [("data.status", "enum", "published | scheduled"), ("data.class_id", "string", "目标班级"), ("data.publish_at", "datetime", "发布时间")], "立即发布时学生任务列表可见；未来时间当前只写 scheduled，项目中没有调度器自动切换为 published，因此不应宣称定时发布已闭环。", "409 已关闭任务；422 无测试用例、课程无知识点或日期非法。", request_json='{\n  "class_id": "class-se1",\n  "publish_at": null,\n  "due_at": "2026-08-30T23:59:00"\n}')
    endpoint(doc, "API-TA-03", "学生读取任务列表", "已实现", "GET", "/api/v1/student/tasks", "学生端", "任务中心", "仅查询当前学生已加入班级", [], [
        ("data[]", "array", "目标班级且 status=published 的任务"), ("data[].test_cases", "array", "只包含公开测试用例；hidden=true 不返回"),
        ("data[].due_at", "datetime", "截止时间"), ("data[].starter_code", "string", "起始代码"),
    ], "按 enrollment 推导 class_id 范围并过滤。学生不传 student_id、class_id 来扩大范围。", "403 非学生。当前接口未分页，且列表 DTO 偏大，建议拆分列表与详情。")
    endpoint(doc, "API-TA-04", "学生提交任务", "已实现", "POST", "/api/v1/student/tasks/{task_id}/submissions", "学生端", "任务详情 > 提交", "任务已发布且学生属于目标班级", [
        ("task_id", "path / string", "是", "任务 ID"), ("source_code", "body / string", "是", "最少 5 字符；生产需限制体积"),
        ("hint_level", "body / integer", "否", "0-3"),
    ], [("data.id", "string", "提交 ID"), ("data.version", "integer", "同学生同任务递增"), ("data.evaluation", "object", "自动评测结果"), ("data.diagnosis", "object", "诊断结果")], "事务内创建 submissions、evaluation_results、diagnosis_results；低置信度时创建 diagnosis_reviews；教师提交监控随后可见。", "404 任务不存在或不可提交；403 不属于目标班级；422 代码过短。当前实现未拒绝逾期提交。", request_json='{\n  "source_code": "int main() { return 0; }",\n  "hint_level": 1\n}')
    endpoint(doc, "API-TA-05", "教师查看任务提交", "已实现", "GET", "/api/v1/teacher/submissions?task_id={task_id}", "教师端", "提交监控 / 批改中心", "任务所属课程归当前教师", [("task_id", "query / string", "是", "任务 ID；代码中存在默认 task-01，不建议生产保留")], [
        ("data[]", "array", "提交及学生、评测、诊断、成绩、反馈"), ("data[].source_code", "string", "学生源码，仅教师可见"),
        ("data[].evaluation.details", "array", "包含完整测试明细"),
    ], "教师端读取学生提交回流；学生端永远不能调用该接口。", "403 非教师；404/403 应统一表现为无权访问对象。")
    endpoint(doc, "API-TA-06", "保存并发布成绩", "部分实现", "PUT + POST", "/api/v1/teacher/submissions/{id}/grade 及 /grade/publish", "教师端", "批改中心 > 保存 / 发布", "提交属于当前教师课程", [
        ("score", "body / integer", "保存必填", "0-100"), ("comment", "body / string", "否", "教师评语"),
        ("dimensions", "body / object", "否", "autoTest、codeQuality、report、participation 等维度"),
    ], [("data.status", "enum", "graded | grade_published"), ("data.score", "integer", "最终成绩"), ("data.dimensions", "object", "维度分")], "保存写 grades；发布设置 published_at、submission.status=grade_published 并为学生创建通知。当前缺学生本人读取发布成绩的 API，因此链路未闭环。", "404 提交不存在；403 非所属教师；409 建议用于未评分直接发布。", request_json='{\n  "score": 93,\n  "comment": "边界处理完整。",\n  "dimensions": {"autoTest": 38, "codeQuality": 28, "report": 18, "participation": 9}\n}')


def add_discussion_apis(doc):
    doc.add_heading("8. 课堂讨论链路 API", level=1)
    add_para(doc, "链路：教师创建并发布讨论 → 后端向班级学生生成通知 → 学生读取并回复 → 后端通知教师 → 教师端查看回复和参与人数。")
    endpoint(doc, "API-DI-01", "创建课堂讨论", "已实现", "POST", "/api/v1/teacher/discussions", "教师端", "课堂讨论 > 新建", "课程归当前教师且 class_id 属于该课程", [
        ("course_id", "body / string", "是", "课程 ID"), ("class_id", "body / string", "是", "目标班级 ID"),
        ("title", "body / string", "是", "1-160 字符"), ("content", "body / string", "是", "1-5000 字符"), ("publish", "body / boolean", "否", "true 时创建后立即发布"),
    ], [("data.status", "enum", "draft | published"), ("data.replies", "array", "创建时为空")], "创建 course_discussions；publish=true 时设置 published_at、为全班学生创建 notifications 并写 audit_logs。", "404 课程不存在；422 班级不属于课程或字段非法。", request_json='{\n  "course_id": "course-ds", "class_id": "class-se1",\n  "title": "链表边界条件讨论", "content": "请说明空链表处理策略。", "publish": true\n}')
    endpoint(doc, "API-DI-02", "发布或结束讨论", "已实现", "POST", "/api/v1/teacher/discussions/{id}/publish 或 /end", "教师端", "课堂讨论 > 发布 / 结束", "讨论所属课程归当前教师", [("discussion_id", "path / string", "是", "讨论 ID")], [("data.status", "enum", "published | ended"), ("data.reply_count", "integer", "回复数")], "发布后学生列表可见；结束后学生列表因当前只返回 published 而消失，且不再允许回复。", "404 讨论不存在；409 只有 published 状态可结束。")
    endpoint(doc, "API-DI-03", "学生读取讨论", "已实现", "GET", "/api/v1/student/discussions", "学生端", "课堂讨论列表", "只读取当前学生所在班级且 published 的讨论", [], [("data[]", "array", "已发布讨论"), ("data[].replies", "array", "当前包含全部回复"), ("data[].participant_count", "integer", "去重参与人数")], "从 enrollment 推导班级范围；草稿和已结束讨论不返回。", "403 非学生。当前未分页，回复过多时响应会膨胀，建议列表不带 replies、详情再分页读取。")
    endpoint(doc, "API-DI-04", "学生回复讨论", "已实现", "POST", "/api/v1/student/discussions/{discussion_id}/replies", "学生端", "讨论详情 > 发送", "讨论已发布且学生属于目标班级", [
        ("discussion_id", "path / string", "是", "讨论 ID"), ("content", "body / string", "是", "1-2000 字符"),
    ], [("data", "object", "更新后的完整讨论对象")], "新增 discussion_replies，为教师创建通知并写审计日志。教师端重新拉取后看到回复、回复数和参与人数变化。", "404 讨论不存在或未发布；403 不属于讨论班级；422 空内容或超长。", request_json='{\n  "content": "空链表应直接返回，并避免访问 head->next。"\n}')


def add_missing_apis(doc):
    doc.add_heading("9. 必须新增的学生端 API", level=1)
    add_callout(doc, "实施口径", "本章接口为目标契约，当前代码中不存在。后端实现、OpenAPI 更新、双端前端接入和自动化测试全部完成后，才能把状态改为“已实现”。", fill="FFF4D6")
    endpoint(doc, "API-GR-01", "学生读取知识图谱列表", "待实现", "GET", "/api/v1/student/knowledge-graphs?course_id={course_id}", "学生端", "知识图谱列表", "学生已加入课程；只返回 status=published 且 target_class_ids 包含其班级", [("course_id", "query / string", "是", "课程 ID"), ("page", "query / integer", "否", "默认 1"), ("page_size", "query / integer", "否", "默认 20，最大 100")], [("data.items[]", "array", "图谱摘要，不返回完整 nodes/edges"), ("data.total", "integer", "总数")], "教师发布图谱后目标学生可见。数据库应将 target_classes 文本名称改为 graph_class_targets(graph_id, class_id) 关联，避免班级改名导致失联。", "403 未加入课程；404 课程不存在。")
    endpoint(doc, "API-GR-02", "学生读取知识图谱详情", "待实现", "GET", "/api/v1/student/knowledge-graphs/{graph_id}", "学生端", "知识图谱详情", "图谱已发布且当前学生在目标班级", [("graph_id", "path / integer", "是", "知识图谱 ID")], [("data.nodes", "array", "发布版本节点"), ("data.edges", "array", "发布版本关系"), ("data.published_at", "datetime", "发布时间"), ("data.version", "integer", "发布版本")], "只读取已发布快照。教师后续编辑草稿不应立即改变学生正在查看的已发布版本。", "404 对不存在和无权访问统一返回，防止枚举。")
    endpoint(doc, "API-RE-01", "学生读取本人提交结果", "待实现", "GET", "/api/v1/student/submissions/{submission_id}/result", "学生端", "任务结果页", "submission.student_id 等于当前学生", [("submission_id", "path / string", "是", "提交 ID")], [("data.evaluation", "object", "允许公开的评测结果"), ("data.grade", "object|null", "仅 grade_published 时返回"), ("data.feedback", "array", "仅 student_visible=true"), ("data.diagnosis", "object|null", "仅允许公开的教师确认内容")], "教师发布成绩后学生获得稳定读取入口。隐藏测试用例只能返回是否通过或汇总，不能返回输入和期望输出。", "404 提交不存在或不属于本人；409 结果仍在处理中可返回 EVALUATION_PENDING。")
    endpoint(doc, "API-AN-01", "学生公告列表与详情", "待实现", "GET", "/api/v1/student/announcements", "学生端", "消息中心 > 公告", "按 enrollment 和公告 audience 过滤", [("course_id", "query / string", "否", "课程筛选"), ("unread_only", "query / boolean", "否", "仅未读"), ("page/page_size", "query / integer", "否", "分页")], [("data.items[]", "array", "公告摘要、发布人、范围、置顶、read"), ("data.unread_count", "integer", "未读数")], "读取 course_announcements，并 LEFT JOIN announcement_reads 计算当前学生 read 状态。教师端当前只有公告读取，没有创建接口，需一并实现 POST /teacher/announcements。", "403 非学生；404 不可见公告详情。")
    endpoint(doc, "API-AN-02", "学生公告标记已读", "待实现", "PATCH", "/api/v1/student/announcements/{announcement_id}/read", "学生端", "打开公告详情", "公告对当前学生可见", [("announcement_id", "path / string", "是", "公告 ID"), ("read", "body / boolean", "否", "默认 true")], [("data.id", "string", "公告 ID"), ("data.read", "boolean", "当前已读状态")], "使用 announcement_reads(announcement_id, user_id) 唯一键 upsert，重复调用不产生重复记录。", "404 公告不可见；409 不需要，接口必须幂等。")
    endpoint(doc, "API-NO-01", "双端统一通知中心", "待实现", "GET", "/api/v1/notifications", "教师端与学生端", "顶部通知入口", "只读取 notification.user_id 等于当前用户", [("type", "query / string", "否", "任务、讨论、成绩、公告等"), ("unread_only", "query / boolean", "否", "仅未读"), ("page/page_size", "query / integer", "否", "分页")], [("data.items[]", "array", "通知列表"), ("data.unread_count", "integer", "未读数")], "统一替代仅教师可用的通知读取方式；通知内容包含 resource_type/resource_id，前端按类型跳转。", "403 未登录；404 不向调用方暴露其他用户通知。")
    endpoint(doc, "API-AU-01", "双端真实登录与当前用户", "待实现", "POST + GET", "/api/v1/auth/login 及 /api/v1/me", "双端", "登录 / 应用启动", "服务端校验凭据并签发 Session/JWT", [("username", "body / string", "登录必填", "账号"), ("password", "body / string", "登录必填", "密码只在 TLS 上传输")], [("data.user.id", "string", "用户 ID"), ("data.user.role", "enum", "teacher | student"), ("data.permissions", "array", "权限码"), ("data.csrf_token", "string|null", "Cookie 会话时使用")], "前端不再传 X-User-Id，也不在 API 方法中硬编码学生。后端从凭据恢复当前用户和租户上下文。", "401 凭据无效；423 账号锁定；429 请求过多。")


def add_data_and_security(doc, bullet_id):
    doc.add_heading("10. 字段与状态字典", level=1)
    add_table(doc, ["对象", "允许状态", "跨端可见性规则"], [
        ("Chapter", "draft → published → draft", "仅 published 对学生可见；撤回后立即隐藏"),
        ("Material", "ready / deleted；visibility=teacher/students", "ready 且 students 才可见"),
        ("Task", "draft → scheduled/published → closed", "仅 published 且班级匹配可见"),
        ("Submission", "submitted → evaluated → grade_published", "学生仅本人；教师仅所属课程"),
        ("Grade", "graded → grade_published", "只有 grade_published 返回学生端"),
        ("Feedback", "draft → published；student_visible", "published 且 student_visible=true"),
        ("Discussion", "draft → published → ended", "仅 published 可列出和回复"),
        ("KnowledgeGraph", "draft → published", "只读发布快照且目标班级匹配"),
    ], [1800, 2850, 4710], font_size=8.4)
    doc.add_heading("11. 错误码契约", level=1)
    add_table(doc, ["HTTP", "业务码示例", "前端处理"], [
        ("400", "BAD_REQUEST", "保留表单内容，显示可操作提示"),
        ("401", "UNAUTHENTICATED", "清理会话并跳转登录；只允许一次刷新令牌重试"),
        ("403", "FORBIDDEN / NOT_ENROLLED", "显示无权限；不得循环重试"),
        ("404", "RESOURCE_NOT_FOUND", "返回列表或 404 页面；对象越权也可统一 404"),
        ("409", "INVALID_STATE / VERSION_CONFLICT", "刷新对象状态，阻止重复发布或过期编辑"),
        ("413", "PAYLOAD_TOO_LARGE", "提示文件或源码体积限制"),
        ("415", "UNSUPPORTED_MEDIA_TYPE", "提示允许格式"),
        ("422", "VALIDATION_ERROR", "映射到具体字段；保留服务端 details"),
        ("429", "RATE_LIMITED", "按 Retry-After 延迟重试"),
        ("500", "INTERNAL_ERROR", "显示 trace_id，不展示堆栈或 SQL"),
    ], [900, 3000, 5460], font_size=8.6)
    doc.add_heading("12. 数据库与权限映射", level=1)
    add_table(doc, ["API 链路", "核心表", "事务/约束"], [
        ("入班", "class_groups, enrollments, users", "UNIQUE(class_id, student_id)；邀请码唯一"),
        ("内容发布", "chapters, materials, knowledge_points", "章节状态与同章节资料可见性同事务提交"),
        ("任务发布", "tasks, test_cases, enrollments", "发布前校验班级归属、知识点和测试用例"),
        ("学生提交", "submissions, evaluation_results, diagnosis_results, diagnosis_reviews", "提交、评测和诊断原子写入；版本号并发唯一"),
        ("批改", "grades, teacher_feedback, notifications", "成绩发布和学生通知同事务；建议 Outbox"),
        ("讨论", "course_discussions, discussion_replies, notifications", "回复写入与教师通知同事务；建议 Outbox"),
        ("图谱", "teacher_knowledge_graphs, graph_class_targets(新增)", "发布快照版本化；班级用外键关联"),
        ("公告", "course_announcements, announcement_reads", "UNIQUE(announcement_id, user_id)，已读 upsert"),
    ], [1600, 4050, 3710], font_size=8.2)
    doc.add_heading("12.1 对象级权限检查顺序", level=2)
    for text in [
        "从 Session/JWT 获取 user_id、role、tenant_id，不接受请求头自报用户。",
        "加载目标对象，并验证 tenant_id 一致。",
        "教师校验 course.teacher_id；学生校验 enrollment 与对象目标班级。",
        "应用状态可见性：draft、未发布成绩、隐藏用例和不可见反馈不得序列化。",
        "写入审计日志；关键发布操作写 Outbox，由异步消费者发送通知。",
    ]:
        add_list_item(doc, text, bullet_id)


def add_frontend(doc, bullet_id):
    doc.add_heading("13. 双端前端接入规范", level=1)
    add_table(doc, ["层级", "要求"], [
        ("HTTP Client", "统一 baseURL、超时、Authorization/Cookie、trace_id、错误解包；禁止页面直接 fetch"),
        ("类型", "由 OpenAPI 生成 DTO，或维护单一契约包；禁止大量 any"),
        ("身份", "删除 request(path, options, userId) 的 userId 参数；删除固定 teacher-01/student-01/student-03"),
        ("缓存键", "包含当前 user_id、course_id、class_id 和资源版本，退出登录时清空"),
        ("写后刷新", "优先使用写接口返回值更新本页，再按资源键失效；另一端通过通知/轮询更新"),
        ("状态 UI", "每个跨端页面必须有 loading、empty、403、404、409、422、500 与离线状态"),
        ("防重复", "发布、提交、回复按钮请求中禁用；创建接口携带 Idempotency-Key"),
    ], [1650, 7710], font_size=8.7)
    doc.add_heading("13.1 当前 src/api.ts 必改项", level=2)
    for text in [
        "删除 _currentUserId='teacher-01' 默认身份及 setCurrentUser 作为安全边界的用法。",
        "joinClass、studentCourseContent、studentDiscussions、replyDiscussion 等方法不再传固定学生 ID。",
        "登录成功后由 Cookie 或内存中的短期 access token 建立会话；不要把长期令牌放 localStorage。",
        "为待新增学生接口建立独立 DTO：StudentSubmissionResult、StudentGraphSummary、StudentAnnouncement、NotificationPage。",
        "开发环境可保留受控的 impersonation 工具，但必须由后端开发开关保护，生产构建完全关闭。",
    ]:
        add_list_item(doc, text, bullet_id)


def add_tests(doc):
    doc.add_heading("14. 双端联调与验收用例", level=1)
    add_para(doc, "以下用例覆盖跨端可见性、状态机、权限和敏感字段。标记“现有自动化”的用例在当前仓库已有测试证据；目标接口用例需实现后执行。")
    rows = [
        ("TS-01", "邀请码首次加入", "学生使用有效码加入", "enrollment 新增；教师成员页可见", "现有链路"),
        ("TS-02", "邀请码重复加入", "同一学生重复请求", "仍返回 joined=true；无重复 enrollment", "补自动化"),
        ("TS-03", "无效邀请码", "提交不存在的码", "404；不写数据库", "补自动化"),
        ("TS-04", "章节发布", "教师 draft→published", "学生内容出现章节与 ready 资料", "现有自动化"),
        ("TS-05", "章节撤回", "教师 published→draft", "学生内容立即隐藏章节", "现有自动化"),
        ("TS-06", "未入班读内容", "其他学生访问课程", "403 NOT_ENROLLED", "补自动化"),
        ("TS-07", "任务发布", "教师发布到 class-se1", "该班学生可见，其他班不可见", "现有自动化部分"),
        ("TS-08", "隐藏用例保护", "学生读取任务", "响应不含 hidden=true 测试用例", "现有自动化"),
        ("TS-09", "学生提交", "合法源码与 hint_level", "生成提交、评测、诊断，教师可见", "现有自动化"),
        ("TS-10", "跨班提交", "非目标班学生提交", "403；不创建提交", "补自动化"),
        ("TS-11", "已关闭/未发布任务", "学生直接调用提交 API", "404/409；不评测", "补自动化"),
        ("TS-12", "多版本提交", "同一学生并发提交两次", "version 唯一递增，无重复版本", "目标"),
        ("TS-13", "教师保存成绩", "提交 0-100 分", "grade=graded；学生暂不可读", "现有自动化"),
        ("TS-14", "教师发布成绩", "发布已保存成绩", "grade_published，通知学生", "现有自动化"),
        ("TS-15", "学生读取本人结果", "请求本人 submission", "只返回已发布成绩与可见反馈", "待实现"),
        ("TS-16", "学生越权读结果", "用他人 submission_id", "统一 404，无数据泄露", "待实现"),
        ("TS-17", "讨论发布", "教师发布到目标班", "该班学生可见并收到通知", "现有链路"),
        ("TS-18", "讨论回复", "目标班学生回复", "教师端回复数、参与人数更新", "现有链路"),
        ("TS-19", "讨论结束后回复", "学生直接调用回复", "404/409；不写回复", "补自动化"),
        ("TS-20", "图谱目标班级", "教师发布图谱", "仅目标班学生列表可见", "待实现"),
        ("TS-21", "公告已读幂等", "学生重复标记已读", "announcement_reads 只有一条", "待实现"),
        ("TS-22", "身份伪造", "客户端提交 X-User-Id", "生产忽略/拒绝，不改变当前会话", "上线门禁"),
        ("TS-23", "教师跨课程越权", "教师 B 访问教师 A 对象", "404/403，无响应差异泄露", "上线门禁"),
        ("TS-24", "通知重复消费", "Outbox 消费重试", "同事件不产生重复通知", "目标"),
    ]
    add_table(doc, ["编号", "场景", "操作", "预期", "状态"], rows, [850, 1600, 2250, 3510, 1150], font_size=7.5)
    doc.add_heading("14.1 代表性端到端步骤", level=2)
    add_table(doc, ["步骤", "调用", "断言"], [
        ("1", "教师 POST /teacher/tasks 创建任务", "201，status=draft，测试用例落库"),
        ("2", "教师 POST /teacher/tasks/{id}/publish", "200，status=published，class_id 正确"),
        ("3", "目标学生 GET /student/tasks", "能看到任务，且没有 hidden=true 用例"),
        ("4", "非目标学生 GET /student/tasks", "看不到任务"),
        ("5", "目标学生 POST /student/tasks/{id}/submissions", "201，生成 evaluation 与 diagnosis"),
        ("6", "教师 GET /teacher/submissions?task_id=...", "看到该学生提交和完整评测"),
        ("7", "教师 PUT /grade 后 POST /grade/publish", "成绩状态 grade_published，产生学生通知"),
        ("8", "学生 GET /student/submissions/{id}/result", "只返回本人已发布成绩、公开评测和可见反馈"),
    ], [900, 4080, 4380], font_size=8.2)


def add_delivery(doc, bullet_id):
    doc.add_heading("15. 实施顺序与发布门禁", level=1)
    add_table(doc, ["阶段", "交付内容", "完成标准"], [
        ("P0-1 身份", "auth/login、me、Session/JWT、移除 X-User-Id", "双端无固定用户；越权测试通过"),
        ("P0-2 成绩", "student submission result、反馈可见性", "教师发布后学生可读，未发布不可读"),
        ("P0-3 图谱", "graph_class_targets、学生列表/详情、发布快照", "只对目标班级可见"),
        ("P1 公告通知", "公告创建/读取/已读、统一通知中心、Outbox", "消息可追踪、幂等、无重复"),
        ("P1 班级治理", "申请审核、拒绝、移除、容量、班级关闭", "完整 enrollment 状态机"),
        ("P2 工程化", "OpenAPI 生成 DTO、契约测试、监控、限流", "CI 门禁和可观测性齐备"),
    ], [1500, 4250, 3610], font_size=8.5)
    doc.add_heading("15.1 上线门禁", level=2)
    for text in [
        "生产构建中不存在 teacher-01、student-01、student-03 或可控 X-User-Id 身份入口。",
        "所有跨端对象均有教师所有权或学生 enrollment 权限测试。",
        "OpenAPI 与前端生成类型无差异；接口状态、枚举、时间格式已冻结。",
        "隐藏测试用例、未发布成绩、草稿反馈和其他学生源码通过响应字段审计。",
        "写接口具备幂等或版本控制；通知使用 Outbox 或等价可靠投递机制。",
        "核心跨端用例在真实数据库环境通过，并完成回滚演练与审计日志抽查。",
    ]:
        add_list_item(doc, text, bullet_id)


def audit(doc):
    audit_document(doc)
    assert len(doc.tables) >= 55
    assert len(doc.inline_shapes) == 1
    assert any(p.style.name == "Heading 1" for p in doc.paragraphs)
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == 9360


def main():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        set_header_footer(section)
    bullet_id = add_numbering_definition(doc, bullet=True)

    doc.core_properties.title = "CodeTrack 教师端—学生端协同 API 接口文档（企业版）"
    doc.core_properties.subject = "教师端与学生端共享后端 API 契约、权限、状态与联调基线"
    doc.core_properties.author = "CodeTrack 项目开发组"
    doc.core_properties.keywords = "CodeTrack, 教师端, 学生端, API, 接口文档, 联调"
    doc.core_properties.comments = f"{DOC_ID} {DOC_VERSION}"
    doc.core_properties.created = datetime(2026, 8, 16)
    doc.core_properties.modified = datetime(2026, 8, 16)

    add_cover(doc)
    add_control_and_toc(doc)
    add_scope(doc, bullet_id)
    add_current_state(doc)
    add_protocol(doc, bullet_id)
    add_matrix(doc)
    add_join_apis(doc)
    add_content_apis(doc)
    add_task_apis(doc)
    add_discussion_apis(doc)
    add_missing_apis(doc)
    add_data_and_security(doc, bullet_id)
    add_frontend(doc, bullet_id)
    add_tests(doc)
    add_delivery(doc, bullet_id)
    audit(doc)
    doc.save(OUTPUT_PATH)
    print(json.dumps({
        "output": str(OUTPUT_PATH),
        "bytes": OUTPUT_PATH.stat().st_size,
        "tables": len(doc.tables),
        "paragraphs": len(doc.paragraphs),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
