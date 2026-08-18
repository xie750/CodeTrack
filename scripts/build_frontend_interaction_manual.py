from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\teacher")
OUT_DIR = ROOT / "artifacts" / "deliverables"
OUT_PATH = OUT_DIR / "CodeTrack教师端前端逻辑交互文档.docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)

GREEN = "2AAD16"
DARK_GREEN = "176B12"
INK = "1F2A24"
MUTED = "5F6B63"
LIGHT_GREEN = "EDF8EA"
LIGHT_BLUE = "E8F1FB"
LIGHT_GRAY = "F4F6F5"
TABLE_HEADER = "E8EEF5"
BORDER = "D8E0DA"
CAUTION = "FFF5D9"
RISK = "FCE8E6"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=None, italic=None, ascii_font="Calibri"):
    run.font.name = ascii_font
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def set_row_cant_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell(cell, bold=False, color=INK, size=8.5):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.12
        for run in p.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, font_size=8.5, header_fill=TABLE_HEADER):
    table = doc.add_table(rows=1, cols=len(headers))
    set_repeat_table_header(table.rows[0])
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = str(value)
        set_cell_shading(table.rows[0].cells[i], header_fill)
        style_cell(table.rows[0].cells[i], bold=True, size=font_size)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
            style_cell(cells[i], size=font_size)
    for row in table.rows:
        set_row_cant_split(row)
    set_table_geometry(table, widths)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return table


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="Compact Bullet")
        p.add_run("• ")
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)


def add_callout(doc, label, text, kind="info"):
    fill = LIGHT_GREEN if kind == "info" else CAUTION if kind == "caution" else RISK
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, bold=True, color=DARK_GREEN if kind == "info" else INK)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def add_h1(doc, text):
    return doc.add_paragraph(text, style="Heading 1")


def add_h2(doc, text):
    return doc.add_paragraph(text, style="Heading 2")


def add_h3(doc, text):
    return doc.add_paragraph(text, style="Heading 3")


def add_interactions(doc, rows):
    return add_table(
        doc,
        ["控件 / 入口", "前置条件", "点击结果 / 目标", "数据动作", "反馈 / 返回"],
        rows,
        [1450, 1350, 2660, 1860, 2040],
    )


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(8)
    title_pr = title._element.get_or_add_pPr()
    border = title_pr.find(qn("w:pBdr"))
    if border is not None:
        title_pr.remove(border)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, GREEN),
        ("Heading 2", 13, 14, 7, DARK_GREEN),
        ("Heading 3", 12, 10, 5, INK),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = name == "Heading 1"

    bullet = styles.add_style("Compact Bullet", 1)
    bullet.base_style = normal
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25


def configure_document(doc):
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    settings = doc.settings.element
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is not None:
        settings.remove(even_odd)

    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("CodeTrack 教师端 | 前端逻辑交互文档")
    set_run_font(r, size=9, bold=True, color=MUTED)

    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


ROUTES = [
    ("全局", "/teacher/dashboard", "工作台首页", "默认落点；未知路径也重定向到此"),
    ("全局", "/teacher/courses", "我的课程", "课程列表、草稿、归档"),
    ("全局", "/teacher/courses/new", "新建课程", "三步创建流程"),
    ("全局", "/teacher/settings", "个人设置", "个人资料与偏好展示"),
    ("课程", "/teacher/courses/:courseId/workspace", "课程工作空间", "课程内总览"),
    ("课程", "/teacher/courses/:courseId/content", "课程章节内容", "章节、知识点、资料、练习"),
    ("课程", "/teacher/courses/:courseId/classes", "班级管理", "班级与加入状态"),
    ("课程", "/teacher/courses/:courseId/invite", "邀请学生", "链接、邀请码、二维码"),
    ("课程", "/teacher/courses/:courseId/tasks", "任务管理", "任务创建、编辑、发布"),
    ("课程", "/teacher/courses/:courseId/materials", "资料管理", "文件夹、资料、回收站"),
    ("课程", "/teacher/courses/:courseId/graph", "课程知识图谱", "多图谱创建、编辑、发布"),
    ("课程", "/teacher/courses/:courseId/monitor", "任务监控", "提交状态与批改入口"),
    ("课程", "/teacher/courses/:courseId/grading", "作业批改", "评分、评语、发布成绩"),
    ("课程", "/teacher/courses/:courseId/analytics", "学情分析", "总览、个体诊断、预警"),
    ("课程", "/teacher/courses/:courseId/reviews", "AI 审核", "接受、驳回、修改建议"),
    ("课程", "/teacher/courses/:courseId/course-settings", "课程设置", "课程管理入口"),
    ("URL 状态", "/teacher/courses/:courseId/workspace?discussion=1", "课堂讨论抽屉", "不是独立页面；关闭时移除查询参数"),
]


def build():
    doc = Document()
    configure_document(doc)

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CODETRACK TEACHER")
    set_run_font(r, size=11, bold=True, color=GREEN)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph("CodeTrack 教师端\n前端逻辑交互文档", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("按钮去向、弹窗状态、数据动作与返回路径的完整说明", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("适用对象：前端开发、后端联调、产品验收、测试与教师培训")
    set_run_font(r, size=11, color=MUTED)
    add_table(doc, ["文档属性", "内容"], [
        ["版本 / 日期", "V1.0 / 2026-08-13"],
        ["编写依据", "当前工作区源代码、路由配置、接口调用与浏览器抽样核验"],
        ["覆盖范围", "登录门户、全局工作台、课程创建及课程内全部教师端模块"],
        ["判定原则", "以当前代码实际行为为准；视觉入口与真实功能严格区分"],
    ], [2200, 7160], font_size=9.2, header_fill=LIGHT_GREEN)

    add_h1(doc, "1. 文档说明与交互判定")
    add_body(doc, "本文档回答的核心问题是：教师在当前页面点击某个按钮、菜单、卡片或列表项后，系统会进入哪个页面或状态，是否调用接口，成功或失败如何反馈，以及如何返回。")
    add_h2(doc, "1.1 状态标识")
    add_table(doc, ["标识", "含义", "测试判定"], [
        ["已实现", "存在真实跳转、弹窗/抽屉、状态更新或接口动作", "应按本文描述产生可观察结果"],
        ["页内状态", "不离开当前路由，仅切换筛选、步骤、选中项或面板", "URL 通常不变，刷新后可能恢复默认"],
        ["展示入口", "控件可见，但当前仅提示、无处理函数或未连接真实流程", "不得按已完成功能验收"],
        ["不可逆", "删除或发布等会永久改变服务端数据/业务状态", "执行前应确认对象和环境"],
    ], [1500, 4360, 3500], font_size=9.2)
    add_h2(doc, "1.2 交互类型")
    add_bullets(doc, [
        "页面跳转：React 路由发生变化，进入另一业务页面。",
        "抽屉/弹窗：保留当前页面，在上层打开详情、编辑、创建或确认界面。",
        "页内状态切换：筛选、步骤、选中项、标签页、展开/收起等本地状态变化。",
        "数据写入/接口动作：创建、修改、发布、删除、上传、批改或重新加载服务端数据。",
        "外部页面/下载：打开新标签页、调用浏览器下载或写入剪贴板。",
        "展示入口：当前版本尚未连接真实动作，单独列入限制清单。",
    ])
    add_callout(doc, "阅读规则", "路径中的 :courseId 代表当前选中课程 ID。进入任何课程内页面前，系统必须已经确定当前课程。", "info")

    add_h1(doc, "2. 系统入口与完整路由")
    add_h2(doc, "2.1 主业务流")
    add_body(doc, "登录门户 → 工作台 → 我的课程/新建课程 → 选中课程 → 课程工作空间 → 章节、班级、任务、资料、知识图谱、学情分析等课程内模块。未知地址会回退到工作台。")
    add_callout(doc, "兼容路由", "旧地址 /teacher/classes、/teacher/tasks、/teacher/materials、/teacher/analytics 当前统一重定向到 /teacher/courses；不能据此进入某门课程的对应模块。", "caution")
    add_h2(doc, "2.2 路由总览")
    add_table(doc, ["范围", "路由", "页面 / 状态", "说明"], ROUTES, [950, 3550, 2100, 2760], font_size=8.7)

    add_h1(doc, "3. 公共导航与全局控件")
    add_h2(doc, "3.1 全局侧边栏与顶栏")
    add_interactions(doc, [
        ["工作台首页", "已登录", "页面跳转 → /teacher/dashboard", "切换全局导航状态", "工作台可继续进入课程"],
        ["我的课程", "已登录", "页面跳转 → /teacher/courses", "加载课程列表", "可返回工作台"],
        ["个人设置", "已登录", "页面跳转 → /teacher/settings", "加载个人资料", "侧边栏返回其他全局页"],
        ["课程工作空间首页", "已选课程", "页面跳转 → 当前课程 workspace", "加载课程总览", "课程侧边栏保持可用"],
        ["章节/班级/任务/资料/图谱/学情", "已选课程", "页面跳转 → 当前课程对应路由", "按模块请求数据", "经课程侧边栏切换"],
        ["课堂讨论", "已选课程", "进入 workspace?discussion=1 并打开讨论抽屉", "URL 查询参数控制抽屉", "关闭后移除 discussion=1"],
        ["通知铃铛", "已登录", "打开通知抽屉", "页内状态", "关闭抽屉回到原页面"],
        ["回到工作台首页", "位于课程内", "页面跳转 → /teacher/dashboard", "保留登录状态", "重新选课后回课程"],
        ["返回我的课程", "位于课程内", "页面跳转 → /teacher/courses", "保留登录状态", "点击课程再次进入"],
        ["AI 助教", "已选课程", "页面跳转 → 当前课程 reviews", "加载 AI 审核列表", "课程侧边栏返回"],
        ["用户菜单/退出登录", "已登录", "退出教师端门户状态", "清理登录态", "返回登录入口"],
        ["面包屑：课程工作空间", "已选课程", "页面跳转 → workspace", "无数据写入", "回到课程总览"],
    ])

    add_h1(doc, "4. 登录门户")
    add_interactions(doc, [
        ["密码眼睛图标", "密码框有无内容均可", "页内切换密码明文/密文", "仅本地状态", "再次点击恢复"],
        ["登录教师端", "输入账号和密码", "校验通过后进入已登录门户；演示密码为 123456", "登录状态写入前端", "错误时显示校验提示"],
        ["演示账号卡片", "无", "快速填充/登录演示教师账号", "设置登录状态", "进入门户"],
        ["进入工作台", "已完成登录", "页面跳转 → /teacher/dashboard", "加载仪表盘数据", "侧边栏可切换"],
        ["进入科研", "已完成登录", "展示入口：仅浏览器 alert，无科研路由", "不写入数据", "关闭提示仍停留当前页"],
    ])

    add_h1(doc, "5. 工作台首页")
    add_interactions(doc, [
        ["刷新", "已登录", "留在当前页并重新加载统计/课程/动态", "调用工作台查询接口", "加载态后刷新内容"],
        ["我的课程指标卡", "无", "页面跳转 → /teacher/courses", "加载课程列表", "可返回工作台"],
        ["待发布任务指标卡", "存在当前课程", "页面跳转 → 当前课程 tasks", "加载任务列表", "课程侧栏返回"],
        ["待批改提交指标卡", "存在当前课程", "页面跳转 → 当前课程 grading", "加载任务与提交", "课程侧栏返回"],
        ["学情提醒指标卡", "存在当前课程", "页面跳转 → 当前课程 analytics", "加载学情数据", "课程侧栏返回"],
        ["创建课程", "无", "页面跳转 → /teacher/courses/new", "进入三步表单", "取消返回我的课程"],
        ["查看全部课程", "无", "页面跳转 → /teacher/courses", "加载课程列表", "可返回工作台"],
        ["课程：进入课程", "课程存在", "选中课程并跳转 → workspace", "更新当前课程上下文", "返回我的课程"],
        ["课程：管理课程", "课程存在", "选中课程并跳转 → course-settings", "加载课程设置", "课程侧栏返回"],
        ["更多：激活/转草稿/归档", "课程存在", "确认后更新课程状态", "调用课程更新接口", "成功提示并刷新列表"],
        ["更多：删除课程", "课程存在", "不可逆：确认后删除课程及关联数据", "调用删除接口，级联清理", "成功后卡片消失"],
        ["今日待办动作", "待办项存在", "页内移除该项", "仅本地状态，无 API", "刷新可能恢复"],
        ["AI 建议：学情", "存在当前课程", "页面跳转 → analytics", "加载分析数据", "课程侧栏返回"],
        ["AI 建议：查看学生", "存在当前课程", "页面跳转 → monitor", "加载提交监控", "课程侧栏返回"],
        ["最近动态/查看全部动态", "无", "展示入口：当前无处理函数", "无", "页面不变"],
    ])

    add_h1(doc, "6. 我的课程")
    add_interactions(doc, [
        ["搜索/学年/专业/状态", "课程列表已加载", "页内过滤课程卡片", "仅本地筛选", "清空条件恢复"],
        ["新建课程", "无", "页面跳转 → /teacher/courses/new", "进入创建表单", "取消返回列表"],
        ["进入课程", "课程非删除状态", "选中课程并跳转 → workspace", "更新当前课程上下文", "返回我的课程"],
        ["管理课程", "课程存在", "页面跳转 → course-settings", "加载课程设置", "课程侧栏返回"],
        ["更多：归档/恢复", "课程存在", "确认后更新课程状态", "调用课程状态接口", "提示后刷新列表"],
        ["更多：删除", "课程存在", "不可逆：确认后删除课程", "调用删除接口", "成功后移除卡片"],
        ["草稿：继续编辑", "浏览器存在创建草稿", "页面跳转 → 新建课程并恢复 localStorage 表单", "读取本地草稿", "可继续下一步"],
        ["草稿：删除", "浏览器存在创建草稿", "删除创建草稿", "移除 codetrack-create-course-draft", "草稿入口消失"],
        ["归档课程：查看", "课程已归档", "页面跳转 → workspace", "加载只读/现有课程数据", "返回课程列表"],
        ["使用技巧上一条/下一条", "技巧区可见", "页内切换提示内容", "仅本地索引", "不改变业务数据"],
        ["快速技巧卡片", "无", "展示入口：无跳转处理", "无", "页面不变"],
    ])

    add_h1(doc, "7. 新建课程")
    add_body(doc, "创建流程分为三步：基本信息 → 教学设置 → 创建完成。章节、知识点和高级设置在最终提交前都只是表单状态。")
    add_interactions(doc, [
        ["保存草稿", "已填写任意表单内容", "留在当前步骤并保存草稿", "写入 localStorage", "提示已保存；课程列表可继续编辑"],
        ["下一步", "当前步骤必填项有效", "页内进入下一步骤", "仅更新表单步骤", "校验失败时停留并提示"],
        ["上一步", "当前步骤大于 0", "页内返回上一表单步骤", "保留已填内容", "可再次继续"],
        ["最终创建", "基本信息/教学设置通过校验", "依次创建课程、班级、章节并进入完成页", "调用多项创建接口", "失败时提示并保留表单"],
        ["取消/返回我的课程", "任意创建步骤", "页面跳转 → /teacher/courses", "不自动提交当前表单", "已保存草稿可恢复"],
        ["学期管理", "教学设置步骤", "打开学期管理弹窗", "弹窗内页内状态", "关闭回到表单"],
        ["AI 高级设置", "教学设置步骤", "打开高级设置弹窗", "更新本地表单配置", "确认/关闭回到表单"],
        ["添加章节/知识点", "教学设置步骤", "在当前表单追加条目", "仅本地状态，最终创建时提交", "可继续编辑/删除"],
        ["进入课程工作空间", "课程创建成功", "当前实际跳转 → 当前课程 invite", "设置新课程为当前课程", "可从课程侧栏进入 workspace"],
        ["完成页后续建议卡片", "课程创建成功", "展示入口：当前未确认真实处理", "无", "使用主按钮或侧边栏继续"],
    ])
    add_callout(doc, "实现差异", "完成页主按钮文案是“进入课程工作空间”，但父级当前实际导航到邀请学生页 /invite。联调与验收应以此实际行为为准，或在后续版本统一文案与路由。", "caution")

    add_h1(doc, "8. 课程工作空间")
    add_interactions(doc, [
        ["课程设置", "已选课程", "页面跳转 → course-settings", "加载课程设置", "课程侧栏返回"],
        ["学生/任务/资料指标卡", "已选课程", "分别跳转 → classes/tasks/materials", "加载对应数据", "课程侧栏返回"],
        ["逾期指标卡", "已选课程", "页面跳转 → monitor", "加载提交状态", "课程侧栏返回"],
        ["进度/优秀率指标卡", "已选课程", "页面跳转 → analytics", "加载学情数据", "课程侧栏返回"],
        ["公告行", "公告存在", "打开公告详情抽屉并标记已读", "读取/写入 localStorage 阅读状态", "关闭回工作空间"],
        ["更多/全部公告", "公告区可见", "打开公告列表抽屉", "页内状态", "关闭回工作空间"],
        ["公告上一条/下一条", "详情抽屉已打开", "页内切换公告详情", "更新选中索引", "可关闭抽屉"],
        ["最近任务", "任务存在", "跳转 tasks 并定位、高亮目标行", "sessionStorage 写 codetrack:focus-task", "任务页清除/消费定位状态"],
        ["最近资料", "资料存在", "页面跳转 → materials", "加载资料列表", "课程侧栏返回"],
        ["快捷：新建作业", "已选课程", "页面跳转 → tasks", "加载任务页", "在任务页点击新建"],
        ["快捷：上传资料", "已选课程", "页面跳转 → materials", "加载资料页", "在资料页选择上传"],
        ["快捷：邀请学生", "已选课程", "页面跳转 → invite", "加载班级邀请码", "课程侧栏返回"],
        ["快捷：课堂讨论", "已选课程", "打开 discussion=1 讨论抽屉", "URL 状态", "关闭移除查询参数"],
        ["班级概览", "已选课程", "页面跳转 → classes", "加载班级列表", "课程侧栏返回"],
        ["AI 完整分析", "已选课程", "页面跳转 → analytics", "加载学情分析", "课程侧栏返回"],
    ])

    add_h1(doc, "9. 课程章节内容")
    add_interactions(doc, [
        ["学生预览", "章节数据已加载", "打开学生视角预览抽屉", "页内状态", "关闭回章节页"],
        ["新建章节（顶部/底部）", "已选课程", "打开新建章节弹窗", "确认后调用章节创建接口", "成功刷新章节列表"],
        ["章节行", "章节存在", "打开章节详情抽屉", "加载/绑定选中章节", "关闭回列表"],
        ["知识点/资料/练习标签", "章节抽屉已打开", "页内切换详情标签", "仅本地状态", "关闭抽屉返回"],
        ["添加知识点", "章节详情-知识点", "打开弹窗；确认后创建知识点", "调用知识点接口", "成功后刷新详情"],
        ["上传章节资料", "章节详情-资料", "选择文件并上传", "调用上传接口", "成功提示并刷新资料"],
        ["创建练习", "章节详情-练习", "打开练习弹窗", "保存草稿或发布接口", "成功后刷新练习列表"],
        ["练习：保存草稿", "练习表单有效", "创建/更新为草稿", "调用练习保存接口", "关闭弹窗或显示成功"],
        ["练习：发布给学生", "练习表单有效", "创建并发布练习", "调用发布接口", "状态变为已发布"],
        ["任务行：任务管理", "关联任务存在", "页面跳转 → tasks", "加载任务列表", "课程侧栏返回"],
        ["学生可见/发布开关", "章节或资料存在", "更新可见/发布状态", "调用对应更新接口", "状态即时刷新"],
    ])

    add_h1(doc, "10. 班级管理与邀请学生")
    add_h2(doc, "10.1 班级管理")
    add_interactions(doc, [
        ["筛选/搜索", "班级已加载", "页内过滤班级", "仅本地状态", "清空恢复"],
        ["新建班级", "已选课程", "打开弹窗；确认后创建班级", "调用班级创建接口", "成功刷新列表"],
        ["班级行", "班级存在", "仅选中该班级", "页内状态", "可继续查看详情"],
        ["查看详情", "班级存在", "打开班级详情抽屉", "绑定当前班级", "关闭返回列表"],
        ["邀请学生", "班级存在", "记录所选班级并跳转 → invite", "更新当前邀请目标", "课程侧栏返回"],
        ["管理班级", "班级存在", "展示入口：当前仅 toast", "无真实管理动作", "关闭提示仍在本页"],
        ["复制邀请码", "详情抽屉已打开", "写入系统剪贴板", "Clipboard API", "成功/失败提示"],
        ["查看全部/管理申请", "详情抽屉已打开", "打开加入状态抽屉", "页内状态", "关闭返回班级页"],
        ["加入状态筛选/搜索", "加入状态抽屉已打开", "页内过滤申请/学生", "仅本地状态", "清空恢复"],
        ["继续邀请", "加入状态抽屉已打开", "页面跳转 → invite", "保留班级选择", "课程侧栏返回"],
        ["学生：查看", "学生项存在", "展示入口：仅 toast", "无", "页面不变"],
        ["详情：编辑", "详情抽屉已打开", "展示入口：无处理函数", "无", "页面不变"],
    ])
    add_h2(doc, "10.2 邀请学生")
    add_interactions(doc, [
        ["选择班级", "至少一个班级", "切换当前邀请目标", "加载/展示对应加入码", "页面不跳转"],
        ["复制邀请链接/邀请码", "加入码有效", "写入剪贴板", "Clipboard API", "成功/失败提示"],
        ["重新生成链接/邀请码/二维码", "班级已选", "确认后生成新的加入码", "调用更新 join code 接口", "旧链接/码失效并刷新显示"],
        ["下载二维码", "二维码已生成", "浏览器下载二维码图片", "Canvas 导出下载", "下载后仍在当前页"],
        ["邀请方式开关", "邀请页已加载", "页内启用/禁用展示方式", "仅本地状态", "刷新可能恢复"],
        ["Excel/CSV 批量上传", "无", "展示入口：尚未接入上传解析", "无", "页面不变"],
        ["模板下载", "无", "展示入口：尚未接入文件下载", "无", "页面不变"],
        ["查看全部/查看更多", "存在加入记录", "写 session 状态并跳转 classes，自动打开加入状态抽屉", "sessionStorage 跨页传递", "关闭抽屉返回班级页"],
    ])

    add_h1(doc, "11. 任务管理")
    add_interactions(doc, [
        ["标签/筛选/排序/搜索", "任务已加载", "页内过滤或排序任务", "仅本地状态", "清空恢复"],
        ["新建任务", "已选课程", "打开右侧任务编辑面板第 0 步", "初始化表单", "关闭/取消回列表"],
        ["学生预览", "任务页可用", "当前实际跳转 → monitor", "加载提交监控", "课程侧栏返回"],
        ["编辑", "任务存在", "打开编辑面板并回填任务", "读取任务数据", "保存/关闭回列表"],
        ["查看成绩", "任务为已发布或已关闭", "页面跳转 → grading", "加载任务提交", "课程侧栏返回"],
        ["发布任务", "草稿任务", "打开编辑面板第 2 步", "回填任务并准备发布", "最终确认后刷新状态"],
        ["更多：提交监控", "任务存在", "页面跳转 → monitor", "加载任务提交", "课程侧栏返回"],
        ["更多：发布设置", "任务存在", "打开任务面板相关步骤", "页内编辑状态", "保存/关闭返回"],
        ["更多：复制为草稿", "任务存在", "复制当前任务为新草稿", "调用创建任务接口", "成功刷新任务列表"],
        ["面板：下一步/上一步", "任务面板已打开", "在第 0-3 步间切换", "仅本地表单状态", "校验失败停留"],
        ["面板：保存草稿", "任务表单满足最小要求", "创建/更新草稿任务", "调用任务保存接口", "成功提示并刷新列表"],
        ["面板：最终确认发布", "发布步骤校验通过", "发布任务", "调用任务发布接口", "状态变更并刷新"],
        ["AI 生成", "任务面板可用", "打开 AI 生成弹窗；生成后进入编辑器", "调用 AI 草稿接口", "失败提示；成功可继续编辑"],
        ["添加测试用例", "任务编辑面板", "展示入口：当前无处理函数", "无", "页面不变"],
        ["分页按钮", "列表底部可见", "展示入口：当前为静态视觉", "无真实分页请求", "页面不变"],
    ])

    add_h1(doc, "12. 任务监控与作业批改")
    add_h2(doc, "12.1 任务监控")
    add_interactions(doc, [
        ["任务选择器", "至少一个任务", "切换当前任务并加载提交", "调用提交查询接口", "列表与统计刷新"],
        ["刷新", "任务已选", "重新加载提交和统计", "调用查询接口", "加载态后更新"],
        ["导出", "任务已选", "展示入口：当前仅 toast", "不生成文件", "关闭提示仍在本页"],
        ["搜索/状态控件", "提交列表已加载", "部分为视觉控件；搜索当前未接筛选", "无或仅本地状态", "验收不应期待请求"],
        ["提交行/详情", "提交存在", "页内选中并展示提交详情", "仅本地选择", "切换其他提交"],
        ["进入批改", "提交存在", "页面跳转 → grading", "加载任务/学生提交", "课程侧栏返回"],
        ["前往任务管理", "当前课程无任务", "页面跳转 → tasks", "加载任务列表", "新建任务后可回监控"],
    ])
    add_h2(doc, "12.2 作业批改")
    add_interactions(doc, [
        ["任务选择器", "至少一个任务", "切换任务并加载提交", "调用提交查询接口", "学生列表刷新"],
        ["学生搜索/学生项", "提交已加载", "筛选或选择当前学生", "仅本地状态", "选择其他学生切换"],
        ["分数输入/四项滑杆", "已选提交", "页内重算最终分数", "仅本地评分状态", "保存前不持久化"],
        ["恢复建议", "存在 AI/评价建议", "恢复建议分数与评价值", "覆盖本地编辑值", "可继续修改"],
        ["保存草稿", "已选提交", "保存成绩与未发布反馈", "调用 grade 保存接口", "成功提示；学生端不可见"],
        ["发布成绩", "评分有效", "保存成绩、发布反馈并发布成绩", "依次调用保存/反馈发布/成绩发布", "状态变为已发布"],
        ["重新评估", "已选提交", "展示入口：当前无处理函数", "无", "页面不变"],
        ["版本对比", "存在版本信息", "展示入口：当前无处理函数", "无", "页面不变"],
    ])

    add_h1(doc, "13. 资料管理")
    add_interactions(doc, [
        ["上传课件/讲义", "已选课程", "选择文件并上传到当前目录", "调用资料上传接口", "成功刷新资料列表"],
        ["添加链接", "无", "创建预设 cppreference 外链资料", "调用资料创建接口", "成功刷新列表"],
        ["新建文件夹", "无", "打开弹窗；确认后创建文件夹", "调用文件夹创建接口", "成功刷新目录"],
        ["文件夹", "文件夹存在", "页内筛选为该文件夹内容", "仅本地过滤", "选择全部恢复"],
        ["文件夹菜单：删除", "文件夹存在", "不可逆确认；空文件夹永久删除，非空移入回收站", "调用删除/回收接口", "成功刷新目录"],
        ["标签/搜索", "资料已加载", "页内筛选资料", "仅本地状态", "清空恢复"],
        ["资料行", "资料存在", "页内选中并显示预览", "加载当前资料状态", "选择其他资料切换"],
        ["更多：导入知识图谱", "资料存在", "打开旧版课程知识点关联弹窗", "确认后调用关联接口", "与新版多图谱编辑器不同"],
        ["更多：移到回收站", "资料非删除状态", "确认后移入回收站", "调用资料删除接口", "列表刷新"],
        ["更多：恢复", "资料在回收站", "恢复资料", "调用恢复接口", "回到原/可用目录"],
        ["已删除文件夹：展开", "回收站有文件夹", "页内展开其资料", "仅本地状态", "可收起"],
        ["恢复整个文件夹", "已删除文件夹存在", "恢复文件夹及内容", "调用恢复接口", "目录刷新"],
        ["单独保留/移动资料", "已删除文件夹内有资料", "保留为独立资料或移动到现有文件夹", "调用资料更新接口", "成功刷新"],
        ["预览：外部打开/下载", "资料 URL 有效", "新标签页打开或浏览器下载", "外部导航/下载", "原页面保留"],
        ["编辑标签/添加标签", "资料已选", "展示入口：当前无真实保存动作", "无", "页面不变"],
        ["AI 大纲", "资料存在", "打开确认弹窗；确认后生成大纲", "调用 AI 大纲接口", "成功展示结果/状态"],
    ])

    add_h1(doc, "14. 课程知识图谱")
    add_body(doc, "该页面是新版多图谱编辑器。图谱列表和服务端保存是持久化边界；画布内节点位置、属性和关系编辑均须点击“保存草稿”后才能保证刷新恢复。")
    add_interactions(doc, [
        ["刷新", "已选课程", "重新加载图谱列表及当前图谱", "GET 图谱列表/详情", "加载后覆盖未保存本地编辑"],
        ["保存草稿", "当前图谱存在", "保存标题、说明、班级、节点、关系和状态", "PUT 当前图谱", "成功提示；刷新可恢复"],
        ["发布", "当前图谱存在", "先保存当前编辑，再发布图谱", "PUT 后 POST publish", "状态变为已发布"],
        ["选择文件", "支持 PDF/DOCX/MD/TXT", "页内显示文件名与大小", "仅本地文件选择", "分析前可重新选择"],
        ["分析并生成图谱", "已选择有效文件", "上传资料并创建草稿图谱", "multipart POST from-files", "成功打开新图谱；失败提示"],
        ["新建空白图谱", "已选课程", "创建含默认核心节点的新图谱", "POST 图谱", "成功加入列表并打开"],
        ["图谱列表项", "图谱存在", "切换并打开所选图谱", "GET 图谱详情", "当前项高亮"],
        ["选择模式", "画布有节点/关系", "点击节点或边并在右侧编辑", "仅本地选择", "点其他元素切换"],
        ["连接模式", "至少两个节点", "先点起点再点终点，创建当前类型的边", "仅本地新增关系", "保存草稿后持久化"],
        ["关系类型", "工具栏可用", "在前驱/后继/相关间切换", "仅本地工具状态", "影响下一条新建关系"],
        ["节点", "当前图谱存在", "新增自定义节点并选中", "仅本地新增节点", "保存草稿后持久化"],
        ["布局", "画布有多个节点", "核心节点居中，其余节点椭圆分布", "批量更新本地坐标", "保存草稿后持久化"],
        ["删除", "已选节点或关系", "删除选中元素；删节点同时删关联边", "仅本地删除", "保存后持久化"],
        ["拖拽节点", "画布可用", "修改节点坐标", "仅本地坐标", "保存后持久化"],
        ["画布平移/缩放", "画布可用", "改变查看视口", "仅前端视图状态", "不修改业务数据"],
        ["属性面板字段", "已选节点/关系", "编辑名称、类型、难度、说明或关系类型", "仅本地更新", "保存后持久化"],
        ["关联节点", "已选节点且存在入边/出边", "切换选中到对应关系", "仅本地选择", "右侧显示关系属性"],
        ["删除当前图谱", "当前图谱存在", "不可逆：确认后删除整张图谱", "DELETE 当前图谱", "列表刷新并选择其他图谱"],
    ])
    add_callout(doc, "未保存风险", "节点新增、拖拽、属性修改、连边、布局和画布删除先发生在前端状态中。直接切换图谱或刷新页面可能丢失这些修改，应先点击“保存草稿”。", "risk")

    add_h1(doc, "15. 学情分析")
    add_interactions(doc, [
        ["班级总览", "分析页已加载", "页内切换至班级统计视图", "加载/展示汇总数据", "同页切换其他分段"],
        ["个体诊断", "分析页已加载", "页内切换至学生诊断视图", "加载个体数据", "同页切换其他分段"],
        ["预警中心", "分析页已加载", "页内切换至预警列表", "加载预警数据", "同页切换其他分段"],
        ["班级/搜索/学生选择", "个体诊断视图", "筛选并切换诊断对象", "查询或本地过滤学生", "诊断内容随选择更新"],
        ["导出报告", "分析数据可见", "展示入口：当前仅 toast", "不生成文件", "关闭提示仍在本页"],
    ])

    add_h1(doc, "16. 课堂讨论")
    add_body(doc, "课堂讨论由 URL 查询参数 discussion=1 控制，是覆盖在课程工作空间上的抽屉流程。内部包含首页、创建、进行中和历史记录状态。")
    add_interactions(doc, [
        ["讨论首页：新建讨论", "抽屉已打开", "页内进入创建表单", "初始化讨论表单", "取消回讨论首页"],
        ["发布开关", "创建表单已打开", "切换保存为草稿或立即发布", "仅本地表单状态", "影响保存后的去向"],
        ["保存讨论", "表单有效", "已发布 → 进行中；草稿 → 讨论首页", "调用讨论保存/发布接口", "成功切换状态，失败提示"],
        ["进行中的讨论", "存在活动讨论", "进入 live 状态并每 2 秒轮询", "周期调用讨论查询接口", "结束/关闭后停止"],
        ["草稿讨论项", "草稿存在", "当前行为：发布该草稿并进入 live", "调用发布接口", "状态变为进行中"],
        ["结束讨论", "讨论进行中", "确认后结束，回首页/历史", "调用结束接口", "轮询停止并刷新列表"],
        ["历史记录", "存在已结束讨论", "页内进入历史状态/查看记录", "加载历史数据", "返回讨论首页"],
        ["学生预览", "讨论抽屉已打开", "新标签页打开 /student/discussions", "外部页面导航", "教师页保留"],
        ["关闭抽屉", "discussion=1 存在", "关闭讨论并移除 URL 查询参数", "更新浏览器地址", "回到 workspace"],
    ])

    add_h1(doc, "17. AI 审核与个人设置")
    add_h2(doc, "17.1 AI 审核")
    add_interactions(doc, [
        ["审核列表项", "存在 AI 建议", "页内选择并显示建议详情", "仅本地选择", "选择其他项切换"],
        ["驳回", "建议未处理", "提交驳回决定", "调用审核 API", "状态刷新"],
        ["接受", "建议未处理", "直接接受建议", "调用审核 API", "状态刷新"],
        ["修改并接受", "建议未处理", "进入编辑状态；确认后提交修改版本", "调用审核 API", "成功退出编辑并刷新"],
        ["取消修改", "正在编辑", "放弃本次本地修改", "仅本地状态", "回到详情"],
    ])
    add_h2(doc, "17.2 个人设置")
    add_interactions(doc, [
        ["左侧设置菜单", "设置页已打开", "展示入口：标签可见，但当前内容始终为个人资料", "无真实分区切换", "页面主体不变"],
        ["个人资料字段", "资料已加载", "只读展示", "不允许编辑", "无保存入口"],
        ["保存偏好", "偏好区可见", "当前仅 toast 提示", "无服务端持久化", "刷新可能恢复默认"],
    ])

    add_h1(doc, "18. 跨页面业务闭环")
    add_table(doc, ["业务目标", "推荐点击路径", "完成判定"], [
        ["创建课程并邀请学生", "工作台/我的课程 → 新建课程 → 三步创建 → 完成页主按钮 → 邀请学生", "新课程、默认班级/章节创建成功，邀请码可复制"],
        ["创建并发布任务", "课程工作空间 → 任务管理 → 新建任务 → 分步填写 → 最终确认发布", "任务状态为已发布，监控页可选择该任务"],
        ["监控并发布成绩", "任务管理 → 提交监控 → 选择提交 → 进入批改 → 发布成绩", "成绩与反馈均为已发布"],
        ["上传资料并生成大纲", "课程工作空间 → 资料管理 → 上传 → 选择资料 → AI 大纲 → 确认", "资料列表存在文件且大纲接口成功"],
        ["资料生成知识图谱并发布", "课程侧栏 → 课程知识图谱 → 选择文件 → 分析生成 → 编辑/连边 → 保存草稿 → 发布", "图谱为已发布，刷新后节点和关系仍存在"],
        ["邀请记录回看", "班级管理 → 邀请学生 → 查看全部/查看更多 → 班级管理加入状态抽屉", "跨页后自动打开正确班级的加入状态"],
        ["发起课堂讨论", "课程工作空间 → 课堂讨论 → 新建讨论 → 打开发布开关 → 保存", "进入 live 状态并开始轮询；结束后进入历史"],
    ], [1900, 4920, 2540], font_size=8.8)

    add_h1(doc, "19. 展示入口与当前限制")
    add_table(doc, ["页面", "入口", "当前实际行为", "建议验收口径"], [
        ["登录门户", "进入科研", "仅 alert，无科研页面", "标记未接入"],
        ["工作台", "最近动态/查看全部动态", "无处理函数", "不作为跳转验收项"],
        ["我的课程", "快速技巧卡片", "无处理函数", "仅视觉验收"],
        ["新建课程", "完成页建议卡片", "未确认真实处理", "使用主按钮验收"],
        ["班级管理", "管理班级/学生查看/详情编辑", "toast 或无处理函数", "标记未接入"],
        ["邀请学生", "批量上传/模板下载", "视觉占位", "标记未接入"],
        ["任务管理", "添加测试用例/分页", "无处理或静态视觉", "标记未接入"],
        ["任务监控", "导出/搜索", "toast 或未接筛选", "不期待文件/搜索结果"],
        ["作业批改", "重新评估/版本对比", "无处理函数", "标记未接入"],
        ["资料管理", "标签编辑/添加标签", "无真实保存", "不期待刷新保持"],
        ["学情分析", "导出报告", "仅 toast", "不期待下载文件"],
        ["个人设置", "设置菜单/保存偏好", "无分区切换/无持久化", "仅展示验收"],
    ], [1560, 2050, 2990, 2760], font_size=8.8)
    add_callout(doc, "测试注意", "展示入口不是缺少测试步骤，而是当前实现边界。测试报告应写“入口可见但未连接真实动作”，不要误写为路由或接口失败。", "caution")

    add_h1(doc, "20. 附录")
    add_h2(doc, "20.1 跨页面状态存储")
    add_table(doc, ["存储位置 / 键", "写入场景", "读取 / 消费场景", "影响"], [
        ["localStorage: codetrack-create-course-draft", "新建课程-保存草稿", "我的课程-继续编辑", "恢复三步创建表单；删除草稿时移除"],
        ["sessionStorage: codetrack:focus-task", "工作空间点击最近任务", "任务管理页加载后", "滚动并高亮目标任务"],
        ["sessionStorage: 班级加入状态上下文", "邀请页查看全部/查看更多", "班级管理页加载后", "自动选择班级并打开加入状态抽屉"],
        ["URL: discussion=1", "工作空间/侧栏点击课堂讨论", "课程工作空间", "打开讨论抽屉；关闭时移除"],
        ["localStorage: 公告已读状态", "打开公告详情", "工作空间公告列表", "控制本浏览器已读显示"],
    ], [2450, 2250, 2400, 2260], font_size=8.8)
    add_h2(doc, "20.2 关键接口动作映射")
    add_table(doc, ["模块", "动作", "接口类型 / 方向", "关键结果"], [
        ["工作台/课程", "查询、状态修改、删除", "GET / PATCH(或 PUT) / DELETE", "刷新统计、课程状态或移除课程"],
        ["新建课程", "创建课程、班级、章节", "POST 多步骤", "完成页出现新课程"],
        ["章节", "章节/知识点/练习/资料", "POST / PUT / 上传", "详情抽屉刷新"],
        ["班级/邀请", "创建班级、更新加入码", "POST / PUT", "班级与邀请码刷新"],
        ["任务", "创建草稿、更新、发布、复制", "POST / PUT / publish", "任务状态刷新"],
        ["批改", "保存成绩、发布反馈、发布成绩", "多次写接口", "学生可见状态改变"],
        ["资料", "上传、创建目录、删除/恢复、AI 大纲", "上传 / POST / DELETE / 恢复", "资料树和预览刷新"],
        ["知识图谱", "列表/详情/创建/文件生成/保存/发布/删除", "GET / POST / PUT / DELETE", "图谱、节点和关系持久化"],
        ["课堂讨论", "保存/发布/轮询/结束", "POST / GET polling / end", "live 与 history 状态变化"],
        ["AI 审核", "接受、驳回、修改接受", "审核写接口", "建议处理状态刷新"],
    ], [1450, 2600, 2460, 2850], font_size=8.7)
    add_h2(doc, "20.3 验收执行顺序")
    add_bullets(doc, [
        "先验证全局导航和课程上下文，再验证课程内页面；避免 courseId 缺失造成误判。",
        "对每个写操作同时检查按钮反馈、列表/状态变化以及刷新后的持久化结果。",
        "对图谱画布先编辑后保存，再刷新验证；未保存编辑丢失属于当前设计边界。",
        "删除课程、删除图谱、重新生成邀请码等不可逆操作只在专用测试数据上执行。",
        "对展示入口只检查可见性、禁用/提示状态，不要求页面跳转、接口请求或下载文件。",
    ])
    add_callout(doc, "结论", "本文档以当前实现为唯一判定依据，既覆盖真实可执行路径，也保留尚未接入的视觉入口。后续按钮行为、路由或接口变化时，应同步更新对应交互矩阵和限制清单。", "info")

    doc.core_properties.title = "CodeTrack 教师端前端逻辑交互文档"
    doc.core_properties.subject = "教师端按钮去向、状态变化与数据动作"
    doc.core_properties.author = "CodeTrack 项目组"
    doc.core_properties.keywords = "CodeTrack, 教师端, 前端交互, 路由, 操作文档"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
