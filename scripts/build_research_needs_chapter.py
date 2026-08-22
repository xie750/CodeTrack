from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"E:\teacher")
OUT_DIR = ROOT / "artifacts"
ASSET_DIR = OUT_DIR / "codetrack_research_needs_assets"
OUT_PATH = OUT_DIR / "CodeTrack_研究需求与总体方案_论文版.docx"

BLUE = "2F6F9F"
GREEN = "3A8D74"
LIGHT_BLUE = "EAF3F8"
LIGHT_GREEN = "EAF5F1"
LIGHT_ORANGE = "FFF5E6"
GRAY = "5B6570"


def set_run_font(run, east_asia="宋体", western="Times New Roman", size=12, bold=None, color=None):
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), east_asia)
    rpr.rFonts.set(qn("w:ascii"), western)
    rpr.rFonts.set(qn("w:hAnsi"), western)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def para_format(p, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if line == 1.5 else WD_LINE_SPACING.SINGLE
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Cm(0.85) if indent else None


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_three_line_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("left", "right", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    for edge, size in (("top", "12"), ("bottom", "12"), ("insideH", "8")):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), "333333")
        borders.append(node)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    para_format(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, east_asia="黑体", size=12, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=12)
    else:
        r = p.add_run(text)
        set_run_font(r, size=12)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=8 if level == 1 else 5, after=6)
    p.paragraph_format.keep_with_next = True
    size = 14 if level == 1 else 12
    east = "黑体" if level == 1 else "宋体"
    r = p.add_run(text)
    set_run_font(r, east_asia=east, size=size, bold=True, color="1E2D38")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=4, line=1.0)
    r = p.add_run(text)
    set_run_font(r, east_asia="宋体", size=10.5)
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, caption, headers, rows, widths=None, page_break_before=False):
    cap = add_caption(doc, caption)
    if page_break_before:
        cap.paragraph_format.page_break_before = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths: cell.width = Cm(widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
        r = p.add_run(h)
        set_run_font(r, east_asia="黑体", size=10.5, bold=True)
    for data in rows:
        row = table.add_row()
        for i, value in enumerate(data):
            cell = row.cells[i]
            if widths: cell.width = Cm(widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0)
            r = p.add_run(str(value))
            set_run_font(r, east_asia="宋体", size=10.5)
    set_three_line_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def img_font(size, bold=False):
    path = r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"
    return ImageFont.truetype(path, size)


def draw_box(draw, box, text, fill, edge, size=28, radius=20):
    if not fill.startswith("#"): fill = "#" + fill
    if not edge.startswith("#"): edge = "#" + edge
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=edge, width=3)
    font = img_font(size)
    bb = draw.multiline_textbbox((0, 0), text, font=font, spacing=8, align="center")
    tw, th = bb[2]-bb[0], bb[3]-bb[1]
    x1, y1, x2, y2 = box
    draw.multiline_text(((x1+x2-tw)/2, (y1+y2-th)/2-2), text, fill="#1E2D38", font=font, spacing=8, align="center")


def arrow(draw, start, end, color="#64737F", width=5):
    if not color.startswith("#"): color = "#" + color
    draw.line([start, end], fill=color, width=width)
    sx, sy = start; ex, ey = end
    if abs(ex-sx) >= abs(ey-sy):
        pts = [(ex, ey), (ex-16 if ex>sx else ex+16, ey-10), (ex-16 if ex>sx else ex+16, ey+10)]
    else:
        pts = [(ex, ey), (ex-10, ey-16 if ey>sy else ey+16), (ex+10, ey-16 if ey>sy else ey+16)]
    draw.polygon(pts, fill=color)


def build_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    # Architecture diagram: layers mirror the prototype's three user-facing terminals and research entry.
    im = Image.new("RGB", (2400, 1360), "white"); d = ImageDraw.Draw(im)
    layers = [("应用层", 60, 280, "#F6FAFC", BLUE), ("智能能力层", 350, 580, LIGHT_GREEN, GREEN), ("知识与数据层", 650, 880, LIGHT_BLUE, BLUE), ("模型与平台支撑层", 950, 1190, "#F7F7F7", "#7D8790")]
    for label, y1, y2, fill, edge in layers:
        d.rounded_rectangle((70,y1,2330,y2), radius=24, outline="#D5DDE2", width=3, fill="#FFFFFF")
        d.text((110,y1+22), label, font=img_font(34, True), fill="#1E2D38")
        d.line((310,y1+12,310,y2-12), fill="#D5DDE2", width=2)
    apps=["教师端\n课程·班级·任务·资料\n知识图谱·学情分析", "学生端\n课程章节·练习·成绩\n分步答疑与个人画像", "管理员端\n权限·版本·日志\n安全审计与运行治理", "科研入口\n文献阅读·代码复现\n实验记录与成果沉淀"]
    for i,t in enumerate(apps): draw_box(d,(370+i*475,115,800+i*475,245),t,"#F6FAFC",BLUE,24)
    agents=["问题分类", "检索增强\nRAG", "步骤拆解", "求解编排", "答案质检"]
    for i,t in enumerate(agents):
        x=390+i*375; draw_box(d,(x,405,x+285,535),t,LIGHT_GREEN,GREEN,24)
        if i<4: arrow(d,(x+285,470),(x+360,470),GREEN)
    data=["课程与教材\n资料库", "结构化题库\n问题-题型-推导-答案", "知识图谱\n课程-章节-知识点", "向量索引与\n版本元数据"]
    for i,t in enumerate(data): draw_box(d,(385+i*480,705,775+i*480,835),t,LIGHT_BLUE,BLUE,22)
    support=["Qwen2.5-7B-Instruct\n+ LoRA适配器", "LangChain + Chroma\n统一API / Gradio", "权限·日志·审计\n人工复核与回溯"]
    for i,t in enumerate(support): draw_box(d,(460+i*620,1005,980+i*620,1145),t,"#F7F7F7","#7D8790",22)
    for y1,y2 in ((280,350),(580,650),(880,950)): arrow(d,(1200,y1+8),(1200,y2-8),"#7D8790")
    im.save(ASSET_DIR/"architecture.png")

    # Business flow diagram.
    im=Image.new("RGB",(2400,760),"white"); d=ImageDraw.Draw(im)
    flow=[("用户进入", "识别角色与课程"), ("提出任务", "题目/资料/科研问题"), ("智能编排", "分类→检索→拆解→求解→质检"), ("结果呈现", "步骤、依据、来源与画像"), ("反馈闭环", "教师确认/纠错回流")]
    for i,(a,b) in enumerate(flow):
        x=70+i*465; draw_box(d,(x,180,x+350,500),a+"\n\n"+b, ["#F6FAFC",LIGHT_BLUE,LIGHT_GREEN,LIGHT_ORANGE,"#FDECEC"][i], [BLUE,BLUE,GREEN,"#C9822B","#B94A48"][i],25)
        if i<4: arrow(d,(x+350,340),(x+440,340),"#64737F")
    d.text((80,65),"CodeTrack 面向三端与科研入口的统一业务流程",font=img_font(38,True),fill="#1E2D38")
    im.save(ASSET_DIR/"business_flow.png")

    # Research framework diagram.
    im=Image.new("RGB",(2400,1120),"white"); d=ImageDraw.Draw(im)
    d.text((90,55),"研究框架：需求牵引、技术实现、验证迭代",font=img_font(38,True),fill="#1E2D38")
    cols=[("问题定义",["通用模型知识偏差","教师答疑难以规模化","工具分散、缺乏闭环"],LIGHT_BLUE,BLUE),
          ("方案设计",["四类用户需求","垂直模型+RAG","多智能体协同"],LIGHT_GREEN,GREEN),
          ("系统实现",["Qwen2.5-7B + LoRA","LangChain + Chroma","三端与科研入口"],LIGHT_ORANGE,"#C9822B"),
          ("验证迭代",["离线对比评测","教师盲评与用户试用","反馈回流与版本审计"],"#FDECEC","#B94A48")]
    for i,(title,items,fill,edge) in enumerate(cols):
        x=100+i*570; draw_box(d,(x,230,x+450,820),title+"\n\n"+"\n".join("· "+s for s in items),fill,edge,26)
        if i<3: arrow(d,(x+450,525),(x+530,525),"#64737F")
    d.text((215,930),"研究目标：形成可复制、可迁移、可审计的学科智能学习平台方案",font=img_font(30,True),fill="#1E2D38")
    im.save(ASSET_DIR/"research_framework.png")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def setup_document():
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54); sec.left_margin=Cm(3.17); sec.right_margin=Cm(3.17)
    sec.header_distance=Cm(1.5); sec.footer_distance=Cm(1.5)
    styles=doc.styles
    normal=styles["Normal"]
    normal.font.name="Times New Roman"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); normal.font.size=Pt(12)
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=header.add_run("CodeTrack：研究需求与总体方案"); set_run_font(r,east_asia="宋体",size=9,color=GRAY)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=footer.add_run("信赛创翼队  ·  "); set_run_font(r,east_asia="宋体",size=9,color=GRAY)
    add_page_field(footer); return doc


def build():
    build_diagrams(); doc=setup_document()
    p=doc.add_paragraph(); para_format(p,indent=False,align=WD_ALIGN_PARAGRAPH.CENTER,before=12,after=8)
    r=p.add_run("CodeTrack：面向学科生态的垂直大模型与自主化学科智能体集成方案"); set_run_font(r,east_asia="黑体",size=18,bold=True,color="1E2D38")
    p=doc.add_paragraph(); para_format(p,indent=False,align=WD_ALIGN_PARAGRAPH.CENTER,after=12)
    r=p.add_run("研究需求与总体方案"); set_run_font(r,east_asia="宋体",size=15,bold=True,color=GREEN)
    p=doc.add_paragraph(); para_format(p,indent=False,align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
    r=p.add_run("信赛创翼队"); set_run_font(r,east_asia="宋体",size=12)
    p=doc.add_paragraph(); para_format(p,indent=False,align=WD_ALIGN_PARAGRAPH.CENTER,after=18)
    r=p.add_run("信阳学院计算机与人工智能学院"); set_run_font(r,east_asia="宋体",size=10.5,color=GRAY)
    p=doc.add_paragraph(); para_format(p,indent=False,align=WD_ALIGN_PARAGRAPH.LEFT,before=8,after=2)
    r=p.add_run("摘  要："); set_run_font(r,east_asia="黑体",size=12,bold=True)
    r=p.add_run("针对高校理工科教学中专业知识分散、复杂习题答疑难以规模化以及通用大模型推理过程不可追溯等问题，本文结合项目申报方案与 CodeTrack 原型界面，提出由教师端、学生端、管理员端和科研入口组成的学科智能学习平台总体方案。平台以 Qwen2.5-7B-Instruct 为计划采用的基座模型，以 LoRA 参数高效微调和检索增强生成（RAG）为核心技术路径，构建问题分类、检索增强、步骤拆解、求解编排和答案质检的协同链路，并通过统一数据对象、权限控制、版本管理和人工复核机制支撑平台治理。本文重点论述四类用户需求、分层系统架构、跨端业务流程及研究框架，明确拟开展工作、评价指标和风险边界，为后续模型实现与系统验证提供规范化依据。")
    set_run_font(r,east_asia="宋体",size=12)
    p=doc.add_paragraph(); para_format(p,indent=False,align=WD_ALIGN_PARAGRAPH.LEFT,after=12)
    r=p.add_run("关键词："); set_run_font(r,east_asia="黑体",size=12,bold=True)
    r=p.add_run("学科垂直大模型；检索增强生成；多智能体协同；智能教学平台；科研入口")
    set_run_font(r,east_asia="宋体",size=12)

    add_heading(doc,"1 研究需求分析",1)
    add_body(doc,"CodeTrack 的需求分析以项目申报表中的教学实践问题为起点，并以现有 UI 原型所呈现的三端工作流与科研入口为产品边界。需求分析不把模型能力假设为已实现功能，而是将其表述为待验证的系统能力：平台首先保证课程、资料、任务、成绩和权限等基本教学业务可用，再将垂类模型、RAG 与多智能体能力嵌入关键节点，形成可解释、可复核的 AI 辅助闭环。")
    add_heading(doc,"1.1 教师端需求",2)
    add_body(doc,"教师端承担课程组织、知识建设、任务发布、学习诊断和 AI 结果确认等职责。原型界面已形成以教学工作台为入口的导航结构，包括课程管理、班级管理、任务管理、教学材料、课程知识图谱和学情分析等模块。由此，教师端需求可归纳为四个层次：一是将课程、章节、知识点和资料组织为可检索的课程知识空间；二是支持作业或练习的创建、发布、批改、成绩查看与版本对比；三是利用学情数据识别高频错误、未完成任务和需要关注的学生；四是对 AI 生成的题目、讲义、反馈或解题步骤进行审核、修改和发布，保留人工教学决策权。")
    add_heading(doc,"1.2 学生端需求",2)
    add_body(doc,"学生端围绕“理解过程而非只获得答案”的学习目标设计。学生进入课程后，应能够按章节查看知识点、教学资料和练习任务，提交作业并查看成绩与教师反馈；在遇到问题时，可直接输入自然语言或 LaTeX 公式，由系统自动识别题型并生成分步推导、依据说明和相关资料来源。学生端还需要形成个人画像，用于记录知识点掌握状态、错误类型、练习轨迹和反馈偏好，但画像仅用于学习支持，不直接替代教师评价或自动生成惩罚性结论。")
    add_heading(doc,"1.3 管理员端需求",2)
    add_body(doc,"管理员端是平台的治理与安全边界，面向用户、课程、知识库、模型服务和运行日志提供统一管理。其核心需求包括：按角色配置最小必要权限；维护课程与知识库的版本状态、来源和审核记录；监测模型调用、检索命中、质检失败和人工转交等关键事件；对异常调用、敏感信息、越权访问和高风险输出进行告警与追溯；支持模型、提示词、检索参数和智能体策略的版本对比与回滚。管理员端不参与具体教学判断，而是保障系统行为可见、可控、可恢复。")
    add_heading(doc,"1.4 科研入口需求",2)
    add_body(doc,"科研入口用于连接平台的教学数据资产与研究工作流，保留为 CodeTrack 的差异化产品入口。该入口面向论文阅读、代码复现、实验记录、数据集管理和结果归档等任务，支持将课程知识点、检索证据、模型版本和实验参数关联到研究项目中。科研入口与教学端共享权限和审计机制，但在数据范围上实行项目级隔离，防止未授权的学生信息或课程资料进入科研空间。由于真实模型代码和实验结果尚未全部完成，当前阶段只定义工作空间、数据对象和操作边界，不宣称已经实现完整科研智能体。")
    add_table(doc,"表1  四类用户需求与验收关注点",["用户端","主要任务","AI辅助点","人工确认/治理要求"],[
        ["教师端","课程、资料、任务、知识图谱、学情分析","题目/讲义草拟；学情诊断；分步解题辅助","教师审核后发布；保留修改记录"],
        ["学生端","章节学习、练习、提交、成绩、反馈","题型识别；RAG检索；分步推导；个人画像","不得以AI结果替代教师评价；可追问与纠错"],
        ["管理员端","用户、权限、版本、日志、安全","异常检测；版本对比；运行监控","最小权限；高风险任务人工转交；可回滚"],
        ["科研入口","文献、代码、实验、数据与成果归档","文献摘要；代码复现提示；实验记录整理","项目级隔离；来源和模型版本可追溯"]], [2.2,4.0,4.1,4.4])

    add_heading(doc,"2 平台总体架构",1)
    add_body(doc,"平台总体架构采用“应用层—智能能力层—知识与数据层—模型与平台支撑层”的分层设计。分层的目的不是增加系统复杂度，而是把用户界面、模型编排、知识资产和基础设施解耦，使教学端和科研入口可以复用同一套知识、模型与审计能力，同时允许不同课程按权限加载不同的知识范围。")
    doc.add_picture(str(ASSET_DIR/"architecture.png"), width=Cm(15.6)); add_caption(doc,"图1  CodeTrack 平台总体架构")
    add_heading(doc,"2.1 应用层：三端与科研入口",2)
    add_body(doc,"应用层承接用户任务与结果呈现。教师端以课程工作台为核心，学生端以课程学习与练习为核心，管理员端以治理控制台为核心，科研入口以研究工作空间为核心。四类入口共享统一身份认证和权限模型，但根据角色限制可见课程、知识库、数据集和操作按钮。原型中的绿色主色、浅色背景、卡片式信息区和清晰的侧边导航体现了平台面向高频教学工作的简洁交互取向；论文中的系统设计沿用这一信息架构，而不把视觉样式误写成技术能力。")
    add_heading(doc,"2.2 智能能力层：垂类模型、RAG 与多智能体",2)
    add_body(doc,"智能能力层由计划采用的 Qwen2.5-7B-Instruct 基座模型、LoRA 适配器、检索增强组件和多智能体调度器组成。项目申报方案计划构建不少于350条“问题—题型—分步推导—答案”四元组数据，并按8:1:1划分训练、验证和测试集；拟采用秩 r=16 的 LoRA 进行参数高效微调，训练环境为单张 RTX 3090 24GB。上述内容属于实施方案与资源规划，最终效果需要通过独立测试集和教师盲评验证。")
    add_body(doc,"多智能体链路由问题分类、检索增强、步骤拆解、求解编排和答案质检五个职责组成。分类智能体判断概念辨析、数值计算、证明推理或实际应用等题型；检索增强智能体在 Chroma 向量库中检索课程资料和相似例题；步骤拆解智能体按题型模板规划推导步骤；求解编排器调用模型完成各步骤生成；答案质检智能体检查术语、公式、步骤依赖和结论一致性。质检未通过时最多回溯两次，连续失败或高风险任务转人工确认。")
    add_heading(doc,"2.3 知识与数据层：课程知识空间与个人画像",2)
    add_body(doc,"知识与数据层保存课程资料、结构化题库、知识点关系、向量索引和版本元数据。课程资料需要经过来源登记、权限确认、格式解析、章节切分、知识点绑定和教师审核后，才进入可检索知识空间。个人画像由学习行为、练习结果、知识点掌握状态和反馈记录构成，使用时应遵循数据最小化和目的限定原则，仅用于个性化学习建议、教师学情分析和研究统计，不将画像直接作为自动评分或纪律处分依据。")
    add_heading(doc,"2.4 模型与平台支撑层：统一接口、安全与审计",2)
    add_body(doc,"支撑层通过统一 API 封装模型推理、检索、智能体调度和业务数据访问，前端可以采用 Gradio 原型或后续正式 Web 客户端。LangChain 用于流程编排，Chroma 用于向量检索，服务端记录请求角色、课程范围、知识库版本、模型版本、检索证据、质检结果和人工处理状态。涉及代码执行的科研任务应放入隔离沙箱，不允许模型直接访问宿主机敏感目录、系统命令或未授权数据库。")

    add_heading(doc,"3 平台业务流程",1)
    add_body(doc,"平台业务流程以用户角色为起点，以可追溯的 AI 辅助结果和人工反馈为终点。不同端的页面操作虽然不同，但在数据层都遵循“任务创建—知识准备—智能处理—结果确认—行为留痕”的统一规则。")
    doc.add_picture(str(ASSET_DIR/"business_flow.png"), width=Cm(15.6)); add_caption(doc,"图2  CodeTrack 三端与科研入口统一业务流程")
    add_heading(doc,"3.1 学生学习与答疑流程",2)
    add_body(doc,"学生首先进入课程并选择章节、知识点或练习任务；系统读取用户身份、课程范围和当前知识库版本，再接收题目文本或公式。调度器执行“分类—检索—拆解—求解—质检”链路，输出最终答案、分步推导、引用来源和可继续追问的子步骤。学生可以标记“看不懂”“步骤错误”或“与课堂方法不一致”，这些反馈进入待复核队列。未经教师或管理员确认的反馈，不得直接改变正式知识库和模型训练数据。")
    add_heading(doc,"3.2 教师教学与审核流程",2)
    add_body(doc,"教师创建课程或进入既有课程工作台，上传讲义、教材、实验指导等资料并绑定章节和知识点；资料经解析与审核后进入课程知识库。教师可在任务管理中生成或编辑练习，设置题型、难度、知识点和截止时间，发布前预览学生视角。系统运行期间，教师查看班级完成情况、错误分布、AI 建议和高风险任务，对 AI 草稿进行确认、修改或驳回；确认后的内容才可作为教学资源或有效反馈。")
    add_heading(doc,"3.3 管理员治理与异常处理流程",2)
    add_body(doc,"管理员配置角色、课程权限和项目空间，维护模型与知识库版本，查看运行日志和异常告警。当出现检索证据不足、质检连续失败、敏感信息风险、越权访问或沙箱异常时，系统暂停自动发布并转人工处理。管理员完成原因标注、版本回滚或策略调整后，形成可审计的处理记录。该流程确保模型的“自动化”始终受到权限和人工复核边界约束。")
    add_heading(doc,"3.4 科研辅助与成果沉淀流程",2)
    add_body(doc,"科研用户在项目空间中导入经过授权的文献、代码、数据和实验记录，系统按项目范围提供检索、摘要、代码解释和实验日志整理。每条研究记录关联数据版本、模型版本、提示词或智能体策略版本及运行时间，便于复现实验。科研辅助输出定位为草稿和分析建议，研究者需要对事实、代码可运行性和实验结论负责。")
    add_table(doc,"表2  核心业务节点及可追溯记录",["节点","系统动作","必须记录的元数据","失败处理"], [
        ["知识入库","解析、切分、向量化、绑定知识点","来源、授权、课程、版本、审核人","退回修改；不进入正式检索库"],
        ["问题处理","分类、检索、拆解、求解、质检","用户角色、题型、证据、模型版本、质检结果","回溯最多2次；失败转人工"],
        ["任务发布","教师编辑并向班级发布","任务版本、知识点、截止时间、发布人","撤回或回滚至上一版本"],
        ["反馈回流","收集学生标记与教师修订","反馈类型、原输出、修订内容、复核状态","复核后再进入补充数据队列"],
        ["科研归档","保存代码、数据、实验与结论","项目、数据版本、运行环境、模型版本","权限隔离；异常记录不可删除"]], [2.4,5.1,4.9,3.0], page_break_before=True)

    add_heading(doc,"4 研究框架",1)
    add_body(doc,"本研究以教学实践中的真实问题为牵引，围绕“知识有效嵌入—复杂推理可控拆解—系统自主决策与闭环迭代”展开。研究框架同时覆盖技术路线和产品验证，避免只评估模型离线指标而忽略教师可用性、学生可跟随性以及系统安全性。")
    doc.add_picture(str(ASSET_DIR/"research_framework.png"), width=Cm(15.6)); add_caption(doc,"图3  CodeTrack 研究框架")
    add_heading(doc,"4.1 研究目标",2)
    add_body(doc,"第一，构建面向学科推理的结构化数据规范，形成可用于微调、检索和评测的四元组样本；第二，完成垂直模型与多智能体链路的原型集成，验证题型分类、知识检索、步骤拆解和答案质检之间的接口契约；第三，建设教师端、学生端、管理员端和科研入口的统一平台原型，验证跨端数据与权限边界；第四，建立离线评测、教师盲评、目标用户试用和反馈回流组成的闭环评价机制。")
    add_heading(doc,"4.2 研究内容与技术路线",2)
    add_body(doc,"研究内容分为四个相互衔接的模块：其一，数据构建与质量控制，围绕题型、推理步骤、答案和来源进行双人交叉校验；其二，LoRA 参数高效微调与基线对比，比较原始 Qwen、Qwen+RAG、LoRA 模型、LoRA+RAG 以及完整多智能体系统的差异；其三，多智能体协同机制与统一 API 集成，明确每个智能体的输入、输出、失败状态和回溯规则；其四，三端产品与科研入口验证，观察系统在课程资料、练习任务、学情分析和科研记录中的可用性。")
    add_heading(doc,"4.3 评价与验证框架",2)
    add_table(doc,"表3  研究评价维度与拟采用方法",["维度","评价问题","拟采用方法","当前状态"], [
        ["答案正确性","结论与计算是否正确","独立测试集；与多组基线对比","待实施，不预设结果"],
        ["推导完整性","步骤是否连续、可理解、可复核","教师5级盲评；一致性检验","待实施，不预设结果"],
        ["方法一致性","是否符合课程与教师约定方法","教师复核与差异归因","待实施，不预设结果"],
        ["检索可靠性","证据是否相关、可定位、版本有效","命中率、来源核对、版本审计","需随知识库建设验证"],
        ["产品可用性","三端与科研入口是否支持核心任务","15—20名学生、2—3名教师试用记录","计划开展"],
        ["安全与治理","是否存在越权、泄露或不可追溯行为","权限测试、日志审计、异常演练","设计阶段"]], [2.5,4.3,5.0,3.2])
    add_heading(doc,"4.4 实施边界与风险控制",2)
    add_body(doc,"本阶段研究不把尚未完成的真实模型训练、完整智能体部署、用户满意度和实验准确率写成既成事实。模型输出必须展示证据来源与版本，知识库内容必须经过授权和审核，个人画像采用最小化字段并限制用途。对于高风险或多次质检失败的任务，系统停止自动化链路并转人工；对于科研代码执行，采用沙箱和资源配额；对于模型迭代，保留独立测试集和旧版本，以防止反馈回流造成性能退化。")
    add_heading(doc,"4.5 本章小结",2)
    add_body(doc,"本章将 CodeTrack 的产品原型、项目申报方案和垂类大模型研究路线统一为一套可执行的研究需求与总体方案。平台以四类入口满足教学、学习、治理和科研任务，以分层架构复用模型、知识与审计能力，以“分类—检索—拆解—求解—质检”链路承载学科推理增强，并以人工确认、版本管理和闭环评价确保系统边界清晰。后续工作应按照本章提出的数据规范、接口契约和评价框架开展实现与验证。")

    add_heading(doc,"参考文献",1)
    refs=[
        "[1] 贾万林. 面向学科生态的垂直大模型构建与自主化学科智能体集成方案研究：项目申报书[R]. 信阳学院, 2026.",
        "[2] CodeTrack. 教学工作台、学生学习视角、管理员治理与科研入口 UI 原型[Z]. 2026.",
        "[3] Hu E J, Shen Y, Wallis P, et al. LoRA: Low-Rank Adaptation of Large Language Models[C]// ICLR. 2022.",
        "[4] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]// NeurIPS. 2020.",
        "[5] Qwen Team. Qwen2.5 Technical Report[R/OL]. 2024.",
        "[6] 刘雪颖, 云静, 李博, 等. 基于大型语言模型的检索增强生成综述[J]. 计算机工程与应用, 2025, 61(13): 1-25.",
        "[7] 赵鑫, 窦志成, 文继荣, 等. 大语言模型智能体综述[J/OL]. 计算机学报, 2025.",
    ]
    for ref in refs:
        p=doc.add_paragraph(); para_format(p,indent=False,align=WD_ALIGN_PARAGRAPH.LEFT,after=2)
        r=p.add_run(ref); set_run_font(r,east_asia="宋体",size=10.5)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__": build()
