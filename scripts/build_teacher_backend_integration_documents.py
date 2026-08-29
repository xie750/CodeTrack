from __future__ import annotations

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_database_integration_document import (
    GREEN, DARK_GREEN, INK, MUTED, LIGHT_GREEN, LIGHT_GRAY,
    add_callout, add_field, add_para, add_table, audit_document,
    configure_section, configure_styles, set_run_font,
    set_table_geometry, style_cell,
)
from build_database_enterprise_document import add_code_block, add_list_item
from build_teacher_student_backend_document import add_numbering_definition


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / 'artifacts' / 'deliverables'
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
CONNECTED_PATH = OUTPUT_DIR / 'CodeTrack教师端已接入后端功能说明_企业版.docx'
RESERVED_PATH = OUTPUT_DIR / 'CodeTrack教师端暂未接入后端与预留接口说明_企业版.docx'
BASELINE = '2026-08-16'


REFERENCES = [
    ('R1', 'CodeTrack教师端详细操作手册_实测版.docx', '操作步骤、页面名称与实测口径'),
    ('R2', 'CodeTrack教师端前端逻辑交互文档.docx', '按钮、路由、抽屉和弹窗交互口径'),
    ('R3', 'CodeTrack教师端数据库内容与前端接入说明_企业版.docx', '表结构、持久化边界与字段来源'),
    ('R4', 'CodeTrack教师端与学生端协同API接口文档_企业版.docx', '跨端契约，仅作为预留基线'),
    ('R5', 'CodeTrack教师端后端API接口文档_企业版.docx', '教师端接口、错误码与数据模型基线'),
]


CONNECTED_MODULES = [
    ('认证与启动', '登录页、应用启动、通知中心', 'GET /teacher/auth/accounts\nPOST /teacher/auth/login\nGET /teacher/bootstrap\nPATCH /teacher/notifications/{id}', 'users, teacher_credentials, notifications', '已接入'),
    ('课程创建草稿', '创建课程向导、恢复/放弃草稿', 'GET/PUT/DELETE /teacher/course-draft', 'course_drafts', '已接入'),
    ('课程与设置', '课程列表、新建、归档、恢复、删除、课程设置', 'POST/PATCH/DELETE /teacher/courses...', 'courses, class_groups, chapters', '已接入'),
    ('公告与偏好', '课程公告列表/已读、个人通知与 AI 设置', 'GET /teacher/courses/{id}/announcements\nPATCH /teacher/announcements/{id}/read\nGET/PUT /teacher/preferences', 'course_announcements, announcement_reads, teacher_preferences', '已接入'),
    ('班级与成员', '班级列表、新建、邀请码、CSV 导入、花名册、加入状态', 'GET/POST /teacher/classes...\nPOST .../join-code\nPOST .../students/import\nGET .../students|join-status', 'class_groups, users, enrollments', '已接入'),
    ('章节与知识点', '章节目录、新建、教学方式、发布/撤回、知识点新增', 'GET/POST /teacher/courses/{id}/chapters\nPATCH /teacher/chapters/{id}\nPOST /teacher/knowledge-points', 'chapters, knowledge_points', '已接入'),
    ('任务管理', '任务列表、AI 草稿、创建、发布', 'GET/POST /teacher/tasks\nPOST /teacher/tasks/ai-draft\nPOST /teacher/tasks/{id}/publish', 'tasks, test_cases', '已接入'),
    ('资料库', '列表、上传、链接、可见性、回收站、恢复、图谱关联', 'GET/POST/PATCH/DELETE /teacher/materials...\nPOST /teacher/materials/upload', 'materials, material_knowledge_links', '已接入'),
    ('资料文件夹', '目录、目录回收站、恢复、保留资料、AI 目录建议', 'GET/POST/DELETE /teacher/material-folders...', 'material_folders, materials', '已接入'),
    ('教师画布图谱', '图谱列表、文件生成、空白创建、编辑保存、发布、删除', 'GET/POST/PUT/DELETE /teacher/knowledge-graphs...', 'teacher_knowledge_graphs', '已接入'),
    ('任务监控与批改', '提交列表、成绩保存/发布、反馈、AI 诊断审核', 'GET /teacher/submissions\nPUT/POST .../grade...\nPOST .../feedback\nGET/POST /teacher/ai-reviews...', 'submissions, grades, teacher_feedback, diagnosis_reviews', '已接入'),
    ('学情与讨论', '学情概览、学生诊断、讨论创建/发布/结束', 'GET /teacher/analytics/overview\nGET/POST /teacher/discussions...', 'submissions, grades, course_discussions, discussion_replies', '已接入'),
]


ACTIVE_ENDPOINTS = [
    ('认证/启动', 'GET', '/teacher/auth/accounts', '登录页加载演示教师账号', 'teacherAccounts'),
    ('认证/启动', 'POST', '/teacher/auth/login', '校验账号密码并进入门户', 'teacherLogin'),
    ('认证/启动', 'GET', '/teacher/bootstrap', '聚合教师、课程、班级、通知', 'bootstrap'),
    ('认证/启动', 'PATCH', '/teacher/notifications/{notification_id}', '通知标记已读', 'markNotification'),
    ('草稿', 'GET', '/teacher/course-draft', '恢复课程创建草稿', 'courseDraft'),
    ('草稿', 'PUT', '/teacher/course-draft', '保存课程创建草稿', 'saveCourseDraft'),
    ('草稿', 'DELETE', '/teacher/course-draft', '放弃课程创建草稿', 'deleteCourseDraft'),
    ('公告/偏好', 'GET', '/teacher/courses/{course_id}/announcements', '课程工作台公告列表', 'announcements'),
    ('公告/偏好', 'PATCH', '/teacher/announcements/{announcement_id}/read', '公告已读', 'markAnnouncementRead'),
    ('公告/偏好', 'GET', '/teacher/preferences', '读取个人设置', 'preferences'),
    ('公告/偏好', 'PUT', '/teacher/preferences', '保存个人设置', 'savePreferences'),
    ('工作台', 'GET', '/teacher/dashboard', '指标、任务、风险与动态', 'dashboard'),
    ('课程', 'POST', '/teacher/courses', '创建课程', 'createCourse'),
    ('课程', 'PATCH', '/teacher/courses/{course_id}', '课程设置、归档和恢复', 'updateCourse'),
    ('课程', 'DELETE', '/teacher/courses/{course_id}', '删除允许删除的课程', 'deleteCourse'),
    ('班级', 'GET', '/teacher/classes', '按课程读取班级', 'classes'),
    ('班级', 'POST', '/teacher/classes', '创建教学班', 'createClass'),
    ('班级', 'POST', '/teacher/classes/{class_id}/join-code', '重新生成邀请码', 'regenerateJoinCode'),
    ('班级', 'POST', '/teacher/classes/{class_id}/students/import', 'CSV 批量导入学生', 'importStudents'),
    ('班级', 'GET', '/teacher/classes/{class_id}/students', '班级学生名单与摘要', 'students'),
    ('班级', 'GET', '/teacher/classes/{class_id}/join-status', '加入状态和容量汇总', 'classJoinStatus'),
    ('章节', 'GET', '/teacher/courses/{course_id}/chapters', '章节和知识点目录', 'chapters'),
    ('章节', 'POST', '/teacher/courses/{course_id}/chapters', '新增章节', 'createChapter'),
    ('章节', 'PATCH', '/teacher/chapters/{chapter_id}', '编辑、发布或撤回章节', 'updateChapter'),
    ('章节', 'POST', '/teacher/knowledge-points', '新增知识点', 'createKnowledgePoint'),
    ('任务', 'GET', '/teacher/tasks', '任务列表', 'tasks'),
    ('任务', 'POST', '/teacher/tasks', '创建任务草稿', 'createTask'),
    ('任务', 'POST', '/teacher/tasks/ai-draft', '生成 AI/规则任务草稿', 'aiTaskDraft'),
    ('任务', 'POST', '/teacher/tasks/{task_id}/publish', '发布任务', 'publishTask'),
    ('资料', 'GET', '/teacher/materials', '资料列表', 'materials'),
    ('资料', 'POST', '/teacher/materials', '新增外部链接或资料元数据', 'createMaterial'),
    ('资料', 'POST', '/teacher/materials/upload', '上传资料文件', 'uploadMaterial'),
    ('资料', 'PATCH', '/teacher/materials/{material_id}', '可见性和状态更新', 'updateMaterial'),
    ('资料', 'DELETE', '/teacher/materials/{material_id}', '资料移入回收站', 'trashMaterial/deleteMaterial'),
    ('资料', 'GET', '/teacher/materials/trash', '读取资料回收站', 'trashMaterials'),
    ('资料', 'POST', '/teacher/materials/{material_id}/restore', '恢复资料', 'restoreMaterial'),
    ('资料', 'POST', '/teacher/materials/{material_id}/knowledge-graph', '资料关联知识点', 'importMaterialToGraph'),
    ('资料目录', 'GET', '/teacher/material-folders', '读取有效目录', 'materialFolders'),
    ('资料目录', 'GET', '/teacher/material-folders/trash', '读取目录回收站', 'trashedMaterialFolders'),
    ('资料目录', 'POST', '/teacher/material-folders', '创建目录', 'createMaterialFolder'),
    ('资料目录', 'DELETE', '/teacher/material-folders/{folder_id}', '目录移入回收站', 'deleteMaterialFolder'),
    ('资料目录', 'POST', '/teacher/material-folders/{folder_id}/restore', '恢复目录', 'restoreMaterialFolder'),
    ('资料目录', 'POST', '/teacher/material-folders/{folder_id}/materials/{material_id}/keep', '保留或迁移资料', 'keepDeletedFolderMaterial'),
    ('资料目录', 'POST', '/teacher/material-folders/ai-outline', '生成目录候选', 'aiMaterialOutline'),
    ('资料目录', 'POST', '/teacher/material-folders/confirm-outline', '确认目录候选', 'confirmMaterialOutline'),
    ('课程图谱', 'GET', '/teacher/knowledge-graph', '资料关联知识点时读取课程图谱', 'graph'),
    ('画布图谱', 'GET', '/teacher/knowledge-graphs', '教师图谱列表', 'teacherGraphs'),
    ('画布图谱', 'GET', '/teacher/knowledge-graphs/{graph_id}', '打开图谱', 'teacherGraph'),
    ('画布图谱', 'POST', '/teacher/knowledge-graphs', '创建空白图谱', 'createTeacherGraph'),
    ('画布图谱', 'POST', '/teacher/knowledge-graphs/from-files', '从文件生成图谱', 'createTeacherGraphFromFiles'),
    ('画布图谱', 'PUT', '/teacher/knowledge-graphs/{graph_id}', '保存节点与关系', 'saveTeacherGraph'),
    ('画布图谱', 'POST', '/teacher/knowledge-graphs/{graph_id}/publish', '发布图谱', 'publishTeacherGraph'),
    ('画布图谱', 'DELETE', '/teacher/knowledge-graphs/{graph_id}', '删除图谱', 'deleteTeacherGraph'),
    ('讨论', 'GET', '/teacher/discussions', '讨论列表和回复监看', 'discussions'),
    ('讨论', 'POST', '/teacher/discussions', '创建讨论', 'createDiscussion'),
    ('讨论', 'POST', '/teacher/discussions/{discussion_id}/publish', '发布讨论', 'publishDiscussion'),
    ('讨论', 'POST', '/teacher/discussions/{discussion_id}/end', '结束讨论', 'endDiscussion'),
    ('监控/批改', 'GET', '/teacher/submissions', '按任务读取提交', 'submissions'),
    ('监控/批改', 'PUT', '/teacher/submissions/{submission_id}/grade', '保存评分', 'saveGrade'),
    ('监控/批改', 'POST', '/teacher/submissions/{submission_id}/grade/publish', '发布成绩', 'publishGrade'),
    ('监控/批改', 'POST', '/teacher/submissions/{submission_id}/feedback', '保存或发布反馈', 'feedback'),
    ('AI 审核', 'GET', '/teacher/ai-reviews', '读取待审诊断', 'reviews'),
    ('AI 审核', 'POST', '/teacher/ai-reviews/{review_id}/action', '接受、修改或拒绝诊断', 'reviewAction'),
    ('学情', 'GET', '/teacher/analytics/overview', '班级和个人学情分析', 'analytics'),
]


TEST_CASES = [
    ('T01', '教师账号加载与正确密码登录', '登录页选择教师并使用 123456', '进入教师门户；身份写入当前会话', '通过'),
    ('T02', '错误密码', '提交错误密码', '返回 401；前端保留登录页并显示错误', '后端覆盖'),
    ('T03', '课程创建草稿', '填写向导并保存/刷新', '草稿按教师恢复；完成创建后删除', '通过'),
    ('T04', '课程创建及初始班级', '完成三步向导', 'courses、class_groups、chapters 持久化', '通过'),
    ('T05', '课程归档/恢复/删除保护', '执行状态菜单', '状态落库；存在关联数据时按规则拒绝删除', '通过'),
    ('T06', '班级新建与邀请码刷新', '新建班级并重新生成邀请码', '班级落库；返回 8 位新邀请码', '通过'),
    ('T07', 'CSV 批量导入', '导入姓名/学号 CSV，包含重复学号', '新增/加入/跳过计数正确；花名册刷新', '通过'),
    ('T08', '章节发布与撤回', '切换章节发布状态', '章节状态和关联资料可见性按规则更新', '通过'),
    ('T09', '教师本地学生视角预览', '点击“查看学生视角”', '只组合教师端已加载数据，不请求 /student/*', '通过'),
    ('T10', '任务创建、AI 草稿和发布', '创建任务并发布', '任务、测试用例、发布时间持久化', '通过'),
    ('T11', '资料上传、回收站与恢复', '上传文件后删除并恢复', '元数据与文件 URL 正确；状态可逆', '通过'),
    ('T12', '资料目录删除与保留资料', '删除含资料目录并选择保留', '目录/资料软删除或迁移结果正确', '通过'),
    ('T13', '教师画布图谱', '文件生成、编辑节点关系、保存、发布', '图谱 JSON、来源文件、状态持久化', '通过'),
    ('T14', '评分与反馈', '修改维度分、保存草稿、发布', 'grades 与 teacher_feedback 更新；发布状态正确', '通过'),
    ('T15', '课堂讨论', '创建、发布、结束讨论', '状态机 draft → published → ended', '通过'),
    ('T16', '学生入口隔离', '访问 /join/{code} 与 /student/discussions', '显示预留页；跨端 API 请求数为 0', '通过'),
    ('T17', '生产构建', 'TypeScript 编译与 Vite build', '构建完成；仅有 bundle 体积提示', '通过'),
    ('T18', '后端回归', 'pytest teacher_backend/tests', '21 passed；无失败用例', '通过'),
]


RESERVED_CROSS_END = [
    ('RS-01', 'POST', '/classes/{join_code}/join', '学生申请加入教师创建的班级', '入口保留；前端不发送请求', '学生认证、幂等、审批策略'),
    ('RS-02', 'GET', '/student/courses/{course_id}/content', '学生读取教师已发布章节/资料/练习', '契约保留；教师预览改为本地组合', '学生课程权限、缓存、内容版本'),
    ('RS-03', 'GET', '/student/tasks', '学生读取可见任务列表', '契约保留；教师端无调用', '班级范围、发布时间、分页'),
    ('RS-04', 'POST', '/student/tasks/{task_id}/submissions', '学生提交代码或作业', '契约保留；模拟提交已禁用', '鉴权、限流、评测队列、幂等键'),
    ('RS-05', 'GET', '/student/discussions', '学生读取已发布讨论', '路由保留；打开后显示未启用页', '课程权限、分页、敏感内容治理'),
    ('RS-06', 'POST', '/student/discussions/{discussion_id}/replies', '学生回复讨论', '契约保留；前端不发送请求', '鉴权、限流、审核、撤回机制'),
]


UNUSED_TEACHER_APIS = [
    ('UT-01', 'GET /health', '服务已实现，教师 SPA 未直接调用', '运维探针，不属于业务页面', 'P2', '由部署平台/网关探活'),
    ('UT-02', 'GET /teacher/courses', '客户端方法存在但页面使用 bootstrap.courses', '避免重复请求', 'P3', '保留为独立刷新或分页接口'),
    ('UT-03', 'POST /teacher/knowledge-graph/ai-candidates', '后端已实现，活动图谱页未调用', '活动页使用独立 teacher_knowledge_graphs 画布体系', 'P2', '产品确认两套图谱是否合并'),
    ('UT-04', 'POST /teacher/knowledge-graph/confirm', '后端已实现，活动图谱页未调用', '同上', 'P2', '合并后再接候选确认流程'),
    ('UT-05', 'PUT /teacher/knowledge-graph/nodes/{node_id}', '后端已实现，活动图谱页未调用', '节点编辑走画布整体 PUT', 'P2', '确定单节点与整体保存策略'),
    ('UT-06', 'GET /teacher/submissions/{submission_id}', '后端已实现，前端列表响应已含详情', '当前无需二次请求', 'P3', '拆分轻量列表后启用详情接口'),
]


MISSING_CAPABILITIES = [
    ('班级管理', '编辑/删除班级', '“管理班级”“编辑”目前只打开详情或无提交', '缺 PATCH/DELETE class API', 'P1'),
    ('成员管理', '学生详情、移除、待审核通过/拒绝', '只读花名册和加入状态', '缺成员详情与 enrollment 状态变更 API', 'P1'),
    ('邀请策略', '链接/二维码/邀请码开关持久化', '开关仅为页面状态', '缺邀请方式策略字段/API', 'P2'),
    ('课程公告', '新增、编辑、删除、置顶公告', '当前仅查询和标记已读', '缺公告写接口', 'P1'),
    ('任务管理', '编辑、删除、关闭任务；真实分页', '创建/发布可用，其余部分按钮未落库', '缺 PATCH/DELETE/close 与分页契约', 'P1'),
    ('测试用例', '增删改测试用例', '创建时可随任务提交，发布后缺独立维护', '缺 test-case CRUD', 'P2'),
    ('监控导出', '导出筛选后的提交数据', '按钮仅提示“已导出”', '缺导出任务/文件接口', 'P1'),
    ('学情导出', '导出分析报告', '按钮仅提示“已导出”', '缺异步报表生成与下载接口', 'P1'),
    ('讨论管理', '编辑/删除讨论、教师回复', '仅创建、发布、结束', '缺更新、删除、教师回复 API', 'P2'),
    ('通知中心', '独立通知列表、分页、全部已读', '通知随 bootstrap 返回，只能单条已读', '缺通知列表与批量操作 API', 'P2'),
    ('资料预览', '统一的文件鉴权下载/在线预览', '已返回 content_url，部分按钮未绑定', '缺短期签名 URL/预览转换服务', 'P1'),
    ('身份安全', 'JWT/Session、刷新、注销、CSRF/权限声明', '当前请求信任 X-User-Id', '缺正式认证授权体系', 'P0'),
    ('计划发布', '到点自动发布 scheduled 任务', '可保存 scheduled，无调度器执行', '缺任务调度与失败重试', 'P0'),
    ('AI 能力', '真实模型服务、队列、配额与审计', '部分 AI 接口为规则回退', '缺模型网关、作业状态和成本治理', 'P1'),
    ('图谱治理', '课程图谱与教师画布图谱统一', '两套模型并存', '缺统一标识、迁移和发布投影', 'P1'),
]


def set_header_footer(section, right_title, doc_id):
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = 'CodeTrack 教师端'
    table.cell(0, 1).text = right_title
    set_table_geometry(table, [3000, 6360], indent_dxa=0)
    style_cell(table.cell(0, 0), bold=True, color=GREEN, size=8.5)
    style_cell(table.cell(0, 1), color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = f'{doc_id}  |  基线 {BASELINE}'
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.add_run('第 '), size=8.5, color=MUTED)
    add_field(right, 'PAGE')
    set_run_font(right.add_run(' 页 / 共 '), size=8.5, color=MUTED)
    add_field(right, 'NUMPAGES')
    set_run_font(right.add_run(' 页'), size=8.5, color=MUTED)
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    style_cell(table.cell(0, 0), size=8.5, color=MUTED)


def setup_doc(title, subject, doc_id, right_title):
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        set_header_footer(section, right_title, doc_id)
    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = 'CodeTrack 项目组'
    doc.core_properties.keywords = 'CodeTrack, React, FastAPI, API, 后端接入, 企业开发'
    return doc


def add_cover(doc, title, subtitle, doc_id, status):
    add_para(doc, 'ENGINEERING DELIVERY / INTERNAL', size=10, bold=True, color=GREEN, before=18, after=12)
    add_para(doc, 'CodeTrack 教学平台', size=14, bold=True, color=DARK_GREEN, after=5)
    add_para(doc, title, size=24, bold=True, color=INK, after=8)
    add_para(doc, subtitle, size=13, color=MUTED, after=24)
    rows = [
        ('文档编号', doc_id), ('版本 / 状态', f'V1.0 / {status}'),
        ('系统范围', 'CodeTrack 教师端 SPA、FastAPI 后端与 SQLite 数据库'),
        ('技术基线', 'React 19 / TypeScript / Vite / FastAPI / SQLAlchemy'),
        ('编制日期', BASELINE), ('编制角色', '产品、前端、后端、测试、架构与安全'),
        ('密级', '内部使用'),
    ]
    add_table(doc, ['元数据', '内容'], rows, [2100, 7260], font_size=9.5)
    add_callout(doc, '阅读结论', '本文按当前代码和实测结果描述，不把“后端路由已存在”等同于“前端已经接入”，也不把教师发布动作误判为学生端已经上线。')
    doc.add_page_break()


def add_sources(doc):
    doc.add_heading('1. 文档范围与判定口径', level=1)
    add_para(doc, '本次分析采用“四份既有业务/技术文档 + 最新教师端 API 文档”的组合基线，并以当前源代码和测试结果作为最终事实来源。', after=8)
    add_table(doc, ['编号', '输入文档', '本次用途'], REFERENCES, [900, 5000, 3460], font_size=9)
    doc.add_heading('1.1 状态定义', level=2)
    bullet = add_numbering_definition(doc, bullet=True)
    for text in [
        '已接入：活动页面存在真实调用，后端返回结果会进入页面状态或持久化流程，并纳入本次验证。',
        '接口已实现但未接入：后端路由和前端客户端方法存在，但活动页面没有调用点。',
        '预留接口：仅保存跨端契约、路由占位或元数据，不在当前教师端发起网络请求。',
        '缺少后端能力：界面存在按钮或业务预期，但没有满足该动作的教师端 API。',
    ]:
        add_list_item(doc, text, bullet)


def build_connected():
    doc = setup_doc('CodeTrack教师端已接入后端功能说明', '教师端前后端接入范围、数据映射、测试与运行说明', 'CT-INT-TCH-001', '已接入后端功能说明')
    add_cover(doc, '教师端已接入后端功能说明', '前端交互、API、数据库、验证证据与运维边界', 'CT-INT-TCH-001', '实施完成版')
    add_sources(doc)

    doc.add_heading('2. 接入结果摘要', level=1)
    add_callout(doc, '当前结论', '教师端 API 客户端共有 71 项教师侧能力，其中 65 项已在活动页面形成调用；6 项后端能力当前未被活动页面调用。学生侧 6 个接口均未接入教师端运行链路。')
    add_table(doc, ['指标', '结果', '说明'], [
        ('已接入教师侧客户端能力', '65', '包含 JSON 请求与 2 个 multipart 文件方法'),
        ('未调用教师侧能力', '6', '健康检查、冗余课程列表、3 个旧课程图谱接口、提交详情'),
        ('跨学生端运行时调用', '0', '浏览器监听未捕获 /student/* 或加入班级 POST'),
        ('后端自动化测试', '21/21 通过', 'pytest teacher_backend/tests -q'),
        ('前端生产构建', '通过', 'TypeScript + Vite；仅有 bundle 体积提示'),
        ('浏览器关键路径', '通过', '登录、章节预览、邀请页、预留学生入口'),
    ], [2800, 1800, 4760], font_size=9.2)

    doc.add_heading('2.1 系统边界', level=2)
    add_code_block(doc, '教师浏览器\n  -> React/TypeScript 页面与 src/api.ts\n  -> /api/v1/teacher/*（当前启用）\n  -> FastAPI + SQLAlchemy\n  -> teacher_backend/codetrack.db + teacher_backend/uploads\n\n/student/* 与 /classes/{join_code}/join\n  -> 仅保留契约，不进入当前教师端网络链路')

    doc.add_heading('3. 已接入功能矩阵', level=1)
    add_para(doc, '下表按教师工作流汇总活动页面、API 和持久化对象。每一行均要求“按钮或页面动作 → API → 数据返回/落库”成立。')
    add_table(doc, ['模块', '活动页面/动作', '主要接口', '主要数据对象', '状态'], CONNECTED_MODULES, [1350, 2360, 2520, 2230, 900], font_size=8.1)

    doc.add_heading('4. 前端调用与接口明细', level=1)
    add_para(doc, '所有路径均省略统一前缀 /api/v1。文件上传使用 multipart/form-data，其余写请求使用 application/json。')
    for module in []:
        pass
    chunks = [ACTIVE_ENDPOINTS[i:i + 13] for i in range(0, len(ACTIVE_ENDPOINTS), 13)]
    for index, chunk in enumerate(chunks, 1):
        if index > 1:
            doc.add_page_break()
        doc.add_heading(f'4.{index} 接口清单（{index}/{len(chunks)}）', level=2)
        add_table(doc, ['域', '方法', '路径', '页面动作/数据用途', '客户端方法'], chunk, [1100, 750, 2920, 2990, 1600], font_size=7.8)

    doc.add_heading('5. 关键业务闭环', level=1)
    workflows = [
        ('5.1 教师登录与启动', ['登录页读取教师账号。', '提交账号密码，后端校验 PBKDF2 凭据。', '登录成功后前端保存本地教师会话标识。', '进入工作台后请求 bootstrap，并按 X-User-Id 做教师数据范围过滤。']),
        ('5.2 课程创建', ['创建向导按教师保存草稿。', '确认创建时写入课程、初始章节和教学设置。', '继续创建初始班级并生成邀请码。', '成功后删除草稿并导航至邀请页。']),
        ('5.3 班级 CSV 导入', ['教师下载 CSV 模板，填写“姓名,学号”。', '前端解析 UTF-8 CSV，并过滤空行和无效行。', 'POST 学生数组到教师端导入接口。', '后端按学号创建学生、建立 enrollment、返回 created/enrolled/skipped/errors。', '前端重新读取花名册、加入状态和 bootstrap。']),
        ('5.4 任务与批改', ['教师创建任务和测试用例草稿。', '发布到指定班级，后端更新状态和可见时间。', '教师端监控读取 submissions。', '保存维度评分与反馈；发布成绩后更新状态。']),
        ('5.5 资料与图谱', ['上传资料或新增外部链接。', '资料可进入文件夹、回收站并恢复。', '资料可关联课程知识点。', '教师画布图谱支持文件生成、节点关系编辑、整体保存与发布。']),
        ('5.6 课堂讨论', ['教师创建讨论草稿或立即发布。', '教师端轮询/刷新读取回复汇总。', '教师结束讨论后状态变为 ended。', '学生读取与回复接口保留但当前不启用。']),
    ]
    for title, steps in workflows:
        doc.add_heading(title, level=2)
        num = add_numbering_definition(doc)
        for step in steps:
            add_list_item(doc, step, num)

    doc.add_heading('6. 数据落库与一致性', level=1)
    add_table(doc, ['业务对象', '主要表/存储', '写入来源', '一致性规则'], [
        ('教师与认证', 'users, teacher_credentials', '账号种子与登录校验', '账号必须为 teacher；当前不签发 token'),
        ('教师配置', 'teacher_preferences, course_drafts', '个人设置、创建课程向导', '按 teacher_id 唯一覆盖'),
        ('课程组织', 'courses, class_groups, chapters', '课程/班级/章节页面', '均校验课程归属'),
        ('成员', 'users, enrollments', 'CSV 导入和加入状态', '学号去重；班级-学生唯一'),
        ('任务评测', 'tasks, test_cases, submissions, evaluation_results', '任务发布与提交结果', '隐藏测试只在教师侧展示'),
        ('批改', 'grades, teacher_feedback, diagnosis_reviews', '评分、反馈、AI 审核', '成绩发布前为草稿'),
        ('资料', 'materials, material_folders, teacher_backend/uploads', '上传、链接、目录、回收站', '软删除与恢复保留审计语义'),
        ('图谱', 'knowledge_points + teacher_knowledge_graphs', '课程图谱关联 + 教师画布编辑', '当前为两套模型，不能混用 ID'),
        ('讨论/公告', 'course_discussions, discussion_replies, course_announcements', '教师讨论动作、公告读取', '讨论遵循状态机；公告目前只读'),
    ], [1500, 2500, 2400, 2960], font_size=8.6)

    doc.add_heading('7. 前端实现位置', level=1)
    add_table(doc, ['文件', '职责', '本次边界'], [
        ('src/api.ts', '统一教师端 API 客户端、类型和 X-User-Id 请求头', '已删除可执行学生端方法'),
        ('src/exact/ExactApp.tsx', '路由装配、bootstrap、通知中心、课程/班级上下文', '仅装配教师端活动页面'),
        ('src/exact/pages-global.tsx', '登录、工作台、课程列表、创建课程', '教师端 API 已接入'),
        ('src/exact/pages-course.tsx', '课程工作空间、邀请、设置和讨论', 'CSV 导入已接入；学生提交模拟已禁用'),
        ('src/exact/ExactCourseContent.tsx', '章节、知识点、资料和练习', '学生视角由教师数据本地组合'),
        ('src/exact/ExactClassesV2.tsx', '班级详情、学生列表和加入状态', '只读成员管理已接入'),
        ('src/exact/ExactTasksV2.tsx', '任务创建、AI 草稿、发布和导航', '教师任务接口已接入'),
        ('src/exact/ExactMaterialsV2.tsx', '资料库、目录、回收站、AI 目录和图谱关联', '教师资料接口已接入'),
        ('src/exact/ExactGraphV2.tsx', '独立教师画布图谱', 'teacher_knowledge_graphs 全流程'),
        ('src/exact/pages-flow.tsx', '监控、批改、学情、AI 审核、设置', '教师查询和写接口已接入'),
        ('src/reserved/studentIntegration.ts', '师生协同契约元数据', 'metadata-only，不执行 fetch'),
        ('src/main.tsx', '应用入口和教师本地会话', '学生路由只展示预留提示页'),
    ], [3100, 3500, 2760], font_size=8.5)

    doc.add_heading('8. 验证记录', level=1)
    add_table(doc, ['用例', '验证目标', '操作/输入', '预期与结果', '结论'], TEST_CASES, [700, 1800, 2500, 3560, 800], font_size=8.0)
    add_callout(doc, '实测证据', '2026-08-16：前端生产构建成功；后端 21 个测试全部通过；Playwright 浏览器检查显示 localPreviewVisible=true、csvButtonVisible=true、reservedPageVisible=true、crossEndRequests=[]、apiErrors=[]。')

    doc.add_heading('9. 配置、启动与排障', level=1)
    doc.add_heading('9.1 本地启动', level=2)
    add_code_block(doc, '前端：node node_modules/vite/bin/vite.js --host 127.0.0.1 --port 5173\n后端：python -m uvicorn teacher_backend.app.main:app --host 127.0.0.1 --port 8001\n健康检查：http://127.0.0.1:8001/api/v1/health')
    doc.add_heading('9.2 前端配置', level=2)
    add_table(doc, ['配置项', '默认值', '说明'], [
        ('VITE_API_BASE', '/api/v1', '生产环境应指向网关统一前缀'),
        ('X-User-Id', '当前登录教师 ID', '原型身份头；正式环境必须由 token 替代'),
        ('上传目录', 'teacher_backend/uploads', '生产环境应迁移到对象存储并使用签名 URL'),
        ('数据库', 'teacher_backend/codetrack.db', 'SQLite 适合本地/演示，不作为高并发生产方案'),
    ], [1900, 2500, 4960], font_size=9)
    doc.add_heading('9.3 常见故障', level=2)
    add_table(doc, ['现象', '检查项', '处置'], [
        ('页面提示无法连接', '8001 端口、/health、Vite proxy', '先恢复后端，再刷新教师端'),
        ('403 Teacher permission required', 'X-User-Id 对应 users.role', '重新登录并核对教师账号'),
        ('课程/班级 404', 'URL 中 courseId/classId 与教师归属', '回到课程列表重新选择'),
        ('CSV 导入 422', '表头、姓名、学号长度和空行', '使用下载模板并确保 UTF-8 CSV'),
        ('上传 413/415', '文件大小与扩展名白名单', '调整文件或后端限制'),
        ('图谱发布 409', '图谱是否为空', '至少创建一个节点后发布'),
    ], [2200, 2900, 4260], font_size=8.8)

    doc.add_heading('10. 已知限制与上线门禁', level=1)
    bullet = add_numbering_definition(doc, bullet=True)
    for text in [
        '认证仍是原型方案：登录成功不签发 JWT/Session，业务请求信任 X-User-Id；正式上线前必须替换。',
        '学情、班级趋势等部分指标包含演示/聚合口径，上线前需要确认指标定义和数据刷新策略。',
        'scheduled 任务缺少自动调度器，不能承诺到点自动发布。',
        '课程知识图谱与教师画布图谱为两套持久化模型，发布对象和节点 ID 不互通。',
        '导出、班级编辑、成员审批等能力仍属于未接入或缺后端范围，详见配套文档。',
        '学生端契约虽在后端存在，但当前教师端没有任何运行时调用；启用必须走独立联调与安全验收。',
    ]:
        add_list_item(doc, text, bullet)

    doc.add_heading('11. 交付验收签字项', level=1)
    add_table(doc, ['角色', '验收重点', '结论/签字'], [
        ('产品负责人', '已接入范围与按钮行为符合需求', '________________'),
        ('前端负责人', '活动页面无学生端运行时调用', '________________'),
        ('后端负责人', '接口、权限、事务与错误码符合基线', '________________'),
        ('测试负责人', '自动化与关键路径证据完整', '________________'),
        ('安全负责人', '原型认证风险已接受或完成整改', '________________'),
    ], [1600, 5000, 2760], font_size=9.2)
    add_para(doc, f'文档生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT, before=12)
    audit_document(doc)
    doc.save(CONNECTED_PATH)


def build_reserved():
    doc = setup_doc('CodeTrack教师端暂未接入后端与预留接口说明', '未接入功能、跨学生端预留契约、风险与后续实施顺序', 'CT-GAP-TCH-001', '暂未接入与预留接口')
    add_cover(doc, '教师端暂未接入后端与预留接口说明', '跨学生端隔离、教师端缺口、优先级与启用门禁', 'CT-GAP-TCH-001', '范围冻结版')
    add_sources(doc)

    doc.add_heading('2. 范围结论', level=1)
    add_callout(doc, '当前发布边界', '教师端自身后端继续启用；所有学生侧接口暂不接入。现有学生 URL 只显示“学生端暂未启用”提示，接口契约集中保存在 metadata-only 模块中，不允许页面直接 fetch。')
    add_table(doc, ['分类', '数量', '处理结论'], [
        ('师生协同预留接口', '6', '保留方法、路径和消费者说明；当前不执行'),
        ('后端已存在但活动页面未调用', '6', '按冗余、待产品决策或未来优化分类'),
        ('界面/业务需要但缺教师后端能力', '15 类', '纳入后续版本，不伪装为已完成'),
        ('浏览器跨端请求', '0', '关键路径实测未捕获'),
    ], [3000, 1400, 4960], font_size=9.2)

    doc.add_heading('3. 师生协同预留接口', level=1)
    add_para(doc, '下列接口属于学生应用消费面。教师端可以产生邀请码、发布章节、任务或讨论，但“产生教师侧数据”不等于“学生端已经接入”。')
    add_table(doc, ['编号', '方法', '预留路径', '业务用途', '当前处理', '启用前置条件'], RESERVED_CROSS_END, [700, 650, 2550, 2000, 1700, 1760], font_size=7.7)

    doc.add_heading('3.1 代码预留方式', level=2)
    add_table(doc, ['位置', '实现', '禁止事项'], [
        ('src/reserved/studentIntegration.ts', '保存 6 条接口的 method/path/teacherTrigger/studentConsumer/status', '不得在该模块封装 fetch'),
        ('src/main.tsx', '/join/{code} 与 /student/discussions 显示预留页', '不得使用固定 student ID 调接口'),
        ('src/api.ts', '仅保留教师端可执行客户端方法', '不得重新加入 studentSubmit/joinClass 等方法'),
        ('ExactCourseContent.tsx', '学生视角由已加载的教师章节/资料/任务本地组合', '不得为预览请求 /student/courses/*'),
        ('旧任务页面', '模拟学生提交按钮禁用', '不得生成假学生提交数据'),
    ], [2600, 4300, 2460], font_size=8.6)
    add_code_block(doc, "export const reservedStudentIntegrations = [\n  { method: 'POST', path: '/api/v1/classes/{join_code}/join', status: 'reserved' },\n  { method: 'GET', path: '/api/v1/student/courses/{course_id}/content', status: 'reserved' },\n  // 其余契约同一模块维护；这里没有 request/fetch。\n]")

    doc.add_heading('4. 后端已实现但活动页面未调用', level=1)
    add_para(doc, '这些接口不属于“后端不存在”。它们暂时没有活动页面调用点，不能在已接入清单中计为完成。')
    add_table(doc, ['编号', '接口', '当前事实', '未接原因', '优先级', '建议'], UNUSED_TEACHER_APIS, [700, 2540, 2100, 1900, 700, 1420], font_size=7.8)

    doc.add_heading('5. 界面需要但缺少教师端后端能力', level=1)
    add_para(doc, '此清单用于避免“按钮能点/显示成功提示”被误认为已完成后端接入。优先级 P0 为上线阻断，P1 为核心业务，P2 为增强项。')
    add_table(doc, ['模块', '缺口功能', '当前界面事实', '后端缺口', '优先级'], MISSING_CAPABILITIES, [1200, 1800, 2600, 2960, 800], font_size=8.0)

    doc.add_heading('6. 风险分级与实施顺序', level=1)
    priorities = [
        ('阶段 0：安全与调度门禁', '正式 JWT/Session、权限声明、注销/刷新、scheduled 发布调度器', '没有这些能力不得面向真实师生发布'),
        ('阶段 1：教师核心闭环', '班级编辑、成员审批/移除、公告写入、任务编辑/关闭、真实导出', '消除当前主要无后端按钮'),
        ('阶段 2：内容与治理', '文件签名预览、讨论编辑删除、通知分页/全部已读、图谱统一', '提升可维护性与数据一致性'),
        ('阶段 3：学生端联调', '学生登录、入班、内容读取、任务提交、讨论回复', '独立环境联调，不直接在教师端开启'),
        ('阶段 4：生产化', '限流、审计、异步作业、对象存储、监控告警、压测和灾备', '完成企业上线验收'),
    ]
    add_table(doc, ['阶段', '实施内容', '完成标准'], priorities, [2200, 4300, 2860], font_size=8.7)

    doc.add_heading('7. 学生端启用门禁', level=1)
    gates = [
        ('身份与授权', '学生 token 可验证；接口不接受前端自报 student_id；课程/班级归属校验完备'),
        ('契约版本', 'OpenAPI 固化版本；字段、错误码、分页、时间格式和幂等策略评审通过'),
        ('隔离环境', '学生端在独立测试环境联调，不能直接使用教师演示数据库'),
        ('任务提交', '提交幂等键、频率限制、文件/源码大小限制、评测队列、超时和重试明确'),
        ('内容可见性', '章节/资料/任务只按 published、班级范围和发布时间返回'),
        ('讨论治理', '敏感词、举报、删除/撤回、教师管理、频控和审计明确'),
        ('隐私与日志', '最小化返回学号/姓名；日志脱敏；审计记录可追溯'),
        ('测试', '单元、契约、权限越权、并发、异常、回归和浏览器端到端测试全部通过'),
        ('灰度与回滚', '功能开关、灰度班级、指标监控、数据库兼容和回滚方案可执行'),
    ]
    add_table(doc, ['门禁项', '验收要求'], gates, [2200, 7160], font_size=9)

    doc.add_heading('8. 预留接口激活流程', level=1)
    num = add_numbering_definition(doc)
    for step in [
        '产品负责人确认学生端版本范围、班级加入策略与数据可见性。',
        '后端负责人将原型身份头替换为正式认证，并补齐权限、幂等、限流和审计。',
        '前后端共同冻结 OpenAPI 契约，生成或校验类型，不手工猜测字段。',
        '学生端在独立环境接入 RS-01 至 RS-06；教师端只保留教师动作和监控页面。',
        '测试执行跨角色越权、状态机、并发提交、重复请求、超时和回滚用例。',
        '通过功能开关按课程/班级灰度，监控错误率、延迟、队列积压和数据一致性。',
        '验收通过后才将 reserved 状态改为 enabled，并更新两份配套说明文档。',
    ]:
        add_list_item(doc, step, num)

    doc.add_heading('9. 未接入验证用例', level=1)
    reserved_tests = [
        ('G01', '静态扫描', '搜索 joinClass/studentCourseContent/studentSubmit/studentDiscussions/replyDiscussion', 'src 活动代码无可执行方法', '通过'),
        ('G02', '章节预览', '教师点击“查看学生视角”', '只读取教师端已有状态；无 /student/* 请求', '通过'),
        ('G03', '加入链接', '访问 /join/ABC12345', '显示预留页；不 POST /classes/{code}/join', '通过'),
        ('G04', '讨论学生入口', '访问 /student/discussions', '显示预留页；不 GET /student/discussions', '通过'),
        ('G05', '旧任务模拟', '检查旧任务页按钮与源代码', '按钮禁用；无 studentSubmit 调用', '通过'),
        ('G06', '教师关键路径', '登录、章节、邀请页', '教师 API 无 4xx/5xx；跨端请求数组为空', '通过'),
        ('G07', '后端保留性', '枚举 FastAPI 路由', '6 个学生契约仍存在，可供后续独立联调', '通过'),
        ('G08', '构建回归', 'TypeScript + Vite build', '预留元数据类型正确，构建通过', '通过'),
    ]
    add_table(doc, ['用例', '检查点', '操作', '预期/结果', '结论'], reserved_tests, [700, 1600, 2300, 3960, 800], font_size=8.2)

    doc.add_heading('10. 变更控制与责任矩阵', level=1)
    add_table(doc, ['变更类型', '提出', '评审', '实施', '验收'], [
        ('启用学生端接口', '产品', '架构/安全/数据', '学生端 + 后端', '产品/测试/安全'),
        ('新增教师端 API', '前端/产品', '后端/架构', '后端 + 前端', '测试'),
        ('修改数据库结构', '后端', '架构/数据', '后端', '测试/运维'),
        ('调整接口字段或错误码', '任一消费方', '所有消费方', '接口所有者', '契约测试'),
        ('解除预留状态', '产品', '发布评审会', '配置/代码所有者', '发布负责人'),
    ], [2400, 1500, 2200, 1800, 1460], font_size=8.8)

    doc.add_heading('11. 发布检查表', level=1)
    bullet = add_numbering_definition(doc, bullet=True)
    for item in [
        '教师端前端仅调用 /teacher/*、资料文件和健康检查允许范围。',
        'src/reserved/studentIntegration.ts 中所有学生契约状态仍为 reserved。',
        '学生 URL 返回预留页，页面无提交按钮和自动轮询。',
        '生产构建、后端回归与浏览器隔离检查全部通过。',
        '未接入按钮不得显示误导性的“已导出/已保存”成功结果，需求方已知晓缺口。',
        '原型认证、调度器、图谱双模型风险已进入上线门禁清单。',
        '任何启用学生端的变更均有独立需求单、接口评审、测试报告和回滚方案。',
    ]:
        add_list_item(doc, item, bullet)

    doc.add_heading('12. 验收签字项', level=1)
    add_table(doc, ['角色', '确认事项', '结论/签字'], [
        ('产品负责人', '未接入范围和优先级无遗漏', '________________'),
        ('教师端负责人', '教师端不存在学生 API 运行时调用', '________________'),
        ('学生端负责人', '预留契约可作为后续联调输入', '________________'),
        ('后端负责人', '现有路由与缺口分类准确', '________________'),
        ('安全负责人', '启用门禁满足身份、权限、限流和隐私要求', '________________'),
    ], [1700, 4900, 2760], font_size=9.2)
    add_para(doc, f'文档生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT, before=12)
    audit_document(doc)
    doc.save(RESERVED_PATH)


if __name__ == '__main__':
    build_connected()
    build_reserved()
    print(CONNECTED_PATH)
    print(RESERVED_PATH)

