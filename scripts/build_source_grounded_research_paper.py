from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import build_research_needs_chapter as b

OUT_PATH = b.OUT_DIR / "CodeTrack_研究需求与总体方案_申报表提取修订版.docx"
ASSET_DIR = b.OUT_DIR / "codetrack_source_research_assets"


def font(size, bold=False):
    p = r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"
    return ImageFont.truetype(p, size)


def box(draw, xy, text, fill, edge, size=26):
    fill = fill if fill.startswith("#") else "#" + fill
    edge = edge if edge.startswith("#") else "#" + edge
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=edge, width=3)
    bb = draw.multiline_textbbox((0, 0), text, font=font(size), spacing=8, align="center")
    tw, th = bb[2]-bb[0], bb[3]-bb[1]
    x1, y1, x2, y2 = xy
    draw.multiline_text(((x1+x2-tw)/2, (y1+y2-th)/2), text, fill="#1E2D38", font=font(size), spacing=8, align="center")


def arrow(draw, start, end, color="#64737F"):
    draw.line([start, end], fill=color, width=5)
    sx, sy = start; ex, ey = end
    pts = [(ex, ey), (ex-16 if ex > sx else ex+16, ey-10), (ex-16 if ex > sx else ex+16, ey+10)]
    draw.polygon(pts, fill=color)


def build_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    # The source proposal explicitly uses a three-layer architecture.
    im = Image.new("RGB", (2400, 1200), "white"); d = ImageDraw.Draw(im)
    layers = [("应用层", 80, 315, "#F6FAFC", b.BLUE), ("能力层", 400, 635, "#EAF5F1", b.GREEN), ("基座层", 720, 1035, "#F7F7F7", "#7D8790")]
    for label, y1, y2, fill, edge in layers:
        d.rounded_rectangle((70, y1, 2330, y2), radius=22, outline="#D5DDE2", width=3, fill="white")
        d.text((110, y1+25), label, font=font(36, True), fill="#1E2D38")
        d.line((310, y1+12, 310, y2-12), fill="#D5DDE2", width=2)
    for i, t in enumerate(["学生端\n输入题目·查看分步推导", "教师端\n资源审核·结果盲评", "教学管理端\n课程与用户管理", "科研入口\n文献·代码·实验记录"]):
        box(d, (370+i*475, 135, 800+i*475, 275), t, "#F6FAFC", b.BLUE, 24)
    for i, t in enumerate(["问题分类", "检索验证\nRAG", "步骤拆解", "主模型求解", "答案质检"]):
        x=390+i*375; box(d,(x,455,x+285,580),t,"#EAF5F1",b.GREEN,24)
        if i<4: arrow(d,(x+285,517),(x+360,517),"#3A8D74")
    for i, t in enumerate(["Qwen2.5-7B-Instruct\n+ LoRA适配器", "LangChain + Chroma\n多智能体调度", "Gradio Web\n统一API与本地部署"]):
        box(d, (390+i*650, 825, 950+i*650, 965), t, "#F7F7F7", "#7D8790", 24)
    for y1,y2 in ((315,400),(635,720)): arrow(d,(1200,y1+10),(1200,y2-10),"#7D8790")
    im.save(ASSET_DIR / "source_architecture.png")

    im = Image.new("RGB", (2500, 820), "white"); d = ImageDraw.Draw(im)
    d.text((90, 45), "产品使用流程：输入—自动处理—查看输出—纠错反馈", font=font(38, True), fill="#1E2D38")
    steps = [("输入题目", "支持自然语言与 LaTeX\n无需手动选题型"), ("自动处理", "分类→检索→拆解\n→求解→质检"), ("查看输出", "最终答案、分步详解\n推理依据与例题来源"), ("纠错反馈", "学生标记问题\n教师复核并回流")]
    for i,(a,c) in enumerate(steps):
        x=110+i*590; box(d,(x,230,x+440,600),a+"\n\n"+c,["#EAF3F8","#EAF5F1","#FFF5E6","#FDECEC"][i],[b.BLUE,b.GREEN,"#C9822B","#B94A48"][i],27)
        if i<3: arrow(d,(x+440,415),(x+545,415))
    im.save(ASSET_DIR / "source_business_flow.png")

    im = Image.new("RGB", (2500, 1180), "white"); d = ImageDraw.Draw(im)
    d.text((90, 50), "研究框架：数据构建—模型适配—智能体协同—评价迭代", font=font(38, True), fill="#1E2D38")
    cols=[("目标", "350条四元组数据\nLoRA垂直适配\n多智能体系统\n对比与用户验证", "#EAF3F8", b.BLUE),
          ("内容", "数据集与质量管控\nLoRA微调与错误归因\n四智能体协同部署\n评测与闭环迭代", "#EAF5F1", b.GREEN),
          ("关键问题", "低资源知识注入\n中间过程可验证\n自主化持续迭代", "#FFF5E6", "#C9822B"),
          ("产出", "数据集·适配器\n原型系统·评测报告\n论文与搭建指南", "#FDECEC", "#B94A48")]
    for i,(title,txt,fill,edge) in enumerate(cols):
        x=110+i*590; box(d,(x,260,x+455,880),title+"\n\n"+txt,fill,edge,27)
        if i<3: arrow(d,(x+455,570),(x+540,570))
    d.text((180, 1010), "技术路线：需求调研 → 数据集构建 → LoRA微调 → 多智能体部署 → 对比验证 → 用户试用 → 数据回流", font=font(29, True), fill="#1E2D38")
    im.save(ASSET_DIR / "source_framework.png")


def body(doc, text, prefix=None):
    return b.add_body(doc, text, bold_prefix=prefix)


def build():
    build_assets()
    doc = b.setup_document()
    p = doc.add_paragraph(); b.para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=8)
    r=p.add_run("面向学科生态的垂直大模型构建与自主化学科智能体集成方案研究"); b.set_run_font(r,east_asia="黑体",size=18,bold=True,color="1E2D38")
    p=doc.add_paragraph(); b.para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    r=p.add_run("研究需求与总体方案"); b.set_run_font(r,east_asia="宋体",size=15,bold=True,color=b.GREEN)
    p=doc.add_paragraph(); b.para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    r=p.add_run("信赛创翼队"); b.set_run_font(r,east_asia="宋体",size=12)
    p=doc.add_paragraph(); b.para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    r=p.add_run("信阳学院计算机与人工智能学院"); b.set_run_font(r,east_asia="宋体",size=10.5,color=b.GRAY)
    p=doc.add_paragraph(); b.para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=2)
    r=p.add_run("摘  要："); b.set_run_font(r,east_asia="黑体",size=12,bold=True)
    r=p.add_run("通用大语言模型在高校理工科学科问答中存在专业术语混淆、推导逻辑跳跃和解题方法与课堂不一致等问题。项目组通过3次习题课课堂观察、3位教师深度访谈和400余条历史答疑记录挖掘，发现学生的核心需求不是直接获得答案，而是获得与课堂方法一致、可跟随的分步推导。基于这一需求，本文从项目申报书和配套研究方案中提取研究事实与技术参数，提出以Qwen2.5-7B-Instruct为基座、LoRA为参数高效适配方法、结构化四元组数据为知识载体、LangChain多智能体为推理编排机制的总体方案，并将教师端、学生端、管理员端和科研入口纳入统一平台边界。本文系统论述研究需求、三层架构、业务流程、研究框架、实施计划和验证方法，为后续模型训练与系统实现提供论文级方案依据。")
    b.set_run_font(r,east_asia="宋体",size=12)
    p=doc.add_paragraph(); b.para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=12)
    r=p.add_run("关键词："); b.set_run_font(r,east_asia="黑体",size=12,bold=True)
    r=p.add_run("垂直大模型；结构化推理数据；LoRA；检索增强生成；多智能体协同")
    b.set_run_font(r,east_asia="宋体",size=12)

    b.add_heading(doc,"1 项目来源与研究需求",1)
    b.add_heading(doc,"1.1 需求调研与问题界定",2)
    body(doc,"本项目来源于高校理工科教学场景中的真实答疑需求。项目组采用课堂观察、教师深度访谈和历史答疑数据挖掘三种方法并行采集第一手资料。课堂观察连续跟踪3次习题课，以每5分钟为一个时间切片，共形成36个时间切片；观察发现，学生在综合题上的平均停留时间超过20分钟，是单一知识点题目的2.5倍，且约60%的学生使用通用AI工具时只是“看一眼答案然后关掉”，没有形成对推导过程的理解。教师访谈覆盖3位授课教师和11个关键答疑场景，教师普遍反感AI解答与课堂教学方法不一致。历史答疑数据挖掘覆盖400余条提问记录，推导中断类问题占48%，综合题占提问总量35%，却贡献了62%的深度答疑。三类证据共同指向同一问题：教学场景需要的是可跟随的分步推导，而不是脱离课程方法的最终答案。")
    b.add_table(doc,"表1  项目需求调研的原始证据",["调研方法","样本与过程","主要发现","对方案的约束"],[
        ["课堂观察","3次习题课；36个5分钟时间切片","综合题停留超20分钟；约60% AI使用停留在看答案","输出必须呈现可跟随的推导过程"],
        ["教师访谈","3位教师；11个关键答疑场景","教师要求解法与课堂方法一致","设置教师审核、盲评与纠错回流"],
        ["答疑数据挖掘","400余条历史提问","推导中断48%；综合题35%却贡献62%深度答疑","优先优化综合题与中间步骤验证"]], [2.7,4.0,4.6,3.8])
    b.add_heading(doc,"1.2 四类平台用户需求",2)
    body(doc,"申报书原始研究对象主要是学生、教师和教学管理者；结合你已确定的最终产品范围，本文将管理员端和科研入口纳入平台集成层，并明确二者属于产品化扩展，不把其功能写成申报书已经完成的实验结果。")
    body(doc,"学生端需求是围绕课程章节、知识点和练习任务获得与课堂方法一致的分步推导，支持自然语言或LaTeX输入，查看推理依据和参考例题来源，并对“看不懂”“步骤错误”“方法不一致”等情况进行反馈。教师端需求是管理课程和资料、审核知识入库、创建和发布练习、查看班级学情、盲评模型输出并确认有效纠错样本。管理员端需求是管理用户角色、课程权限、模型和知识库版本、运行日志及异常告警，保证高风险任务可暂停、可转人工、可回滚。科研入口需求是管理文献、代码、数据和实验记录，关联模型与知识库版本，形成可复现实验档案；该入口的智能功能在当前阶段属于待实现范围。")
    b.add_table(doc,"表2  四类端的研究需求边界",["端/入口","原始材料中的依据","本研究拟解决的任务","完成状态"],[
        ["学生端","课后习题与课程答疑调研；四步产品流程","题型识别、检索、分步推导、来源展示、反馈","原型与方案设计；模型能力待验证"],
        ["教师端","教师访谈、教师盲评和资料审核要求","课程资料、任务发布、知识审核、学情分析、人工确认","原型已有交互；算法联调待完成"],
        ["管理员端","平台安全、权限、版本和日志要求","角色权限、版本治理、异常告警、审计回滚","产品化扩展设计"],
        ["科研入口","项目范围内的研究工作空间需求","文献、代码、实验记录和成果归档","产品化扩展设计；不宣称已实现"]], [2.4,4.5,5.0,3.3], page_break_before=True)

    b.add_heading(doc,"2 平台总体方案",1)
    body(doc,"项目申报书将系统总体架构划分为基座层、能力层和应用层。本文沿用这一三层架构，避免把产品页面、模型能力和实验计划混为同一层。基座层提供语言理解与生成能力；能力层完成知识检索、问题分类、步骤拆解、求解和质检；应用层向学生、教师、教学管理者以及科研入口提供交互服务。")
    doc.add_picture(str(ASSET_DIR/"source_architecture.png"), width=Cm(15.6)); b.add_caption(doc,"图1  基于申报方案提取的三层系统总体架构")
    b.add_heading(doc,"2.1 基座层：Qwen2.5-7B-Instruct与LoRA适配",2)
    body(doc,"基座层计划采用Qwen2.5-7B-Instruct作为开源基础模型，通过LoRA进行参数高效微调。项目方案设定LoRA秩r=16，在注意力层与全连接层注入低秩增量矩阵，冻结基础模型原始权重，仅训练增量参数；训练计划使用单张RTX 3090（24GB）完成，并以验证集损失早停策略选择较优checkpoint。上述参数是项目实施计划，不代表模型训练和效果验证已经完成。")
    b.add_heading(doc,"2.2 能力层：结构化知识与多智能体协同",2)
    body(doc,"能力层以“问题—题型—分步推导—答案”四元组作为主要数据组织形式。项目计划从本科计算机与人工智能方向核心课程近五年的教材习题和考试真题中收集数据，经去重、格式统一、质量筛选和双人交叉校验后，形成不少于350条样本，题型包括概念辨析、数值计算、证明推理和实际应用，并按8:1:1划分训练集、验证集和测试集。")
    body(doc,"多智能体部分基于LangChain构建四个专用智能体：问题分类智能体判断题型；检索增强智能体在Chroma向量库中检索同类型例题及其推理链并注入提示词；步骤拆解智能体按题型策略将复杂问题分解为子步骤，并依次调用主模型求解；答案质检智能体基于规则检查推理链的逻辑连贯性和结论可导出性。四者形成“分类→检索→拆解→求解→质检”链路，质检不通过时最大回溯2次，连续失败或高风险任务转人工确认。")
    b.add_heading(doc,"2.3 应用层：三端与科研入口",2)
    body(doc,"应用层计划采用Gradio开发Web界面并通过统一API封装后端功能。学生端面向题目输入、分步推导和反馈；教师端面向课程资料、任务发布、知识审核和学情分析；管理员端面向权限、版本、日志和安全；科研入口面向文献、代码、实验记录和成果归档。申报书明确的交付物是面向本科计算机与人工智能方向核心课程的自动化垂直问答原型系统，管理员端和科研入口的完整实现应在产品化阶段单独验收。")

    b.add_heading(doc,"3 平台业务流程",1)
    body(doc,"项目方案将产品使用流程明确为四步闭环：输入题目、系统自动处理、查看输出和纠错反馈。该流程以用户不需要手动选择题型或多次提问为前提，将多智能体协同隐藏在统一任务流中，同时保留教师对结果的最终确认权。")
    doc.add_picture(str(ASSET_DIR/"source_business_flow.png"), width=Cm(15.6)); b.add_caption(doc,"图2  申报方案中的四步产品使用流程")
    b.add_heading(doc,"3.1 学生端流程",2)
    body(doc,"学生在Web界面输入问题或LaTeX公式后，系统读取用户角色、课程范围和当前知识库版本，自动执行“分类→检索→拆解→求解→质检”。结果页面展示最终答案、各步骤详解、每一步的推理依据和参考例题来源，学生可以继续追问子步骤或提交“看不懂”“步骤错误”“方法不一致”等反馈。未经教师复核的反馈不能直接写入正式训练集。")
    b.add_heading(doc,"3.2 教师端流程",2)
    body(doc,"教师上传教材、讲义、实验指导等课程资料，完成章节和知识点绑定后提交审核；教师创建练习时设置课程目标、知识点、题型和截止时间，并在发布前预览学生视角。系统运行后，教师查看班级完成情况、错误类型和AI输出，对步骤进行盲评、修改或驳回。教师确认的纠错样本进入更新队列，达到计划阈值后触发增量微调或策略调整。")
    b.add_heading(doc,"3.3 管理员端与科研入口流程",2)
    body(doc,"管理员端负责角色权限、课程与知识库版本、模型及提示词版本、运行日志和异常告警。遇到检索证据不足、质检连续失败、敏感信息风险、越权访问或代码沙箱异常时，系统暂停自动发布并转人工处理。科研入口在项目空间内管理文献、代码、数据和实验记录，每条记录关联数据版本、模型版本和运行时间；其输出定位为研究草稿与分析建议，不替代研究者对代码、事实和结论的责任。")

    b.add_heading(doc,"4 研究框架与技术路线",1)
    body(doc,"本项目围绕三个关键科学与技术问题展开：低资源条件下如何将学科知识有效嵌入模型参数空间；复杂推理任务的中间过程如何实现可控拆解与可追溯检验；如何建立全流程自动化和可持续的数据驱动迭代机制。")
    doc.add_picture(str(ASSET_DIR/"source_framework.png"), width=Cm(15.6)); b.add_caption(doc,"图3  从申报书提取的研究框架")
    b.add_heading(doc,"4.1 研究目标",2)
    body(doc,"目标一是完成不少于350条结构化推理数据集构建和Qwen2.5-7B-Instruct的LoRA微调；目标二是基于LangChain部署四个专用智能体，形成自动化问答链路并完成Gradio端到端联调；目标三是以原始Qwen2.5-7B-Instruct、Qwen+RAG和DeepSeek-V3为主要基线，从答案正确率、推导完整性和方法一致性三个维度开展评测，并通过15—20名学生和2—3名教师试用验证；目标四是交付数据集、微调模型权重、系统源代码、搭建指南和学术论文，并探索向数学、物理等理工科课程迁移。")
    b.add_heading(doc,"4.2 研究内容",2)
    body(doc,"研究内容包括四个方面：第一，学科结构化推理数据集的构建规范与质量管控，重点记录题型、推理依据、来源和版本；第二，LoRA垂直大模型微调与错误归因，区分术语混淆、公式误用、推导跳跃和计算失误，为质检规则提供依据；第三，自主化多智能体协同机制和产品化集成，明确各智能体输入、输出、失败状态和最大回溯次数；第四，多维对比评测和闭环迭代，按照“评测→试用→数据回流→迭代”更新系统。")
    b.add_heading(doc,"4.3 实施安排与评价指标",2)
    b.add_table(doc,"表3  项目实施阶段与拟验证指标",["阶段","时间与主要工作","拟交付/验证","状态边界"],[
        ["需求调研与原型开发","2026年7—8月；调研、数据清洗标注、模型与多智能体原型、Gradio联调","约350条四元组；可运行原型","计划任务"],
        ["系统验证与闭环迭代","2026年9—12月；基线对比、教师盲评、两周用户试用、反馈回流","正确率、推导完整性、方法一致性；50条有效样本触发增量微调","待实施"],
        ["成果沉淀与推广","2027年1—6月；课程群拓展、1—2个专业迁移、开源与论文","模型权重、数据集、代码、搭建指南、论文","预期成果"]], [2.8,5.6,4.1,3.2])
    body(doc,"申报方案提出的预期验收目标包括：测试集答案正确率较原始基线提升不少于10个百分点，教师盲评推导完整性均分不低于4.0/5.0，单课程试点进入条件为测试集正确率不低于70%，方法一致性评分不低于4.0/5.0，推导可读性评分不低于3.5/5.0。以上均为计划指标，须在实际实验与用户试用完成后报告真实结果。")

    b.add_heading(doc,"5 可行性、创新点与边界",1)
    b.add_heading(doc,"5.1 可行性与风险控制",2)
    body(doc,"技术上，Qwen2.5-7B-Instruct、LoRA、LangChain、Chroma和Gradio均为成熟开源技术；数据上，教材习题和考试真题可作为公开资源，经授权、去重和双人校验后使用；团队具备Python与PyTorch基础，学院实验室提供服务器和开发环境。针对小样本过拟合风险，方案设置轻量数据增强、验证集早停、较低LoRA秩和独立测试集；针对检索错误，采用来源、版本和教师审核机制；针对智能体协作失败，设置质检回溯和人工转交；针对反馈回流造成退化，保留独立测试集和旧版本进行A/B对比。")
    b.add_heading(doc,"5.2 项目创新点",2)
    body(doc,"第一，将检索增强与分步拆解融合为“分类→检索→拆解→求解→质检”的多智能体协同推理机制，并通过质检失败回溯增强过程可观察性和容错性。第二，提出面向学科习题求解的“问题—题型—分步推导—答案”四元组组织方式，将每一步推理依据编码为可训练、可校验的数据。第三，采用“7B级开源模型+LoRA+消费级GPU”的低成本适配路线，降低普通高校开展垂直模型研究的门槛。第四，将教师和学生的真实反馈纳入“评测—试用—数据回流—迭代”闭环，避免系统一次性交付后失去改进依据。")
    b.add_heading(doc,"5.3 研究边界",2)
    body(doc,"截至当前，真实模型代码、完整训练过程、对比实验结果和用户试用数据尚未全部完成，本文不将其写成已实现事实。管理员端和科研入口属于在三端教学原型基础上的产品化扩展，其功能需求已纳入总体方案，但需在后续开发中单独实现和验收。个人画像只用于学习支持和学情分析，不直接生成成绩或处分结论；模型输出必须提供证据来源并保留人工确认通道。")
    b.add_heading(doc,"6 结论",1)
    body(doc,"本文依据项目申报书及配套研究方案，提取真实调研数据、技术参数、实施计划和预期成果，形成了CodeTrack研究需求与总体方案。研究以“与课堂方法一致的分步推导”为核心需求，以Qwen2.5-7B-Instruct和LoRA完成学科知识注入，以结构化四元组数据承载推理过程，以LangChain组织分类、检索、拆解、求解和质检智能体，并通过三层架构和四步业务流程连接学生端、教师端、管理员端和科研入口。后续应严格按照计划指标开展训练、评测和用户试用，依据真实结果补充论文中的实验章节。")

    b.add_heading(doc,"参考文献",1)
    refs=[
        "[1] 贾万林. 面向学科生态的垂直大模型构建与自主化学科智能体集成方案研究：项目申报书[R]. 信阳学院, 2026.",
        "[2] 籍欣萌, 昝红英, 崔婷婷, 张坤丽. 大模型在垂直领域应用的现状与挑战[J]. 计算机工程与应用, 2025, 61(12): 1-11.",
        "[3] 周明, 刘群, 张伟, 等. 低资源场景下大模型高效微调技术比较研究[J]. 计算机科学, 2025, 52(5): 1-15.",
        "[4] 刘雪颖, 云静, 李博, 等. 基于大型语言模型的检索增强生成综述[J]. 计算机工程与应用, 2025, 61(13): 1-25.",
        "[5] 张高飞, 李欢, 池云仙, 等. 一种基于领域知识的检索增强生成方法[J]. 河北工业科技, 2025, 42(2): 103-110.",
        "[6] 赵鑫, 窦志成, 文继荣, 等. 大语言模型智能体综述[J/OL]. 计算机学报, 2025.",
        "[7] 朱孟潇, 陆文灏, 王禧玥, 等. 智能教育中的大模型智能体：基础、模式和挑战[J]. 教师教育论坛, 2025, 38(2): 49-60.",
    ]
    for ref in refs:
        p=doc.add_paragraph(); b.para_format(p,indent=False,align=WD_ALIGN_PARAGRAPH.LEFT,after=2)
        r=p.add_run(ref); b.set_run_font(r,east_asia="宋体",size=10.5)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__": build()
