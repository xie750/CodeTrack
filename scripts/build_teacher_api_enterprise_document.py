from __future__ import annotations

from datetime import datetime
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

from build_database_integration_document import (
    GREEN,
    INK,
    LIGHT_GREEN,
    MUTED,
    ROOT,
    add_callout as _add_callout,
    add_field,
    add_para,
    add_table,
    audit_document,
    configure_section,
    configure_styles,
    set_run_font,
    set_row_cant_split,
    set_table_geometry,
    style_cell,
)
from build_database_enterprise_document import add_code_block, add_list_item
from build_teacher_student_backend_document import add_numbering_definition


OUTPUT_DIR = ROOT / "artifacts" / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "CodeTrack教师端后端API接口文档_企业版.docx"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

DOC_ID = "CT-API-TCH-001"
DOC_VERSION = "V1.0"
BASELINE_DATE = "2026-08-16"


def add_callout(doc, title, text, fill=LIGHT_GREEN):
    _add_callout(doc, title, text, fill=fill)
    set_row_cant_split(doc.tables[-1].rows[0])


MODULES = [
    {
        "name": "平台、认证与教师端持久化",
        "scope": "健康检查、原型登录、首屏聚合、仪表盘、课程创建草稿、课程公告和个人偏好。",
        "tables": "users, teacher_credentials, teacher_preferences, course_drafts, course_announcements, announcement_reads, notifications",
        "rows": [
            ("SYS-01", "GET", "/api/v1/health", "服务健康检查", "应用启动/运维探针", "无", "status=ok", "200", "稳定"),
            ("AUTH-01", "GET", "/api/v1/teacher/auth/accounts", "列出教师原型账号", "登录页账号选择", "无", "教师摘要数组", "200", "原型"),
            ("AUTH-02", "POST", "/api/v1/teacher/auth/login", "校验教师账号密码", "登录按钮", "TeacherLogin", "教师摘要", "200/401/422", "原型"),
            ("APP-01", "GET", "/api/v1/teacher/bootstrap", "聚合首屏基础数据", "应用启动", "course_id?, class_id?", "教师、课程、班级、通知", "200/403/404", "已实现"),
            ("APP-02", "GET", "/api/v1/teacher/dashboard", "读取教师工作台数据", "工作台", "course_id, class_id", "指标、任务、动态、风险学生", "200/403/404", "部分原型"),
            ("DRAFT-01", "GET", "/api/v1/teacher/course-draft", "读取本人课程创建草稿", "创建课程恢复", "无", "草稿对象或 null", "200/403", "已实现"),
            ("DRAFT-02", "PUT", "/api/v1/teacher/course-draft", "保存/覆盖课程创建草稿", "创建课程自动保存", "CourseDraftUpsert", "草稿及 savedAt", "200/403/422", "已实现"),
            ("DRAFT-03", "DELETE", "/api/v1/teacher/course-draft", "删除本人课程草稿", "放弃草稿", "无", "deleted=true", "200/403", "已实现"),
            ("ANN-01", "GET", "/api/v1/teacher/courses/{course_id}/announcements", "读取课程公告", "课程工作台/公告", "course_id", "公告数组及 read 状态", "200/403/404", "已实现"),
            ("ANN-02", "PATCH", "/api/v1/teacher/announcements/{announcement_id}/read", "公告标记已读", "打开公告", "announcement_id", "id, read=true", "200/403/404", "已实现"),
            ("PREF-01", "GET", "/api/v1/teacher/preferences", "读取教师偏好", "设置页", "无", "通知、AI、邮件摘要设置", "200/403", "已实现"),
            ("PREF-02", "PUT", "/api/v1/teacher/preferences", "更新教师偏好", "设置页保存", "TeacherPreferenceUpdate", "更新后的偏好", "200/403/422", "已实现"),
        ],
    },
    {
        "name": "课程管理",
        "scope": "课程列表、新建、编辑、归档/恢复和删除。",
        "tables": "courses, class_groups, chapters, tasks, materials, audit_logs",
        "rows": [
            ("COURSE-01", "GET", "/api/v1/teacher/courses", "列出当前教师课程", "课程列表", "无", "课程摘要数组", "200/403", "已实现"),
            ("COURSE-02", "POST", "/api/v1/teacher/courses", "创建课程及初始章节", "创建课程提交", "CourseCreate", "新课程摘要", "201/403/409/422", "已实现"),
            ("COURSE-03", "PATCH", "/api/v1/teacher/courses/{course_id}", "修改课程与状态", "课程设置/归档/恢复", "CourseUpdate", "更新后的课程", "200/403/404/422", "已实现"),
            ("COURSE-04", "DELETE", "/api/v1/teacher/courses/{course_id}", "删除课程", "课程列表删除", "course_id", "id, deleted=true", "200/403/404/409", "已实现"),
        ],
    },
    {
        "name": "班级与成员管理",
        "scope": "教学班列表、新建、邀请码、批量导入、学生花名册与加入状态。",
        "tables": "class_groups, enrollments, users, submissions, audit_logs",
        "rows": [
            ("CLASS-01", "GET", "/api/v1/teacher/classes", "按课程列出教学班", "班级管理", "course_id", "班级摘要数组", "200/403/404", "已实现"),
            ("CLASS-02", "POST", "/api/v1/teacher/classes", "创建教学班", "新建班级", "ClassCreate", "班级与邀请码", "201/403/404/422", "已实现"),
            ("CLASS-03", "POST", "/api/v1/teacher/classes/{class_id}/join-code", "刷新班级邀请码", "班级邀请码操作", "class_id", "class_id, join_code", "200/403/404", "已实现"),
            ("CLASS-04", "POST", "/api/v1/teacher/classes/{class_id}/students/import", "批量导入学生", "导入学生", "StudentImportPayload", "created/enrolled/skipped/errors", "200/403/404/422", "已实现"),
            ("CLASS-05", "GET", "/api/v1/teacher/classes/{class_id}/students", "读取班级学生与学情摘要", "成员/学情列表", "class_id", "学生、进度、成绩、提交", "200/403/404", "部分原型"),
            ("CLASS-06", "GET", "/api/v1/teacher/classes/{class_id}/join-status", "读取入班进度", "加入进度", "class_id", "容量汇总与成员明细", "200/403/404", "部分原型"),
        ],
    },
    {
        "name": "章节与知识点",
        "scope": "课程章节、教学方式、发布状态和课程知识点。",
        "tables": "chapters, knowledge_points, materials, audit_logs",
        "rows": [
            ("CHAPTER-01", "GET", "/api/v1/teacher/courses/{course_id}/chapters", "列出章节和知识点", "课程内容", "course_id", "章节数组", "200/403/404", "已实现"),
            ("CHAPTER-02", "POST", "/api/v1/teacher/courses/{course_id}/chapters", "创建章节", "新增章节", "ChapterCreate", "新章节", "201/403/404/422", "已实现"),
            ("CHAPTER-03", "PATCH", "/api/v1/teacher/chapters/{chapter_id}", "编辑、发布或撤回章节", "章节保存/发布", "ChapterUpdate", "更新后的章节", "200/403/404/422", "已实现"),
            ("KP-01", "POST", "/api/v1/teacher/knowledge-points", "创建课程知识点", "章节知识点新增", "KnowledgePointCreate", "知识点对象", "201/403/404/422", "已实现"),
        ],
    },
    {
        "name": "资料与文件",
        "scope": "教学资料列表、元数据创建、真实文件上传下载、可见性、回收站、恢复和图谱关联。",
        "tables": "materials, material_knowledge_links, knowledge_points, audit_logs; backend/uploads",
        "rows": [
            ("MAT-01", "GET", "/api/v1/teacher/materials", "列出课程有效资料", "教学资料", "course_id?", "资料数组及关联知识点", "200/403/404", "已实现"),
            ("MAT-02", "POST", "/api/v1/teacher/materials", "创建资料元数据", "新增链接/占位资料", "MaterialCreate", "资料对象", "201/403/404/422", "已实现"),
            ("MAT-03", "POST", "/api/v1/teacher/materials/upload", "上传文件并创建资料", "上传资料", "multipart upload form", "id/title/status/size/content_url", "201/403/404/413/415", "已实现"),
            ("MAT-04", "GET", "/api/v1/material-files/{stored_name}", "下载资料文件", "资料打开/下载", "stored_name", "二进制文件流", "200/401/403/404", "已实现"),
            ("MAT-05", "PATCH", "/api/v1/teacher/materials/{material_id}", "修改可见性和状态", "资料发布/停用", "MaterialUpdate", "id/visibility/status", "200/403/404/422", "已实现"),
            ("MAT-06", "DELETE", "/api/v1/teacher/materials/{material_id}", "将资料移入回收站", "资料删除", "material_id", "资料对象 status=deleted", "200/403/404", "已实现"),
            ("MAT-07", "GET", "/api/v1/teacher/materials/trash", "列出资料回收站", "资料回收站", "course_id", "已删除资料数组", "200/403/404", "已实现"),
            ("MAT-08", "POST", "/api/v1/teacher/materials/{material_id}/restore", "恢复已删除资料", "回收站恢复", "material_id", "资料对象 status=ready", "200/403/404", "已实现"),
            ("MAT-09", "POST", "/api/v1/teacher/materials/{material_id}/knowledge-graph", "将资料关联至课程知识点", "导入知识图谱", "MaterialGraphImport", "关联数、创建节点、知识点", "200/403/404/409/422", "已实现"),
        ],
    },
    {
        "name": "资料文件夹",
        "scope": "手工文件夹、文件夹回收站、文件保留/迁移和 AI 目录建议。",
        "tables": "material_folders, materials, chapters, audit_logs",
        "rows": [
            ("FOLDER-01", "GET", "/api/v1/teacher/material-folders", "列出有效文件夹", "资料目录", "course_id", "文件夹数组", "200/403/404", "已实现"),
            ("FOLDER-02", "POST", "/api/v1/teacher/material-folders", "创建资料文件夹", "新建文件夹", "FolderCreate", "文件夹对象", "201/403/404/422", "已实现"),
            ("FOLDER-03", "GET", "/api/v1/teacher/material-folders/trash", "列出文件夹回收站", "文件夹回收站", "course_id", "文件夹及关联资料", "200/403/404", "已实现"),
            ("FOLDER-04", "DELETE", "/api/v1/teacher/material-folders/{folder_id}", "软删除文件夹及标记资料", "删除文件夹", "folder_id", "删除结果与资料数量", "200/403/404", "已实现"),
            ("FOLDER-05", "POST", "/api/v1/teacher/material-folders/{folder_id}/restore", "恢复文件夹及资料", "恢复文件夹", "folder_id", "恢复后的文件夹", "200/403/404", "已实现"),
            ("FOLDER-06", "POST", "/api/v1/teacher/material-folders/{folder_id}/materials/{material_id}/keep", "保留或迁移被删文件夹中的资料", "保留资料", "DeletedFolderMaterialKeep", "资料与目标文件夹", "200/403/404/409/422", "已实现"),
            ("FOLDER-07", "POST", "/api/v1/teacher/material-folders/ai-outline", "生成 AI 目录候选", "AI 整理资料", "OutlineRequest", "候选文件夹数组", "200/403/404", "规则回退"),
            ("FOLDER-08", "POST", "/api/v1/teacher/material-folders/confirm-outline", "确认 AI 目录", "确认目录建议", "OutlineConfirm", "最终文件夹数组", "200/403/404/422", "已实现"),
        ],
    },
    {
        "name": "任务与发布",
        "scope": "任务列表、手工创建、规则型 AI 草稿和面向班级发布。",
        "tables": "tasks, test_cases, chapters, knowledge_points, audit_logs",
        "rows": [
            ("TASK-01", "GET", "/api/v1/teacher/tasks", "列出课程任务", "任务管理", "course_id?", "任务数组", "200/403/404", "已实现"),
            ("TASK-02", "POST", "/api/v1/teacher/tasks", "创建任务草稿与测试用例", "新建任务", "TaskCreate", "完整任务", "201/403/404/422", "已实现"),
            ("TASK-03", "POST", "/api/v1/teacher/tasks/ai-draft", "生成规则型 AI 任务草稿", "AI 创建任务", "AITaskDraftRequest", "草稿摘要", "201/403/404/422", "规则回退"),
            ("TASK-04", "POST", "/api/v1/teacher/tasks/{task_id}/publish", "发布或设为 scheduled", "任务发布", "TaskPublish", "发布后的任务", "200/403/404/409/422", "部分实现"),
        ],
    },
    {
        "name": "提交、批改、反馈与 AI 审核",
        "scope": "学生提交监控、详情、评分、成绩发布、教师反馈和低置信度诊断审核。",
        "tables": "submissions, evaluation_results, grades, teacher_feedback, diagnosis_results, diagnosis_reviews, notifications",
        "rows": [
            ("SUB-01", "GET", "/api/v1/teacher/submissions", "按任务列出提交", "提交监控", "task_id?", "提交数组", "200/403/404", "已实现"),
            ("SUB-02", "GET", "/api/v1/teacher/submissions/{submission_id}", "读取提交详情", "批改详情", "submission_id", "源码、评测、诊断、成绩、反馈", "200/403/404", "已实现"),
            ("GRADE-01", "PUT", "/api/v1/teacher/submissions/{submission_id}/grade", "保存或覆盖评分", "保存批改", "GradeUpsert", "成绩对象", "200/403/404/422", "已实现"),
            ("GRADE-02", "POST", "/api/v1/teacher/submissions/{submission_id}/grade/publish", "发布成绩并通知学生", "发布成绩", "submission_id", "发布后的成绩", "200/403/404/409", "已实现"),
            ("FEED-01", "POST", "/api/v1/teacher/submissions/{submission_id}/feedback", "创建草稿或公开反馈", "提交反馈", "FeedbackCreate", "反馈对象", "201/403/404/422", "已实现"),
            ("REVIEW-01", "GET", "/api/v1/teacher/ai-reviews", "列出待审诊断", "AI 审核", "无", "审核任务数组", "200/403", "已实现"),
            ("REVIEW-02", "POST", "/api/v1/teacher/ai-reviews/{review_id}/action", "接受、修改或拒绝诊断", "审核提交", "ReviewAction", "审核结果", "200/403/404/422", "已实现"),
        ],
    },
    {
        "name": "学情分析与通知",
        "scope": "课程/班级分析汇总与教师通知已读。",
        "tables": "classes, enrollments, tasks, submissions, evaluation_results, grades, notifications",
        "rows": [
            ("ANA-01", "GET", "/api/v1/teacher/analytics/overview", "读取学情分析概览", "学情诊断", "course_id, class_id", "指标、趋势、知识点、风险", "200/403/404", "部分原型"),
            ("NOTICE-01", "PATCH", "/api/v1/teacher/notifications/{notification_id}", "修改本人通知已读状态", "通知中心", "NotificationRead", "id, read", "200/403/404/422", "已实现"),
        ],
    },
    {
        "name": "课程知识图谱",
        "scope": "基于 chapters/knowledge_points 的课程图谱、AI 候选确认和节点更新。",
        "tables": "chapters, knowledge_points, materials, material_knowledge_links, audit_logs",
        "rows": [
            ("CKG-01", "GET", "/api/v1/teacher/knowledge-graph", "读取课程知识图谱", "课程知识图谱", "course_id?", "nodes/edges/materials", "200/403/404", "已实现"),
            ("CKG-02", "POST", "/api/v1/teacher/knowledge-graph/ai-candidates", "生成知识点候选", "AI 生成节点", "CandidateRequest", "候选节点数组", "200/403/404/422", "规则回退"),
            ("CKG-03", "POST", "/api/v1/teacher/knowledge-graph/confirm", "确认候选并落库", "确认 AI 节点", "ConfirmCandidates", "新建知识点数组", "200/403/404/422", "已实现"),
            ("CKG-04", "PUT", "/api/v1/teacher/knowledge-graph/nodes/{node_id}", "更新课程知识点节点", "编辑节点", "NodeUpdate", "更新后的节点", "200/403/404/422", "已实现"),
        ],
    },
    {
        "name": "教师画布知识图谱",
        "scope": "独立画布图谱的 CRUD、文件生成和发布。该模块使用 teacher_knowledge_graphs，不等同于课程知识图谱。",
        "tables": "teacher_knowledge_graphs",
        "rows": [
            ("TGRAPH-01", "GET", "/api/v1/teacher/knowledge-graphs", "列出本人画布图谱", "知识图谱列表", "无", "图谱摘要数组", "200/403", "已实现"),
            ("TGRAPH-02", "POST", "/api/v1/teacher/knowledge-graphs", "创建空白画布图谱", "新建图谱", "GraphCreate", "完整图谱", "201/403/422", "已实现"),
            ("TGRAPH-03", "GET", "/api/v1/teacher/knowledge-graphs/{graph_id}", "读取画布图谱详情", "打开图谱", "graph_id", "nodes/edges/source files", "200/403/404", "已实现"),
            ("TGRAPH-04", "PUT", "/api/v1/teacher/knowledge-graphs/{graph_id}", "保存图谱节点与关系", "保存图谱", "GraphUpdate", "完整图谱", "200/403/404/422", "已实现"),
            ("TGRAPH-05", "DELETE", "/api/v1/teacher/knowledge-graphs/{graph_id}", "删除画布图谱", "删除图谱", "graph_id", "id, deleted=true", "200/403/404", "已实现"),
            ("TGRAPH-06", "POST", "/api/v1/teacher/knowledge-graphs/from-files", "从 txt/md/docx/pdf 生成图谱", "上传资料生成", "multipart graph form", "完整图谱与来源", "201/403/413/415/422", "AI/规则回退"),
            ("TGRAPH-07", "POST", "/api/v1/teacher/knowledge-graphs/{graph_id}/publish", "发布非空图谱", "发布图谱", "graph_id", "status=published", "200/403/404/409", "已实现"),
        ],
    },
    {
        "name": "课堂讨论",
        "scope": "教师查看、创建、发布和结束课程讨论。",
        "tables": "course_discussions, discussion_replies, notifications, audit_logs",
        "rows": [
            ("DISC-01", "GET", "/api/v1/teacher/discussions", "列出课程/班级讨论", "课堂讨论", "course_id, class_id?", "讨论及回复数组", "200/403/404", "已实现"),
            ("DISC-02", "POST", "/api/v1/teacher/discussions", "创建讨论草稿或立即发布", "新建讨论", "DiscussionCreate", "讨论对象", "201/403/404/422", "已实现"),
            ("DISC-03", "POST", "/api/v1/teacher/discussions/{discussion_id}/publish", "发布讨论", "讨论发布", "discussion_id", "status=published", "200/403/404", "已实现"),
            ("DISC-04", "POST", "/api/v1/teacher/discussions/{discussion_id}/end", "结束已发布讨论", "结束讨论", "discussion_id", "status=ended", "200/403/404/409", "已实现"),
        ],
    },
]


SCHEMAS = [
    ("TeacherLogin", "username:string(1-80), password:string(1-128)", "原型登录；成功不签发 token/session"),
    ("CourseDraftUpsert", "payload:object", "按教师唯一覆盖保存"),
    ("TeacherPreferenceUpdate", "notifications_enabled:boolean, ai_assistant_enabled:boolean, email_digest:boolean", "全量 PUT"),
    ("CourseCreate", "name:string(2-160), code:string(2-40), term:string, description:string?, student_visible:boolean?, chapter_titles:string[]?", "code 唯一"),
    ("CourseUpdate", "name?, term?, description?, status:active|preparing|archived?, student_visible?", "PATCH 仅更新非空字段"),
    ("ClassCreate", "course_id, name(2-120), grade?, major?, schedule?, mentor?, status:active|closed|pending?", "自动生成 8 位邀请码"),
    ("StudentImportPayload", "students:[{name(1-80), number(1-40)}]", "按学号去重，返回逐行错误"),
    ("ChapterCreate", "title(2-160), description?, teaching_mode?", "position 自动追加"),
    ("ChapterUpdate", "title?, description?, teaching_mode?, status:draft|published?", "发布/撤回会联动资料可见性"),
    ("KnowledgePointCreate", "chapter_id, name(1-120), description?, difficulty?, position_x/y:0-100", "章节必须属于教师课程"),
    ("MaterialCreate", "course_id, title, type, chapter_label, size?, visibility?, content_url?", "仅创建元数据"),
    ("multipart upload form", "course_id, chapter_label?, visibility?, file", "允许 pdf/ppt/pptx/doc/docx/csv/xlsx/mp4/txt；最大 200 MB"),
    ("MaterialUpdate", "visibility:teacher|students?, status:ready|disabled?", "删除/恢复使用独立接口"),
    ("MaterialGraphImport", "knowledge_point_ids:string[]?, create_from_material:boolean?", "至少选择或创建一个节点"),
    ("FolderCreate", "course_id, name(1-160), parent_id?", "同课程名称应由业务层避免重复"),
    ("DeletedFolderMaterialKeep", "target_folder_id:string|null", "为空表示保留在未分类"),
    ("OutlineRequest", "course_id", "根据章节和资料生成目录候选"),
    ("OutlineConfirm", "course_id, folders:string[]", "确认后新增/恢复并重排目录"),
    ("AITaskDraftRequest", "course_id, class_id?, prompt:string(4-1000)", "当前 generator=rule_fallback"),
    ("TaskCreate", "course_id, class_id?, title(2-200), type?, chapter_label, description?, starter_code?, difficulty?, total_score(1-1000), due_at, allow_hints?, test_cases[]", "至少一个测试用例"),
    ("TestCaseCreate", "name, input_data?, expected_output?, hidden?, weight:0-100", "隐藏用例仅教师端可见"),
    ("TaskPublish", "class_id, publish_at?, due_at", "due_at 必须晚于 publish_at/当前时间"),
    ("GradeUpsert", "score:0-100, comment?, dimensions:object?", "同一提交 upsert"),
    ("FeedbackCreate", "content:string(>=1), publish:boolean?", "publish=true 时 student_visible=true"),
    ("ReviewAction", "status:accepted|modified|rejected, reviewed_explanation?, comment?", "modified 必须填写 reviewed_explanation"),
    ("CandidateRequest", "course_id", "生成课程知识点候选"),
    ("ConfirmCandidates", "course_id, candidates:CandidateNode[]", "写 knowledge_points"),
    ("NodeUpdate", "name?, description?, difficulty?, mastery?, position_x?, position_y?", "仅允许课程所属教师"),
    ("GraphCreate", "title, description?, target_classes:string[]?, nodes[]?, edges[]?", "无节点时创建默认核心节点"),
    ("GraphUpdate", "title, description?, target_classes[], status:draft|published, nodes[], edges[]", "全量 PUT 图谱画布"),
    ("multipart graph form", "files[], title, description?, target_classes?", "txt/md/docx/pdf；单文件最大 20 MB"),
    ("DiscussionCreate", "course_id, class_id, title(1-160), content(1-5000), publish:boolean?", "班级必须属于课程"),
    ("NotificationRead", "read:boolean=true", "通知必须属于当前教师"),
]


DTO_ROWS = [
    ("Teacher", "id, name, number, email, department", "不应包含密码盐或密码哈希"),
    ("Course", "id, name, code, term, description, status, student_visible, progress, classes, students, task_count", "课程统计字段随列表返回"),
    ("Class", "id, course_id, name, grade, major, schedule, mentor, join_code, status, students, completion, active_rate, risk_count, capacity", "部分指标当前为原型值"),
    ("Chapter", "id, title, description, position, teaching_mode, status, knowledge_points[]", "发布状态影响学生可见性"),
    ("Material", "id, course_id, title, type, chapter, size, visibility, status, citations, content_url, updated_at, knowledge_points[]", "content_url 受下载接口权限保护"),
    ("Task", "id, course_id, class_id, title, type, chapter, description, starter_code, status, difficulty, total_score, publish_at, due_at, test_cases[]", "教师端包含隐藏测试用例"),
    ("Submission", "id, task_id, student, version, source_code, status, hint_level, submitted_at, evaluation, diagnosis, grade, feedback[]", "源码与隐藏明细仅教师可读"),
    ("Discussion", "id, course_id, class_id, class_name, title, content, status, participant_count, reply_count, published_at, replies[]", "列表当前携带完整回复"),
    ("CourseGraph", "nodes[{id,name,description,difficulty,mastery,x,y,materials}], edges[]", "基于 knowledge_points 动态拼装"),
    ("TeacherGraph", "id, title, description, status, target_classes, source_files, source_summary, nodes, edges, counts, timestamps", "target_classes 当前为名称数组"),
]


def total_api_count():
    return sum(len(module["rows"]) for module in MODULES)


def set_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = "CodeTrack 教师端 API"
    table.cell(0, 1).text = "后端接口契约与联调基线"
    set_table_geometry(table, [3600, 5760], indent_dxa=0)
    style_cell(table.cell(0, 0), bold=True, color=GREEN, size=8.5)
    style_cell(table.cell(0, 1), color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.cell(0, 0).text = f"{DOC_ID}  |  {DOC_VERSION}  |  内部开发使用"
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.add_run("第 "), size=8.5, color=MUTED)
    add_field(right, "PAGE")
    set_run_font(right.add_run(" 页 / 共 "), size=8.5, color=MUTED)
    add_field(right, "NUMPAGES")
    set_run_font(right.add_run(" 页"), size=8.5, color=MUTED)
    set_table_geometry(table, [5000, 4360], indent_dxa=0)
    style_cell(table.cell(0, 0), size=8.5, color=MUTED)


def add_cover(doc):
    add_para(doc, "CODETRACK  /  TEACHER API REFERENCE", size=10, bold=True, color=GREEN, after=16)
    add_para(doc, "教师端后端", size=28, bold=True, color=INK, after=2)
    add_para(doc, "API 接口文档", size=28, bold=True, color=INK, after=12)
    add_para(doc, "企业版：接口清单、契约、权限、数据与验收标准", size=14, color=MUTED, after=24)
    add_callout(doc, "文档定位", f"本文件只描述教师端前端调用的后端 API，不展开学生端接口。当前规范路径共 {total_api_count()} 个，覆盖教师端现有全部功能模块；另附 7 个知识图谱兼容别名。", fill=LIGHT_GREEN)
    add_table(doc, ["文档属性", "内容"], [
        ("文档编号", DOC_ID),
        ("版本 / 基线", f"{DOC_VERSION} / {BASELINE_DATE}"),
        ("技术栈", "React/Vite/TypeScript + FastAPI/SQLAlchemy + SQLite"),
        ("规范基础路径", "/api/v1"),
        ("接口规模", f"{total_api_count()} 个规范调用接口 + 7 个兼容别名"),
        ("数据库快照", "27 张业务表，1,191 行记录"),
        ("适用对象", "教师端前端、后端、测试、架构、产品和运维"),
    ], [1900, 7460], font_size=9.1)


def add_front_matter(doc):
    heading = doc.add_heading("文档控制", level=1)
    heading.paragraph_format.page_break_before = True
    add_table(doc, ["角色", "评审重点", "签署输出"], [
        ("教师端前端", "调用时机、字段类型、加载/空/错状态、兼容性", "页面—API 映射确认"),
        ("后端", "请求模型、权限、事务、状态机、错误码、审计", "OpenAPI 与实现确认"),
        ("测试", "正向、负向、越权、幂等、并发、回归", "接口测试报告"),
        ("架构/安全", "认证、对象级权限、敏感数据、限流与日志", "上线门禁确认"),
    ], [1800, 5200, 2360], font_size=8.8)
    doc.add_heading("阅读导航", level=1)
    add_table(doc, ["章节", "内容"], [
        ("1-4", "范围、系统基线、统一协议和完整接口清单"),
        ("5-16", "按教师端功能模块查看全部接口契约"),
        ("17-20", "请求模型、响应 DTO、错误码、数据与安全"),
        ("21-23", "前端接入、测试用例、发布门禁与兼容接口"),
    ], [1700, 7660], font_size=9.2)
    add_callout(doc, "状态定义", "“已实现”表示路由、数据库操作和前端调用已存在；“部分原型”表示接口可用但含演示指标或不完整生产能力；“规则回退”表示当前不是外部大模型生成；“原型”表示不得直接作为生产认证能力。")


def add_scope(doc, bullet_id):
    heading = doc.add_heading("1. 范围与原则", level=1)
    heading.paragraph_format.page_break_before = True
    add_para(doc, "本文以当前代码为事实来源，收录教师端页面直接调用、应用启动依赖或教师资料下载依赖的全部规范 API。学生加入班级、读取内容、提交任务和回复讨论等 /student 或学生行为接口不在本文详细范围内。")
    doc.add_heading("1.1 收录口径", level=2)
    for text in [
        "收录 /api/v1/teacher/** 规范路径。",
        "收录教师端必须调用的 /api/v1/health 和 /api/v1/material-files/{stored_name}。",
        "不把同一路径在 src/api.ts 中的重复方法计算为不同接口。",
        "知识图谱 /api/teacher/** 兼容别名放入附录，不作为新开发规范路径。",
        "所有请求、响应和错误说明以当前实现为准，目标生产改造单独标注。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("1.2 设计原则", level=2)
    add_table(doc, ["原则", "要求"], [
        ("单一事实源", "以服务端数据库和状态机为准，前端不长期保存业务副本。"),
        ("对象级权限", "除 role=teacher 外，还必须验证 course.teacher_id 或对象所属课程。"),
        ("契约优先", "OpenAPI、DTO、枚举、错误码和示例在发布前冻结。"),
        ("敏感最小化", "密码哈希、其他教师对象、学生隐藏评测数据不得越权返回。"),
        ("可追踪", "关键创建、发布、删除、评分和审核操作写 audit_logs 与 trace_id。"),
    ], [1800, 7560], font_size=8.9)


def add_baseline(doc):
    doc.add_heading("2. 系统基线与关键判断", level=1)
    add_table(doc, ["项目", "当前值"], [
        ("前端", "React / Vite / TypeScript，统一调用入口 src/api.ts"),
        ("后端", "FastAPI + Pydantic + SQLAlchemy"),
        ("数据库", "SQLite codetrack.db；27 张业务表，1,191 行"),
        ("成功响应", "大多数接口返回 {\"data\": ...}"),
        ("当前身份", "X-User-Id 请求头；多个依赖默认 teacher-01"),
        ("规范前缀", "/api/v1；/api/teacher 仅为画布图谱历史兼容"),
        ("文件存储", "backend/uploads 本地目录"),
    ], [2100, 7260], font_size=9.0, keep_together=True)
    add_callout(doc, "生产阻断项", "POST /teacher/auth/login 只验证 PBKDF2 密码并返回教师资料，没有签发 Cookie、JWT 或服务端 Session；后续请求仍由 X-User-Id 决定身份。GET /teacher/auth/accounts 还会公开列出教师账号。生产前必须完成真实会话、限流、账号锁定和 CSRF/Token 策略。", fill="FDECEC")
    doc.add_heading("2.1 两套知识图谱必须区分", level=2)
    add_table(doc, ["模块", "规范路径", "数据来源", "适用页面"], [
        ("课程知识图谱", "/api/v1/teacher/knowledge-graph", "chapters + knowledge_points + material links", "课程内容/资料关联图谱"),
        ("教师画布知识图谱", "/api/v1/teacher/knowledge-graphs", "teacher_knowledge_graphs 的 JSON 画布", "独立图谱编辑器"),
    ], [1800, 3200, 2600, 1760], font_size=8.5)


def add_protocol(doc, bullet_id):
    doc.add_heading("3. 通用接口协议", level=1)
    add_table(doc, ["项目", "当前约定", "企业目标"], [
        ("基础地址", "开发 http://127.0.0.1:8001/api/v1", "按环境注入，不在代码写死"),
        ("认证", "X-User-Id", "HttpOnly Session Cookie 或短期 JWT"),
        ("成功响应", "{data: ...}", "保留并补 meta/trace_id"),
        ("失败响应", "{detail: string|array}", "稳定 code/message/details/trace_id"),
        ("时间", "ISO 8601，部分接口无时区", "数据库 UTC，响应带时区"),
        ("分页", "多数列表一次性返回", "page/page_size/total，最大 100"),
        ("幂等", "部分接口业务幂等", "创建/发布支持 Idempotency-Key"),
        ("并发", "最后写入覆盖", "version 或 ETag/If-Match"),
    ], [1550, 3650, 4160], font_size=8.4)
    doc.add_heading("3.1 请求示例", level=2)
    add_code_block(doc, 'PATCH /api/v1/teacher/courses/course-ds\nContent-Type: application/json\nX-User-Id: teacher-01\n\n{"status":"archived"}')
    doc.add_heading("3.2 响应示例", level=2)
    add_code_block(doc, '{\n  "data": {\n    "id": "course-ds",\n    "status": "archived"\n  }\n}')
    doc.add_heading("3.3 权限检查顺序", level=2)
    for text in [
        "解析当前教师身份；生产环境不得信任客户端传入 user_id。",
        "加载目标对象并验证 tenant_id（目标架构）与 course.teacher_id。",
        "验证对象状态，例如 deleted、closed、published 或 scheduled。",
        "校验请求字段与跨对象引用，例如班级、章节、知识点必须属于同一课程。",
        "事务提交后写审计记录；向学生发通知的操作建议使用 Outbox。",
    ]:
        add_list_item(doc, text, bullet_id)


def add_inventory(doc):
    doc.add_heading("4. 教师端 API 完整清单", level=1)
    add_para(doc, f"规范接口共 {total_api_count()} 个。下表按模块汇总；第 5-16 章给出逐项契约。")
    summary = []
    for index, module in enumerate(MODULES, 1):
        counts = {}
        for row in module["rows"]:
            counts[row[1]] = counts.get(row[1], 0) + 1
        methods = ", ".join(f"{key} {value}" for key, value in counts.items())
        summary.append((index, module["name"], len(module["rows"]), methods, module["scope"]))
    add_table(doc, ["序", "模块", "接口数", "方法分布", "说明"], summary, [650, 2100, 850, 2000, 3760], font_size=8.0)
    add_callout(doc, "计数说明", "71 个规范接口包含健康检查和资料文件下载；不包含任何学生行为接口。附录中的 7 个 /api/teacher/knowledge-graphs 兼容别名与规范路径执行同一组处理函数，不应在新前端继续使用。")


def add_module(doc, number, module):
    doc.add_heading(f"{number}. {module['name']}", level=1)
    add_para(doc, module["scope"], size=10.2)
    add_para(doc, f"主要数据：{module['tables']}", size=9.2, color=MUTED, italic=True, after=6)
    add_table(doc, ["ID", "方法", "路径", "功能", "页面触发", "请求", "响应", "状态码", "成熟度"], module["rows"], [650, 550, 2350, 1200, 950, 900, 1050, 850, 860], font_size=7.0)
    doc.add_heading(f"{number}.1 模块业务规则", level=2)
    rules = {
        "平台、认证与教师端持久化": [
            "auth/accounts 与 auth/login 当前无真实会话，必须视为演示认证。",
            "course-draft 按 teacher_id 唯一保存，PUT 为覆盖语义，DELETE 重复调用仍返回成功。",
            "公告仅实现教师读取与已读；教师发布公告接口当前不存在。",
            "bootstrap 为首屏聚合，不应替代各业务列表的分页接口。",
        ],
        "课程管理": [
            "课程 code 唯一；冲突应返回 409，而不是 500。",
            "status 只允许 active、preparing、archived；恢复归档使用 PATCH status=preparing。",
            "DELETE 当前删除课程实体及级联关系，企业环境应增加二次确认、引用检查与可恢复策略。",
        ],
        "班级与成员管理": [
            "所有班级操作先验证所属课程由当前教师拥有。",
            "学生批量导入以学号识别已有用户，并分别统计 created、enrolled、skipped 和 errors。",
            "join-status 的 pending/invited 当前固定为 0；join_method、last_active 和部分容量指标为原型展示值。",
        ],
        "章节与知识点": [
            "章节发布会把同标题章节下 ready 资料 visibility 改为 students；撤回则改回 teacher。",
            "教学方式仅允许服务端白名单中的六类值。",
            "知识点坐标范围 0-100，所属章节决定课程权限。",
        ],
        "资料与文件": [
            "DELETE 是软删除，物理文件保留；恢复将状态改回 ready。",
            "文件上传限制 200 MB，扩展名使用白名单；下载前校验用户与资料状态。",
            "资料图谱关联会先删除旧关联再按本次请求重建，并更新 citations。",
            "MaterialUpdate 不允许直接写 deleted，避免绕过删除审计。",
        ],
        "资料文件夹": [
            "文件夹删除通过 source=deleted|... 标记；相关资料状态使用 folder_deleted|... 标记。",
            "保留资料可恢复原状态或迁移到目标文件夹，目标必须属于同一课程。",
            "AI 目录当前按章节/资料规则生成，不是外部大模型结果。",
        ],
        "任务与发布": [
            "任务草稿至少包含一个测试用例；发布前还要求课程存在知识点。",
            "publish_at 未来时间只把状态写成 scheduled；当前项目没有调度器自动切换 published。",
            "AI 草稿当前 generator=rule_fallback，必须由教师确认后发布。",
        ],
        "提交、批改、反馈与 AI 审核": [
            "教师只能读取自己课程的提交、源码和完整测试明细。",
            "成绩保存状态为 graded；发布后改为 grade_published 并创建学生通知。",
            "反馈 publish=true 时 student_visible=true；草稿反馈不得下发。",
            "modified 审核必须提供 reviewed_explanation。",
        ],
        "学情分析与通知": [
            "分析接口按课程与班级授权，部分进度/风险展示仍含原型计算。",
            "通知 PATCH 只能操作 notification.user_id 等于当前教师的记录。",
        ],
        "课程知识图谱": [
            "该模块以关系表和知识点实体为事实源，节点更新会写回 knowledge_points。",
            "AI 候选和资料匹配包含规则回退；候选确认后才落库。",
        ],
        "教师画布知识图谱": [
            "画布节点与边以 JSON 保存在 teacher_knowledge_graphs，不与 knowledge_points 自动同步。",
            "文件生成允许 txt、md、docx、pdf；单文件最大 20 MB。",
            "空图谱不能发布；target_classes 当前是班级名称数组，生产应改 class_id 关联。",
        ],
        "课堂讨论": [
            "讨论状态为 draft、published、ended；只有 published 可结束。",
            "发布会为目标班级学生创建通知；重复发布不会重复生成通知。",
            "教师列表当前返回完整 replies，数据量增长后应拆为详情分页。",
        ],
    }[module["name"]]
    for rule in rules:
        p = doc.add_paragraph()
        p.style = doc.styles["List Bullet"]
        p.paragraph_format.space_after = doc.styles["Normal"].paragraph_format.space_after
        set_run_font(p.add_run(rule), size=9.6, color=INK)


def add_schemas(doc):
    doc.add_heading("17. 请求模型字典", level=1)
    add_para(doc, "字段后的问号表示可选；未单独说明时字符串由 Pydantic 做基础类型校验。")
    add_table(doc, ["模型", "字段", "规则/备注"], SCHEMAS, [2000, 4700, 2660], font_size=7.7)
    doc.add_heading("18. 核心响应 DTO", level=1)
    add_table(doc, ["DTO", "主要字段", "安全/业务说明"], DTO_ROWS, [1700, 5200, 2460], font_size=7.8)


def add_errors_and_data(doc, bullet_id):
    doc.add_heading("19. 错误码与前端处理", level=1)
    add_table(doc, ["HTTP", "建议业务码", "典型场景", "前端处理"], [
        ("400", "BAD_REQUEST", "请求语义错误", "保留输入并显示可操作提示"),
        ("401", "UNAUTHENTICATED", "登录失败、文件下载未登录", "登录页或单次刷新会话"),
        ("403", "FORBIDDEN", "非教师或角色不匹配", "停止重试并显示无权限"),
        ("404", "RESOURCE_NOT_FOUND", "对象不存在或不属于当前教师", "返回列表或 404 页面"),
        ("409", "INVALID_STATE/CONFLICT", "关闭任务、空图谱、已删除资料、唯一冲突", "刷新状态并阻止重复操作"),
        ("413", "PAYLOAD_TOO_LARGE", "文件超限", "提示 200 MB/20 MB 对应限制"),
        ("415", "UNSUPPORTED_MEDIA_TYPE", "文件类型不支持", "显示允许格式"),
        ("422", "VALIDATION_ERROR", "字段、枚举、日期和跨对象引用非法", "映射字段错误，保留 details"),
        ("429", "RATE_LIMITED", "登录/AI/上传限流（目标）", "按 Retry-After 重试"),
        ("500", "INTERNAL_ERROR", "未处理异常", "显示 trace_id，不展示堆栈"),
    ], [800, 2300, 2900, 3360], font_size=8.0, keep_together=True)
    add_callout(doc, "当前差异", "现有实现主要返回 FastAPI detail 文本，且中英文错误信息并存。企业发布前应统一 error.code、message、details 和 trace_id，同时保留 HTTP 状态语义。")
    doc.add_heading("20. 数据、事务与安全", level=1)
    add_table(doc, ["业务域", "核心表/存储", "事务与约束"], [
        ("账号与偏好", "users, teacher_credentials, teacher_preferences", "密码仅存 PBKDF2 salt/hash；偏好按 teacher_id 唯一"),
        ("课程班级", "courses, class_groups, enrollments", "课程 code、邀请码、班级学生组合唯一"),
        ("内容资料", "chapters, knowledge_points, materials, material_folders", "章节发布与资料可见性同事务；删除使用状态标记"),
        ("任务批改", "tasks, test_cases, submissions, evaluations, grades, feedback", "提交相关数据对象级授权；成绩发布与通知同事务"),
        ("知识图谱", "knowledge_points/material links；teacher_knowledge_graphs", "两套模型隔离；发布状态和版本需审计"),
        ("讨论通知", "course_discussions, replies, notifications", "发布/回复与通知建议 Outbox"),
        ("文件", "backend/uploads + materials.content_url", "路径净化、白名单、体积限制；生产应接对象存储"),
        ("审计", "audit_logs", "记录 actor/action/resource/detail；目标补 IP、trace_id、before/after"),
    ], [1600, 3600, 4160], font_size=8.0)
    doc.add_heading("20.1 安全上线要求", level=2)
    for text in [
        "删除生产环境中的 X-User-Id 信任链和 teacher-01 默认值。",
        "auth/accounts 仅在受控开发模式开放；生产登录不得暴露教师账号目录。",
        "登录使用 PBKDF2/Argon2/bcrypt，增加失败次数、账号锁定、验证码或 MFA 策略。",
        "所有对象访问统一执行 tenant、role、owner 三层校验，并对无权对象统一 404。",
        "上传文件做 MIME 与扩展名双校验、病毒扫描、随机对象键、下载审计和过期签名。",
        "日志不得记录密码、完整源码、令牌和敏感个人信息。",
    ]:
        add_list_item(doc, text, bullet_id)


def add_frontend_and_tests(doc, bullet_id):
    doc.add_heading("21. 教师端前端接入规范", level=1)
    add_table(doc, ["层级", "企业要求"], [
        ("HTTP Client", "统一 baseURL、超时、认证、trace_id、错误解包；页面组件不得直接 fetch。"),
        ("认证", "登录成功后建立服务端会话；API 方法不再接收 userId，也不保留默认教师。"),
        ("类型", "从 OpenAPI 生成 DTO，替换 src/api.ts 中的 unknown 和 any。"),
        ("缓存", "查询键包含 user/course/class/filter；归档、发布、删除后精准失效。"),
        ("写操作", "请求中禁用按钮；创建/发布携带 Idempotency-Key；409 后刷新对象。"),
        ("列表", "为任务、资料、提交、讨论、通知补分页和筛选，避免一次返回全部数据。"),
        ("文件", "上传显示进度、取消、413/415 处理；下载使用认证请求或短期签名 URL。"),
        ("状态 UI", "每个页面具备 loading、empty、403、404、409、422、500、离线和重试状态。"),
    ], [1600, 7760], font_size=8.7)
    doc.add_heading("21.1 页面—API 映射", level=2)
    add_table(doc, ["教师端页面", "主要接口 ID", "写后刷新资源"], [
        ("登录", "AUTH-01/02", "当前用户与全局缓存"),
        ("工作台", "APP-01/02, ANN-01/02, NOTICE-01", "bootstrap/dashboard/notifications"),
        ("课程列表/创建", "COURSE-01~04, DRAFT-01~03", "courses, draft"),
        ("班级管理", "CLASS-01~06", "classes, students, join-status"),
        ("课程内容", "CHAPTER-01~03, KP-01", "chapters, student-visible content"),
        ("教学资料", "MAT-01~09, FOLDER-01~08", "materials, folders, graph"),
        ("任务管理", "TASK-01~04", "tasks"),
        ("提交监控/批改", "SUB-01/02, GRADE-01/02, FEED-01", "submissions, notifications"),
        ("AI 审核", "REVIEW-01/02", "reviews, submission detail"),
        ("学情诊断", "ANA-01", "analytics"),
        ("课程知识图谱", "CKG-01~04", "course graph"),
        ("教师画布图谱", "TGRAPH-01~07", "teacher graphs"),
        ("课堂讨论", "DISC-01~04", "discussions, notifications"),
        ("设置", "PREF-01/02", "preferences"),
    ], [2300, 3600, 3460], font_size=8.2)

    doc.add_heading("22. 接口测试与验收", level=1)
    add_para(doc, "当前仓库后端测试作为现状证据；以下是企业级补充验收基线。")
    cases = [
        ("T-01", "教师登录成功/失败", "正确密码 200；错误密码 401；不泄露账号存在性", "认证"),
        ("T-02", "伪造 X-User-Id", "生产忽略/拒绝，不改变会话", "安全门禁"),
        ("T-03", "跨教师课程访问", "课程、班级、资料、任务统一 404/403", "权限"),
        ("T-04", "课程 code 冲突", "返回 409，事务不写入", "课程"),
        ("T-05", "课程归档与恢复", "状态正确，列表展示一致", "课程"),
        ("T-06", "课程删除引用检查", "明确可删/不可删策略，失败不部分删除", "课程"),
        ("T-07", "批量导入重复学生", "created/enrolled/skipped 统计准确", "班级"),
        ("T-08", "邀请码刷新", "新码可用、旧码失效、审计存在", "班级"),
        ("T-09", "章节发布/撤回", "资料 visibility 同事务切换", "内容"),
        ("T-10", "上传格式与体积", "非法 415，超限 413，无残留文件", "资料"),
        ("T-11", "资料软删除/恢复", "下载 404 后恢复 200，关联不丢失", "资料"),
        ("T-12", "文件路径穿越", "stored_name 净化后拒绝目录逃逸", "安全"),
        ("T-13", "文件夹删除/恢复", "资料状态与文件夹状态一致", "资料目录"),
        ("T-14", "AI 目录确认幂等", "重复确认不创建重复目录", "资料目录"),
        ("T-15", "任务无测试用例", "422，不创建任务", "任务"),
        ("T-16", "任务未来发布", "scheduled；调度到时才 published", "任务目标"),
        ("T-17", "已关闭任务发布", "409，不修改状态", "任务"),
        ("T-18", "保存/覆盖成绩", "同一 submission 只有一条 grade", "批改"),
        ("T-19", "发布成绩", "grade_published，通知与成绩同事务", "批改"),
        ("T-20", "草稿反馈", "student_visible=false", "反馈"),
        ("T-21", "修改 AI 审核无说明", "422；有说明时成功", "AI 审核"),
        ("T-22", "画布图谱空发布", "409", "知识图谱"),
        ("T-23", "图谱文件格式/体积", "415/413；无脏数据", "知识图谱"),
        ("T-24", "讨论重复发布", "不重复通知", "讨论"),
        ("T-25", "结束非 published 讨论", "409", "讨论"),
        ("T-26", "修改他人通知", "404，不能推断记录存在", "通知"),
        ("T-27", "并发保存图谱", "版本冲突返回 409，不静默覆盖", "并发目标"),
        ("T-28", "统一错误响应", "code/message/details/trace_id 完整", "契约目标"),
    ]
    add_table(doc, ["编号", "场景", "预期", "模块/状态"], cases, [850, 2300, 4750, 1460], font_size=7.7)
    doc.add_heading("22.1 自动化分层", level=2)
    for text in [
        "单元测试：Pydantic 约束、状态机、序列化、权限辅助函数。",
        "API 集成测试：真实事务、唯一约束、文件上传、对象级权限和错误契约。",
        "契约测试：OpenAPI 与前端生成类型差异为零。",
        "E2E：从教师登录到课程、班级、资料、任务、批改、图谱和讨论的主流程。",
        "安全测试：越权、路径穿越、上传绕过、暴力登录、敏感字段和限流。",
    ]:
        add_list_item(doc, text, bullet_id)


def add_delivery_and_aliases(doc, bullet_id):
    doc.add_heading("23. 实施优先级与发布门禁", level=1)
    add_table(doc, ["阶段", "交付内容", "完成标准"], [
        ("P0 认证", "Session/JWT、/me、移除 X-User-Id、关闭账号枚举", "生产无默认教师或可伪造身份"),
        ("P0 权限", "统一教师所有权依赖、越权自动化", "全部对象级权限测试通过"),
        ("P0 契约", "统一错误码、OpenAPI 响应模型、生成 DTO", "前后端契约无差异"),
        ("P1 数据", "分页、筛选、真实统计、定时发布调度", "无展示原型值，scheduled 闭环"),
        ("P1 可靠性", "幂等键、版本控制、Outbox、审计增强", "重复/并发/通知测试通过"),
        ("P2 基础设施", "对象存储、病毒扫描、限流、监控告警", "安全与可观测性门禁通过"),
    ], [1500, 4300, 3560], font_size=8.4)
    doc.add_heading("23.1 上线检查清单", level=2)
    for text in [
        "规范路径全部使用 /api/v1；新前端不调用兼容别名。",
        "登录建立真实会话，GET /me 可恢复教师和权限。",
        "OpenAPI 为所有 71 个接口声明请求、响应、错误和鉴权。",
        "全部列表接口明确分页策略，时间和枚举已冻结。",
        "关键写操作具备幂等、版本控制、事务、审计和回滚方案。",
        "后端回归、契约、E2E、安全和性能测试均达到门禁。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("附录 A. 知识图谱兼容别名", level=1)
    add_callout(doc, "兼容策略", "以下 7 个 /api/teacher 路径与 /api/v1/teacher/knowledge-graphs 对应接口共用处理函数，仅用于历史测试和兼容。新开发必须使用 /api/v1。")
    aliases = [
        ("GET", "/api/teacher/knowledge-graphs", "GET /api/v1/teacher/knowledge-graphs"),
        ("POST", "/api/teacher/knowledge-graphs", "POST /api/v1/teacher/knowledge-graphs"),
        ("GET", "/api/teacher/knowledge-graphs/{graph_id}", "GET /api/v1/teacher/knowledge-graphs/{graph_id}"),
        ("PUT", "/api/teacher/knowledge-graphs/{graph_id}", "PUT /api/v1/teacher/knowledge-graphs/{graph_id}"),
        ("DELETE", "/api/teacher/knowledge-graphs/{graph_id}", "DELETE /api/v1/teacher/knowledge-graphs/{graph_id}"),
        ("POST", "/api/teacher/knowledge-graphs/from-files", "POST /api/v1/teacher/knowledge-graphs/from-files"),
        ("POST", "/api/teacher/knowledge-graphs/{graph_id}/publish", "POST /api/v1/teacher/knowledge-graphs/{graph_id}/publish"),
    ]
    add_table(doc, ["方法", "兼容路径", "规范替代"], aliases, [1100, 4130, 4130], font_size=8.2)


def audit(doc):
    audit_document(doc)
    assert total_api_count() == 71
    assert len(doc.tables) >= 30
    assert any(p.style.name == "Heading 1" for p in doc.paragraphs)
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == 9360


def main():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        set_header_footer(section)
    bullet_id = add_numbering_definition(doc, bullet=True)

    doc.core_properties.title = "CodeTrack 教师端后端 API 接口文档（企业版）"
    doc.core_properties.subject = "教师端全部后端接口、契约、权限、数据、测试与发布基线"
    doc.core_properties.author = "CodeTrack 项目开发组"
    doc.core_properties.keywords = "CodeTrack, 教师端, API, FastAPI, 接口文档, 企业开发"
    doc.core_properties.comments = f"{DOC_ID} {DOC_VERSION}"
    doc.core_properties.created = datetime(2026, 8, 16)
    doc.core_properties.modified = datetime(2026, 8, 16)

    add_cover(doc)
    add_front_matter(doc)
    add_scope(doc, bullet_id)
    add_baseline(doc)
    add_protocol(doc, bullet_id)
    add_inventory(doc)
    for number, module in enumerate(MODULES, 5):
        add_module(doc, number, module)
    add_schemas(doc)
    add_errors_and_data(doc, bullet_id)
    add_frontend_and_tests(doc, bullet_id)
    add_delivery_and_aliases(doc, bullet_id)
    audit(doc)
    doc.save(OUTPUT_PATH)
    print(json.dumps({
        "output": str(OUTPUT_PATH),
        "bytes": OUTPUT_PATH.stat().st_size,
        "canonical_endpoints": total_api_count(),
        "legacy_aliases": 7,
        "tables": len(doc.tables),
        "paragraphs": len(doc.paragraphs),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
